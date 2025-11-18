from flask_login import login_required, current_user, LoginManager, login_user, logout_user
from flask import Blueprint, render_template, request, redirect, url_for, flash, jsonify, session, Flask, send_file, current_app, abort, render_template_string
import pdfkit
from datetime import datetime, timezone, timedelta, date, time
from sqlalchemy.orm import joinedload
from sqlalchemy import or_, func
import logging
import traceback
import re
import json
import unicodedata
import chardet
import os
import random
import subprocess
import base64
from tempfile import NamedTemporaryFile
from docxtpl import DocxTemplate
from docx import Document
from io import BytesIO
from app import db
from app.models import Funcionario,Chamado,PrescricaoEnfermagemTemplate,Leito,PrescricaoEmergencia,MedicacoesPadrao,AdmissaoEnfermagem, AtendimentosGestante ,ListaInternacao, ListaObservacao, Paciente, Atendimento, InternacaoSae, Internacao, EvolucaoAtendimentoClinica, PrescricaoClinica, EvolucaoEnfermagem, PrescricaoEnfermagem, InternacaoEspecial, Aprazamento, ReceituarioClinica, AtestadoClinica, PacienteRN, now_brasilia, FichaReferencia, EvolucaoFisioterapia,EvolucaoAssistenteSocial, EvolucaoNutricao, AtendimentosGestante, FluxoDisp, FluxoPaciente
from app.timezone_helper import formatar_datetime_br_completo, formatar_datetime_br, converter_para_brasilia, formatar_data_br
from zoneinfo import ZoneInfo

# Cria o Blueprint principal
bp = Blueprint('main', __name__)
# Utilitário para normalizar strings removendo acentos/cedilhas
def normalize_status(value: str) -> str:
    try:
        txt = (value or '').strip().upper()
        # Normaliza para NFD e remove marcas combinantes (acentos)
        decomposed = unicodedata.normalize('NFD', txt)
        return ''.join(ch for ch in decomposed if unicodedata.category(ch)[0] != 'M')
    except Exception:
        return (value or '').strip().upper()


# Cria o Blueprint para internações especiais
internacoes_especiais_bp = Blueprint('internacoes_especiais', __name__)

# Login required decorator personalizado
def login_required(f):
    from functools import wraps
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            flash('Por favor, faça login para acessar esta página', 'warning')
            return redirect(url_for('main.index'))
        return f(*args, **kwargs)
    return decorated_function

# Helper para obter o usuário atual
def get_current_user():
    if 'user_id' not in session:
        return None
    return Funcionario.query.get(session['user_id'])

# ====== API: Prescrições de Emergência (porta) ======
@bp.route('/api/prescricoes-emergencia/<string:atendimento_id>', methods=['GET'])
@login_required
def listar_prescricoes_emergencia(atendimento_id):
    try:
        atendimento = Atendimento.query.get(atendimento_id)
        if not atendimento:
            return jsonify({'success': False, 'message': 'Atendimento não encontrado'}), 404

        prescs = (PrescricaoEmergencia.query
                  .filter_by(atendimento_id=atendimento_id)
                  .order_by(PrescricaoEmergencia.horario_prescricao.desc())
                  .all())

        def serialize(p):
            return {
                'id': p.id,
                'atendimento_id': p.atendimento_id,
                'medico_id': p.medico_id,
                'enfermeiro_id': p.enfermeiro_id,
                'texto_dieta': p.texto_dieta,
                'texto_procedimento_medico': p.texto_procedimento_medico,
                'texto_procedimento_multi': p.texto_procedimento_multi,
                'horario_prescricao': p.horario_prescricao.isoformat() if p.horario_prescricao else None,
                'medicamentos': p.medicamentos or []
            }

        return jsonify({'success': True, 'prescricoes': [serialize(p) for p in prescs]})
    except Exception as e:
        logging.error(f"Erro ao listar prescrições de emergência: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro interno'}), 500


@bp.route('/api/prescricoes-emergencia', methods=['POST'])
@login_required
def criar_prescricao_emergencia():
    try:
        dados = request.get_json(force=True) or {}
        atendimento_id = (dados.get('atendimento_id') or '').strip()
        medico_id = dados.get('medico_id') or session.get('user_id')
        if not atendimento_id:
            return jsonify({'success': False, 'message': 'atendimento_id é obrigatório'}), 400
        if not medico_id:
            return jsonify({'success': False, 'message': 'medico_id é obrigatório'}), 400

        atendimento = Atendimento.query.get(atendimento_id)
        if not atendimento:
            return jsonify({'success': False, 'message': 'Atendimento não encontrado'}), 404

        # Campos opcionais
        texto_dieta = (dados.get('texto_dieta') or None)
        texto_proc_med = (dados.get('texto_procedimento_medico') or None)
        texto_proc_multi = (dados.get('texto_procedimento_multi') or None)
        medicamentos = dados.get('medicamentos') or []
        if not isinstance(medicamentos, list):
            return jsonify({'success': False, 'message': 'medicamentos deve ser uma lista'}), 400

        nova = PrescricaoEmergencia(
            atendimento_id=atendimento_id,
            medico_id=medico_id,
            enfermeiro_id=dados.get('enfermeiro_id'),
            texto_dieta=texto_dieta,
            texto_procedimento_medico=texto_proc_med,
            texto_procedimento_multi=texto_proc_multi,
            horario_prescricao=converter_para_brasilia(datetime.now(ZoneInfo('America/Sao_Paulo'))),
            medicamentos=medicamentos
        )
        db.session.add(nova)
        db.session.commit()

        return jsonify({'success': True, 'id': nova.id}), 201
    except Exception as e:
        logging.error(f"Erro ao criar prescrição de emergência: {str(e)}")
        logging.error(traceback.format_exc())
        db.session.rollback()
        return jsonify({'success': False, 'message': 'Erro interno ao salvar'}), 500


@bp.route('/api/imprimir-prescricao-emergencia-html/<int:prescricao_id>', methods=['GET'])
@login_required
def imprimir_prescricao_emergencia_html(prescricao_id):
    try:
        p = PrescricaoEmergencia.query.get(prescricao_id)
        if not p:
            abort(404, 'Prescrição não encontrada.')

        atendimento = Atendimento.query.get(p.atendimento_id)
        if not atendimento:
            abort(404, 'Atendimento não encontrado.')

        paciente = atendimento.paciente
        medico = Funcionario.query.get(p.medico_id) if p.medico_id else None

        # Renderização simples em HTML para impressão sem depender de internação
        html = render_template_string(
            '''<!DOCTYPE html>
            <html lang="pt-br">
            <head>
              <meta charset="utf-8" />
              <title>Prescrição de Emergência #{{ p.id }}</title>
              <style>
                body { font-family: Arial, sans-serif; margin: 24px; }
                h1 { margin: 0 0 8px 0; }
                .muted { color: #666; font-size: 12px; }
                .box { border: 1px solid #ddd; border-radius: 8px; padding: 12px; margin-bottom: 12px; }
                ul { margin: 6px 0 0 18px; }
              </style>
            </head>
            <body>
              <h1>Prescrição de Emergência</h1>
              <div class="muted">ID: {{ p.id }} • Data: {{ p.horario_prescricao.strftime('%d/%m/%Y %H:%M') if p.horario_prescricao else '' }}</div>
              <div class="box">
                <strong>Paciente:</strong> {{ paciente.nome or '-' }}<br/>
                <strong>Atendimento:</strong> {{ atendimento.id }}<br/>
                <strong>Médico:</strong> {{ (medico.nome if medico else '-') }}
              </div>
              {% if p.texto_dieta %}
              <div class="box">
                <strong>Dieta</strong>
                <div>{{ p.texto_dieta|safe }}</div>
              </div>
              {% endif %}
              {% if p.texto_procedimento_medico %}
              <div class="box">
                <strong>Procedimentos Médicos</strong>
                <div>{{ p.texto_procedimento_medico|safe }}</div>
              </div>
              {% endif %}
              {% if p.texto_procedimento_multi %}
              <div class="box">
                <strong>Procedimentos Multiprofissionais</strong>
                <div>{{ p.texto_procedimento_multi|safe }}</div>
              </div>
              {% endif %}
              {% if p.medicamentos %}
              <div class="box">
                <strong>Medicamentos</strong>
                <ul>
                  {% for m in p.medicamentos %}
                    <li>{{ m.get('nome_medicamento','') }}{% if m.get('descricao_uso') %} - {{ m.get('descricao_uso') }}{% endif %}</li>
                  {% endfor %}
                </ul>
              </div>
              {% endif %}
            </body>
            </html>''',
            p=p, atendimento=atendimento, paciente=paciente, medico=medico
        )
        return html
    except Exception as e:
        logging.error(f"Erro ao imprimir prescrição de emergência: {str(e)}")
        logging.error(traceback.format_exc())
        abort(500, f'Erro interno: {str(e)}')

# Rota de teste simples para diagnóstico
@bp.route('/api/status')
def api_status():
    try:
        # Tentar fazer uma consulta simples ao banco de dados
        count = Funcionario.query.count()
        return jsonify({
            'status': 'online',
            'database': 'connected',
            'funcionarios_count': count,
            'timestamp': datetime.now(ZoneInfo("America/Sao_Paulo")).isoformat()
        })
    except Exception as e:
        logging.error(f"Erro na rota de status: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({
            'status': 'error',
            'error': str(e),
            'timestamp': datetime.now(ZoneInfo("America/Sao_Paulo")).isoformat()
        }), 500
        
# API para adicionar paciente em observação

@bp.route('/api/observacao-paciente', methods=['POST'])
@login_required
def observacao_paciente():
            """
            Registra um novo paciente e cria uma observação associada.
            """
            try:
                current_user = get_current_user()
                if current_user.cargo.lower() != 'medico':
                    return jsonify({
                        'success': False,
                        'message': 'Apenas médicos podem adicionar pacientes em observação'
                    }), 403
        
                dados = request.get_json()

                # Campos obrigatórios de observação (sempre exigidos)
                campos_observacao_obrigatorios = ['hda', 'diagnostico_inicial']
                for campo in campos_observacao_obrigatorios:
                    if campo not in dados or not dados[campo]:
                        return jsonify({
                            'success': False,
                            'message': f'Campo obrigatório não informado: {campo}'
                        }), 400

                # Tentar usar atendimento existente, se fornecido
                atendimento_id = (dados.get('atendimento_id') or '').strip()
                atendimento = Atendimento.query.get(atendimento_id) if atendimento_id else None
                agora = datetime.now(ZoneInfo("America/Sao_Paulo"))

                if atendimento is not None:
                    logging.info(f"📋 Usando atendimento existente: {atendimento_id}")
                    paciente = Paciente.query.get(atendimento.paciente_id)
                    atendimento.status = 'Em Observação'
                    atendimento.horario_observacao = agora
                    db.session.add(atendimento)
                else:
                    # Não há atendimento existente: criar paciente e atendimento novos
                    campos_paciente_obrigatorios = ['nome', 'cpf', 'data_nascimento', 'sexo']
                    for campo in campos_paciente_obrigatorios:
                        if campo not in dados or not dados[campo]:
                            return jsonify({
                                'success': False,
                                'message': f'Campo obrigatório não informado: {campo}'
                            }), 400

                    # Normalizar CPF (remover pontos, traços e quaisquer não dígitos)
                    cpf_raw = (dados.get('cpf') or '').strip()
                    cpf_limpo = re.sub(r'\D', '', cpf_raw)
                    dados['cpf'] = cpf_limpo

                    # Verificar se o paciente já existe com o CPF normalizado
                    paciente_existente = Paciente.query.filter_by(cpf=cpf_limpo).first()

                    if paciente_existente:
                        paciente = paciente_existente
                        # Atualizar alergias do paciente existente se fornecidas
                        if dados.get('alergias'):
                            paciente.alergias = dados.get('alergias')
                    else:
                        try:
                            data_nascimento = datetime.strptime(dados['data_nascimento'], '%Y-%m-%d').date()
                        except ValueError:
                            return jsonify({
                                'success': False,
                                'message': 'Formato de data de nascimento inválido. Use YYYY-MM-DD.'
                            }), 400

                        cartao_sus = dados.get('cartao_sus')
                        if dados.get('sem_cartao_sus', False) or not cartao_sus:
                            cartao_sus = None

                        paciente = Paciente(
                            nome=dados['nome'],
                            filiacao=dados.get('filiacao', 'Não informado'),
                            cpf=dados['cpf'],
                            data_nascimento=data_nascimento,
                            sexo=dados['sexo'],
                            telefone=dados.get('telefone', 'Não informado'),
                            endereco=dados.get('endereco', 'Não informado'),
                            municipio=dados.get('municipio', 'Não informado'),
                            bairro=dados.get('bairro', 'Não informado'),
                            cartao_sus=cartao_sus,
                            nome_social=dados.get('nome_social', ''),
                            cor=dados.get('cor', 'Não informada'),
                            alergias=dados.get('alergias', ''),
                            identificado=True
                        )
                        db.session.add(paciente)
                        db.session.flush()

                    agora = datetime.now(ZoneInfo("America/Sao_Paulo"))
                    atendimento_id = (dados.get('atendimento_id') or '').strip()
                    if atendimento_id:
                        if len(atendimento_id) > 8:
                            return jsonify({
                                'success': False,
                                'message': 'ID de atendimento excede o limite máximo de 8 dígitos.'
                            }), 400
                        if Atendimento.query.get(atendimento_id):
                            return jsonify({
                                'success': False,
                                'message': 'ID de atendimento já existe. Tente novamente.'
                            }), 400
                    else:
                        prefixo_data = agora.strftime('%y%m%d')
                        numero_unico = str(paciente.id)[-2:].zfill(2)
                        atendimento_id = f"{prefixo_data}{numero_unico}"

                    # Atualizar alergias do paciente se fornecidas
                    if dados.get('alergias'):
                        paciente.alergias = dados.get('alergias')

                    atendimento = Atendimento(
                        id=atendimento_id,
                        paciente_id=paciente.id,
                        funcionario_id=current_user.id,
                        medico_id=current_user.id,
                        data_atendimento=date.today(),
                        hora_atendimento=time(agora.hour, agora.minute, agora.second),
                        status='Em Observação',
                        horario_observacao=agora
                    )
                    db.session.add(atendimento)
                    db.session.flush()
        
                # CRÍTICO: SEMPRE garantir que existe um registro de Internacao
                # Verificar se já existe uma internação para este atendimento
                internacao = Internacao.query.filter_by(atendimento_id=atendimento_id).first()
                
                if internacao:
                    # Já existe - atualizar dados se necessário
                    logging.info(f"✅ Internação já existe (ID: {internacao.id}) - Atualizando dados")
                    
                    # Atualizar campos importantes se foram fornecidos
                    if dados.get('hda'):
                        internacao.hda = dados.get('hda')
                    if dados.get('diagnostico_inicial'):
                        internacao.diagnostico_inicial = dados.get('diagnostico_inicial')
                    if dados.get('exame_fisico'):
                        internacao.folha_anamnese = dados.get('exame_fisico')
                    if dados.get('cid_principal'):
                        internacao.cid_principal = dados.get('cid_principal')
                    
                    # Garantir que está marcado como observação
                    if not internacao.leito or internacao.leito == '':
                        internacao.leito = 'Observação'
                    if not internacao.carater_internacao:
                        internacao.carater_internacao = 'Observação'
                    
                    db.session.add(internacao)
                else:
                    # Não existe - CRIAR OBRIGATORIAMENTE
                    logging.info(f"🆕 Criando novo registro de Internacao para atendimento: {atendimento_id}")
                    
                    internacao = Internacao(
                        atendimento_id=atendimento_id,
                        paciente_id=paciente.id,
                        medico_id=current_user.id,
                        enfermeiro_id=None,
                        hda=dados.get('hda', ''),
                        diagnostico_inicial=dados.get('diagnostico_inicial', ''),
                        folha_anamnese=dados.get('exame_fisico', ''),
                        cid_principal=dados.get('cid_principal', ''),
                        cid_10_secundario=dados.get('cid_secundario', ''),
                        data_internacao=agora,
                        leito='Observação',
                        carater_internacao='Observação',
                        dieta=None  # Dieta NULL durante a internação/observação
                    )
                    db.session.add(internacao)
                    db.session.flush()
                    
                    logging.info(f"✅ Internacao criada com sucesso (ID: {internacao.id})")

                # Criar primeira evolução na tabela EvolucaoAtendimentoClinica (se fornecida)
                if dados.get('primeira_evolucao'):
                    try:
                        primeira_evolucao = EvolucaoAtendimentoClinica(
                            atendimentos_clinica_id=internacao.id,
                            funcionario_id=current_user.id,
                            data_evolucao=agora,
                            hda=dados.get('hda', ''),
                            evolucao=dados.get('primeira_evolucao', ''),
                            conduta='Primeira evolução médica - Observação'
                        )
                        db.session.add(primeira_evolucao)
                        db.session.flush()
                        logging.info(f"✅ Primeira evolução criada")
                    except Exception as e:
                        logging.error(f"⚠️ Erro ao criar primeira evolução (não crítico): {str(e)}")
        
                # CRÍTICO: SEMPRE garantir que existe um registro em ListaObservacao
                observacao = ListaObservacao.query.filter_by(id_atendimento=atendimento_id).first()
                
                if observacao:
                    # Já existe - atualizar se necessário
                    logging.info(f"✅ Registro em ListaObservacao já existe - Atualizando")
                    observacao.medico_entrada = current_user.nome
                    observacao.data_entrada = agora
                    # Resetar campos de saída caso esteja reabrindo observação
                    observacao.medico_conduta = None
                    observacao.data_saida = None
                    observacao.conduta_final = None
                    db.session.add(observacao)
                else:
                    # Não existe - CRIAR OBRIGATORIAMENTE
                    logging.info(f"🆕 Criando novo registro em ListaObservacao")
                    observacao = ListaObservacao(
                        id_atendimento=atendimento_id,
                        id_paciente=paciente.id,
                        medico_entrada=current_user.nome,
                        data_entrada=agora,
                        medico_conduta=None,
                        data_saida=None,
                        conduta_final=None
                    )
                    db.session.add(observacao)
                    db.session.flush()
                    logging.info(f"✅ ListaObservacao criada com sucesso")
        
                # Commit final
                db.session.commit()
                
                logging.info(f"💾 ✅ Observação registrada com sucesso - Atendimento: {atendimento_id}, Internacao ID: {internacao.id}")
        
                return jsonify({
                    'success': True,
                    'message': 'Paciente adicionado à observação com sucesso',
                    'paciente_id': paciente.id,
                    'internacao_id': internacao.id,
                    'observacao_id': observacao.id,
                    'atendimento_id': atendimento_id
                }), 201
        
            except Exception as e:
                db.session.rollback()
                logging.error(f'❌ ERRO ao adicionar paciente em observação: {str(e)}')
                logging.error(traceback.format_exc())
                
                # Tentar identificar o tipo de erro para dar uma mensagem mais específica
                erro_msg = str(e)
                if 'IntegrityError' in str(type(e)):
                    erro_msg = 'Erro de integridade no banco de dados. Verifique se os dados não estão duplicados.'
                elif 'OperationalError' in str(type(e)):
                    erro_msg = 'Erro de conexão com o banco de dados. Tente novamente em alguns instantes.'
                elif 'DataError' in str(type(e)):
                    erro_msg = 'Dados inválidos fornecidos. Verifique os campos e tente novamente.'
                else:
                    erro_msg = f'Erro ao processar observação: {str(e)}'
                
                return jsonify({
                    'success': False,
                    'message': erro_msg,
                    'details': 'Verifique os logs do servidor para mais informações'
                }), 500


# GET /api/medico/nome
@bp.route('/api/medico/nome', methods=['GET'])
@login_required
def get_nome_medico():
    """
    Retorna o nome do médico logado.
    """
    try:
        # Usar o sistema de autenticação personalizado
        current_user = get_current_user()
        if not current_user:
            return jsonify({'success': False, 'message': 'Usuário não autenticado.'}), 401
        
        if current_user.cargo.lower() != 'medico':
            return jsonify({'success': False, 'message': 'Usuário não é médico.'}), 403

        return jsonify({
            'success': True,
            'nome': current_user.nome
        })

    except Exception as e:
        logging.error(f"Erro ao buscar nome do médico: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': 'Erro ao buscar nome do médico.',
            'error': str(e)
        }), 500


# ===================== FICHA GERAL DO PACIENTE =====================

@bp.route('/ficha_geral_paciente')
@login_required
def ficha_geral_paciente():
    """
    Página visual da ficha geral do paciente.
    Recebe query param `paciente_id` no frontend.
    """
    try:
        # Apenas valida sessão; a lógica de carregamento é feita via JS
        return render_template('ficha_geral_paciente.html')
    except Exception:
        logging.error('Erro ao renderizar ficha_geral_paciente:\n' + traceback.format_exc())
        abort(500)


@bp.route('/api/enfermeiro/ficha_geral', methods=['GET'])
@login_required
def api_ficha_geral_paciente():
    """
    Retorna dados consolidados para a Ficha do Paciente:
    - Dados cadastrais do paciente
    - Último atendimento (triagem/conduta) se houver
    - (Opcional) dados do enfermeiro que realizou a triagem
    """
    try:
        paciente_id = request.args.get('paciente_id', type=int)
        if not paciente_id:
            return jsonify({'success': False, 'message': 'paciente_id é obrigatório'}), 400

        paciente = Paciente.query.get(paciente_id)
        if not paciente:
            return jsonify({'success': False, 'message': 'Paciente não encontrado'}), 404

        # Buscar o último atendimento do paciente (por data/hora)
        ultimo_atendimento = (
            Atendimento.query
            .filter_by(paciente_id=paciente.id)
            .order_by(Atendimento.data_atendimento.desc(), Atendimento.hora_atendimento.desc())
            .first()
        )

        triagem = None
        enfermeiro_info = None
        if ultimo_atendimento:
            triagem = {
                'id': ultimo_atendimento.id,
                'data_atendimento': ultimo_atendimento.data_atendimento.strftime('%Y-%m-%d') if ultimo_atendimento.data_atendimento else None,
                'hora_atendimento': ultimo_atendimento.hora_atendimento.strftime('%H:%M') if ultimo_atendimento.hora_atendimento else None,
                'pressao': ultimo_atendimento.pressao,
                'pulso': ultimo_atendimento.pulso,
                'sp02': ultimo_atendimento.sp02,
                'temp': ultimo_atendimento.temp,
                'peso': ultimo_atendimento.peso,
                'altura': ultimo_atendimento.altura,
                'fr': ultimo_atendimento.fr,
                'dx': ultimo_atendimento.dx,
                'triagem': ultimo_atendimento.triagem,
                'alergias': ultimo_atendimento.paciente.alergias if ultimo_atendimento.paciente else '',
                'classificacao_risco': ultimo_atendimento.classificacao_risco,
                'anamnese_exame_fisico': ultimo_atendimento.anamnese_exame_fisico,
                'observacao': ultimo_atendimento.observacao,
                'reavaliacao': ultimo_atendimento.reavaliacao,
                'prescricao_medica': ultimo_atendimento.prescricao_medica,
                'conduta_final': ultimo_atendimento.conduta_final,
                'status': ultimo_atendimento.status,
            }

            if ultimo_atendimento.enfermeiro_id:
                enf = Funcionario.query.get(ultimo_atendimento.enfermeiro_id)
                if enf:
                    enfermeiro_info = {
                        'id': enf.id,
                        'nome': enf.nome,
                        'numero_profissional': enf.numero_profissional,
                    }

        paciente_dict = {
            'id': paciente.id,
            'nome': paciente.nome,
            'nome_social': paciente.nome_social,
            'filiacao': paciente.filiacao,
            'data_nascimento': paciente.data_nascimento.strftime('%Y-%m-%d') if paciente.data_nascimento else None,
            'sexo': paciente.sexo,
            'cpf': paciente.cpf,
            'cartao_sus': paciente.cartao_sus,
            'telefone': paciente.telefone,
            'endereco': paciente.endereco,
            'municipio': paciente.municipio,
            'bairro': paciente.bairro,
            'cor': paciente.cor,
            'identificado': paciente.identificado,
        }

        return jsonify({
            'success': True,
            'paciente': paciente_dict,
            'triagem': triagem,
            'enfermeiro': enfermeiro_info,
        })
    except Exception:
        logging.error('Erro na API ficha_geral:\n' + traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro interno do servidor'}), 500


@bp.route('/api/pacientes/<int:paciente_id>/atendimentos', methods=['GET'])
@login_required
def listar_atendimentos_por_paciente(paciente_id):
    """
    Lista atendimentos de um paciente, ordenados por data/hora desc.
    """
    try:
        paciente = Paciente.query.get(paciente_id)
        if not paciente:
            return jsonify({'success': False, 'message': 'Paciente não encontrado'}), 404

        atendimentos = (
            Atendimento.query
            .filter_by(paciente_id=paciente_id)
            .order_by(Atendimento.data_atendimento.desc(), Atendimento.hora_atendimento.desc())
            .all()
        )

        resultado = []
        for a in atendimentos:
            resultado.append({
                'id': a.id,
                'data_atendimento': a.data_atendimento.strftime('%Y-%m-%d') if a.data_atendimento else None,
                'hora_atendimento': a.hora_atendimento.strftime('%H:%M') if a.hora_atendimento else None,
                'status': a.status,
                'classificacao_risco': a.classificacao_risco,
            })

        return jsonify({'success': True, 'atendimentos': resultado})
    except Exception:
        logging.error('Erro ao listar atendimentos do paciente:\n' + traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro interno do servidor'}), 500


@bp.route('/api/usuario/cargo', methods=['GET'])
@login_required
def api_usuario_cargo():
    """Retorna o cargo do usuário logado para decisões de navegação no frontend."""
    try:
        current = get_current_user()
        if not current:
            return jsonify({'success': False, 'message': 'Não autenticado'}), 401
        return jsonify({'success': True, 'cargo': (current.cargo or '').lower()})
    except Exception:
        logging.error('Erro ao obter cargo do usuário:\n' + traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro interno do servidor'}), 500

# LISTA PARA MÉDICO: ATENDIMENTOS AGUARDANDO MÉDICO
@bp.route('/api/medico/atendimentos/aguardando', methods=['GET'])
@login_required
def api_medico_atendimentos_aguardando():
    """
    Retorna lista de atendimentos com status "Aguardando Medico" contendo
    nome do paciente, idade, classificação de risco e texto da triagem.
    """
    try:
        current_user = get_current_user()
        if not current_user:
            return jsonify({'success': False, 'message': 'Usuário não autenticado.'}), 401

        if current_user.cargo.lower() not in ['medico', 'multi', 'admin']:
            return jsonify({'success': False, 'message': 'Acesso não autorizado'}), 403

        atendimentos = Atendimento.query.filter(
            db.or_(
                Atendimento.status.ilike('%aguardando%medico%'),
                Atendimento.status.ilike('%aguardando%médico%')
            )
        ).order_by(Atendimento.horario_triagem.desc().nullslast(), Atendimento.data_atendimento.desc()).all()

        def calcular_idade(data_nascimento):
            if not data_nascimento:
                return None
            hoje = date.today()
            anos = hoje.year - data_nascimento.year - ((hoje.month, hoje.day) < (data_nascimento.month, data_nascimento.day))
            return anos

        resultado = []
        for a in atendimentos:
            paciente = Paciente.query.get(a.paciente_id)
            if not paciente:
                continue

            # Verifica se há chamado ativo para este atendimento
            chamado_ativo = Chamado.query.filter(
                Chamado.id_atendimento == a.id,
                Chamado.status == 'ativo'
            ).order_by(Chamado.data.desc(), Chamado.hora.desc()).first()

            chamado_info = None
            if chamado_ativo:
                medico = Funcionario.query.get(chamado_ativo.id_medico) if chamado_ativo.id_medico else None
                chamado_info = {
                    'medico_nome': medico.nome if medico else None,
                    'local': chamado_ativo.local,
                    'hora': chamado_ativo.hora.strftime('%H:%M') if chamado_ativo.hora else None
                }

            resultado.append({
                'atendimento_id': a.id,
                'paciente_id': paciente.id,
                'nome': paciente.nome,
                'data_nascimento': paciente.data_nascimento.strftime('%Y-%m-%d') if paciente.data_nascimento else None,
                'idade': calcular_idade(paciente.data_nascimento) if paciente.data_nascimento else None,
                'classificacao_risco': a.classificacao_risco,
                'triagem': a.triagem,
                'horario_triagem': a.horario_triagem.strftime('%Y-%m-%d %H:%M:%S') if a.horario_triagem else None,
                'prioridade': paciente.prioridade if hasattr(paciente, 'prioridade') else False,
                'desc_prioridade': paciente.desc_prioridade if hasattr(paciente, 'desc_prioridade') else None,
                'chamado_ativo': chamado_info
            })

        return jsonify({'success': True, 'atendimentos': resultado, 'total': len(resultado)})

    except Exception as e:
        logging.error(f"Erro ao listar atendimentos aguardando médico: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro interno do servidor'}), 500


# Médico assume um atendimento (registro explícito de início de consulta)
@bp.route('/api/atendimento/<string:atendimento_id>/assumir', methods=['POST'])
@login_required
def api_medico_assumir_atendimento(atendimento_id):
    try:
        current_user = get_current_user()
        if current_user.cargo.strip().lower() != 'medico':
            return jsonify({'success': False, 'message': 'Acesso não autorizado'}), 403

        atendimento = Atendimento.query.get(atendimento_id)
        if not atendimento:
            return jsonify({'success': False, 'message': 'Atendimento não encontrado'}), 404

        # Horário do servidor menos 3 horas
        agora = datetime.utcnow() - timedelta(hours=3)

        # Se ainda não havia horário de consulta, registra agora
        if not atendimento.horario_consulta_medica:
            atendimento.horario_consulta_medica = agora

        # Não alterar o status aqui; permanecer em "Aguardando Médico" até definir conduta

        # Garante médico responsável
        atendimento.medico_id = atendimento.medico_id or current_user.id

        # Registrar no FluxoPaciente
        try:
            fluxo = FluxoPaciente(
                id_atendimento=atendimento.id,
                id_medico=current_user.id,
                id_enfermeiro=None,
                nome_paciente=atendimento.paciente.nome if atendimento.paciente else '',
                mudanca_status='INICIO CONSULTA',
                mudanca_hora=agora
            )
            db.session.add(fluxo)
        except Exception as e:
            logging.error(f"Falha ao registrar INICIO CONSULTA no FluxoPaciente: {str(e)}")

        db.session.add(atendimento)
        db.session.commit()

        return jsonify({'success': True, 'message': 'Atendimento assumido pelo médico.'})

    except Exception as e:
        db.session.rollback()
        logging.error(f"Erro ao assumir atendimento: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro interno do servidor'}), 500


@bp.route('/api/medico/atendimentos/alta-medicacao', methods=['GET'])
@login_required
def api_medico_atendimentos_alta_medicacao():
    """
    Retorna lista de atendimentos com status relacionado a "Alta após medicação".
    """
    try:
        current_user = get_current_user()
        if not current_user:
            return jsonify({'success': False, 'message': 'Usuário não autenticado.'}), 401

        if current_user.cargo.lower() not in ['medico', 'multi', 'admin']:
            return jsonify({'success': False, 'message': 'Acesso não autorizado'}), 403

        # Após normalização, o status é salvo sem acentos: 'ALTA APOS MEDICACAO'
        atendimentos = Atendimento.query.filter(
            Atendimento.status.ilike('%ALTA%APOS%MEDICACAO%')
        ).order_by(Atendimento.horario_medicacao.desc().nullslast(), Atendimento.data_atendimento.desc()).all()

        def calcular_idade(data_nascimento):
            if not data_nascimento:
                return None
            hoje = date.today()
            anos = hoje.year - data_nascimento.year - ((hoje.month, hoje.day) < (data_nascimento.month, data_nascimento.day))
            return anos

        resultado = []
        for a in atendimentos:
            paciente = Paciente.query.get(a.paciente_id)
            if not paciente:
                continue
            resultado.append({
                'atendimento_id': a.id,
                'paciente_id': paciente.id,
                'nome': paciente.nome,
                'data_nascimento': paciente.data_nascimento.strftime('%Y-%m-%d') if paciente.data_nascimento else None,
                'idade': calcular_idade(paciente.data_nascimento) if paciente.data_nascimento else None,
                'classificacao_risco': a.classificacao_risco,
                'triagem': a.triagem,
                'horario_triagem': a.horario_triagem.strftime('%Y-%m-%d %H:%M:%S') if a.horario_triagem else None,
            })

        return jsonify({'success': True, 'atendimentos': resultado, 'total': len(resultado)})

    except Exception as e:
        logging.error(f"Erro ao listar atendimentos de alta após medicação: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro interno do servidor'}), 500


@bp.route('/api/medico/atendimentos/reavaliacao-medicacao', methods=['GET'])
@login_required
def api_medico_atendimentos_reavaliacao_medicacao():
    """
    Retorna lista de atendimentos marcados para "Reavaliação após medicação".
    Inclui informações sobre ciclos de reavaliação (até 3 ciclos de 2 horas cada).
    """
    try:
        current_user = get_current_user()
        if not current_user:
            return jsonify({'success': False, 'message': 'Usuário não autenticado.'}), 401

        if current_user.cargo.lower() not in ['medico', 'multi', 'admin']:
            return jsonify({'success': False, 'message': 'Acesso não autorizado'}), 403

        # Após normalização, o status para reavaliação é salvo como 'REAVALIACAO'
        atendimentos = Atendimento.query.filter(
            Atendimento.status.ilike('%REAVALIACAO%')
        ).order_by(Atendimento.horario_medicacao.desc().nullslast(), Atendimento.data_atendimento.desc()).all()

        def calcular_idade(data_nascimento):
            if not data_nascimento:
                return None
            hoje = date.today()
            anos = hoje.year - data_nascimento.year - ((hoje.month, hoje.day) < (data_nascimento.month, data_nascimento.day))
            return anos

        def contar_ciclos_reavaliacao(medicacao_utilizada_text):
            """Conta quantas vezes o valor '1' aparece em medicacao_utilizada"""
            if not medicacao_utilizada_text:
                return 0
            try:
                # Tenta parsear como JSON primeiro
                import json
                medicacoes = json.loads(medicacao_utilizada_text)
                if isinstance(medicacoes, list):
                    # Conta quantos registros existem (cada dispensação é um ciclo)
                    return len(medicacoes)
            except:
                pass
            
            # Se não for JSON válido, conta ocorrências de "1"
            return medicacao_utilizada_text.count('1')

        resultado = []
        agora = now_brasilia()
        
        for a in atendimentos:
            paciente = Paciente.query.get(a.paciente_id)
            if not paciente:
                continue
            
            # Contar ciclos baseado em medicacao_utilizada
            ciclo_atual = contar_ciclos_reavaliacao(a.medicacao_utilizada) + 1  # +1 porque estamos no próximo ciclo
            
            # Calcular tempo desde a consulta médica ou última medicação
            horario_referencia = a.horario_medicacao or a.horario_consulta_medica
            horas_decorridas = None
            minutos_para_reavaliacao = None
            precisa_observacao = ciclo_atual > 3
            
            if horario_referencia:
                # Garantir que horario_referencia tem timezone antes de subtrair
                horario_referencia_tz = converter_para_brasilia(horario_referencia)
                diferenca = agora - horario_referencia_tz
                horas_decorridas = diferenca.total_seconds() / 3600
                minutos_para_reavaliacao = (120 - (diferenca.total_seconds() / 60))  # 120 minutos = 2 horas
            
            resultado.append({
                'atendimento_id': a.id,
                'paciente_id': paciente.id,
                'nome': paciente.nome,
                'data_nascimento': paciente.data_nascimento.strftime('%Y-%m-%d') if paciente.data_nascimento else None,
                'idade': calcular_idade(paciente.data_nascimento) if paciente.data_nascimento else None,
                'classificacao_risco': a.classificacao_risco,
                'triagem': a.triagem,
                'horario_triagem': a.horario_triagem.strftime('%Y-%m-%d %H:%M:%S') if a.horario_triagem else None,
                'horario_consulta_medica': a.horario_consulta_medica.strftime('%Y-%m-%d %H:%M:%S') if a.horario_consulta_medica else None,
                'horario_medicacao': a.horario_medicacao.strftime('%Y-%m-%d %H:%M:%S') if a.horario_medicacao else None,
                'ciclo_atual': min(ciclo_atual, 3),  # Limita a 3 para exibição
                'horas_decorridas': round(horas_decorridas, 1) if horas_decorridas is not None else None,
                'minutos_para_reavaliacao': round(minutos_para_reavaliacao) if minutos_para_reavaliacao is not None else None,
                'precisa_observacao': precisa_observacao,
                'passou_2_horas': horas_decorridas >= 2 if horas_decorridas is not None else False,
                'prioridade': paciente.prioridade if hasattr(paciente, 'prioridade') else False,
                'desc_prioridade': paciente.desc_prioridade if hasattr(paciente, 'desc_prioridade') else None
            })

        return jsonify({'success': True, 'atendimentos': resultado, 'total': len(resultado)})

    except Exception as e:
        logging.error(f"Erro ao listar atendimentos de reavaliação após medicação: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro interno do servidor'}), 500


@bp.route('/api/medico/atendimentos/aguardando-cirurgia', methods=['GET'])
@login_required
def api_medico_atendimentos_aguardando_cirurgia():
    """
    Retorna lista de atendimentos com status "aguardando cirurgia" para a aba de cirurgia.
    """
    try:
        current_user = get_current_user()
        if not current_user:
            return jsonify({'success': False, 'message': 'Usuário não autenticado.'}), 401

        if current_user.cargo.lower() not in ['medico', 'multi', 'admin']:
            return jsonify({'success': False, 'message': 'Acesso não autorizado'}), 403

        atendimentos = Atendimento.query.filter(
            Atendimento.status.ilike('%aguardando%cirurgia%')
        ).order_by(Atendimento.horario_consulta_medica.desc().nullslast(), Atendimento.data_atendimento.desc()).all()

        def calcular_idade(data_nascimento):
            if not data_nascimento:
                return None
            hoje = date.today()
            anos = hoje.year - data_nascimento.year - ((hoje.month, hoje.day) < (data_nascimento.month, data_nascimento.day))
            return anos

        resultado = []
        for a in atendimentos:
            paciente = Paciente.query.get(a.paciente_id)
            if not paciente:
                continue

            # Verifica se há chamado ativo para este atendimento
            chamado_ativo = Chamado.query.filter(
                Chamado.id_atendimento == a.id,
                Chamado.status == 'ativo'
            ).order_by(Chamado.data.desc(), Chamado.hora.desc()).first()

            chamado_info = None
            if chamado_ativo:
                medico = Funcionario.query.get(chamado_ativo.id_medico) if chamado_ativo.id_medico else None
                chamado_info = {
                    'medico_nome': medico.nome if medico else None,
                    'local': chamado_ativo.local,
                    'hora': chamado_ativo.hora.strftime('%H:%M') if chamado_ativo.hora else None
                }

            resultado.append({
                'atendimento_id': a.id,
                'paciente_id': paciente.id,
                'nome': paciente.nome,
                'data_nascimento': paciente.data_nascimento.strftime('%Y-%m-%d') if paciente.data_nascimento else None,
                'idade': calcular_idade(paciente.data_nascimento) if paciente.data_nascimento else None,
                'classificacao_risco': a.classificacao_risco,
                'triagem': a.triagem,
                'horario_triagem': a.horario_triagem.strftime('%Y-%m-%d %H:%M:%S') if a.horario_triagem else None,
                'horario_consulta_medica': a.horario_consulta_medica.strftime('%Y-%m-%d %H:%M:%S') if a.horario_consulta_medica else None,
                'conduta_final': a.conduta_final,
                'prioridade': paciente.prioridade if hasattr(paciente, 'prioridade') else False,
                'desc_prioridade': paciente.desc_prioridade if hasattr(paciente, 'desc_prioridade') else None,
                'chamado_ativo': chamado_info
            })

        return jsonify({'success': True, 'atendimentos': resultado, 'total': len(resultado)})

    except Exception as e:
        logging.error(f"Erro ao listar atendimentos aguardando cirurgia: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro interno do servidor'}), 500


@bp.route('/api/medico/atendimentos/gestantes', methods=['GET'])
@login_required
def api_medico_atendimentos_gestantes():
    """
    Lista atendimentos de pacientes gestantes (possuem registro em AtendimentosGestante),
    priorizando os que estão aguardando (médico ou triagem).
    Retorna campos compatíveis com o frontend da página do médico.
    """
    try:
        current_user = get_current_user()
        if not current_user:
            return jsonify({'success': False, 'message': 'Usuário não autenticado.'}), 401

        if current_user.cargo.lower() not in ['medico', 'multi', 'admin']:
            return jsonify({'success': False, 'message': 'Acesso não autorizado'}), 403

        # Seleciona atendimentos com registro em AtendimentosGestante
        # e prioriza os "aguardando" (médico/triagem) na ordenação
        query = db.session.query(Atendimento).join(
            AtendimentosGestante, AtendimentosGestante.id_atendimentos == Atendimento.id
        )

        # Mantém apenas atendimentos em estados de espera, para foco operacional
        atendimentos = query.filter(
            Atendimento.status.ilike('%aguardando%')
        ).order_by(
            Atendimento.horario_triagem.desc().nullslast(),
            Atendimento.data_atendimento.desc()
        ).all()

        def calcular_idade(data_nascimento):
            if not data_nascimento:
                return None
            hoje = date.today()
            anos = hoje.year - data_nascimento.year - ((hoje.month, hoje.day) < (data_nascimento.month, data_nascimento.day))
            return anos

        resultado = []
        for a in atendimentos:
            paciente = Paciente.query.get(a.paciente_id)
            if not paciente:
                continue
            resultado.append({
                'atendimento_id': a.id,
                'paciente_id': paciente.id,
                'nome': paciente.nome,
                'data_nascimento': paciente.data_nascimento.strftime('%Y-%m-%d') if paciente.data_nascimento else None,
                'idade': calcular_idade(paciente.data_nascimento) if paciente.data_nascimento else None,
                'classificacao_risco': a.classificacao_risco,
                'triagem': a.triagem,
                'horario_triagem': a.horario_triagem.strftime('%Y-%m-%d %H:%M:%S') if a.horario_triagem else None,
                'prioridade': paciente.prioridade if hasattr(paciente, 'prioridade') else False,
                'desc_prioridade': paciente.desc_prioridade if hasattr(paciente, 'desc_prioridade') else None
            })

        return jsonify({'success': True, 'atendimentos': resultado, 'total': len(resultado)})

    except Exception as e:
        logging.error(f"Erro ao listar atendimentos gestantes (médico): {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro interno do servidor'}), 500


@bp.route('/api/medico/atendimento/buscar', methods=['GET'])
@login_required
def api_medico_buscar_atendimento():
    """
    Busca atendimentos por número (ID) para reabrir prontuário.
    Retorna informações do atendimento incluindo status.
    """
    try:
        current_user = get_current_user()
        if not current_user:
            return jsonify({'success': False, 'message': 'Usuário não autenticado.'}), 401

        if current_user.cargo.lower() not in ['medico', 'multi', 'admin']:
            return jsonify({'success': False, 'message': 'Acesso não autorizado'}), 403

        numero = request.args.get('numero', '').strip()
        if len(numero) < 2:
            return jsonify({'success': False, 'message': 'Digite pelo menos 2 caracteres'}), 400

        # Busca atendimentos que contenham o número digitado
        atendimentos = Atendimento.query.filter(
            Atendimento.id.like(f'%{numero}%')
        ).order_by(Atendimento.data_atendimento.desc()).limit(20).all()

        resultado = []
        for a in atendimentos:
            paciente = Paciente.query.get(a.paciente_id)
            if not paciente:
                continue
            
            # Verificar se está finalizado
            status_lower = (a.status or '').lower()
            status_finalizado = (
                'finalizado' in status_lower or 
                'alta' in status_lower
            )
            
            resultado.append({
                'id': a.id,
                'paciente_nome': paciente.nome,
                'data_atendimento': a.data_atendimento.strftime('%Y-%m-%d %H:%M:%S') if a.data_atendimento else None,
                'status': a.status,
                'status_finalizado': status_finalizado
            })

        return jsonify({'success': True, 'atendimentos': resultado, 'total': len(resultado)})

    except Exception as e:
        logging.error(f"Erro ao buscar atendimentos: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro interno do servidor'}), 500


@bp.route('/medico/aguardando-medico', methods=['GET'])
@login_required
def pagina_medico_aguardando_medico():
    try:
        current_user = get_current_user()
        if not current_user:
            abort(401)
        if current_user.cargo.lower() not in ['medico', 'multi', 'admin']:
            abort(403)
        return render_template('medico_aguardando_medico.html')
    except Exception as e:
        logging.error(f"Erro ao renderizar página aguardando médico: {str(e)}")
        logging.error(traceback.format_exc())
        return render_template_string('<h3>Erro ao carregar a página</h3>'), 500


# Lista pacientes aguardando triagem (visão do médico)
@bp.route('/api/medico/atendimentos/aguardando-triagem', methods=['GET'])
@login_required
def api_medico_aguardando_triagem():
    try:
        current_user = get_current_user()
        if not current_user:
            return jsonify({'success': False, 'message': 'Usuário não autenticado'}), 401
        if current_user.cargo.lower() not in ['medico', 'multi', 'admin']:
            return jsonify({'success': False, 'message': 'Acesso não autorizado'}), 403

        atendimentos = Atendimento.query.filter(
            Atendimento.status.ilike('%aguardando triagem%')
        ).all()

        def calc_idade(dn):
            if not dn:
                return None
            hoje = date.today()
            return hoje.year - dn.year - ((hoje.month, hoje.day) < (dn.month, dn.day))

        resultado = []
        for a in atendimentos:
            paciente = Paciente.query.get(a.paciente_id)
            if not paciente:
                continue
            # Usar data/hora de criação da ficha como referência de "entrada"
            criado_em = None
            try:
                if a.data_atendimento and a.hora_atendimento:
                    criado_em = datetime.combine(a.data_atendimento, a.hora_atendimento).strftime('%Y-%m-%d %H:%M:%S')
            except Exception:
                criado_em = None

            resultado.append({
                'atendimento_id': a.id,
                'paciente_id': paciente.id,
                'nome': paciente.nome,
                'data_nascimento': paciente.data_nascimento.strftime('%Y-%m-%d') if paciente.data_nascimento else None,
                'idade': calc_idade(paciente.data_nascimento) if paciente.data_nascimento else None,
                'classificacao_risco': a.classificacao_risco,
                'triagem': a.triagem,
                'horario_triagem': a.horario_triagem.strftime('%Y-%m-%d %H:%M:%S') if a.horario_triagem else criado_em,
                'prioridade': paciente.prioridade if hasattr(paciente, 'prioridade') else False,
                'desc_prioridade': paciente.desc_prioridade if hasattr(paciente, 'desc_prioridade') else None
            })

        return jsonify({'success': True, 'atendimentos': resultado, 'total': len(resultado)})

    except Exception as e:
        logging.error(f"Erro ao listar aguardando triagem (médico): {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro interno do servidor'}), 500

# FICHA DE ATENDIMENTO MÉDICO (visualização principal do atendimento)
@bp.route('/medico/atendimento/<string:atendimento_id>', methods=['GET'])
@login_required
def pagina_medico_ficha_atendimento(atendimento_id):
    try:
        current_user = get_current_user()
        if not current_user:
            abort(401)
        if current_user.cargo.lower() not in ['medico', 'multi', 'admin']:
            abort(403)

        atendimento = Atendimento.query.get(atendimento_id)
        if not atendimento:
            flash('Atendimento não encontrado.', 'warning')
            return redirect(url_for('main.pagina_medico_aguardando_medico'))

        paciente = Paciente.query.get(atendimento.paciente_id)
        if not paciente:
            flash('Paciente não encontrado.', 'danger')
            return redirect(url_for('main.pagina_medico_aguardando_medico'))

        return render_template('medico_ficha_atendimento.html', paciente=paciente, atendimento=atendimento)
    except Exception as e:
        logging.error(f"Erro ao renderizar ficha de atendimento: {str(e)}")
        logging.error(traceback.format_exc())
        return render_template_string('<h3>Erro ao carregar a ficha</h3>'), 500


# API de dados básicos do atendimento (para a ficha)
@bp.route('/api/atendimento/<string:atendimento_id>', methods=['GET'])
@login_required
def api_dados_atendimento(atendimento_id):
    try:
        current_user = get_current_user()
        if not current_user:
            return jsonify({'success': False, 'message': 'Usuário não autenticado'}), 401
        if current_user.cargo.lower() not in ['medico', 'multi', 'admin', 'enfermeiro']:
            return jsonify({'success': False, 'message': 'Acesso não autorizado'}), 403

        atendimento = Atendimento.query.get(atendimento_id)
        if not atendimento:
            return jsonify({'success': False, 'message': 'Atendimento não encontrado'}), 404

        paciente = Paciente.query.get(atendimento.paciente_id)
        if not paciente:
            return jsonify({'success': False, 'message': 'Paciente não encontrado'}), 404

        def calc_idade(dn):
            if not dn:
                return None
            hoje = date.today()
            return hoje.year - dn.year - ((hoje.month, hoje.day) < (dn.month, dn.day))

        data = {
            'atendimento': {
                'id': atendimento.id,
                'status': atendimento.status,
                'classificacao_risco': atendimento.classificacao_risco,
                'triagem': atendimento.triagem,
                'pressao': atendimento.pressao,
                'pulso': atendimento.pulso,
                'sp02': atendimento.sp02,
                'temp': atendimento.temp,
                'fr': atendimento.fr,
                'peso': atendimento.peso,
                'altura': atendimento.altura,
                'dx': atendimento.dx,
                'alergias': atendimento.paciente.alergias if atendimento.paciente else '',
                'anamnese_exame_fisico': atendimento.anamnese_exame_fisico,
                'conduta_final': atendimento.conduta_final,
                'reavaliacao': atendimento.reavaliacao,
                'observacao': atendimento.observacao,
                'medicacao_utilizada': atendimento.medicacao_utilizada,
                'data_atendimento': atendimento.data_atendimento.isoformat() if atendimento.data_atendimento else None,
                'hora_atendimento': atendimento.hora_atendimento.strftime('%H:%M') if atendimento.hora_atendimento else None,
                'horario_triagem': atendimento.horario_triagem.strftime('%Y-%m-%d %H:%M:%S') if atendimento.horario_triagem else None
            },
            'paciente': {
                'id': paciente.id,
                'nome': paciente.nome,
                'cpf': paciente.cpf,
                'cartao_sus': paciente.cartao_sus,
                'data_nascimento': paciente.data_nascimento.strftime('%Y-%m-%d') if paciente.data_nascimento else None,
                'idade': calc_idade(paciente.data_nascimento) if paciente.data_nascimento else None,
                'sexo': paciente.sexo,
                'telefone': paciente.telefone,
                'endereco': paciente.endereco,
                'bairro': paciente.bairro,
                'municipio': paciente.municipio,
                'cor': paciente.cor,
                'filiacao': paciente.filiacao
            }
        }

        return jsonify({'success': True, 'data': data})
    except Exception as e:
        logging.error(f"Erro ao obter dados do atendimento: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro interno do servidor'}), 500


# API para registrar ciclo de reavaliação
@bp.route('/api/atendimento/<string:atendimento_id>/registrar-ciclo-reavaliacao', methods=['POST'])
@login_required
def registrar_ciclo_reavaliacao(atendimento_id):
    """
    Registra um novo ciclo de reavaliação, incrementando o contador de medicação utilizada.
    """
    try:
        current_user = get_current_user()
        if not current_user:
            return jsonify({'success': False, 'message': 'Usuário não autenticado'}), 401
        
        if current_user.cargo.lower() not in ['medico', 'multi', 'admin']:
            return jsonify({'success': False, 'message': 'Acesso não autorizado. Apenas médicos podem registrar ciclos de reavaliação.'}), 403

        # Buscar atendimento
        atendimento = Atendimento.query.get(atendimento_id)
        if not atendimento:
            return jsonify({'success': False, 'message': 'Atendimento não encontrado'}), 404

        # Verificar se está em reavaliação
        if not atendimento.status or 'REAVALIACAO' not in atendimento.status.upper():
            return jsonify({'success': False, 'message': 'Atendimento não está em status de reavaliação'}), 400

        # Contar ciclos atuais
        ciclos_atuais = 0
        if atendimento.medicacao_utilizada:
            try:
                medicacoes = json.loads(atendimento.medicacao_utilizada)
                if isinstance(medicacoes, list):
                    ciclos_atuais = len(medicacoes)
            except:
                ciclos_atuais = atendimento.medicacao_utilizada.count('1')
        
        # Registrar novo ciclo
        novo_ciclo = {
            'ciclo': ciclos_atuais + 1,
            'horario': now_brasilia().strftime('%Y-%m-%d %H:%M:%S'),
            'medico_id': current_user.id,
            'medico_nome': current_user.nome
        }
        
        # Atualizar medicacao_utilizada
        if atendimento.medicacao_utilizada:
            try:
                medicacoes = json.loads(atendimento.medicacao_utilizada)
                if isinstance(medicacoes, list):
                    medicacoes.append(novo_ciclo)
                else:
                    medicacoes = [novo_ciclo]
            except:
                medicacoes = [novo_ciclo]
        else:
            medicacoes = [novo_ciclo]
        
        atendimento.medicacao_utilizada = json.dumps(medicacoes, ensure_ascii=False)
        atendimento.horario_medicacao = now_brasilia()
        
        # Se completou 3 ciclos, alertar sobre necessidade de observação
        ciclo_atual = len(medicacoes)
        precisa_observacao = ciclo_atual > 3
        
        db.session.commit()
        
        logging.info(f"Ciclo de reavaliação registrado - Atendimento: {atendimento_id} - Ciclo: {ciclo_atual} - Médico: {current_user.nome}")
        
        return jsonify({
            'success': True,
            'message': 'Ciclo de reavaliação registrado com sucesso',
            'ciclo_atual': ciclo_atual,
            'precisa_observacao': precisa_observacao,
            'horario_medicacao': atendimento.horario_medicacao.strftime('%Y-%m-%d %H:%M:%S') if atendimento.horario_medicacao else None
        })

    except Exception as e:
        db.session.rollback()
        logging.error(f"Erro ao registrar ciclo de reavaliação: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro interno do servidor'}), 500


# API para buscar histórico de atendimentos do paciente
@bp.route('/api/paciente/<int:paciente_id>/historico-atendimentos', methods=['GET'])
@login_required
def api_historico_atendimentos_paciente(paciente_id):
    try:
        current_user = get_current_user()
        if not current_user:
            return jsonify({'success': False, 'message': 'Usuário não autenticado'}), 401
        if current_user.cargo.lower() not in ['medico', 'multi', 'admin', 'enfermeiro']:
            return jsonify({'success': False, 'message': 'Acesso não autorizado'}), 403

        # Pegar o ID do atendimento atual para excluí-lo do histórico
        atendimento_atual_id = request.args.get('atendimento_atual_id')
        
        # Buscar todos os atendimentos do paciente, ordenados por data decrescente, excluindo o atual
        query = Atendimento.query.filter_by(paciente_id=paciente_id)
        if atendimento_atual_id:
            query = query.filter(Atendimento.id != atendimento_atual_id)
        
        atendimentos = query.order_by(Atendimento.data_atendimento.desc(), Atendimento.hora_atendimento.desc())\
            .limit(20)\
            .all()
        
        # Calcular atendimentos nos últimos 7 dias
        data_limite = date.today() - timedelta(days=7)
        atendimentos_recentes = Atendimento.query.filter(
            Atendimento.paciente_id == paciente_id,
            Atendimento.data_atendimento >= data_limite
        ).count()
        
        historico = []
        for atend in atendimentos:
            medico = Funcionario.query.get(atend.medico_id) if atend.medico_id else None
            
            # Buscar prescrições de emergência do atendimento
            prescricoes = PrescricaoEmergencia.query.filter_by(atendimento_id=atend.id).all()
            prescricao_texto = []
            for presc in prescricoes:
                if presc.texto_dieta:
                    prescricao_texto.append(f"Dieta: {presc.texto_dieta}")
                if presc.medicamentos:
                    meds = json.loads(presc.medicamentos) if isinstance(presc.medicamentos, str) else presc.medicamentos
                    if meds:
                        prescricao_texto.append("Medicamentos:")
                        for med in meds:
                            prescricao_texto.append(f"  • {med.get('nome_medicamento', '')} - {med.get('descricao_uso', '')}")
                if presc.texto_procedimento_medico:
                    prescricao_texto.append(f"Procedimentos: {presc.texto_procedimento_medico}")
            
            # Verificar se existe observação ou internação
            observacao = ListaObservacao.query.filter_by(id_atendimento=atend.id).first()
            internacao = Internacao.query.filter_by(atendimento_id=atend.id).first()
            
            historico.append({
                'id': atend.id,
                'data': atend.data_atendimento.strftime('%d/%m/%Y') if atend.data_atendimento else '-',
                'hora': atend.hora_atendimento.strftime('%H:%M') if atend.hora_atendimento else '-',
                'status': atend.status or '-',
                'classificacao_risco': atend.classificacao_risco or '-',
                'triagem': atend.triagem or '-',
                'medico': medico.nome if medico else '-',
                'conduta_final': atend.conduta_final or '-',
                'anamnese_exame_fisico': atend.anamnese_exame_fisico or '-',
                'prescricao': '\n'.join(prescricao_texto) if prescricao_texto else '-',
                'tem_observacao': observacao is not None,
                'tem_internacao': internacao is not None
            })
        
        return jsonify({
            'success': True, 
            'historico': historico,
            'total': len(historico),
            'atendimentos_ultimos_7_dias': atendimentos_recentes,
            'alerta_frequente': atendimentos_recentes > 3
        })
    except Exception as e:
        logging.error(f"Erro ao buscar histórico de atendimentos: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro interno do servidor'}), 500


# API de dados completos do atendimento (complemento automático de informações)
@bp.route('/api/atendimento/<string:atendimento_id>/completo', methods=['GET'])
@login_required
def api_dados_atendimento_completo(atendimento_id):
    try:
        current_user = get_current_user()
        if not current_user:
            return jsonify({'success': False, 'message': 'Usuário não autenticado'}), 401
        if current_user.cargo.lower() not in ['medico', 'multi', 'admin', 'enfermeiro']:
            return jsonify({'success': False, 'message': 'Acesso não autorizado'}), 403

        atendimento = Atendimento.query.get(atendimento_id)
        if not atendimento:
            return jsonify({'success': False, 'message': 'Atendimento não encontrado'}), 404

        paciente = Paciente.query.get(atendimento.paciente_id)
        if not paciente:
            return jsonify({'success': False, 'message': 'Paciente não encontrado'}), 404

        # Profissionais
        medico = Funcionario.query.get(atendimento.medico_id) if atendimento.medico_id else None
        enfermeiro = Funcionario.query.get(atendimento.enfermeiro_id) if atendimento.enfermeiro_id else None

        # Internação (quando existir)
        internacao = Internacao.query.filter_by(atendimento_id=atendimento_id).first()

        # Helpers
        def calc_idade(dn):
            if not dn:
                return None
            hoje = date.today()
            return hoje.year - dn.year - ((hoje.month, hoje.day) < (dn.month, dn.day))

        def fmt_date(dt):
            try:
                return dt.strftime('%d/%m/%Y') if dt else None
            except Exception:
                return None

        def fmt_datetime(dt):
            try:
                return dt.strftime('%d/%m/%Y %H:%M') if dt else None
            except Exception:
                return None

        def fmt_time(tm):
            try:
                return tm.strftime('%H:%M') if tm else None
            except Exception:
                return None

        # Normalização de prescrições clínicas (quando houver internação)
        prescricoes_serializadas = []
        if internacao:
            try:
                prescricoes = PrescricaoClinica.query.filter_by(
                    atendimentos_clinica_id=internacao.id
                ).order_by(PrescricaoClinica.horario_prescricao.desc()).all()

                for p in prescricoes:
                    prescricoes_serializadas.append({
                        'id': p.id,
                        'horario_prescricao_iso': p.horario_prescricao.isoformat() if p.horario_prescricao else None,
                        'horario_prescricao_pretty': fmt_datetime(p.horario_prescricao),
                        'texto_dieta': p.texto_dieta,
                        'texto_procedimento_medico': p.texto_procedimento_medico,
                        'texto_procedimento_multi': p.texto_procedimento_multi,
                        'medicamentos_json': p.medicamentos_json
                    })
            except Exception as e:
                logging.error(f"Falha ao serializar prescrições clínicas: {str(e)}")
                logging.error(traceback.format_exc())

        data = {
            'paciente': {
                'id': paciente.id,
                'nome': paciente.nome,
                'cpf': paciente.cpf,
                'cartao_sus': paciente.cartao_sus,
                'data_nascimento_iso': paciente.data_nascimento.strftime('%Y-%m-%d') if paciente.data_nascimento else None,
                'data_nascimento_pretty': fmt_date(paciente.data_nascimento),
                'idade': calc_idade(paciente.data_nascimento),
                'sexo': paciente.sexo,
                'telefone': paciente.telefone,
                'endereco': paciente.endereco,
                'bairro': paciente.bairro,
                'municipio': paciente.municipio,
                'cor': paciente.cor,
                'filiacao': paciente.filiacao
            },
            'atendimento': {
                'id': atendimento.id,
                'status': atendimento.status,
                'classificacao_risco': atendimento.classificacao_risco,
                'triagem': atendimento.triagem,
                'pressao': atendimento.pressao,
                'pulso': atendimento.pulso,
                'sp02': atendimento.sp02,
                'temp': atendimento.temp,
                'fr': atendimento.fr,
                'peso': atendimento.peso,
                'altura': atendimento.altura,
                'dx': atendimento.dx,
                'alergias': atendimento.paciente.alergias if atendimento.paciente else '',
                'anamnese_exame_fisico': atendimento.anamnese_exame_fisico,
                'conduta_final': atendimento.conduta_final,
                'reavaliacao': atendimento.reavaliacao,
                'observacao': atendimento.observacao,
                'data_atendimento_iso': atendimento.data_atendimento.isoformat() if atendimento.data_atendimento else None,
                'data_atendimento_pretty': fmt_date(atendimento.data_atendimento) if atendimento.data_atendimento else None,
                'hora_atendimento': fmt_time(atendimento.hora_atendimento),
                'horario_triagem': fmt_datetime(atendimento.horario_triagem),
                'horario_consulta_medica': fmt_datetime(atendimento.horario_consulta_medica),
                'horario_medicacao': fmt_datetime(atendimento.horario_medicacao),
                'horario_alta': fmt_datetime(atendimento.horario_alta)
            },
            'medico': ({
                'id': medico.id,
                'nome': medico.nome,
                'numero_profissional': getattr(medico, 'numero_profissional', None)
            } if medico else None),
            'enfermeiro': ({
                'id': enfermeiro.id,
                'nome': enfermeiro.nome,
                'numero_profissional': getattr(enfermeiro, 'numero_profissional', None)
            } if enfermeiro else None),
            'internacao': ({
                'id': internacao.id,
                'leito': internacao.leito,
                'carater_internacao': internacao.carater_internacao,
                'data_internacao': fmt_datetime(internacao.data_internacao),
                'data_alta': fmt_datetime(internacao.data_alta)
            } if internacao else None),
            'prescricoes_clinica': prescricoes_serializadas
        }

        return jsonify({'success': True, 'data': data})
    except Exception as e:
        logging.error(f"Erro ao obter dados completos do atendimento: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro interno do servidor'}), 500

# Atualiza Anamnese e Conduta do atendimento
@bp.route('/api/atendimento/<string:atendimento_id>/anamnese-conduta', methods=['PUT'])
@login_required
def api_atendimento_anamnese_conduta(atendimento_id):
    try:
        current_user = get_current_user()
        if not current_user:
            return jsonify({'success': False, 'message': 'Usuário não autenticado'}), 401
        if current_user.cargo.lower() not in ['medico', 'multi', 'admin']:
            return jsonify({'success': False, 'message': 'Acesso não autorizado'}), 403

        dados = request.get_json() or {}
        atendimento = Atendimento.query.get(atendimento_id)
        if not atendimento:
            return jsonify({'success': False, 'message': 'Atendimento não encontrado'}), 404

        # Atualiza campos
        if 'anamnese_exame_fisico' in dados:
            atendimento.anamnese_exame_fisico = dados.get('anamnese_exame_fisico')
            # Sempre que a anamnese/exame físico for alterada, registra o horário da consulta
            try:
                # Registrar horário do servidor menos 3 horas
                atendimento.horario_consulta_medica = datetime.utcnow() - timedelta(hours=3)
            except Exception:
                pass
        if 'reavaliacao' in dados:
            atendimento.reavaliacao = dados.get('reavaliacao')
        if 'conduta_final' in dados:
            conduta_final = dados.get('conduta_final')
            atendimento.conduta_final = conduta_final
            # Sincronizar status com conduta_final (normalizado sem acento/cedilha)
            if conduta_final:
                atendimento.status = normalize_status(conduta_final)

        db.session.commit()

        return jsonify({'success': True, 'message': 'Dados atualizados com sucesso'})
    except Exception as e:
        db.session.rollback()
        logging.error(f"Erro ao atualizar anamnese/conduta do atendimento: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro interno do servidor'}), 500


# API para atualizar alergias do paciente
@bp.route('/api/paciente/<int:paciente_id>/alergias', methods=['PUT'])
@login_required
def api_atualizar_alergias_paciente(paciente_id):
    """
    Atualiza as alergias do paciente.
    Permite que enfermeiros e médicos atualizem as alergias conhecidas.
    """
    try:
        current_user = get_current_user()
        if not current_user:
            return jsonify({'success': False, 'message': 'Usuário não autenticado'}), 401

        # Permitir que enfermeiros, médicos e multi atualizem alergias
        if current_user.cargo.lower() not in ['enfermeiro', 'medico', 'multi', 'admin']:
            return jsonify({'success': False, 'message': 'Acesso não autorizado'}), 403

        dados = request.get_json() or {}
        novas_alergias = dados.get('alergias')

        # Converter string vazia em None
        if isinstance(novas_alergias, str) and novas_alergias.strip() == '':
            novas_alergias = None

        paciente = Paciente.query.get(paciente_id)
        if not paciente:
            return jsonify({'success': False, 'message': 'Paciente não encontrado'}), 404

        # Atualizar alergias do paciente
        paciente.alergias = novas_alergias

        db.session.commit()

        logging.info(f'Alergias atualizadas para paciente {paciente_id} por {current_user.nome} (ID: {current_user.id})')

        return jsonify({
            'success': True,
            'message': 'Alergias atualizadas com sucesso',
            'alergias': novas_alergias
        })
    except Exception as e:
        db.session.rollback()
        logging.error(f"Erro ao atualizar alergias do paciente: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro interno do servidor'}), 500

# API LEGADA - Mantida para compatibilidade, redireciona para atualizar alergias do paciente
@bp.route('/api/atendimento/<string:atendimento_id>/alergias', methods=['PUT'])
@login_required
def api_atualizar_alergias(atendimento_id):
    """
    [LEGADO] Endpoint mantido para compatibilidade. Atualiza alergias do paciente vinculado ao atendimento.
    """
    try:
        atendimento = Atendimento.query.get(atendimento_id)
        if not atendimento:
            return jsonify({'success': False, 'message': 'Atendimento não encontrado'}), 404

        # Redireciona para o novo endpoint que atualiza o paciente
        return api_atualizar_alergias_paciente(atendimento.paciente_id)
    except Exception as e:
        logging.error(f"Erro ao processar atualização de alergias (legado): {str(e)}")
        return jsonify({'success': False, 'message': 'Erro interno do servidor'}), 500


# Encerrar atendimento (alta/alta após medicação/evasão)
@bp.route('/api/atendimento/<string:atendimento_id>/encerrar', methods=['PUT'])
@login_required
def api_encerrar_atendimento(atendimento_id):
    try:
        current_user = get_current_user()
        if not current_user:
            return jsonify({'success': False, 'message': 'Usuário não autenticado'}), 401
        if current_user.cargo.lower() not in ['medico', 'multi', 'admin']:
            return jsonify({'success': False, 'message': 'Acesso não autorizado'}), 403

        dados = request.get_json() or {}
        conduta_final = (dados.get('conduta_final') or '').strip().upper()
        conduta_norm = normalize_status(conduta_final)
        status_customizado = dados.get('status')  # Aceitar status personalizado do frontend

        atendimento = Atendimento.query.get(atendimento_id)
        if not atendimento:
            return jsonify({'success': False, 'message': 'Atendimento não encontrado'}), 404

        # Atualiza conduta final registrada
        if conduta_final:
            atendimento.conduta_final = conduta_final
            # Se foi enviado um status personalizado, usar ele; senão, usar conduta normalizada
            if status_customizado:
                atendimento.status = status_customizado
                logging.info(f"Status personalizado aplicado: {status_customizado}")
            else:
                # Sincronizar status normalizado (sem acentos) no campo status
                atendimento.status = conduta_norm

        # Horário do servidor menos 3 horas
        agora = datetime.utcnow() - timedelta(hours=3)

        # Casos especiais que mantém atendimento em aberto: REAVALIACAO e ENCAMINHAR PARA CIRURGIA
        if conduta_norm == 'REAVALIACAO':
            db.session.commit()
            return jsonify({'success': True, 'message': 'Atendimento marcado para reavaliação; permanece em aberto'})

        if 'CIRURGIA' in conduta_norm or (status_customizado and 'cirurgia' in status_customizado.lower()):
            # Registrar fluxo do paciente
            try:
                fluxo = FluxoPaciente(
                    id_atendimento=atendimento.id,
                    id_medico=current_user.id if current_user else None,
                    id_enfermeiro=None,
                    nome_paciente=atendimento.paciente.nome if atendimento.paciente else '',
                    mudanca_status='Aguardando Cirurgia',
                    mudanca_hora=datetime.utcnow() - timedelta(hours=3)
                )
                db.session.add(fluxo)
                logging.info(f"FluxoPaciente registrado: Aguardando Cirurgia")
            except Exception as e:
                logging.error(f"Falha ao registrar FluxoPaciente no encaminhamento para cirurgia: {str(e)}")
            db.session.commit()
            return jsonify({'success': True, 'message': 'Paciente encaminhado para cirurgia; atendimento permanece em aberto'})

        # Para condutas de encerramento (alta, evasão, etc.), registrar horário de alta
        if conduta_final and conduta_norm != 'REAVALIACAO':
            atendimento.horario_alta = agora
            # Registrar fluxo do paciente com status personalizado ou conduta final
            status_para_fluxo = status_customizado if status_customizado else conduta_final
            try:
                fluxo = FluxoPaciente(
                    id_atendimento=atendimento.id,
                    id_medico=current_user.id if current_user else None,
                    id_enfermeiro=None,
                    nome_paciente=atendimento.paciente.nome if atendimento.paciente else '',
                    mudanca_status=status_para_fluxo,
                    mudanca_hora=agora
                )
                db.session.add(fluxo)
                logging.info(f"FluxoPaciente registrado: {status_para_fluxo}")
            except Exception as e:
                logging.error(f"Falha ao registrar FluxoPaciente no encerramento: {str(e)}")
            db.session.commit()
            
            mensagem_sucesso = f'Atendimento encerrado com conduta: {conduta_final}'
            if status_customizado:
                mensagem_sucesso += f' (Status: {status_customizado})'
            
            return jsonify({'success': True, 'message': mensagem_sucesso})

        # Para OBSERVAÇÃO e INTERNAMENTO o frontend usará endpoints dedicados
        return jsonify({'success': False, 'message': 'Use os fluxos de Observação ou Internamento para concluir esta conduta.'}), 400

    except Exception as e:
        db.session.rollback()
        logging.error(f"Erro ao encerrar atendimento: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro interno do servidor'}), 500

# POST /api/medico/mudar-senha
@bp.route('/api/medico/mudar-senha', methods=['POST'])
@login_required
def mudar_senha_medico():
    """
    Permite que o médico altere a própria senha.
    """
    try:
        # Usar o sistema de autenticação personalizado
        current_user = get_current_user()
        if not current_user:
            return jsonify({'success': False, 'message': 'Usuário não autenticado.'}), 401
        
        if current_user.cargo.lower() != 'medico':
            return jsonify({'success': False, 'message': 'Usuário não é médico.'}), 403

        dados = request.get_json()
        if not dados:
            return jsonify({'success': False, 'message': 'Dados não fornecidos.'}), 400

        senha_atual = dados.get('senha_atual')
        nova_senha = dados.get('nova_senha')

        if not senha_atual or not nova_senha:
            return jsonify({'success': False, 'message': 'Campos obrigatórios não preenchidos.'}), 400

        # Verificar se a senha atual confere
        if not current_user.check_password(senha_atual):
            return jsonify({'success': False, 'message': 'Senha atual incorreta.'}), 400

        # Atualizar a senha
        current_user.set_password(nova_senha)
        db.session.commit()

        return jsonify({'success': True, 'message': 'Senha alterada com sucesso.'})

    except Exception as e:
        db.session.rollback()
        logging.error(f"Erro ao alterar senha do médico: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': 'Erro ao alterar a senha.',
            'error': str(e)
        }), 500

# Rota protegida para médicos redirecionarem
@bp.route('/medico')
@login_required
def painel_medico():
    """
    Renderiza o painel do médico.
    """
    try:
        # Usar o sistema de autenticação personalizado
        current_user = get_current_user()
        if not current_user:
            flash('Sessão expirada. Por favor, faça login novamente.', 'warning')
            return redirect(url_for('main.index'))

        if current_user.cargo.strip().lower() != 'medico':
            flash('Acesso restrito a médicos.', 'danger')
            return redirect(url_for('main.index'))
        
        return render_template('medico.html')
        
    except Exception as e:
        logging.error(f"Erro ao acessar painel do médico: {str(e)}")
        logging.error(traceback.format_exc())
        flash('Erro ao acessar o painel. Por favor, tente novamente.', 'danger')
        return redirect(url_for('main.index'))
    
@bp.route('/enfermeiro')
@login_required
def painel_enfermeiro():
    try:
        # Verificar se o usuário é enfermeiro
        current_user = get_current_user()
        if current_user.cargo.lower() != 'enfermeiro':
            flash('Acesso restrito a enfermeiros.', 'danger')
            return redirect(url_for('main.index'))
        
        # Renderizar o template do painel do enfermeiro
        return render_template('enfermeiro.html')
        
    except Exception as e:
        logging.error(f"Erro ao acessar painel do enfermeiro: {str(e)}")
        logging.error(traceback.format_exc())
        flash('Erro ao acessar o painel. Por favor, tente novamente.', 'danger')
        return redirect(url_for('main.index'))

@bp.route('/multi')
@login_required
def painel_multi():
    try:
        # Verificar se o usuário é enfermeiro
        current_user = get_current_user()
        if current_user.cargo.lower() != 'multi':
            flash('Acesso restrito a enfermeiros.', 'danger')
            return redirect(url_for('main.index'))
        
        # Renderizar o template do painel do enfermeiro
        return render_template('multi.html')
        
    except Exception as e:
        logging.error(f"Erro ao acessar painel da Equipe multiprofissional : {str(e)}")
        logging.error(traceback.format_exc())
        flash('Erro ao acessar o painel. Por favor, tente novamente.', 'danger')
        return redirect(url_for('main.index'))

@bp.route('/administrador')
@login_required
def painel_administrador():
    try:
        user = get_current_user()
        if not user:
            flash('Sessão expirada. Por favor, faça login novamente.', 'warning')
            return redirect(url_for('main.index'))

        if user.cargo.strip().lower() != 'administrador':
            flash('Acesso restrito a administradores.', 'danger')
            return redirect(url_for('main.index'))

        # Cálculo dos indicadores
        agora_br = datetime.now(ZoneInfo("America/Sao_Paulo"))
        inicio_dia = datetime(agora_br.year, agora_br.month, agora_br.day, tzinfo=ZoneInfo("America/Sao_Paulo"))
        fim_dia = inicio_dia + timedelta(days=1)

        # 1) Internações hoje
        total_internacoes_hoje = db.session.query(func.count(Internacao.id)).filter(
            Internacao.data_internacao >= inicio_dia,
            Internacao.data_internacao < fim_dia
        ).scalar() or 0

        # 2) Taxa de ocupação de leitos = soma(ocupacao_atual)/soma(capacidade_maxima)
        soma_capacidade, soma_ocupacao = db.session.query(
            func.coalesce(func.sum(Leito.capacidade_maxima), 0),
            func.coalesce(func.sum(Leito.ocupacao_atual), 0)
        ).filter(Leito.status != 'Interditado').first()

        taxa_ocupacao = 0
        try:
            if soma_capacidade and soma_capacidade > 0:
                taxa_ocupacao = round((float(soma_ocupacao) / float(soma_capacidade)) * 100, 1)
        except Exception:
            taxa_ocupacao = 0

        # 3) Pacientes em observação (status "Em Observação")
        pacientes_observacao = db.session.query(func.count(Atendimento.id)).filter(
            Atendimento.status == 'Em Observação'
        ).scalar() or 0

        # 4) Tempo médio de permanência (em dias) para altas concluídas
        # Usa média do intervalo em segundos dividido por 86400, removendo outliers
        tempo_medio_permanencia = 0
        try:
            # Calcula tempos em dias, removendo outliers (< 0.1 dias ou > 180 dias)
            internacoes_com_tempo = db.session.query(
                func.extract('epoch', Internacao.data_alta - Internacao.data_internacao).label('tempo_segundos')
            ).filter(
                Internacao.data_internacao.isnot(None),
                Internacao.data_alta.isnot(None)
            ).subquery()
            
            media_segundos = db.session.query(
                func.avg(internacoes_com_tempo.c.tempo_segundos)
            ).filter(
                internacoes_com_tempo.c.tempo_segundos >= 8640,  # >= 0.1 dia (2.4 horas)
                internacoes_com_tempo.c.tempo_segundos <= 15552000  # <= 180 dias
            ).scalar()
            
            if media_segundos is not None:
                tempo_medio_permanencia = round(float(media_segundos) / 86400.0, 1)
        except Exception as e:
            logging.warning(f"Erro ao calcular tempo médio de permanência: {str(e)}")
            tempo_medio_permanencia = 0

        # 5) Atendimentos na porta hoje por classificação de risco
        hoje_data = inicio_dia.date()
        total_atendimentos_hoje = db.session.query(func.count(Atendimento.id)).filter(
            Atendimento.data_atendimento == hoje_data
        ).scalar() or 0

        grupos = db.session.query(
            Atendimento.classificacao_risco,
            func.count(Atendimento.id)
        ).filter(
            Atendimento.data_atendimento == hoje_data,
            Atendimento.classificacao_risco.isnot(None),
            Atendimento.classificacao_risco != ''
        ).group_by(
            Atendimento.classificacao_risco
        ).all()

        porta_risco_labels = [g[0] for g in grupos]
        porta_risco_valores = [int(g[1]) for g in grupos]

        # 6) Número de pacientes internados por dia (últimos 7 dias)
        fluxo_labels = []
        fluxo_entradas = []
        for i in range(6, -1, -1):
            dia = (inicio_dia - timedelta(days=i)).date()
            if i == 0:
                fluxo_labels.append('Hoje')
            else:
                fluxo_labels.append(f'D-{i}')
            
            # Contar internações no dia
            entradas_dia = db.session.query(func.count(Internacao.id)).filter(
                func.date(Internacao.data_internacao) == dia
            ).scalar() or 0
            fluxo_entradas.append(entradas_dia)

        # 7) Atendimentos por hora do dia (só mostrar até a hora atual)
        hora_atual = agora_br.hour
        atend_hora_labels = []
        atend_hora_valores = []
        
        # Criar labels e valores apenas até a hora atual
        for h in range(hora_atual + 1):
            atend_hora_labels.append(f'{h:02d}h')
            atend_hora_valores.append(0)
        
        atendimentos_hora = db.session.query(
            func.extract('hour', Atendimento.hora_atendimento).label('hora'),
            func.count(Atendimento.id).label('total')
        ).filter(
            Atendimento.data_atendimento == hoje_data
        ).group_by('hora').all()
        
        for hora, total in atendimentos_hora:
            if hora is not None and 0 <= int(hora) <= hora_atual:
                atend_hora_valores[int(hora)] = int(total)

        # 8) Calcular tempos médios reais usando os campos de horário
        # Removendo outliers: tempos < 1 min ou > 600 min (10 horas)
        
        # Tempo médio de TRIAGEM (hora_atendimento até horario_triagem)
        tempo_medio_triagem = 0
        qtd_com_triagem = 0
        try:
            # Criar subquery para calcular tempos em minutos
            tempos_triagem = db.session.query(
                (func.extract('epoch', Atendimento.horario_triagem - 
                 func.cast(func.concat(Atendimento.data_atendimento, ' ', Atendimento.hora_atendimento), db.DateTime)) / 60.0).label('tempo_min')
            ).filter(
                Atendimento.horario_triagem.isnot(None),
                Atendimento.hora_atendimento.isnot(None)
            ).subquery()
            
            # Calcular média removendo outliers
            resultado_triagem = db.session.query(
                func.avg(tempos_triagem.c.tempo_min)
            ).filter(
                tempos_triagem.c.tempo_min >= 1,    # >= 1 minuto
                tempos_triagem.c.tempo_min <= 600   # <= 10 horas
            ).scalar()
            
            if resultado_triagem:
                tempo_medio_triagem = round(float(resultado_triagem), 1)
            
            qtd_com_triagem = db.session.query(func.count(Atendimento.id)).filter(
                Atendimento.horario_triagem.isnot(None)
            ).scalar() or 0
        except Exception as e:
            logging.warning(f"Erro ao calcular tempo médio de triagem: {str(e)}")
            tempo_medio_triagem = 0

        # Tempo médio de CONSULTA MÉDICA (horario_atendimento até horario_consulta_medica)
        tempo_medio_consulta = 0
        qtd_com_consulta = 0
        try:
            # Criar subquery para calcular tempos em minutos
            tempos_consulta = db.session.query(
                (func.extract('epoch', Atendimento.horario_consulta_medica - 
                 func.cast(func.concat(Atendimento.data_atendimento, ' ', Atendimento.hora_atendimento), db.DateTime)) / 60.0).label('tempo_min')
            ).filter(
                Atendimento.hora_atendimento.isnot(None),
                Atendimento.horario_consulta_medica.isnot(None)
            ).subquery()
            
            # Calcular média removendo outliers
            resultado_consulta = db.session.query(
                func.avg(tempos_consulta.c.tempo_min)
            ).filter(
                tempos_consulta.c.tempo_min >= 1,    # >= 1 minuto
                tempos_consulta.c.tempo_min <= 600   # <= 10 horas
            ).scalar()
            
            if resultado_consulta:
                tempo_medio_consulta = round(float(resultado_consulta), 1)
            
            qtd_com_consulta = db.session.query(func.count(Atendimento.id)).filter(
                Atendimento.horario_consulta_medica.isnot(None)
            ).scalar() or 0
        except Exception as e:
            logging.warning(f"Erro ao calcular tempo médio de consulta: {str(e)}")
            tempo_medio_consulta = 0

        # Tempo médio de OBSERVAÇÃO (entrada até saída de observação)
        tempo_medio_observacao = 0
        qtd_com_observacao = 0
        try:
            tempos_observacao = db.session.query(
                (func.extract('epoch', Atendimento.horario_alta - Atendimento.horario_observacao) / 60.0).label('tempo_min')
            ).filter(
                Atendimento.horario_observacao.isnot(None),
                Atendimento.horario_alta.isnot(None)
            ).subquery()
            
            resultado_observacao = db.session.query(
                func.avg(tempos_observacao.c.tempo_min)
            ).filter(
                tempos_observacao.c.tempo_min >= 1,
                tempos_observacao.c.tempo_min <= 600
            ).scalar()
            
            if resultado_observacao:
                tempo_medio_observacao = round(float(resultado_observacao), 1)
            
            qtd_com_observacao = db.session.query(func.count(Atendimento.id)).filter(
                Atendimento.horario_observacao.isnot(None)
            ).scalar() or 0
        except Exception as e:
            logging.warning(f"Erro ao calcular tempo médio de observação: {str(e)}")
            tempo_medio_observacao = 0

        # Tempo médio de INTERNAÇÃO (entrada até saída)
        tempo_medio_internacao = 0
        qtd_com_internacao = 0
        try:
            tempos_internacao = db.session.query(
                (func.extract('epoch', Atendimento.horario_alta - Atendimento.horario_internacao) / 60.0).label('tempo_min')
            ).filter(
                Atendimento.horario_internacao.isnot(None),
                Atendimento.horario_alta.isnot(None)
            ).subquery()
            
            resultado_internacao = db.session.query(
                func.avg(tempos_internacao.c.tempo_min)
            ).filter(
                tempos_internacao.c.tempo_min >= 1,
                tempos_internacao.c.tempo_min <= 600
            ).scalar()
            
            if resultado_internacao:
                tempo_medio_internacao = round(float(resultado_internacao), 1)
            
            qtd_com_internacao = db.session.query(func.count(Atendimento.id)).filter(
                Atendimento.horario_internacao.isnot(None)
            ).scalar() or 0
        except Exception as e:
            logging.warning(f"Erro ao calcular tempo médio de internação: {str(e)}")
            tempo_medio_internacao = 0

        # Tempo TOTAL (chegada até alta)
        tempo_medio_total = 0
        qtd_finalizados = 0
        try:
            tempos_total = db.session.query(
                (func.extract('epoch', Atendimento.horario_alta - 
                 func.cast(func.concat(Atendimento.data_atendimento, ' ', Atendimento.hora_atendimento), db.DateTime)) / 60.0).label('tempo_min')
            ).filter(
                Atendimento.hora_atendimento.isnot(None),
                Atendimento.horario_alta.isnot(None)
            ).subquery()
            
            resultado_total = db.session.query(
                func.avg(tempos_total.c.tempo_min)
            ).filter(
                tempos_total.c.tempo_min >= 1,
                tempos_total.c.tempo_min <= 600
            ).scalar()
            
            if resultado_total:
                tempo_medio_total = round(float(resultado_total), 1)
            
            qtd_finalizados = db.session.query(func.count(Atendimento.id)).filter(
                Atendimento.horario_alta.isnot(None)
            ).scalar() or 0
        except Exception as e:
            logging.warning(f"Erro ao calcular tempo médio total: {str(e)}")
            tempo_medio_total = 0

        # 9) Tempo médio de espera POR CLASSIFICAÇÃO DE RISCO (horario_atendimento até horario_consulta)
        # Removendo outliers: tempos < 1 min ou > 600 min (10 horas)
        tempos_por_risco = []
        try:
            # Criar subquery com tempos calculados
            tempos_risco_sq = db.session.query(
                func.coalesce(Atendimento.classificacao_risco, 'Não classificado').label('classificacao'),
                (func.extract('epoch', Atendimento.horario_consulta_medica - 
                 func.cast(func.concat(Atendimento.data_atendimento, ' ', Atendimento.hora_atendimento), db.DateTime)) / 60.0).label('tempo_min'),
                Atendimento.id
            ).filter(
                Atendimento.hora_atendimento.isnot(None),
                Atendimento.horario_consulta_medica.isnot(None)
            ).subquery()
            
            # Calcular média por classificação removendo outliers
            resultados_risco = db.session.query(
                tempos_risco_sq.c.classificacao,
                func.avg(tempos_risco_sq.c.tempo_min).label('tempo_medio'),
                func.count(tempos_risco_sq.c.id).label('quantidade')
            ).filter(
                tempos_risco_sq.c.tempo_min >= 1,
                tempos_risco_sq.c.tempo_min <= 600
            ).group_by(
                tempos_risco_sq.c.classificacao
            ).all()
            
            for classificacao, tempo, qtd in resultados_risco:
                tempos_por_risco.append({
                    'classificacao': classificacao,
                    'tempo_medio': round(float(tempo), 1) if tempo else 0,
                    'quantidade': int(qtd)
                })
        except Exception as e:
            logging.warning(f"Erro ao calcular tempos por risco: {str(e)}")
            tempos_por_risco = []

        # 9) Pacientes aguardando triagem (mesmo critério usado no enfermeiro.html)
        aguardando_triagem = db.session.query(func.count(Atendimento.id)).filter(
            Atendimento.status.ilike('%aguardando triagem%')
        ).scalar() or 0

        # 10) Pacientes aguardando médico (status "Aguardando Medico")
        aguardando_medico = db.session.query(func.count(Atendimento.id)).filter(
            Atendimento.status == 'Aguardando Medico'
        ).scalar() or 0

        # 11) Últimas consultas realizadas (últimas 10 com horario_consulta_medica preenchido)
        ultimas_consultas = []
        try:
            consultas_query = db.session.query(
                Atendimento,
                Paciente,
                Funcionario.nome.label('medico_nome'),
                Funcionario
            ).join(
                Paciente, Atendimento.paciente_id == Paciente.id
            ).outerjoin(
                Funcionario, Atendimento.medico_id == Funcionario.id
            ).filter(
                Atendimento.horario_consulta_medica.isnot(None),
                Atendimento.hora_atendimento.isnot(None)
            ).order_by(
                Atendimento.horario_consulta_medica.desc()
            ).limit(10).all()
            
            for atendimento, paciente, medico_nome, medico_obj in consultas_query:
                # Calcular idade
                idade = 0
                if paciente.data_nascimento:
                    hoje = date.today()
                    idade = hoje.year - paciente.data_nascimento.year - (
                        (hoje.month, hoje.day) < (paciente.data_nascimento.month, paciente.data_nascimento.day)
                    )
                
                # Calcular tempo da triagem até consulta médica em minutos
                tempo_consulta = 0
                try:
                    if atendimento.horario_consulta_medica and atendimento.horario_triagem:
                        # Horário da triagem
                        if isinstance(atendimento.horario_triagem, datetime):
                            dt_triagem = atendimento.horario_triagem
                        else:
                            dt_triagem = datetime.combine(atendimento.data_atendimento, atendimento.horario_triagem)
                        
                        # Horário da consulta médica
                        if isinstance(atendimento.horario_consulta_medica, datetime):
                            dt_consulta = atendimento.horario_consulta_medica
                        else:
                            dt_consulta = datetime.combine(atendimento.data_atendimento, atendimento.horario_consulta_medica)
                        
                        # Calcular diferença
                        diferenca = dt_consulta - dt_triagem
                        tempo_consulta = round(diferenca.total_seconds() / 60.0)
                except Exception as e:
                    logging.warning(f"Erro ao calcular tempo de consulta: {str(e)}")
                    tempo_consulta = 0
                
                # Nome do enfermeiro
                enfermeiro_nome = "Não atribuído"
                if atendimento.enfermeiro_id:
                    enfermeiro = db.session.query(Funcionario).get(atendimento.enfermeiro_id)
                    if enfermeiro:
                        enfermeiro_nome = enfermeiro.nome
                
                # Calcular tempo até triagem (hora_atendimento até horario_triagem)
                tempo_ate_triagem = '-'
                try:
                    if atendimento.horario_triagem and atendimento.hora_atendimento:
                        # Combinar data e hora do atendimento (chegada)
                        dt_atendimento = datetime.combine(atendimento.data_atendimento, atendimento.hora_atendimento)
                        
                        # Hora da triagem
                        if isinstance(atendimento.horario_triagem, datetime):
                            dt_triagem = atendimento.horario_triagem
                        else:
                            dt_triagem = datetime.combine(atendimento.data_atendimento, atendimento.horario_triagem)
                        
                        # Calcular diferença
                        diferenca = dt_triagem - dt_atendimento
                        tempo_ate_triagem = round(diferenca.total_seconds() / 60.0)
                except:
                    tempo_ate_triagem = '-'
                
                ultimas_consultas.append({
                    'id_atendimento': atendimento.id,
                    'paciente_nome': paciente.nome,
                    'classificacao_risco': atendimento.classificacao_risco or 'Não informado',
                    'idade': idade,
                    'medico': medico_nome or 'Não atribuído',
                    'enfermeiro': enfermeiro_nome,
                    'tempo_consulta': tempo_consulta,
                    'tempo_ate_triagem': tempo_ate_triagem,
                    'horario_consulta': atendimento.horario_consulta_medica.strftime('%d/%m/%Y %H:%M') if atendimento.horario_consulta_medica else '-'
                })
        except Exception as e:
            logging.error(f"Erro ao buscar últimas consultas: {str(e)}")
            logging.error(traceback.format_exc())
            ultimas_consultas = []

        # 12) Últimos pacientes internados (apenas com status "Internado" no atendimento)
        ultimas_internacoes = []
        try:
            internacoes_query = db.session.query(
                Internacao,
                Paciente,
                Funcionario.nome.label('medico_nome')
            ).join(
                Paciente, Internacao.paciente_id == Paciente.id
            ).join(
                Atendimento, Internacao.atendimento_id == Atendimento.id
            ).outerjoin(
                Funcionario, Internacao.medico_id == Funcionario.id
            ).filter(
                Internacao.data_internacao.isnot(None),
                Atendimento.status == 'Internado'
            ).order_by(
                Internacao.data_internacao.desc()
            ).limit(10).all()
            
            for internacao, paciente, medico_nome in internacoes_query:
                # Verificar se estava em observação antes da internação
                estava_em_observacao = False
                if internacao.atendimento_id:
                    # Verificar na tabela lista_observacao
                    obs = db.session.query(ListaObservacao).filter(
                        ListaObservacao.id_atendimento == internacao.atendimento_id
                    ).first()
                    if obs:
                        estava_em_observacao = True
                    else:
                        # Verificar pelo status do atendimento
                        atend = db.session.query(Atendimento).filter(
                            Atendimento.id == internacao.atendimento_id
                        ).first()
                        if atend and atend.horario_observacao:
                            estava_em_observacao = True
                
                # Pegar diagnóstico (pode ser de vários campos)
                diagnostico = internacao.diagnostico_inicial or internacao.diagnostico or internacao.justificativa_internacao_sinais_e_sintomas or 'Não informado'
                # Limitar tamanho do diagnóstico para exibição
                if len(diagnostico) > 100:
                    diagnostico = diagnostico[:97] + '...'
                
                # Formatar data de internação
                data_internacao_fmt = '-'
                if internacao.data_internacao:
                    if isinstance(internacao.data_internacao, datetime):
                        data_internacao_fmt = internacao.data_internacao.strftime('%d/%m/%Y %H:%M')
                    else:
                        data_internacao_fmt = str(internacao.data_internacao)
                
                # Calcular idade do paciente
                idade_paciente = '-'
                if paciente.data_nascimento:
                    try:
                        if isinstance(paciente.data_nascimento, str):
                            data_nasc = datetime.strptime(paciente.data_nascimento, '%Y-%m-%d').date()
                        else:
                            data_nasc = paciente.data_nascimento
                        idade_paciente = (agora_br.date() - data_nasc).days // 365
                    except:
                        idade_paciente = '-'
                
                # Calcular tempo internado
                tempo_internado = '-'
                if internacao.data_internacao:
                    try:
                        if isinstance(internacao.data_internacao, datetime):
                            data_int = internacao.data_internacao
                        else:
                            data_int = datetime.strptime(str(internacao.data_internacao), '%Y-%m-%d %H:%M:%S')
                        
                        if data_int.tzinfo is None:
                            data_int = data_int.replace(tzinfo=ZoneInfo("America/Sao_Paulo"))
                        
                        delta = agora_br - data_int
                        dias = delta.days
                        horas = delta.seconds // 3600
                        
                        if dias > 0:
                            tempo_internado = f"{dias}d {horas}h"
                        else:
                            tempo_internado = f"{horas}h"
                    except:
                        tempo_internado = '-'
                
                # Buscar CID
                cid = internacao.cid_principal or 'Sem CID'
                
                # Buscar última evolução médica (EvolucaoAtendimentoClinica)
                ultima_evolucao_medica = 'Sem evolução'
                try:
                    from app.models import EvolucaoAtendimentoClinica
                    evo_med = db.session.query(EvolucaoAtendimentoClinica).filter(
                        EvolucaoAtendimentoClinica.atendimentos_clinica_id == internacao.id
                    ).order_by(EvolucaoAtendimentoClinica.data_evolucao.desc()).first()
                    
                    if evo_med and evo_med.data_evolucao:
                        if isinstance(evo_med.data_evolucao, datetime):
                            ultima_evolucao_medica = evo_med.data_evolucao.strftime('%d/%m %H:%M')
                        else:
                            ultima_evolucao_medica = str(evo_med.data_evolucao)
                except Exception as e:
                    logging.warning(f"Erro ao buscar evolução médica: {str(e)}")
                    ultima_evolucao_medica = 'Sem evolução'
                
                # Buscar última evolução de enfermagem (EvolucaoEnfermagem)
                ultima_evolucao_enfermagem = 'Sem evolução'
                try:
                    from app.models import EvolucaoEnfermagem
                    evo_enf = db.session.query(EvolucaoEnfermagem).filter(
                        EvolucaoEnfermagem.atendimentos_clinica_id == internacao.id
                    ).order_by(EvolucaoEnfermagem.data_evolucao.desc()).first()
                    
                    if evo_enf and evo_enf.data_evolucao:
                        if isinstance(evo_enf.data_evolucao, datetime):
                            ultima_evolucao_enfermagem = evo_enf.data_evolucao.strftime('%d/%m %H:%M')
                        else:
                            ultima_evolucao_enfermagem = str(evo_enf.data_evolucao)
                except Exception as e:
                    logging.warning(f"Erro ao buscar evolução de enfermagem: {str(e)}")
                    ultima_evolucao_enfermagem = 'Sem evolução'
                
                ultimas_internacoes.append({
                    'id_internacao': internacao.id,
                    'paciente_nome': paciente.nome,
                    'diagnostico': diagnostico,
                    'medico': medico_nome or 'Não atribuído',
                    'estava_em_observacao': estava_em_observacao,
                    'data_internacao': data_internacao_fmt,
                    'leito': internacao.leito or 'Não definido',
                    'idade': idade_paciente,
                    'cid': cid,
                    'tempo_internado': tempo_internado,
                    'ultima_evolucao_medica': ultima_evolucao_medica,
                    'ultima_evolucao_enfermagem': ultima_evolucao_enfermagem
                })
        except Exception as e:
            logging.error(f"Erro ao buscar últimas internações: {str(e)}")
            logging.error(traceback.format_exc())
            ultimas_internacoes = []

        # Dados adicionais para o modo TV
        hoje_date = agora_br.date()
        logging.info(f"=== INÍCIO CÁLCULOS MODO TV - Data: {hoje_date} ===")
        
        # 1) Média de idade dos pacientes de hoje
        media_idade_hoje = 0
        try:
            # Usar join para garantir que os dados do paciente sejam carregados
            atendimentos_hoje = db.session.query(Atendimento).join(
                Paciente, Atendimento.paciente_id == Paciente.id
            ).filter(
                Atendimento.data_atendimento == hoje_date
            ).all()
            
            logging.info(f"Total de atendimentos hoje ({hoje_date}): {len(atendimentos_hoje)}")
            
            idades = []
            pacientes_sem_data = 0
            for atend in atendimentos_hoje:
                paciente = db.session.query(Paciente).get(atend.paciente_id)
                if paciente and paciente.data_nascimento:
                    try:
                        if isinstance(paciente.data_nascimento, str):
                            data_nasc = datetime.strptime(paciente.data_nascimento, '%Y-%m-%d').date()
                        else:
                            data_nasc = paciente.data_nascimento
                        idade = (agora_br.date() - data_nasc).days // 365
                        if 0 <= idade <= 120:
                            idades.append(idade)
                    except Exception as ex:
                        logging.warning(f"Erro ao calcular idade do paciente #{paciente.id}: {str(ex)}")
                        continue
                else:
                    pacientes_sem_data += 1
            
            if pacientes_sem_data > 0:
                logging.info(f"Pacientes sem data de nascimento: {pacientes_sem_data}")
            
            if idades:
                media_idade_hoje = round(sum(idades) / len(idades))
            
            logging.info(f"Idade média calculada: {media_idade_hoje} (baseado em {len(idades)} pacientes)")
        except Exception as e:
            logging.error(f"Erro ao calcular média de idade: {str(e)}")
            logging.error(traceback.format_exc())
            media_idade_hoje = 0

        # 2) Divisão por sexo dos pacientes de hoje
        divisao_sexo_masculino = 0
        divisao_sexo_feminino = 0
        sexos_nao_reconhecidos = []
        try:
            # Usar join para garantir que os dados do paciente sejam carregados
            atendimentos_hoje = db.session.query(Atendimento).join(
                Paciente, Atendimento.paciente_id == Paciente.id
            ).filter(
                Atendimento.data_atendimento == hoje_date
            ).all()
            
            pacientes_sem_sexo = 0
            for atend in atendimentos_hoje:
                paciente = db.session.query(Paciente).get(atend.paciente_id)
                if paciente and paciente.sexo:
                    sexo_original = str(paciente.sexo).strip()
                    sexo_lower = sexo_original.lower()
                    
                    # Suporta: M, Masculino, Masc, Homem
                    if sexo_lower in ['m', 'masculino', 'masc', 'homem']:
                        divisao_sexo_masculino += 1
                    # Suporta: F, Feminino, Fem, Mulher
                    elif sexo_lower in ['f', 'feminino', 'fem', 'mulher']:
                        divisao_sexo_feminino += 1
                    else:
                        # Registrar valores não reconhecidos
                        if sexo_original not in sexos_nao_reconhecidos:
                            sexos_nao_reconhecidos.append(sexo_original)
                            logging.warning(f"Valor de sexo não reconhecido: '{sexo_original}'")
                else:
                    pacientes_sem_sexo += 1
            
            if pacientes_sem_sexo > 0:
                logging.info(f"Pacientes sem sexo informado: {pacientes_sem_sexo}")
            
            logging.info(f"Divisão por sexo calculada - Masculino: {divisao_sexo_masculino}, Feminino: {divisao_sexo_feminino}")
            if sexos_nao_reconhecidos:
                logging.warning(f"Valores de sexo não reconhecidos encontrados: {sexos_nao_reconhecidos}")
        except Exception as e:
            logging.error(f"Erro ao calcular divisão por sexo: {str(e)}")
            divisao_sexo_masculino = 0
            divisao_sexo_feminino = 0

        # 3) Bairros mais comuns (top 6)
        bairros_labels = []
        bairros_valores = []
        try:
            bairros_query = db.session.query(
                Paciente.bairro,
                func.count(Atendimento.id).label('total')
            ).join(
                Atendimento, Atendimento.paciente_id == Paciente.id
            ).filter(
                Atendimento.data_atendimento == hoje_date,
                Paciente.bairro.isnot(None),
                Paciente.bairro != ''
            ).group_by(
                Paciente.bairro
            ).order_by(
                func.count(Atendimento.id).desc()
            ).limit(6).all()
            
            for bairro, total in bairros_query:
                bairros_labels.append(bairro.strip())
                bairros_valores.append(total)
        except Exception as e:
            logging.error(f"Erro ao buscar bairros mais comuns: {str(e)}")
            bairros_labels = []
            bairros_valores = []

        # 4) Tempo médio de atendimento do dia (hora_atendimento até horario_consulta_medica)
        tempo_medio_atendimento_dia = 0
        try:
            # Buscar atendimentos de hoje que tenham hora_atendimento E horario_consulta_medica preenchidos
            atendimentos_com_consulta = db.session.query(Atendimento).filter(
                Atendimento.data_atendimento == hoje_date,
                Atendimento.hora_atendimento.isnot(None),
                Atendimento.horario_consulta_medica.isnot(None)
            ).all()
            
            logging.info(f"Atendimentos com consulta médica hoje: {len(atendimentos_com_consulta)}")
            
            if len(atendimentos_com_consulta) == 0:
                logging.warning("Nenhum atendimento encontrado com hora_atendimento E horario_consulta_medica preenchidos")
                tempo_medio_atendimento_dia = 0
            else:
                tempos_atendimento = []
                erros_calculo = 0
                
                for idx, atend in enumerate(atendimentos_com_consulta):
                    try:
                        # Log detalhado para os primeiros 3 atendimentos
                        if idx < 3:
                            logging.info(f"Processando atendimento #{atend.id}:")
                            logging.info(f"  - data_atendimento: {atend.data_atendimento} (tipo: {type(atend.data_atendimento)})")
                            logging.info(f"  - hora_atendimento: {atend.hora_atendimento} (tipo: {type(atend.hora_atendimento)})")
                            logging.info(f"  - horario_consulta_medica: {atend.horario_consulta_medica} (tipo: {type(atend.horario_consulta_medica)})")
                        
                        # PASSO 1: Converter hora_atendimento para datetime
                        if isinstance(atend.hora_atendimento, datetime):
                            dt_inicio = atend.hora_atendimento
                        elif isinstance(atend.hora_atendimento, time):
                            # Se for time, combinar com data_atendimento
                            dt_inicio = datetime.combine(atend.data_atendimento, atend.hora_atendimento)
                            # Adicionar timezone se necessário
                            if dt_inicio.tzinfo is None:
                                dt_inicio = dt_inicio.replace(tzinfo=ZoneInfo("America/Sao_Paulo"))
                        else:
                            logging.warning(f"Atendimento #{atend.id}: hora_atendimento tem tipo inesperado: {type(atend.hora_atendimento)}")
                            erros_calculo += 1
                            continue
                        
                        # PASSO 2: Converter horario_consulta_medica para datetime
                        if isinstance(atend.horario_consulta_medica, datetime):
                            dt_fim = atend.horario_consulta_medica
                        elif isinstance(atend.horario_consulta_medica, time):
                            # Se for time, combinar com data_atendimento
                            dt_fim = datetime.combine(atend.data_atendimento, atend.horario_consulta_medica)
                            # Adicionar timezone se necessário
                            if dt_fim.tzinfo is None:
                                dt_fim = dt_fim.replace(tzinfo=ZoneInfo("America/Sao_Paulo"))
                        else:
                            logging.warning(f"Atendimento #{atend.id}: horario_consulta_medica tem tipo inesperado: {type(atend.horario_consulta_medica)}")
                            erros_calculo += 1
                            continue
                        
                        # PASSO 3: Garantir que ambos tenham ou não tenham timezone
                        if dt_inicio.tzinfo is None and dt_fim.tzinfo is not None:
                            dt_inicio = dt_inicio.replace(tzinfo=ZoneInfo("America/Sao_Paulo"))
                        elif dt_inicio.tzinfo is not None and dt_fim.tzinfo is None:
                            dt_fim = dt_fim.replace(tzinfo=ZoneInfo("America/Sao_Paulo"))
                        
                        # PASSO 4: Calcular diferença em minutos
                        tempo_min = (dt_fim - dt_inicio).total_seconds() / 60.0
                        
                        if idx < 3:
                            logging.info(f"  - dt_inicio: {dt_inicio}")
                            logging.info(f"  - dt_fim: {dt_fim}")
                            logging.info(f"  - tempo calculado: {tempo_min:.1f} min")
                        
                        # PASSO 5: Filtrar outliers (tempo entre 1 min e 600 min = 10 horas)
                        if 1 <= tempo_min <= 600:
                            tempos_atendimento.append(tempo_min)
                        else:
                            if idx < 3:
                                logging.info(f"  - REJEITADO: tempo fora do range válido (1-600 min)")
                            logging.debug(f"Atendimento #{atend.id}: tempo fora do range ({tempo_min:.1f} min)")
                        
                    except Exception as ex:
                        logging.warning(f"Erro ao calcular tempo do atendimento #{atend.id}: {str(ex)}")
                        logging.warning(traceback.format_exc())
                        erros_calculo += 1
                        continue
                
                # Calcular média
                if tempos_atendimento:
                    tempo_medio_atendimento_dia = round(sum(tempos_atendimento) / len(tempos_atendimento))
                    logging.info(f"✓ Tempo médio de atendimento: {tempo_medio_atendimento_dia} min (baseado em {len(tempos_atendimento)} atendimentos válidos)")
                else:
                    logging.warning(f"Nenhum tempo válido calculado! Erros: {erros_calculo}")
                    tempo_medio_atendimento_dia = 0
                
                if erros_calculo > 0:
                    logging.warning(f"Houve {erros_calculo} erros ao calcular tempos de atendimento")
                    
        except Exception as e:
            logging.error(f"Erro ao calcular tempo médio de atendimento: {str(e)}")
            logging.error(traceback.format_exc())
            tempo_medio_atendimento_dia = 0
        
        # 5) Taxas de desfecho dos atendimentos de hoje
        taxa_alta = 0
        taxa_transferencia = 0
        taxa_obito = 0
        taxa_evasao = 0
        try:
            # Buscar todos os atendimentos de hoje
            total_atend_hoje = db.session.query(func.count(Atendimento.id)).filter(
                Atendimento.data_atendimento == hoje_date
            ).scalar() or 0
            
            logging.info(f"Total de atendimentos para cálculo de taxas: {total_atend_hoje}")
            
            if total_atend_hoje > 0:
                # Contar altas
                count_altas = db.session.query(func.count(Atendimento.id)).filter(
                    Atendimento.data_atendimento == hoje_date,
                    db.or_(
                        Atendimento.status == 'Alta',
                        Atendimento.status.ilike('%alta%')
                    )
                ).scalar() or 0
                taxa_alta = round((count_altas / total_atend_hoje) * 100, 1)
                
                # Contar transferências
                count_transfer = db.session.query(func.count(Atendimento.id)).filter(
                    Atendimento.data_atendimento == hoje_date,
                    db.or_(
                        Atendimento.status == 'Transferido',
                        Atendimento.status.ilike('%transfer%')
                    )
                ).scalar() or 0
                taxa_transferencia = round((count_transfer / total_atend_hoje) * 100, 1)
                
                # Contar óbitos
                count_obitos = db.session.query(func.count(Atendimento.id)).filter(
                    Atendimento.data_atendimento == hoje_date,
                    db.or_(
                        Atendimento.status == 'Óbito',
                        Atendimento.status.ilike('%obito%'),
                        Atendimento.status.ilike('%óbito%')
                    )
                ).scalar() or 0
                taxa_obito = round((count_obitos / total_atend_hoje) * 100, 1)
                
                # Contar evasões
                count_evasoes = db.session.query(func.count(Atendimento.id)).filter(
                    Atendimento.data_atendimento == hoje_date,
                    db.or_(
                        Atendimento.status == 'Evasão',
                        Atendimento.status.ilike('%evas%')
                    )
                ).scalar() or 0
                taxa_evasao = round((count_evasoes / total_atend_hoje) * 100, 1)
                
                logging.info(f"Taxas calculadas - Altas: {taxa_alta}%, Transf: {taxa_transferencia}%, Óbitos: {taxa_obito}%, Evasões: {taxa_evasao}%")
        except Exception as e:
            logging.error(f"Erro ao calcular taxas de desfecho: {str(e)}")
            taxa_alta = 0
            taxa_transferencia = 0
            taxa_obito = 0
            taxa_evasao = 0

        contexto = dict(
            total_internacoes_hoje=total_internacoes_hoje,
            taxa_ocupacao=taxa_ocupacao,
            tempo_medio_permanencia=tempo_medio_permanencia,
            pacientes_observacao=pacientes_observacao,
            total_atendimentos_hoje=total_atendimentos_hoje,
            porta_risco_labels=porta_risco_labels,
            porta_risco_valores=porta_risco_valores,
            fluxo_labels=fluxo_labels,
            fluxo_entradas=fluxo_entradas,
            atend_hora_labels=atend_hora_labels,
            atend_hora_valores=atend_hora_valores,
            # Tempos médios por etapa
            tempo_medio_triagem=tempo_medio_triagem,
            tempo_medio_consulta=tempo_medio_consulta,
            tempo_medio_observacao=tempo_medio_observacao,
            tempo_medio_internacao=tempo_medio_internacao,
            tempo_medio_total=tempo_medio_total,
            # Quantidades por etapa
            qtd_com_triagem=qtd_com_triagem,
            qtd_com_consulta=qtd_com_consulta,
            qtd_com_observacao=qtd_com_observacao,
            qtd_com_internacao=qtd_com_internacao,
            qtd_finalizados=qtd_finalizados,
            # Tempos por classificação de risco
            tempos_por_risco=tempos_por_risco,
            # Aguardando
            aguardando_triagem=aguardando_triagem,
            aguardando_medico=aguardando_medico,
            # Listas
            ultimas_consultas=ultimas_consultas,
            ultimas_internacoes=ultimas_internacoes,
            # Dados para modo TV
            media_idade_hoje=media_idade_hoje,
            divisao_sexo_masculino=divisao_sexo_masculino,
            divisao_sexo_feminino=divisao_sexo_feminino,
            bairros_labels=bairros_labels,
            bairros_valores=bairros_valores,
            # Tempo médio de atendimento do dia
            tempo_medio_atendimento_dia=tempo_medio_atendimento_dia,
            # Taxas de desfecho
            taxa_alta=taxa_alta,
            taxa_transferencia=taxa_transferencia,
            taxa_obito=taxa_obito,
            taxa_evasao=taxa_evasao,
        )
        return render_template('administrador.html', **contexto)
    except Exception as e:
        logging.error(f"Erro ao acessar painel do administrador: {str(e)}")
        logging.error(traceback.format_exc())
        flash('Erro ao acessar o painel de administração.', 'danger')
        return redirect(url_for('main.index'))

@bp.route('/administrador/relatorios')
@login_required
def relatorios_admin():
    try:
        user = get_current_user()
        if not user or user.cargo.strip().lower() != 'administrador':
            flash('Acesso restrito a administradores.', 'danger')
            return redirect(url_for('main.index'))
        return render_template('administrador_relatorios.html')
    except Exception as e:
        logging.error(f"Erro ao acessar gerador de relatórios: {str(e)}")
        logging.error(traceback.format_exc())
        flash('Erro ao acessar o gerador de relatórios.', 'danger')
        return redirect(url_for('main.index'))

@bp.route('/administrador/relatorios/preview')
@login_required
def preview_relatorio():
    """Rota de exemplo para visualizar o template de relatório com dados fictícios"""
    try:
        user = get_current_user()
        if not user or user.cargo.strip().lower() != 'administrador':
            flash('Acesso restrito a administradores.', 'danger')
            return redirect(url_for('main.index'))
        
        # Dados de exemplo para demonstração
        agora = datetime.now(ZoneInfo("America/Sao_Paulo"))
        
        # Exemplo de lista de atendimentos
        atendimentos_exemplo = [
            {'numero': '25102401', 'cpf': '123.456.789-00', 'paciente': 'João da Silva Santos', 
             'medico': 'Dr. Carlos Eduardo', 'enfermeiro': 'Maria José Silva', 
             'status': 'Alta', 'badge_status': 'success', 'data_hora': '24/10/2025 08:30'},
            {'numero': '25102402', 'cpf': '987.654.321-00', 'paciente': 'Ana Paula Oliveira', 
             'medico': 'Dra. Fernanda Lima', 'enfermeiro': 'José Carlos Santos', 
             'status': 'Internado', 'badge_status': 'danger', 'data_hora': '24/10/2025 09:15'},
            {'numero': '25102403', 'cpf': '456.789.123-00', 'paciente': 'Pedro Henrique Costa', 
             'medico': 'Dr. Roberto Alves', 'enfermeiro': 'Maria José Silva', 
             'status': 'Alta após Medicação', 'badge_status': 'warning', 'data_hora': '24/10/2025 10:00'},
            {'numero': '25102501', 'cpf': '321.654.987-00', 'paciente': 'Mariana Souza Lima', 
             'medico': 'Dr. Carlos Eduardo', 'enfermeiro': 'Ana Paula Ferreira', 
             'status': 'Alta', 'badge_status': 'success', 'data_hora': '25/10/2025 07:45'},
            {'numero': '25102502', 'cpf': '789.123.456-00', 'paciente': 'Francisco José Pereira', 
             'medico': 'Dra. Fernanda Lima', 'enfermeiro': 'José Carlos Santos', 
             'status': 'Transferência', 'badge_status': 'warning', 'data_hora': '25/10/2025 11:20'},
            {'numero': '25102503', 'cpf': '147.258.369-00', 'paciente': 'Juliana Rodrigues', 
             'medico': 'Dr. Roberto Alves', 'enfermeiro': 'Maria José Silva', 
             'status': 'Alta', 'badge_status': 'success', 'data_hora': '25/10/2025 14:30'},
            {'numero': '25102601', 'cpf': '258.369.147-00', 'paciente': 'Carlos Alberto Mendes', 
             'medico': 'Dr. Carlos Eduardo', 'enfermeiro': 'Ana Paula Ferreira', 
             'status': 'Observação', 'badge_status': 'secondary', 'data_hora': '26/10/2025 08:15'},
            {'numero': '25102602', 'cpf': '369.147.258-00', 'paciente': 'Beatriz Santos Costa', 
             'medico': 'Dra. Fernanda Lima', 'enfermeiro': 'José Carlos Santos', 
             'status': 'Alta', 'badge_status': 'success', 'data_hora': '26/10/2025 13:00'},
            {'numero': '25102701', 'cpf': '741.852.963-00', 'paciente': 'Roberto Silva Oliveira', 
             'medico': 'Dr. Roberto Alves', 'enfermeiro': 'Maria José Silva', 
             'status': 'Alta após Medicação', 'badge_status': 'warning', 'data_hora': '27/10/2025 09:30'},
            {'numero': '25102801', 'cpf': '852.963.741-00', 'paciente': 'Patrícia Lima Santos', 
             'medico': 'Dr. Carlos Eduardo', 'enfermeiro': 'Ana Paula Ferreira', 
             'status': 'Alta', 'badge_status': 'success', 'data_hora': '28/10/2025 10:45'},
        ]
        
        # Exemplo de dados diários (para série histórica)
        dados_diarios_exemplo = [
            {'data': '24/10/2025', 'total': 18, 'altas': 12, 'internacoes': 3, 'outros': 3},
            {'data': '25/10/2025', 'total': 22, 'altas': 15, 'internacoes': 4, 'outros': 3},
            {'data': '26/10/2025', 'total': 20, 'altas': 13, 'internacoes': 5, 'outros': 2},
            {'data': '27/10/2025', 'total': 25, 'altas': 16, 'internacoes': 6, 'outros': 3},
            {'data': '28/10/2025', 'total': 19, 'altas': 14, 'internacoes': 2, 'outros': 3},
            {'data': '29/10/2025', 'total': 23, 'altas': 17, 'internacoes': 3, 'outros': 3},
            {'data': '30/10/2025', 'total': 21, 'altas': 15, 'internacoes': 4, 'outros': 2},
        ]
        
        # Exemplo de distribuição por status
        dados_status_exemplo = [
            {'status': 'Alta', 'quantidade': 102, 'percentual': 65.4, 'badge_class': 'success'},
            {'status': 'Alta após Medicação', 'quantidade': 25, 'percentual': 16.0, 'badge_class': 'warning'},
            {'status': 'Internado', 'quantidade': 18, 'percentual': 11.5, 'badge_class': 'danger'},
            {'status': 'Transferência', 'quantidade': 7, 'percentual': 4.5, 'badge_class': 'warning'},
            {'status': 'Observação', 'quantidade': 4, 'percentual': 2.6, 'badge_class': 'secondary'},
        ]
        
        # Calcular taxa de conclusão (sem evasão)
        total_exemplo = len(atendimentos_exemplo)
        sem_evasao = sum(1 for a in atendimentos_exemplo if 'evasao' not in a['status'].lower())
        taxa_conclusao_exemplo = f"{round((sem_evasao / total_exemplo * 100), 1)}%" if total_exemplo > 0 else '0%'
        
        contexto = {
            'titulo_relatorio': 'Relatório de Atendimento Emergência (Série Histórica)',
            'modelo_relatorio': 'Série Histórica',
            'periodo_inicio': '24/10/2025',
            'periodo_fim': '30/10/2025',
            'tipo_relatorio_nome': 'Atendimento Emergência',
            'icone_relatorio': 'fas fa-ambulance',
            'data_geracao': agora.strftime('%d/%m/%Y às %H:%M'),
            'ano_atual': agora.year,
            
            # Lista de atendimentos
            'atendimentos_lista': atendimentos_exemplo,
            
            # Métricas
            'total_atendimentos': 156,
            'media_diaria': 22.3,
            'pacientes_unicos': 148,
            'taxa_conclusao': taxa_conclusao_exemplo,
            
            # Dados para gráfico e tabela
            'dados_diarios': dados_diarios_exemplo,
            'dados_status': dados_status_exemplo,
            
            'mostrar_assinaturas': True,
            'assinatura_1_nome': 'Dr. João Silva',
            'assinatura_1_cargo': 'Diretor Técnico',
        }
        
        return render_template('relatorio_emergencia.html', **contexto)
        
    except Exception as e:
        logging.error(f"Erro ao gerar preview de relatório: {str(e)}")
        logging.error(traceback.format_exc())
        flash('Erro ao gerar preview.', 'danger')
        return redirect(url_for('main.relatorios_admin'))

@bp.route('/administrador/relatorios/preview-analitico')
@login_required
def preview_relatorio_analitico():
    """Preview do relatório analítico com dados de exemplo"""
    try:
        user = get_current_user()
        if not user or user.cargo.strip().lower() != 'administrador':
            flash('Acesso restrito a administradores.', 'danger')
            return redirect(url_for('main.index'))
        
        agora = datetime.now(ZoneInfo("America/Sao_Paulo"))
        
        # Dados de exemplo para modo Analítico
        dados_medicos_exemplo = [
            {'medico': 'Dr. Carlos Eduardo Silva', 'quantidade': 45, 'percentual': 28.8, 'media_dia': 6.4},
            {'medico': 'Dra. Fernanda Lima Santos', 'quantidade': 38, 'percentual': 24.4, 'media_dia': 5.4},
            {'medico': 'Dr. Roberto Alves Costa', 'quantidade': 32, 'percentual': 20.5, 'media_dia': 4.6},
            {'medico': 'Dra. Ana Paula Ferreira', 'quantidade': 25, 'percentual': 16.0, 'media_dia': 3.6},
            {'medico': 'Dr. José Carlos Mendes', 'quantidade': 16, 'percentual': 10.3, 'media_dia': 2.3},
        ]
        
        dados_genero_exemplo = [
            {'genero': 'Feminino', 'quantidade': 82, 'percentual': 52.6, 'idade_media': 42.3},
            {'genero': 'Masculino', 'quantidade': 74, 'percentual': 47.4, 'idade_media': 38.7},
        ]
        
        dados_tempo_diario_exemplo = [
            {'data': '24/10/2025', 'quantidade': 18, 'tempo_medio': '1h25min', 'tempo_minimo': '35min', 'tempo_maximo': '3h15min'},
            {'data': '25/10/2025', 'quantidade': 22, 'tempo_medio': '1h32min', 'tempo_minimo': '28min', 'tempo_maximo': '2h50min'},
            {'data': '26/10/2025', 'quantidade': 20, 'tempo_medio': '1h18min', 'tempo_minimo': '42min', 'tempo_maximo': '3h05min'},
            {'data': '27/10/2025', 'quantidade': 25, 'tempo_medio': '1h45min', 'tempo_minimo': '50min', 'tempo_maximo': '3h30min'},
            {'data': '28/10/2025', 'quantidade': 19, 'tempo_medio': '1h22min', 'tempo_minimo': '38min', 'tempo_maximo': '2h45min'},
            {'data': '29/10/2025', 'quantidade': 23, 'tempo_medio': '1h28min', 'tempo_minimo': '45min', 'tempo_maximo': '3h00min'},
            {'data': '30/10/2025', 'quantidade': 21, 'tempo_medio': '1h35min', 'tempo_minimo': '40min', 'tempo_maximo': '2h55min'},
        ]
        
        dados_status_exemplo = [
            {'status': 'Alta', 'quantidade': 102, 'percentual': 65.4, 'badge_class': 'success'},
            {'status': 'Alta após Medicação', 'quantidade': 25, 'percentual': 16.0, 'badge_class': 'warning'},
            {'status': 'Internado', 'quantidade': 18, 'percentual': 11.5, 'badge_class': 'danger'},
            {'status': 'Transferência', 'quantidade': 7, 'percentual': 4.5, 'badge_class': 'warning'},
            {'status': 'Observação', 'quantidade': 4, 'percentual': 2.6, 'badge_class': 'secondary'},
        ]
        
        # Dados diários para o gráfico
        dados_diarios_exemplo = [
            {'data': '24/10/2025', 'total': 18, 'altas': 12, 'internacoes': 3, 'outros': 3},
            {'data': '25/10/2025', 'total': 22, 'altas': 15, 'internacoes': 4, 'outros': 3},
            {'data': '26/10/2025', 'total': 20, 'altas': 13, 'internacoes': 5, 'outros': 2},
            {'data': '27/10/2025', 'total': 25, 'altas': 16, 'internacoes': 6, 'outros': 3},
            {'data': '28/10/2025', 'total': 19, 'altas': 14, 'internacoes': 2, 'outros': 3},
            {'data': '29/10/2025', 'total': 23, 'altas': 17, 'internacoes': 3, 'outros': 3},
            {'data': '30/10/2025', 'total': 21, 'altas': 15, 'internacoes': 4, 'outros': 2},
        ]
        
        # Dados de exemplo de distribuição por idade (0-100 anos)
        # Simulando uma distribuição realista
        import random
        random.seed(42)  # Para resultados consistentes
        dados_idade_exemplo = []
        for idade in range(0, 101):
            # Distribuição mais realista - picos em idades comuns
            if idade < 10:
                qtd = random.choice([0, 0, 1, 2, 1, 0])
            elif idade < 20:
                qtd = random.choice([0, 1, 2, 3, 2, 1])
            elif idade < 40:
                qtd = random.choice([1, 2, 3, 4, 5, 6, 4, 3])
            elif idade < 60:
                qtd = random.choice([2, 3, 4, 5, 6, 5, 4])
            elif idade < 80:
                qtd = random.choice([1, 2, 3, 4, 3, 2])
            else:
                qtd = random.choice([0, 0, 1, 1, 2, 0])
            
            dados_idade_exemplo.append({
                'idade': idade,
                'quantidade': qtd
            })
        
        # Dados de tempo por médico
        dados_tempo_medico_exemplo = [
            {'medico': 'Dr. Carlos Eduardo Silva', 'quantidade': 45, 'tempo_medio': '1h32min', 'tempo_minimo': '25min', 'tempo_maximo': '3h15min'},
            {'medico': 'Dra. Fernanda Lima Santos', 'quantidade': 38, 'tempo_medio': '1h28min', 'tempo_minimo': '30min', 'tempo_maximo': '2h50min'},
            {'medico': 'Dr. Roberto Alves Costa', 'quantidade': 32, 'tempo_medio': '1h45min', 'tempo_minimo': '35min', 'tempo_maximo': '3h30min'},
            {'medico': 'Dra. Ana Paula Ferreira', 'quantidade': 25, 'tempo_medio': '1h22min', 'tempo_minimo': '28min', 'tempo_maximo': '2h45min'},
            {'medico': 'Dr. José Carlos Mendes', 'quantidade': 16, 'tempo_medio': '1h38min', 'tempo_minimo': '32min', 'tempo_maximo': '3h00min'},
        ]
        
        # Dados de tempo por enfermeiro
        dados_tempo_enfermeiro_exemplo = [
            {'enfermeiro': 'Maria José Silva', 'quantidade': 52, 'tempo_medio': '18min', 'tempo_minimo': '8min', 'tempo_maximo': '45min'},
            {'enfermeiro': 'José Carlos Santos', 'quantidade': 48, 'tempo_medio': '22min', 'tempo_minimo': '10min', 'tempo_maximo': '38min'},
            {'enfermeiro': 'Ana Paula Ferreira', 'quantidade': 34, 'tempo_medio': '20min', 'tempo_minimo': '12min', 'tempo_maximo': '42min'},
            {'enfermeiro': 'Roberto Lima Costa', 'quantidade': 22, 'tempo_medio': '25min', 'tempo_minimo': '15min', 'tempo_maximo': '50min'},
        ]
        
        contexto = {
            'titulo_relatorio': 'Relatório de Atendimento Emergência (Analítico)',
            'modelo_relatorio': 'Analítico',
            'periodo_inicio': '24/10/2025',
            'periodo_fim': '30/10/2025',
            'tipo_relatorio_nome': 'Atendimento Emergência',
            'icone_relatorio': 'fas fa-ambulance',
            'data_geracao': agora.strftime('%d/%m/%Y às %H:%M'),
            'ano_atual': agora.year,
            
            # Métricas
            'total_atendimentos': 156,
            'media_diaria': 22.3,
            'pacientes_unicos': 148,
            'taxa_conclusao': '97.4%',
            'idade_media': 40.5,
            
            # Dados específicos do analítico
            'dados_medicos': dados_medicos_exemplo,
            'dados_genero': dados_genero_exemplo,
            'dados_tempo_diario': dados_tempo_diario_exemplo,
            'dados_tempo_medico': dados_tempo_medico_exemplo,
            'dados_tempo_enfermeiro': dados_tempo_enfermeiro_exemplo,
            'dados_idade_detalhada': dados_idade_exemplo,
            'dados_status': dados_status_exemplo,
            'dados_diarios': dados_diarios_exemplo,  # Para o gráfico
            
            'mostrar_assinaturas': True,
            'assinatura_1_nome': 'Dr. João Silva',
            'assinatura_1_cargo': 'Diretor Técnico',
        }
        
        return render_template('relatorio_emergencia.html', **contexto)
        
    except Exception as e:
        logging.error(f"Erro ao gerar preview analítico: {str(e)}")
        logging.error(traceback.format_exc())
        flash('Erro ao gerar preview.', 'danger')
        return redirect(url_for('main.relatorios_admin'))

@bp.route('/administrador/relatorios/preview-produtividade')
@login_required
def preview_relatorio_produtividade():
    """Preview do relatório de produtividade clínica com dados de exemplo"""
    try:
        user = get_current_user()
        if not user or user.cargo.strip().lower() != 'administrador':
            flash('Acesso restrito a administradores.', 'danger')
            return redirect(url_for('main.index'))
        
        agora = datetime.now(ZoneInfo("America/Sao_Paulo"))
        
        # Dados de exemplo
        dados_internacoes_medico_exemplo = [
            {'medico': 'Dr. Carlos Eduardo Silva', 'quantidade': 23, 'percentual': 28.8, 'media_dia': 3.3},
            {'medico': 'Dra. Fernanda Lima Santos', 'quantidade': 19, 'percentual': 23.8, 'media_dia': 2.7},
            {'medico': 'Dr. Roberto Alves Costa', 'quantidade': 16, 'percentual': 20.0, 'media_dia': 2.3},
            {'medico': 'Dra. Ana Paula Ferreira', 'quantidade': 14, 'percentual': 17.5, 'media_dia': 2.0},
            {'medico': 'Dr. José Carlos Mendes', 'quantidade': 8, 'percentual': 10.0, 'media_dia': 1.1},
        ]
        
        dados_leitos_exemplo = [
            {'nome': 'Leito 01 - Clínica', 'capacidade': 2, 'ocupados': 2, 'disponiveis': 0, 'taxa_ocupacao': 100.0, 'status': 'Ocupado', 'badge_status': 'danger'},
            {'nome': 'Leito 02 - Clínica', 'capacidade': 2, 'ocupados': 1, 'disponiveis': 1, 'taxa_ocupacao': 50.0, 'status': 'Parcial', 'badge_status': 'warning'},
            {'nome': 'Leito 03 - Clínica', 'capacidade': 2, 'ocupados': 2, 'disponiveis': 0, 'taxa_ocupacao': 100.0, 'status': 'Ocupado', 'badge_status': 'danger'},
            {'nome': 'Leito 04 - Observação', 'capacidade': 1, 'ocupados': 0, 'disponiveis': 1, 'taxa_ocupacao': 0.0, 'status': 'Disponível', 'badge_status': 'success'},
            {'nome': 'Leito 05 - Observação', 'capacidade': 1, 'ocupados': 1, 'disponiveis': 0, 'taxa_ocupacao': 100.0, 'status': 'Ocupado', 'badge_status': 'danger'},
        ]
        
        dados_cid_exemplo = [
            {'cid': 'A09.9', 'descricao': 'Diarreia e gastroenterite', 'quantidade': 12, 'percentual': 15.0},
            {'cid': 'J18.9', 'descricao': 'Pneumonia não especificada', 'quantidade': 10, 'percentual': 12.5},
            {'cid': 'I10', 'descricao': 'Hipertensão essencial', 'quantidade': 8, 'percentual': 10.0},
            {'cid': 'K35.8', 'descricao': 'Apendicite aguda', 'quantidade': 7, 'percentual': 8.8},
            {'cid': 'N39.0', 'descricao': 'Infecção do trato urinário', 'quantidade': 6, 'percentual': 7.5},
            {'cid': 'E11.9', 'descricao': 'Diabetes mellitus', 'quantidade': 5, 'percentual': 6.3},
            {'cid': 'J44.0', 'descricao': 'DPOC com infecção', 'quantidade': 4, 'percentual': 5.0},
        ]
        
        dados_bairro_exemplo = [
            {'bairro': 'Centro', 'municipio': 'Cedro', 'quantidade': 18, 'percentual': 22.5},
            {'bairro': 'Bairro Novo', 'municipio': 'Cedro', 'quantidade': 12, 'percentual': 15.0},
            {'bairro': 'Vila Rica', 'municipio': 'Cedro', 'quantidade': 10, 'percentual': 12.5},
            {'bairro': 'Jardim América', 'municipio': 'Lavras da Mangabeira', 'quantidade': 8, 'percentual': 10.0},
            {'bairro': 'Centro', 'municipio': 'Lavras da Mangabeira', 'quantidade': 7, 'percentual': 8.8},
            {'bairro': 'Cohab', 'municipio': 'Cedro', 'quantidade': 6, 'percentual': 7.5},
            {'bairro': 'Alto da Boa Vista', 'municipio': 'Cedro', 'quantidade': 5, 'percentual': 6.3},
        ]
        
        dados_readmissao_exemplo = [
            {'paciente': 'João da Silva Santos', 'cpf': '123.456.789-00', 
             'data_primeira': '20/10/2025', 'data_alta': '23/10/2025', 
             'data_segunda': '26/10/2025', 'dias_diferenca': 3},
            {'paciente': 'Maria Oliveira Costa', 'cpf': '987.654.321-00', 
             'data_primeira': '19/10/2025', 'data_alta': '22/10/2025', 
             'data_segunda': '27/10/2025', 'dias_diferenca': 5},
        ]
        
        contexto = {
            'titulo_relatorio': 'Relatório de Produtividade Clínica (Analítico)',
            'modelo_relatorio': 'Analítico',
            'periodo_inicio': '24/10/2025',
            'periodo_fim': '30/10/2025',
            'tipo_relatorio_nome': 'Produtividade Clínica',
            'icone_relatorio': 'fas fa-hospital',
            'data_geracao': agora.strftime('%d/%m/%Y às %H:%M'),
            'ano_atual': agora.year,
            
            # Métricas principais
            'total_internacoes': 80,
            'total_altas': 65,
            'tempo_medio_alta': 4.2,
            'taxa_mortalidade': '1.3%',
            'taxa_readmissao': '3.1%',
            'taxa_ocupacao': '75.0%',
            'total_readmissoes': 2,
            
            # Dados das tabelas
            'dados_internacoes_medico': dados_internacoes_medico_exemplo,
            'dados_leitos': dados_leitos_exemplo,
            'capacidade_total': 8,
            'ocupados_total': 6,
            'disponiveis_total': 2,
            'taxa_ocupacao_geral': 75.0,
            'dados_cid': dados_cid_exemplo,
            'dados_bairro': dados_bairro_exemplo,
            'dados_readmissao': dados_readmissao_exemplo,
            
            'mostrar_assinaturas': True,
            'assinatura_1_nome': 'Dr. João Silva',
            'assinatura_1_cargo': 'Diretor Técnico',
        }
        
        return render_template('relatorio_produtividade.html', **contexto)
        
    except Exception as e:
        logging.error(f"Erro ao gerar preview de produtividade: {str(e)}")
        logging.error(traceback.format_exc())
        flash('Erro ao gerar preview.', 'danger')
        return redirect(url_for('main.relatorios_admin'))

@bp.route('/administrador/relatorios/preview-produtividade-historica')
@login_required
def preview_relatorio_produtividade_historica():
    """Preview do relatório de produtividade clínica (série histórica) com dados de exemplo"""
    try:
        user = get_current_user()
        if not user or user.cargo.strip().lower() != 'administrador':
            flash('Acesso restrito a administradores.', 'danger')
            return redirect(url_for('main.index'))
        
        agora = datetime.now(ZoneInfo("America/Sao_Paulo"))
        
        # Lista de exemplo de internações
        internacoes_exemplo = [
            {'numero': '25102401', 'cpf': '123.456.789-00', 'paciente': 'João da Silva Santos',
             'cid': 'A09.9', 'diagnostico': 'Diarreia e gastroenterite de origem...', 
             'leito': 'Leito 01', 'data_internacao': '24/10/2025', 'data_alta': '27/10/2025',
             'dias_internacao': 3, 'status': 'Alta', 'badge_status': 'success'},
            {'numero': '25102402', 'cpf': '987.654.321-00', 'paciente': 'Ana Paula Oliveira',
             'cid': 'J18.9', 'diagnostico': 'Pneumonia não especificada',
             'leito': 'Leito 02', 'data_internacao': '24/10/2025', 'data_alta': '30/10/2025',
             'dias_internacao': 6, 'status': 'Alta', 'badge_status': 'success'},
            {'numero': '25102501', 'cpf': '456.789.123-00', 'paciente': 'Pedro Henrique Costa',
             'cid': 'K35.8', 'diagnostico': 'Apendicite aguda',
             'leito': 'Leito 03', 'data_internacao': '25/10/2025', 'data_alta': '28/10/2025',
             'dias_internacao': 3, 'status': 'Alta', 'badge_status': 'success'},
            {'numero': '25102601', 'cpf': '321.654.987-00', 'paciente': 'Mariana Souza Lima',
             'cid': 'I10', 'diagnostico': 'Hipertensão essencial primária',
             'leito': 'Leito 01', 'data_internacao': '26/10/2025', 'data_alta': 'Internado',
             'dias_internacao': '4*', 'status': 'Internado', 'badge_status': 'secondary'},
            {'numero': '25102701', 'cpf': '789.123.456-00', 'paciente': 'Francisco José Pereira',
             'cid': 'N39.0', 'diagnostico': 'Infecção do trato urinário',
             'leito': 'Leito 02', 'data_internacao': '27/10/2025', 'data_alta': '29/10/2025',
             'dias_internacao': 2, 'status': 'Alta', 'badge_status': 'success'},
            {'numero': '25102801', 'cpf': '147.258.369-00', 'paciente': 'Juliana Rodrigues Silva',
             'cid': 'E11.9', 'diagnostico': 'Diabetes mellitus tipo 2',
             'leito': 'Leito 03', 'data_internacao': '28/10/2025', 'data_alta': 'Internado',
             'dias_internacao': '2*', 'status': 'Internado', 'badge_status': 'secondary'},
        ]
        
        dados_diarios_exemplo = [
            {'data': '24/10/2025', 'internacoes': 2, 'altas': 0, 'transferencias': 0, 'obitos': 0, 'saldo': 2},
            {'data': '25/10/2025', 'internacoes': 1, 'altas': 0, 'transferencias': 0, 'obitos': 0, 'saldo': 1},
            {'data': '26/10/2025', 'internacoes': 1, 'altas': 0, 'transferencias': 0, 'obitos': 0, 'saldo': 1},
            {'data': '27/10/2025', 'internacoes': 1, 'altas': 1, 'transferencias': 0, 'obitos': 0, 'saldo': 0},
            {'data': '28/10/2025', 'internacoes': 1, 'altas': 1, 'transferencias': 0, 'obitos': 0, 'saldo': 0},
            {'data': '29/10/2025', 'internacoes': 0, 'altas': 1, 'transferencias': 0, 'obitos': 0, 'saldo': -1},
            {'data': '30/10/2025', 'internacoes': 0, 'altas': 1, 'transferencias': 0, 'obitos': 0, 'saldo': -1},
        ]
        
        dados_status_exemplo = [
            {'status': 'Alta', 'quantidade': 4, 'percentual': 66.7, 'badge_class': 'success'},
            {'status': 'Internado', 'quantidade': 2, 'percentual': 33.3, 'badge_class': 'secondary'},
        ]
        
        contexto = {
            'titulo_relatorio': 'Relatório de Produtividade Clínica (Série Histórica)',
            'modelo_relatorio': 'Série Histórica',
            'periodo_inicio': '24/10/2025',
            'periodo_fim': '30/10/2025',
            'tipo_relatorio_nome': 'Produtividade Clínica',
            'icone_relatorio': 'fas fa-hospital',
            'data_geracao': agora.strftime('%d/%m/%Y às %H:%M'),
            'ano_atual': agora.year,
            
            # Lista completa
            'internacoes_lista': internacoes_exemplo,
            
            # Métricas
            'total_internacoes': 6,
            'media_diaria': 0.9,
            'total_altas': 4,
            'tempo_medio_permanencia': 3.5,
            'taxa_mortalidade': '0%',
            'taxa_ocupacao': '75.0%',
            
            # Dados para gráfico e tabelas
            'dados_diarios': dados_diarios_exemplo,
            'dados_status': dados_status_exemplo,
            
            'mostrar_assinaturas': True,
            'assinatura_1_nome': 'Dr. João Silva',
            'assinatura_1_cargo': 'Diretor Técnico',
        }
        
        return render_template('relatorio_produtividade.html', **contexto)
        
    except Exception as e:
        logging.error(f"Erro ao gerar preview de produtividade histórica: {str(e)}")
        logging.error(traceback.format_exc())
        flash('Erro ao gerar preview.', 'danger')
        return redirect(url_for('main.relatorios_admin'))

@bp.route('/administrador/relatorios/preview-observacao')
@login_required
def preview_relatorio_observacao():
    """Preview do relatório de observação com dados de exemplo"""
    try:
        user = get_current_user()
        if not user or user.cargo.strip().lower() != 'administrador':
            flash('Acesso restrito a administradores.', 'danger')
            return redirect(url_for('main.index'))
        
        agora = datetime.now(ZoneInfo("America/Sao_Paulo"))
        
        # Dados de exemplo
        dados_condutas_exemplo = [
            {'conduta': 'Alta Melhorada', 'quantidade': 45, 'percentual': 56.3, 'tempo_medio': '8.5h', 'badge_class': 'success'},
            {'conduta': 'Internação', 'quantidade': 22, 'percentual': 27.5, 'tempo_medio': '12.3h', 'badge_class': 'danger'},
            {'conduta': 'Transferência', 'quantidade': 8, 'percentual': 10.0, 'tempo_medio': '10.2h', 'badge_class': 'warning'},
            {'conduta': 'Evasão', 'quantidade': 5, 'percentual': 6.3, 'tempo_medio': '4.5h', 'badge_class': 'danger'},
        ]
        
        pacientes_internados_exemplo = [
            {'numero': '25102401', 'paciente': 'João da Silva Santos', 'cpf': '123.456.789-00',
             'data_obs': '24/10/2025 08:30', 'data_internacao': '24/10/2025 18:45',
             'tempo_obs': '10h15min', 'cid': 'J18.9'},
            {'numero': '25102502', 'paciente': 'Ana Paula Oliveira', 'cpf': '987.654.321-00',
             'data_obs': '25/10/2025 10:15', 'data_internacao': '25/10/2025 22:30',
             'tempo_obs': '12h15min', 'cid': 'A09.9'},
            {'numero': '25102603', 'paciente': 'Pedro Henrique Costa', 'cpf': '456.789.123-00',
             'data_obs': '26/10/2025 14:00', 'data_internacao': '27/10/2025 02:00',
             'tempo_obs': '12h00min', 'cid': 'K35.8'},
        ]
        
        dados_genero_exemplo = [
            {'genero': 'Feminino', 'quantidade': 42, 'percentual': 52.5, 'idade_media': 38.5},
            {'genero': 'Masculino', 'quantidade': 38, 'percentual': 47.5, 'idade_media': 42.3},
        ]
        
        dados_tempo_permanencia_exemplo = [
            {'faixa': 'Menos de 6h', 'quantidade': 18, 'percentual': 22.5},
            {'faixa': '6h a 12h', 'quantidade': 32, 'percentual': 40.0},
            {'faixa': '12h a 24h', 'quantidade': 22, 'percentual': 27.5},
            {'faixa': '24h a 48h', 'quantidade': 6, 'percentual': 7.5},
            {'faixa': 'Mais de 48h', 'quantidade': 2, 'percentual': 2.5},
        ]
        
        dados_cid_exemplo = [
            {'cid': 'A09.9', 'descricao': 'Diarreia e gastroenterite', 'quantidade': 15, 'percentual': 18.8},
            {'cid': 'R10.0', 'descricao': 'Abdome agudo', 'quantidade': 12, 'percentual': 15.0},
            {'cid': 'J06.9', 'descricao': 'Infecção aguda vias aéreas', 'quantidade': 10, 'percentual': 12.5},
            {'cid': 'R50.9', 'descricao': 'Febre não especificada', 'quantidade': 8, 'percentual': 10.0},
            {'cid': 'K59.0', 'descricao': 'Constipação', 'quantidade': 6, 'percentual': 7.5},
        ]
        
        dados_bairro_exemplo = [
            {'bairro': 'Centro', 'municipio': 'Cedro', 'quantidade': 20, 'percentual': 25.0},
            {'bairro': 'Bairro Novo', 'municipio': 'Cedro', 'quantidade': 14, 'percentual': 17.5},
            {'bairro': 'Vila Rica', 'municipio': 'Cedro', 'quantidade': 11, 'percentual': 13.8},
            {'bairro': 'Jardim América', 'municipio': 'Lavras da Mangabeira', 'quantidade': 9, 'percentual': 11.3},
            {'bairro': 'Centro', 'municipio': 'Lavras da Mangabeira', 'quantidade': 8, 'percentual': 10.0},
        ]
        
        contexto = {
            'titulo_relatorio': 'Relatório de Observação (Analítico)',
            'modelo_relatorio': 'Analítico',
            'periodo_inicio': '24/10/2025',
            'periodo_fim': '30/10/2025',
            'tipo_relatorio_nome': 'Observação',
            'icone_relatorio': 'fas fa-eye',
            'data_geracao': agora.strftime('%d/%m/%Y às %H:%M'),
            'ano_atual': agora.year,
            
            # Métricas
            'total_observacao': 80,
            'tempo_medio_observacao': 10.5,
            'taxa_internacao': '27.5%',
            'taxa_alta': '56.3%',
            'idade_media': 40.4,
            'media_diaria': 11.4,
            
            # Dados das tabelas
            'dados_condutas': dados_condutas_exemplo,
            'pacientes_internados': pacientes_internados_exemplo,
            'dados_genero': dados_genero_exemplo,
            'dados_tempo_permanencia': dados_tempo_permanencia_exemplo,
            'dados_cid': dados_cid_exemplo,
            'dados_bairro': dados_bairro_exemplo,
            
            'mostrar_assinaturas': True,
            'assinatura_1_nome': 'Dr. João Silva',
            'assinatura_1_cargo': 'Diretor Técnico',
        }
        
        return render_template('relatorio_observacao.html', **contexto)
        
    except Exception as e:
        logging.error(f"Erro ao gerar preview de observação: {str(e)}")
        logging.error(traceback.format_exc())
        flash('Erro ao gerar preview.', 'danger')
        return redirect(url_for('main.relatorios_admin'))

@bp.route('/administrador/profissionais')
@login_required
def gestao_profissionais():
    try:
        user = get_current_user()
        if not user or user.cargo.strip().lower() != 'administrador':
            flash('Acesso restrito a administradores.', 'danger')
            return redirect(url_for('main.index'))
        
        # Buscar todos os profissionais
        profissionais = db.session.query(Funcionario).order_by(Funcionario.nome).all()
        
        # Contar por cargo
        cargos_count = db.session.query(
            Funcionario.cargo,
            func.count(Funcionario.id)
        ).group_by(Funcionario.cargo).all()
        
        cargos_stats = {cargo: count for cargo, count in cargos_count}
        
        return render_template('gestao_profissionais.html', 
                             profissionais=profissionais,
                             cargos_stats=cargos_stats)
    except Exception as e:
        logging.error(f"Erro ao acessar gestão de profissionais: {str(e)}")
        logging.error(traceback.format_exc())
        flash('Erro ao acessar gestão de profissionais.', 'danger')
        return redirect(url_for('main.index'))

@bp.route('/administrador/profissionais/cadastrar', methods=['POST'])
@login_required
def cadastrar_profissional():
    try:
        user = get_current_user()
        if not user or user.cargo.strip().lower() != 'administrador':
            return jsonify({'success': False, 'message': 'Acesso negado'}), 403
        
        data = request.get_json()
        
        # Validar dados obrigatórios
        if not data.get('nome') or not data.get('cpf') or not data.get('cargo'):
            return jsonify({'success': False, 'message': 'Nome, CPF e cargo são obrigatórios'}), 400
        
        # Verificar se CPF já existe
        cpf_exists = db.session.query(Funcionario).filter_by(cpf=data['cpf']).first()
        if cpf_exists:
            return jsonify({'success': False, 'message': 'CPF já cadastrado'}), 400
        
        # Criar novo profissional
        novo_profissional = Funcionario(
            nome=data['nome'],
            cpf=data['cpf'],
            cargo=data['cargo'],
            crm=data.get('crm'),
            coren=data.get('coren'),
            especialidade=data.get('especialidade'),
            telefone=data.get('telefone'),
            email=data.get('email')
        )
        
        # Gerar senha inicial (CPF sem pontuação)
        senha_inicial = data['cpf'].replace('.', '').replace('-', '')
        novo_profissional.set_password(senha_inicial)
        
        db.session.add(novo_profissional)
        db.session.commit()
        
        return jsonify({
            'success': True, 
            'message': 'Profissional cadastrado com sucesso',
            'profissional_id': novo_profissional.id
        })
    except Exception as e:
        db.session.rollback()
        logging.error(f"Erro ao cadastrar profissional: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro ao cadastrar profissional'}), 500

@bp.route('/administrador/profissionais/<int:id>/editar', methods=['PUT'])
@login_required
def editar_profissional(id):
    try:
        user = get_current_user()
        if not user or user.cargo.strip().lower() != 'administrador':
            return jsonify({'success': False, 'message': 'Acesso negado'}), 403
        
        profissional = db.session.query(Funcionario).get(id)
        if not profissional:
            return jsonify({'success': False, 'message': 'Profissional não encontrado'}), 404
        
        data = request.get_json()
        
        # Atualizar dados
        if data.get('nome'):
            profissional.nome = data['nome']
        if data.get('cargo'):
            profissional.cargo = data['cargo']
        if data.get('crm'):
            profissional.crm = data['crm']
        if data.get('coren'):
            profissional.coren = data['coren']
        if data.get('especialidade'):
            profissional.especialidade = data['especialidade']
        if data.get('telefone'):
            profissional.telefone = data['telefone']
        if data.get('email'):
            profissional.email = data['email']
        
        db.session.commit()
        
        return jsonify({'success': True, 'message': 'Profissional atualizado com sucesso'})
    except Exception as e:
        db.session.rollback()
        logging.error(f"Erro ao editar profissional: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro ao editar profissional'}), 500

@bp.route('/administrador/profissionais/<int:id>/resetar-senha', methods=['POST'])
@login_required
def resetar_senha_profissional(id):
    try:
        user = get_current_user()
        if not user or user.cargo.strip().lower() != 'administrador':
            return jsonify({'success': False, 'message': 'Acesso negado'}), 403
        
        profissional = db.session.query(Funcionario).get(id)
        if not profissional:
            return jsonify({'success': False, 'message': 'Profissional não encontrado'}), 404
        
        # Resetar senha para CPF
        senha_inicial = profissional.cpf.replace('.', '').replace('-', '')
        profissional.set_password(senha_inicial)
        
        db.session.commit()
        
        return jsonify({'success': True, 'message': 'Senha resetada com sucesso para o CPF'})
    except Exception as e:
        db.session.rollback()
        logging.error(f"Erro ao resetar senha: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro ao resetar senha'}), 500

@bp.route('/administrador/pacientes')
@login_required
def gestao_pacientes():
    try:
        user = get_current_user()
        if not user or user.cargo.strip().lower() != 'administrador':
            flash('Acesso restrito a administradores.', 'danger')
            return redirect(url_for('main.index'))
        
        # Buscar pacientes com paginação (últimos 100)
        pacientes = db.session.query(Paciente).order_by(Paciente.id.desc()).limit(100).all()
        
        # Estatísticas
        total_pacientes = db.session.query(func.count(Paciente.id)).scalar() or 0
        pacientes_com_sus = db.session.query(func.count(Paciente.id)).filter(
            Paciente.cartao_sus.isnot(None),
            Paciente.cartao_sus != ''
        ).scalar() or 0
        pacientes_com_cpf = db.session.query(func.count(Paciente.id)).filter(
            Paciente.cpf.isnot(None),
            Paciente.cpf != ''
        ).scalar() or 0
        
        return render_template('gestao_pacientes.html',
                             pacientes=pacientes,
                             total_pacientes=total_pacientes,
                             pacientes_com_sus=pacientes_com_sus,
                             pacientes_com_cpf=pacientes_com_cpf,
                             today=datetime.now(ZoneInfo("America/Sao_Paulo")).date())
    except Exception as e:
        logging.error(f"Erro ao acessar gestão de pacientes: {str(e)}")
        logging.error(traceback.format_exc())
        flash('Erro ao acessar gestão de pacientes.', 'danger')
        return redirect(url_for('main.index'))

@bp.route('/administrador/pacientes/buscar')
@login_required
def buscar_pacientes_admin():
    try:
        user = get_current_user()
        if not user or user.cargo.strip().lower() != 'administrador':
            return jsonify({'success': False, 'message': 'Acesso negado'}), 403
        
        termo = request.args.get('termo', '').strip()
        if not termo or len(termo) < 3:
            return jsonify({'success': False, 'message': 'Digite ao menos 3 caracteres'}), 400
        
        # Buscar por nome, CPF ou cartão SUS
        pacientes = db.session.query(Paciente).filter(
            db.or_(
                Paciente.nome.ilike(f'%{termo}%'),
                Paciente.cpf.ilike(f'%{termo}%'),
                Paciente.cartao_sus.ilike(f'%{termo}%')
            )
        ).order_by(Paciente.nome).limit(50).all()
        
        resultado = [{
            'id': p.id,
            'nome': p.nome,
            'cpf': p.cpf or 'Não informado',
            'cartao_sus': p.cartao_sus or 'Não informado',
            'data_nascimento': p.data_nascimento.strftime('%d/%m/%Y') if p.data_nascimento else 'Não informado',
            'sexo': p.sexo or 'Não informado',
            'telefone': p.telefone or 'Não informado'
        } for p in pacientes]
        
        return jsonify({'success': True, 'pacientes': resultado})
    except Exception as e:
        logging.error(f"Erro ao buscar pacientes: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro ao buscar pacientes'}), 500

@bp.route('/administrador/pacientes/<int:id>')
@login_required
def obter_paciente_admin(id):
    try:
        user = get_current_user()
        if not user or user.cargo.strip().lower() != 'administrador':
            return jsonify({'success': False, 'message': 'Acesso negado'}), 403
        
        paciente = db.session.query(Paciente).get(id)
        if not paciente:
            return jsonify({'success': False, 'message': 'Paciente não encontrado'}), 404
        
        dados = {
            'id': paciente.id,
            'nome': paciente.nome,
            'nome_social': paciente.nome_social,
            'cpf': paciente.cpf,
            'cartao_sus': paciente.cartao_sus,
            'data_nascimento': paciente.data_nascimento.strftime('%Y-%m-%d') if paciente.data_nascimento else '',
            'sexo': paciente.sexo,
            'cor': paciente.cor,
            'filiacao': paciente.filiacao,
            'endereco': paciente.endereco,
            'bairro': paciente.bairro,
            'municipio': paciente.municipio,
            'telefone': paciente.telefone,
            'identificado': paciente.identificado,
            'descricao_nao_identificado': paciente.descricao_nao_identificado
        }
        
        return jsonify({'success': True, 'paciente': dados})
    except Exception as e:
        logging.error(f"Erro ao obter paciente: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro ao obter dados do paciente'}), 500

@bp.route('/administrador/pacientes/<int:id>/editar', methods=['PUT'])
@login_required
def editar_paciente(id):
    try:
        user = get_current_user()
        if not user or user.cargo.strip().lower() != 'administrador':
            return jsonify({'success': False, 'message': 'Acesso negado'}), 403
        
        paciente = db.session.query(Paciente).get(id)
        if not paciente:
            return jsonify({'success': False, 'message': 'Paciente não encontrado'}), 404
        
        data = request.get_json()
        
        # Registrar alterações para log
        alteracoes = []
        
        # Atualizar campos
        if 'nome' in data and data['nome'] != paciente.nome:
            alteracoes.append(f"Nome: {paciente.nome} → {data['nome']}")
            paciente.nome = data['nome']
        
        if 'nome_social' in data:
            paciente.nome_social = data['nome_social']
        
        if 'cpf' in data and data['cpf'] != paciente.cpf:
            # Verificar se CPF já existe em outro paciente
            if data['cpf']:
                cpf_existe = db.session.query(Paciente).filter(
                    Paciente.cpf == data['cpf'],
                    Paciente.id != id
                ).first()
                if cpf_existe:
                    return jsonify({'success': False, 'message': 'CPF já cadastrado para outro paciente'}), 400
            alteracoes.append(f"CPF: {paciente.cpf} → {data['cpf']}")
            paciente.cpf = data['cpf']
        
        if 'cartao_sus' in data and data['cartao_sus'] != paciente.cartao_sus:
            # Verificar se cartão SUS já existe
            if data['cartao_sus']:
                sus_existe = db.session.query(Paciente).filter(
                    Paciente.cartao_sus == data['cartao_sus'],
                    Paciente.id != id
                ).first()
                if sus_existe:
                    return jsonify({'success': False, 'message': 'Cartão SUS já cadastrado para outro paciente'}), 400
            alteracoes.append(f"Cartão SUS: {paciente.cartao_sus} → {data['cartao_sus']}")
            paciente.cartao_sus = data['cartao_sus']
        
        if 'data_nascimento' in data:
            from datetime import datetime
            if data['data_nascimento']:
                nova_data = datetime.strptime(data['data_nascimento'], '%Y-%m-%d').date()
                if nova_data != paciente.data_nascimento:
                    alteracoes.append(f"Data Nascimento alterada")
                paciente.data_nascimento = nova_data
        
        if 'sexo' in data:
            paciente.sexo = data['sexo']
        
        if 'cor' in data:
            paciente.cor = data['cor']
        
        if 'filiacao' in data:
            paciente.filiacao = data['filiacao']
        
        if 'endereco' in data:
            paciente.endereco = data['endereco']
        
        if 'bairro' in data:
            paciente.bairro = data['bairro']
        
        if 'municipio' in data:
            paciente.municipio = data['municipio']
        
        if 'telefone' in data:
            paciente.telefone = data['telefone']
        
        db.session.commit()
        
        # Log das alterações
        if alteracoes:
            logging.info(f"Paciente {id} editado por {user.nome}. Alterações: {', '.join(alteracoes)}")
        
        return jsonify({'success': True, 'message': 'Paciente atualizado com sucesso'})
    except Exception as e:
        db.session.rollback()
        logging.error(f"Erro ao editar paciente: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro ao editar paciente'}), 500

@bp.route('/administrador/atendimentos')
@login_required
def gestao_atendimentos():
    try:
        user = get_current_user()
        if not user or user.cargo.strip().lower() != 'administrador':
            flash('Acesso restrito a administradores.', 'danger')
            return redirect(url_for('main.index'))
        
        # Buscar atendimentos recentes (últimos 100)
        atendimentos = db.session.query(Atendimento).join(
            Paciente, Atendimento.paciente_id == Paciente.id
        ).order_by(Atendimento.data_atendimento.desc(), Atendimento.hora_atendimento.desc()).limit(100).all()
        
        # Estatísticas por status
        stats_status = db.session.query(
            Atendimento.status,
            func.count(Atendimento.id)
        ).filter(
            Atendimento.status.isnot(None)
        ).group_by(Atendimento.status).all()
        
        status_count = {status: count for status, count in stats_status}
        total_atendimentos = sum(status_count.values())
        
        # Atendimentos abertos (não finalizados)
        atendimentos_abertos = db.session.query(func.count(Atendimento.id)).filter(
            Atendimento.status.notin_(['Alta', 'Evasão', 'Óbito', 'Transferência', 'Finalizado'])
        ).scalar() or 0
        
        return render_template('gestao_atendimentos.html',
                             atendimentos=atendimentos,
                             status_count=status_count,
                             total_atendimentos=total_atendimentos,
                             atendimentos_abertos=atendimentos_abertos,
                             today=datetime.now(ZoneInfo("America/Sao_Paulo")).date())
    except Exception as e:
        logging.error(f"Erro ao acessar gestão de atendimentos: {str(e)}")
        logging.error(traceback.format_exc())
        flash('Erro ao acessar gestão de atendimentos.', 'danger')
        return redirect(url_for('main.index'))

@bp.route('/administrador/atendimentos/buscar')
@login_required
def buscar_atendimentos_admin():
    try:
        user = get_current_user()
        if not user or user.cargo.strip().lower() != 'administrador':
            return jsonify({'success': False, 'message': 'Acesso negado'}), 403
        
        termo = request.args.get('termo', '').strip()
        if not termo or len(termo) < 2:
            return jsonify({'success': False, 'message': 'Digite ao menos 2 caracteres'}), 400
        
        # Buscar por ID do atendimento ou nome do paciente
        atendimentos = db.session.query(Atendimento).join(
            Paciente, Atendimento.paciente_id == Paciente.id
        ).filter(
            db.or_(
                Atendimento.id.ilike(f'%{termo}%'),
                Paciente.nome.ilike(f'%{termo}%'),
                Paciente.cpf.ilike(f'%{termo}%')
            )
        ).order_by(Atendimento.data_atendimento.desc()).limit(50).all()
        
        resultado = []
        for atend in atendimentos:
            paciente = atend.paciente
            resultado.append({
                'id': atend.id,
                'paciente_nome': paciente.nome,
                'paciente_cpf': paciente.cpf or 'Não informado',
                'data_atendimento': atend.data_atendimento.strftime('%d/%m/%Y') if atend.data_atendimento else 'Não informado',
                'hora_atendimento': atend.hora_atendimento.strftime('%H:%M') if atend.hora_atendimento else 'Não informado',
                'status': atend.status or 'Sem status',
                'classificacao_risco': atend.classificacao_risco or 'Não classificado'
            })
        
        return jsonify({'success': True, 'atendimentos': resultado})
    except Exception as e:
        logging.error(f"Erro ao buscar atendimentos: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro ao buscar atendimentos'}), 500

@bp.route('/administrador/atendimentos/<string:id>')
@login_required
def obter_atendimento(id):
    try:
        user = get_current_user()
        if not user or user.cargo.strip().lower() != 'administrador':
            return jsonify({'success': False, 'message': 'Acesso negado'}), 403
        
        atendimento = db.session.query(Atendimento).get(id)
        if not atendimento:
            return jsonify({'success': False, 'message': 'Atendimento não encontrado'}), 404
        
        paciente = atendimento.paciente
        
        dados = {
            'id': atendimento.id,
            'paciente_id': paciente.id,
            'paciente_nome': paciente.nome,
            'paciente_cpf': paciente.cpf or 'Não informado',
            'data_atendimento': atendimento.data_atendimento.strftime('%d/%m/%Y') if atendimento.data_atendimento else '',
            'hora_atendimento': atendimento.hora_atendimento.strftime('%H:%M') if atendimento.hora_atendimento else '',
            'status': atendimento.status or '',
            'classificacao_risco': atendimento.classificacao_risco or '',
            'triagem': atendimento.triagem or '',
            'conduta_final': atendimento.conduta_final or '',
            'horario_triagem': atendimento.horario_triagem.strftime('%d/%m/%Y %H:%M') if atendimento.horario_triagem else 'Não realizado',
            'horario_consulta_medica': atendimento.horario_consulta_medica.strftime('%d/%m/%Y %H:%M') if atendimento.horario_consulta_medica else 'Não realizado',
            'horario_alta': atendimento.horario_alta.strftime('%d/%m/%Y %H:%M') if atendimento.horario_alta else 'Não realizado'
        }
        
        return jsonify({'success': True, 'atendimento': dados})
    except Exception as e:
        logging.error(f"Erro ao obter atendimento: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro ao obter dados do atendimento'}), 500

@bp.route('/administrador/atendimentos/<string:id>/alterar-status', methods=['PUT'])
@login_required
def alterar_status_atendimento(id):
    try:
        user = get_current_user()
        if not user or user.cargo.strip().lower() != 'administrador':
            return jsonify({'success': False, 'message': 'Acesso negado'}), 403
        
        atendimento = db.session.query(Atendimento).get(id)
        if not atendimento:
            return jsonify({'success': False, 'message': 'Atendimento não encontrado'}), 404
        
        data = request.get_json()
        novo_status = data.get('status', '').strip()
        motivo = data.get('motivo', '').strip()
        
        if not novo_status:
            return jsonify({'success': False, 'message': 'Status é obrigatório'}), 400
        
        # Registrar alteração no log
        status_anterior = atendimento.status
        atendimento.status = novo_status
        
        # Se o status for finalizado, registrar horário de alta se não tiver
        if novo_status in ['Alta', 'Evasão', 'Óbito', 'Transferência', 'Finalizado']:
            if not atendimento.horario_alta:
                # Horário do servidor menos 3 horas
                atendimento.horario_alta = datetime.utcnow() - timedelta(hours=3)
        
        db.session.commit()
        
        # Log completo da alteração
        logging.warning(f"STATUS ALTERADO MANUALMENTE - Atendimento: {id} | Paciente: {atendimento.paciente.nome} | "
                       f"Status: {status_anterior} → {novo_status} | "
                       f"Alterado por: {user.nome} (ID: {user.id}) | "
                       f"Motivo: {motivo or 'Não informado'}")
        
        return jsonify({
            'success': True, 
            'message': 'Status atualizado com sucesso',
            'novo_status': novo_status
        })
    except Exception as e:
        db.session.rollback()
        logging.error(f"Erro ao alterar status do atendimento: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro ao alterar status'}), 500

@bp.route('/administrador/relatorios/internacoes')
@login_required
def relatorio_internacoes():
    try:
        user = get_current_user()
        if not user or user.cargo.strip().lower() != 'administrador':
            flash('Acesso restrito a administradores.', 'danger')
            return redirect(url_for('main.index'))

        periodo = (request.args.get('periodo') or '7d').lower()
        inicio_param = request.args.get('inicio')
        fim_param = request.args.get('fim')

        agora_br = datetime.now(ZoneInfo("America/Sao_Paulo"))
        inicio = datetime(agora_br.year, agora_br.month, agora_br.day, tzinfo=ZoneInfo("America/Sao_Paulo"))
        fim = inicio + timedelta(days=1)

        if periodo == 'hoje':
            pass
        elif periodo == 'ontem':
            inicio = inicio - timedelta(days=1)
            fim = inicio + timedelta(days=1)
        elif periodo == '7d':
            inicio = inicio - timedelta(days=6)
            fim = agora_br.replace(tzinfo=ZoneInfo("America/Sao_Paulo")) + timedelta(days=1)
        elif periodo == '30d':
            inicio = inicio - timedelta(days=29)
            fim = agora_br.replace(tzinfo=ZoneInfo("America/Sao_Paulo")) + timedelta(days=1)
        elif periodo == 'custom' and inicio_param and fim_param:
            try:
                ini_date = datetime.strptime(inicio_param, '%Y-%m-%d').date()
                fim_date = datetime.strptime(fim_param, '%Y-%m-%d').date()
                inicio = datetime(ini_date.year, ini_date.month, ini_date.day, tzinfo=ZoneInfo("America/Sao_Paulo"))
                fim = datetime(fim_date.year, fim_date.month, fim_date.day, tzinfo=ZoneInfo("America/Sao_Paulo")) + timedelta(days=1)
            except ValueError:
                flash('Período personalizado inválido. Use YYYY-MM-DD.', 'warning')
        else:
            # fallback para 7 dias
            inicio = inicio - timedelta(days=6)
            fim = agora_br.replace(tzinfo=ZoneInfo("America/Sao_Paulo")) + timedelta(days=1)

        fim_exibicao = fim - timedelta(days=1)

        # KPIs
        # Taxa de ocupação
        soma_capacidade, soma_ocupacao = db.session.query(
            func.coalesce(func.sum(Leito.capacidade_maxima), 0),
            func.coalesce(func.sum(Leito.ocupacao_atual), 0)
        ).filter(Leito.status != 'Interditado').first()
        taxa_ocupacao = 0
        if soma_capacidade and soma_capacidade > 0:
            try:
                taxa_ocupacao = round((float(soma_ocupacao) / float(soma_capacidade)) * 100, 1)
            except Exception:
                taxa_ocupacao = 0

        # Altas no período
        total_altas = db.session.query(func.count(Internacao.id)).filter(
            Internacao.data_alta >= inicio,
            Internacao.data_alta < fim
        ).scalar() or 0

        # Internações iniciadas no período
        total_internacoes_periodo = db.session.query(func.count(Internacao.id)).filter(
            Internacao.data_internacao >= inicio,
            Internacao.data_internacao < fim
        ).scalar() or 0

        # Tempo médio de permanência (altas no período)
        tempo_medio_permanencia = 0
        try:
            media_segundos = db.session.query(
                func.avg(func.extract('epoch', Internacao.data_alta - Internacao.data_internacao))
            ).filter(
                Internacao.data_internacao.isnot(None),
                Internacao.data_alta.isnot(None),
                Internacao.data_alta >= inicio,
                Internacao.data_alta < fim
            ).scalar()
            if media_segundos is not None:
                tempo_medio_permanencia = round(float(media_segundos) / 86400.0, 1)
        except Exception:
            tempo_medio_permanencia = 0

        # Lista de pacientes atualmente internados
        internados_rows = db.session.query(Internacao, Paciente, Atendimento).join(
            Paciente, Internacao.paciente_id == Paciente.id
        ).join(
            Atendimento, Internacao.atendimento_id == Atendimento.id
        ).filter(
            Internacao.data_alta.is_(None)
        ).order_by(Internacao.data_internacao.desc()).all()

        def fmt_dt(dt):
            try:
                return dt.strftime('%d/%m/%Y %H:%M') if dt else ''
            except Exception:
                return ''

        lista_internados = []
        for it, pac, at in internados_rows:
            dias = 0
            try:
                base_fim = it.data_alta or agora_br
                if it.data_internacao:
                    dias = max(0, int((base_fim - it.data_internacao).total_seconds() // 86400))
            except Exception:
                dias = 0
            lista_internados.append({
                'atendimento_id': it.atendimento_id,
                'paciente': pac.nome,
                'leito': it.leito,
                'data_internacao': fmt_dt(it.data_internacao),
                'data_alta': fmt_dt(it.data_alta),
                'dias': dias,
                'status_atendimento': at.status
            })

        return render_template(
            'relatorio_internacoes.html',
            periodo=periodo,
            inicio=inicio,
            fim=fim,
            fim_exibicao=fim_exibicao,
            taxa_ocupacao=taxa_ocupacao,
            tempo_medio_permanencia=tempo_medio_permanencia,
            total_altas=total_altas,
            total_internacoes_periodo=total_internacoes_periodo,
            internados=lista_internados
        )
    except Exception as e:
        logging.error(f"Erro ao gerar Relatório de Internações: {str(e)}")
        logging.error(traceback.format_exc())
        flash('Erro ao gerar relatório.', 'danger')
        return redirect(url_for('main.relatorios_admin'))

@bp.route('/administrador/relatorios/visualizar')
@login_required
def visualizar_relatorio():
    try:
        user = get_current_user()
        if not user or user.cargo.strip().lower() != 'administrador':
            flash('Acesso restrito a administradores.', 'danger')
            return redirect(url_for('main.index'))

        # Obter parâmetros do formulário
        modelo = request.args.get('modelo', 'analitico')  # 'analitico' ou 'serie_historica'
        tipo_relatorio = request.args.get('tipo_relatorio', 'emergencia')  # 'emergencia', 'produtividade', 'observacao', 'perfil'
        data_inicio_str = request.args.get('data_inicio')
        data_fim_str = request.args.get('data_fim')
        
        # Processar datas
        agora_br = datetime.now(ZoneInfo("America/Sao_Paulo"))
        
        if data_inicio_str and data_fim_str:
            try:
                data_inicio = datetime.strptime(data_inicio_str, '%Y-%m-%d')
                data_fim = datetime.strptime(data_fim_str, '%Y-%m-%d')
            except ValueError:
                flash('Formato de data inválido', 'danger')
                return redirect(url_for('main.relatorios_admin'))
        else:
            # Padrão: últimos 7 dias
            data_fim = agora_br
            data_inicio = agora_br - timedelta(days=7)
        
        # Selecionar template apropriado
        templates_map = {
            'emergencia': 'relatorio_emergencia.html',
            'produtividade': 'relatorio_produtividade.html',
            'observacao': 'relatorio_observacao.html',
            'perfil': 'relatorio_perfil_paciente.html'
        }
        
        template_name = templates_map.get(tipo_relatorio, 'relatorio_exemplo.html')
        
        # Nomes dos relatórios
        nomes_relatorios = {
            'emergencia': 'Atendimento Emergência',
            'produtividade': 'Produtividade Clínica',
            'observacao': 'Observação',
            'perfil': 'Perfil de Paciente'
        }
        
        # Ícones dos relatórios
        icones_relatorios = {
            'emergencia': 'fas fa-ambulance',
            'produtividade': 'fas fa-hospital',
            'observacao': 'fas fa-eye',
            'perfil': 'fas fa-user-md'
        }
        
        # Preparar contexto base
        contexto = {
            'titulo_relatorio': f'Relatório de {nomes_relatorios.get(tipo_relatorio, "Dados")}',
            'modelo_relatorio': 'Analítico' if modelo == 'analitico' else 'Série Histórica',
            'periodo_inicio': data_inicio.strftime('%d/%m/%Y'),
            'periodo_fim': data_fim.strftime('%d/%m/%Y'),
            'tipo_relatorio_nome': nomes_relatorios.get(tipo_relatorio, 'Não definido'),
            'icone_relatorio': icones_relatorios.get(tipo_relatorio, 'fas fa-chart-bar'),
            'data_geracao': agora_br.strftime('%d/%m/%Y às %H:%M'),
            'ano_atual': agora_br.year,
            'mostrar_assinaturas': True,
            'assinatura_1_nome': 'Responsável Técnico',
            'assinatura_1_cargo': 'Diretor(a) Técnico(a)'
        }
        
        # Processar dados conforme o tipo de relatório
        if tipo_relatorio == 'emergencia':
            # Buscar atendimentos do período
            atendimentos = db.session.query(
                Atendimento, Paciente, Funcionario.nome.label('medico_nome'), 
                db.func.coalesce(Funcionario.nome, 'Não informado').label('enfermeiro_nome')
            ).join(
                Paciente, Atendimento.paciente_id == Paciente.id
            ).outerjoin(
                Funcionario, Atendimento.medico_id == Funcionario.id
            ).filter(
                Atendimento.data_atendimento >= data_inicio.date(),
                Atendimento.data_atendimento <= data_fim.date()
            ).order_by(Atendimento.data_atendimento.desc(), Atendimento.hora_atendimento.desc()).all()
            
            # Processar lista de atendimentos
            atendimentos_lista = []
            status_count = {}
            pacientes_set = set()
            
            for atend, pac, medico_nome, enfermeiro_nome in atendimentos:
                # Buscar enfermeiro
                enfermeiro_obj = None
                if atend.enfermeiro_id:
                    enfermeiro_obj = Funcionario.query.get(atend.enfermeiro_id)
                    enfermeiro_nome_final = enfermeiro_obj.nome if enfermeiro_obj else 'Não informado'
                else:
                    enfermeiro_nome_final = 'Não informado'
                
                # Formatar CPF
                cpf_formatado = pac.cpf or 'Não informado'
                if cpf_formatado and cpf_formatado != 'Não informado' and len(cpf_formatado) == 11:
                    cpf_formatado = f'{cpf_formatado[:3]}.{cpf_formatado[3:6]}.{cpf_formatado[6:9]}-{cpf_formatado[9:]}'
                
                # Data e hora formatada
                data_hora_str = ''
                if atend.data_atendimento and atend.hora_atendimento:
                    data_hora_str = f'{atend.data_atendimento.strftime("%d/%m/%Y")} {atend.hora_atendimento.strftime("%H:%M")}'
                elif atend.data_atendimento:
                    data_hora_str = atend.data_atendimento.strftime('%d/%m/%Y')
                else:
                    data_hora_str = 'Não informado'
                
                # Badge do status
                status = atend.status or 'Indefinido'
                badge_map = {
                    'aguardando': 'warning',
                    'triagem': 'primary',
                    'medico': 'primary',
                    'medicacao': 'warning',
                    'observacao': 'secondary',
                    'internado': 'danger',
                    'alta': 'success',
                    'evasao': 'danger',
                    'transferencia': 'warning'
                }
                
                badge_class = 'secondary'
                for key, value in badge_map.items():
                    if key in status.lower():
                        badge_class = value
                        break
                
                atendimentos_lista.append({
                    'numero': atend.id,
                    'cpf': cpf_formatado,
                    'paciente': pac.nome or 'Não informado',
                    'medico': medico_nome or 'Não informado',
                    'enfermeiro': enfermeiro_nome_final,
                    'status': status,
                    'badge_status': badge_class,
                    'data_hora': data_hora_str
                })
                
                # Contabilizar status
                status_count[status] = status_count.get(status, 0) + 1
                
                # Contabilizar pacientes únicos
                pacientes_set.add(pac.id)
            
            # Processar dados de status
            total_atend = len(atendimentos_lista)
            dados_status = []
            for status, qtd in sorted(status_count.items(), key=lambda x: x[1], reverse=True):
                percentual = round((qtd / total_atend * 100), 1) if total_atend > 0 else 0
                
                badge_class = 'secondary'
                for key, value in badge_map.items():
                    if key in status.lower():
                        badge_class = value
                        break
                
                dados_status.append({
                    'status': status,
                    'quantidade': qtd,
                    'percentual': percentual,
                    'badge_class': badge_class
                })
            
            # Calcular média diária
            dias_periodo = (data_fim.date() - data_inicio.date()).days + 1
            media_diaria = round(total_atend / dias_periodo, 1) if dias_periodo > 0 else 0
            
            # Calcular taxa de conclusão (atendimentos sem evasão)
            atendimentos_sem_evasao = sum(1 for item in atendimentos_lista if 'evasao' not in item['status'].lower() and 'evas' not in item['status'].lower())
            taxa_conclusao = f"{round((atendimentos_sem_evasao / total_atend * 100), 1)}%" if total_atend > 0 else '0%'
            
            # Processar dados diários para série histórica
            dados_diarios = []
            if modelo == 'serie_historica':
                from sqlalchemy import cast, Date, func
                from collections import defaultdict
                
                # Agrupar atendimentos por dia
                atendimentos_por_dia = defaultdict(lambda: {'total': 0, 'altas': 0, 'internacoes': 0, 'outros': 0})
                
                for atend, pac, medico_nome, enfermeiro_nome in atendimentos:
                    if atend.data_atendimento:
                        data_key = atend.data_atendimento.strftime('%d/%m/%Y')
                        atendimentos_por_dia[data_key]['total'] += 1
                        
                        status_lower = (atend.status or '').lower()
                        if 'alta' in status_lower:
                            atendimentos_por_dia[data_key]['altas'] += 1
                        elif 'internado' in status_lower or 'internacao' in status_lower or 'internamento' in status_lower:
                            atendimentos_por_dia[data_key]['internacoes'] += 1
                        else:
                            atendimentos_por_dia[data_key]['outros'] += 1
                
                # Gerar lista para todos os dias do período
                current_date = data_inicio.date()
                while current_date <= data_fim.date():
                    data_str = current_date.strftime('%d/%m/%Y')
                    dados_dia = atendimentos_por_dia.get(data_str, {'total': 0, 'altas': 0, 'internacoes': 0, 'outros': 0})
                    
                    dados_diarios.append({
                        'data': data_str,
                        'total': dados_dia['total'],
                        'altas': dados_dia['altas'],
                        'internacoes': dados_dia['internacoes'],
                        'outros': dados_dia['outros']
                    })
                    
                    current_date += timedelta(days=1)
            
            # Processar dados específicos do modo Analítico
            if modelo == 'analitico':
                from datetime import date as date_class
                from collections import defaultdict
                
                # 1. Produtividade Médica (médicos e quantidade de atendimentos)
                medico_count = {}
                for atend, pac, medico_nome, enfermeiro_nome in atendimentos:
                    medico = medico_nome or 'Não informado'
                    medico_count[medico] = medico_count.get(medico, 0) + 1
                
                dados_medicos = []
                for medico, qtd in sorted(medico_count.items(), key=lambda x: x[1], reverse=True):
                    percentual = round((qtd / total_atend * 100), 1) if total_atend > 0 else 0
                    media_dia = round(qtd / dias_periodo, 1) if dias_periodo > 0 else 0
                    dados_medicos.append({
                        'medico': medico,
                        'quantidade': qtd,
                        'percentual': percentual,
                        'media_dia': media_dia
                    })
                
                # 2. Distribuição por Gênero e Idade Média
                genero_data = defaultdict(lambda: {'count': 0, 'idades': []})
                idade_total = []
                
                for atend, pac, medico_nome, enfermeiro_nome in atendimentos:
                    # Calcular idade
                    if pac.data_nascimento:
                        hoje = date_class.today()
                        idade = hoje.year - pac.data_nascimento.year
                        if (hoje.month, hoje.day) < (pac.data_nascimento.month, pac.data_nascimento.day):
                            idade -= 1
                        idade_total.append(idade)
                        
                        # Agrupar por gênero
                        genero = 'Masculino' if pac.sexo and pac.sexo.upper() in ['M', 'MASCULINO'] else 'Feminino' if pac.sexo and pac.sexo.upper() in ['F', 'FEMININO'] else 'Não informado'
                        genero_data[genero]['count'] += 1
                        genero_data[genero]['idades'].append(idade)
                
                # Calcular idade média geral
                idade_media = round(sum(idade_total) / len(idade_total), 1) if idade_total else 0
                
                # Processar dados de gênero
                dados_genero = []
                for genero, data in genero_data.items():
                    qtd = data['count']
                    percentual = round((qtd / total_atend * 100), 1) if total_atend > 0 else 0
                    idade_media_genero = round(sum(data['idades']) / len(data['idades']), 1) if data['idades'] else 0
                    dados_genero.append({
                        'genero': genero,
                        'quantidade': qtd,
                        'percentual': percentual,
                        'idade_media': idade_media_genero
                    })
                
                # Ordenar por quantidade
                dados_genero.sort(key=lambda x: x['quantidade'], reverse=True)
                
                # 3. Tempo Médio de Consulta por Dia
                tempo_por_dia = defaultdict(lambda: [])
                
                for atend, pac, medico_nome, enfermeiro_nome in atendimentos:
                    if atend.data_atendimento and atend.hora_atendimento and atend.horario_consulta_medica:
                        # Combinar data e hora de atendimento
                        dt_entrada = datetime.combine(atend.data_atendimento, atend.hora_atendimento)
                        dt_consulta = atend.horario_consulta_medica
                        
                        # Calcular diferença em minutos
                        if isinstance(dt_consulta, datetime):
                            diferenca = (dt_consulta - dt_entrada).total_seconds() / 60
                            if diferenca >= 0:  # Apenas se a consulta foi após a entrada
                                data_key = atend.data_atendimento.strftime('%d/%m/%Y')
                                tempo_por_dia[data_key].append(diferenca)
                
                # Processar dados de tempo por dia
                dados_tempo_diario = []
                current_date = data_inicio.date()
                while current_date <= data_fim.date():
                    data_str = current_date.strftime('%d/%m/%Y')
                    tempos = tempo_por_dia.get(data_str, [])
                    
                    if tempos:
                        tempo_medio_min = sum(tempos) / len(tempos)
                        tempo_minimo_min = min(tempos)
                        tempo_maximo_min = max(tempos)
                        
                        # Converter para formato legível (horas e minutos)
                        def formatar_tempo(minutos):
                            horas = int(minutos // 60)
                            mins = int(minutos % 60)
                            if horas > 0:
                                return f'{horas}h{mins:02d}min'
                            else:
                                return f'{mins}min'
                        
                        dados_tempo_diario.append({
                            'data': data_str,
                            'quantidade': len(tempos),
                            'tempo_medio': formatar_tempo(tempo_medio_min),
                            'tempo_minimo': formatar_tempo(tempo_minimo_min),
                            'tempo_maximo': formatar_tempo(tempo_maximo_min)
                        })
                    
                    current_date += timedelta(days=1)
                
                # 4. Tempo Médio de Atendimento por Médico
                tempo_por_medico = defaultdict(lambda: [])
                
                for atend, pac, medico_nome, enfermeiro_nome in atendimentos:
                    if atend.hora_atendimento and atend.horario_consulta_medica:
                        dt_entrada = datetime.combine(atend.data_atendimento, atend.hora_atendimento) if atend.data_atendimento else None
                        dt_consulta = atend.horario_consulta_medica
                        
                        if dt_entrada and isinstance(dt_consulta, datetime):
                            diferenca = (dt_consulta - dt_entrada).total_seconds() / 60
                            if diferenca >= 0:
                                medico = medico_nome or 'Não informado'
                                tempo_por_medico[medico].append(diferenca)
                
                # Processar dados de tempo por médico
                dados_tempo_medico = []
                for medico, tempos in sorted(tempo_por_medico.items(), key=lambda x: len(x[1]), reverse=True):
                    if tempos:
                        tempo_medio_min = sum(tempos) / len(tempos)
                        tempo_minimo_min = min(tempos)
                        tempo_maximo_min = max(tempos)
                        
                        def formatar_tempo(minutos):
                            horas = int(minutos // 60)
                            mins = int(minutos % 60)
                            if horas > 0:
                                return f'{horas}h{mins:02d}min'
                            else:
                                return f'{mins}min'
                        
                        dados_tempo_medico.append({
                            'medico': medico,
                            'quantidade': len(tempos),
                            'tempo_medio': formatar_tempo(tempo_medio_min),
                            'tempo_minimo': formatar_tempo(tempo_minimo_min),
                            'tempo_maximo': formatar_tempo(tempo_maximo_min)
                        })
                
                # 5. Tempo Médio de Triagem por Enfermeiro
                tempo_triagem_por_enfermeiro = defaultdict(lambda: [])
                
                for atend, pac, medico_nome, enfermeiro_nome in atendimentos:
                    if atend.hora_atendimento and atend.horario_triagem:
                        dt_entrada = datetime.combine(atend.data_atendimento, atend.hora_atendimento) if atend.data_atendimento else None
                        dt_triagem = atend.horario_triagem
                        
                        if dt_entrada and isinstance(dt_triagem, datetime):
                            diferenca = (dt_triagem - dt_entrada).total_seconds() / 60
                            if diferenca >= 0:
                                # Buscar enfermeiro
                                enfermeiro = 'Não informado'
                                if atend.enfermeiro_id:
                                    enfermeiro_obj = Funcionario.query.get(atend.enfermeiro_id)
                                    if enfermeiro_obj:
                                        enfermeiro = enfermeiro_obj.nome
                                tempo_triagem_por_enfermeiro[enfermeiro].append(diferenca)
                
                # Processar dados de tempo de triagem por enfermeiro
                dados_tempo_enfermeiro = []
                for enfermeiro, tempos in sorted(tempo_triagem_por_enfermeiro.items(), key=lambda x: len(x[1]), reverse=True):
                    if tempos:
                        tempo_medio_min = sum(tempos) / len(tempos)
                        tempo_minimo_min = min(tempos)
                        tempo_maximo_min = max(tempos)
                        
                        def formatar_tempo_triagem(minutos):
                            horas = int(minutos // 60)
                            mins = int(minutos % 60)
                            if horas > 0:
                                return f'{horas}h{mins:02d}min'
                            else:
                                return f'{mins}min'
                        
                        dados_tempo_enfermeiro.append({
                            'enfermeiro': enfermeiro,
                            'quantidade': len(tempos),
                            'tempo_medio': formatar_tempo_triagem(tempo_medio_min),
                            'tempo_minimo': formatar_tempo_triagem(tempo_minimo_min),
                            'tempo_maximo': formatar_tempo_triagem(tempo_maximo_min)
                        })
                
                # 6. Distribuição Detalhada por Idade (0-100 anos)
                idade_count = {}
                for idade in range(0, 101):
                    idade_count[idade] = 0
                
                # Contar pacientes por idade
                for atend, pac, medico_nome, enfermeiro_nome in atendimentos:
                    if pac.data_nascimento:
                        hoje = date_class.today()
                        idade = hoje.year - pac.data_nascimento.year
                        if (hoje.month, hoje.day) < (pac.data_nascimento.month, pac.data_nascimento.day):
                            idade -= 1
                        
                        # Garantir que a idade está no range
                        if 0 <= idade <= 100:
                            idade_count[idade] += 1
                
                # Criar lista com todas as idades
                dados_idade_detalhada = []
                for idade in range(0, 101):
                    dados_idade_detalhada.append({
                        'idade': idade,
                        'quantidade': idade_count[idade]
                    })
                
                # Atualizar contexto com dados analíticos
                contexto.update({
                    'total_atendimentos': total_atend,
                    'media_diaria': media_diaria,
                    'pacientes_unicos': len(pacientes_set),
                    'taxa_conclusao': taxa_conclusao,
                    'dados_status': dados_status,
                    'idade_media': idade_media,
                    'dados_medicos': dados_medicos,
                    'dados_genero': dados_genero,
                    'dados_tempo_diario': dados_tempo_diario,
                    'dados_tempo_medico': dados_tempo_medico,
                    'dados_tempo_enfermeiro': dados_tempo_enfermeiro,
                    'dados_idade_detalhada': dados_idade_detalhada,
                    'dados_diarios': dados_diarios  # Para o gráfico
                })
            else:
                # Atualizar contexto para série histórica
                contexto.update({
                    'atendimentos_lista': atendimentos_lista,
                    'total_atendimentos': total_atend,
                    'media_diaria': media_diaria,
                    'pacientes_unicos': len(pacientes_set),
                    'taxa_conclusao': taxa_conclusao,  # Percentual sem evasão
                    'dados_status': dados_status,
                    'dados_diarios': dados_diarios if modelo == 'serie_historica' else None
                })
        
        elif tipo_relatorio == 'produtividade' and modelo == 'analitico':
            # RELATÓRIO DE PRODUTIVIDADE CLÍNICA - MODO ANALÍTICO
            from datetime import date as date_class
            from collections import defaultdict
            
            # Buscar internações do período
            internacoes = db.session.query(
                Internacao, Paciente, Atendimento, Funcionario.nome.label('medico_nome')
            ).join(
                Paciente, Internacao.paciente_id == Paciente.id
            ).join(
                Atendimento, Internacao.atendimento_id == Atendimento.id
            ).outerjoin(
                Funcionario, Atendimento.medico_id == Funcionario.id
            ).filter(
                Internacao.data_internacao >= data_inicio.date(),
                Internacao.data_internacao <= data_fim.date()
            ).all()
            
            total_internacoes = len(internacoes)
            dias_periodo = (data_fim.date() - data_inicio.date()).days + 1
            
            # 1. Internamentos por Médico
            medico_count = {}
            for intern, pac, atend, medico_nome in internacoes:
                medico = medico_nome or 'Não informado'
                medico_count[medico] = medico_count.get(medico, 0) + 1
            
            dados_internacoes_medico = []
            for medico, qtd in sorted(medico_count.items(), key=lambda x: x[1], reverse=True):
                percentual = round((qtd / total_internacoes * 100), 1) if total_internacoes > 0 else 0
                media_dia = round(qtd / dias_periodo, 1) if dias_periodo > 0 else 0
                dados_internacoes_medico.append({
                    'medico': medico,
                    'quantidade': qtd,
                    'percentual': percentual,
                    'media_dia': media_dia
                })
            
            # 2. Tempo Médio de Alta (permanência)
            tempos_permanencia = []
            total_altas = 0
            total_obitos = 0
            
            for intern, pac, atend, medico_nome in internacoes:
                if intern.data_alta and intern.data_internacao:
                    diferenca_dias = (intern.data_alta.date() - intern.data_internacao.date()).days
                    if diferenca_dias >= 0:
                        tempos_permanencia.append(diferenca_dias)
                        total_altas += 1
                
                # Contar óbitos (verificar se o status ou conduta indica óbito)
                if atend.status and 'obito' in atend.status.lower():
                    total_obitos += 1
                elif intern.relatorio_alta and 'obito' in intern.relatorio_alta.lower():
                    total_obitos += 1
            
            tempo_medio_alta = round(sum(tempos_permanencia) / len(tempos_permanencia), 1) if tempos_permanencia else 0
            
            # 3. Taxa de Mortalidade
            taxa_mortalidade = f'{round((total_obitos / total_internacoes * 100), 1)}%' if total_internacoes > 0 else '0%'
            
            # 4. Taxa de Ocupação de Leitos
            leitos = Leito.query.filter(Leito.status != 'Interditado').all()
            dados_leitos = []
            capacidade_total = 0
            ocupados_total = 0
            
            for leito in leitos:
                capacidade = leito.capacidade_maxima or 1
                ocupados = leito.ocupacao_atual or 0
                disponiveis = capacidade - ocupados
                taxa_ocp = round((ocupados / capacidade * 100), 1) if capacidade > 0 else 0
                
                badge_status = 'success' if leito.status == 'Disponível' else 'danger' if leito.status == 'Ocupado' else 'secondary'
                
                dados_leitos.append({
                    'nome': leito.nome,
                    'capacidade': capacidade,
                    'ocupados': ocupados,
                    'disponiveis': disponiveis,
                    'taxa_ocupacao': taxa_ocp,
                    'status': leito.status,
                    'badge_status': badge_status
                })
                
                capacidade_total += capacidade
                ocupados_total += ocupados
            
            disponiveis_total = capacidade_total - ocupados_total
            taxa_ocupacao_geral = round((ocupados_total / capacidade_total * 100), 1) if capacidade_total > 0 else 0
            
            # 5. Taxa de Readmissão (pacientes com alta que retornaram em até 7 dias)
            readmissoes = []
            pacientes_com_alta = {}
            
            # Primeiro, mapear todas as altas
            for intern, pac, atend, medico_nome in internacoes:
                if intern.data_alta:
                    paciente_id = pac.id
                    if paciente_id not in pacientes_com_alta:
                        pacientes_com_alta[paciente_id] = []
                    pacientes_com_alta[paciente_id].append({
                        'internacao': intern,
                        'paciente': pac,
                        'data_internacao': intern.data_internacao,
                        'data_alta': intern.data_alta
                    })
            
            # Verificar readmissões
            for paciente_id, historico in pacientes_com_alta.items():
                # Ordenar por data de internação
                historico_ordenado = sorted(historico, key=lambda x: x['data_internacao'])
                
                for i in range(len(historico_ordenado) - 1):
                    alta_atual = historico_ordenado[i]['data_alta']
                    proxima_internacao = historico_ordenado[i + 1]['data_internacao']
                    
                    if alta_atual and proxima_internacao:
                        diferenca = (proxima_internacao.date() - alta_atual.date()).days
                        
                        if 0 < diferenca <= 7:
                            pac = historico_ordenado[i]['paciente']
                            cpf_formatado = pac.cpf or 'Não informado'
                            if cpf_formatado and cpf_formatado != 'Não informado' and len(cpf_formatado) == 11:
                                cpf_formatado = f'{cpf_formatado[:3]}.{cpf_formatado[3:6]}.{cpf_formatado[6:9]}-{cpf_formatado[9:]}'
                            
                            readmissoes.append({
                                'paciente': pac.nome,
                                'cpf': cpf_formatado,
                                'data_primeira': historico_ordenado[i]['data_internacao'].strftime('%d/%m/%Y'),
                                'data_alta': alta_atual.strftime('%d/%m/%Y'),
                                'data_segunda': proxima_internacao.strftime('%d/%m/%Y'),
                                'dias_diferenca': diferenca
                            })
            
            total_readmissoes = len(readmissoes)
            taxa_readmissao = f'{round((total_readmissoes / total_altas * 100), 1)}%' if total_altas > 0 else '0%'
            
            # 6. CID Mais Comum
            cid_count = {}
            for intern, pac, atend, medico_nome in internacoes:
                cid = intern.cid_principal or 'Não informado'
                if cid and cid != 'Não informado':
                    cid_count[cid] = cid_count.get(cid, 0) + 1
            
            dados_cid = []
            for cid, qtd in sorted(cid_count.items(), key=lambda x: x[1], reverse=True)[:15]:  # Top 15
                percentual = round((qtd / total_internacoes * 100), 1) if total_internacoes > 0 else 0
                dados_cid.append({
                    'cid': cid,
                    'descricao': 'Descrição do CID',  # Pode adicionar lookup de descrição
                    'quantidade': qtd,
                    'percentual': percentual
                })
            
            # 7. Taxa por Bairro
            bairro_count = defaultdict(lambda: {'count': 0, 'municipio': ''})
            for intern, pac, atend, medico_nome in internacoes:
                bairro = pac.bairro or 'Não informado'
                municipio = pac.municipio or 'Não informado'
                bairro_count[bairro]['count'] += 1
                bairro_count[bairro]['municipio'] = municipio
            
            dados_bairro = []
            for bairro, data in sorted(bairro_count.items(), key=lambda x: x[1]['count'], reverse=True)[:20]:  # Top 20
                qtd = data['count']
                percentual = round((qtd / total_internacoes * 100), 1) if total_internacoes > 0 else 0
                dados_bairro.append({
                    'bairro': bairro,
                    'municipio': data['municipio'],
                    'quantidade': qtd,
                    'percentual': percentual
                })
            
            # Atualizar contexto com dados de produtividade clínica
            contexto.update({
                'total_internacoes': total_internacoes,
                'total_altas': total_altas,
                'tempo_medio_alta': tempo_medio_alta,
                'taxa_mortalidade': taxa_mortalidade,
                'taxa_readmissao': taxa_readmissao,
                'total_readmissoes': total_readmissoes,
                'taxa_ocupacao': f'{taxa_ocupacao_geral}%',
                'dados_internacoes_medico': dados_internacoes_medico,
                'dados_leitos': dados_leitos,
                'capacidade_total': capacidade_total,
                'ocupados_total': ocupados_total,
                'disponiveis_total': disponiveis_total,
                'taxa_ocupacao_geral': taxa_ocupacao_geral,
                'dados_cid': dados_cid,
                'dados_bairro': dados_bairro,
                'dados_readmissao': readmissoes
            })
        
        elif tipo_relatorio == 'produtividade' and modelo == 'serie_historica':
            # RELATÓRIO DE PRODUTIVIDADE CLÍNICA - MODO SÉRIE HISTÓRICA
            from datetime import date as date_class
            from collections import defaultdict
            
            # Buscar internações do período
            internacoes = db.session.query(
                Internacao, Paciente, Atendimento
            ).join(
                Paciente, Internacao.paciente_id == Paciente.id
            ).join(
                Atendimento, Internacao.atendimento_id == Atendimento.id
            ).filter(
                Internacao.data_internacao >= data_inicio.date(),
                Internacao.data_internacao <= data_fim.date()
            ).order_by(Internacao.data_internacao.desc()).all()
            
            total_internacoes = len(internacoes)
            dias_periodo = (data_fim.date() - data_inicio.date()).days + 1
            
            # Processar lista completa de internações
            internacoes_lista = []
            total_altas = 0
            total_obitos = 0
            tempo_permanencias = []
            status_count = {}
            
            for intern, pac, atend in internacoes:
                # Formatar CPF
                cpf_formatado = pac.cpf or 'Não informado'
                if cpf_formatado and cpf_formatado != 'Não informado' and len(cpf_formatado) == 11:
                    cpf_formatado = f'{cpf_formatado[:3]}.{cpf_formatado[3:6]}.{cpf_formatado[6:9]}-{cpf_formatado[9:]}'
                
                # Calcular dias de internação
                dias_internacao = '-'
                if intern.data_alta and intern.data_internacao:
                    diferenca = (intern.data_alta.date() - intern.data_internacao.date()).days
                    dias_internacao = diferenca
                    tempo_permanencias.append(diferenca)
                    total_altas += 1
                elif intern.data_internacao:
                    # Se ainda está internado, calcular até hoje
                    diferenca = (date_class.today() - intern.data_internacao.date()).days
                    dias_internacao = f'{diferenca}*'  # Asterisco indica ainda internado
                
                # Determinar status
                status = 'Internado'
                badge_status = 'danger'
                if intern.data_alta:
                    if atend.status and 'obito' in atend.status.lower():
                        status = 'Óbito'
                        badge_status = 'danger'
                        total_obitos += 1
                    elif intern.relatorio_alta and 'obito' in intern.relatorio_alta.lower():
                        status = 'Óbito'
                        badge_status = 'danger'
                        total_obitos += 1
                    elif atend.status and 'transferencia' in atend.status.lower():
                        status = 'Transferência'
                        badge_status = 'warning'
                    else:
                        status = 'Alta'
                        badge_status = 'success'
                
                # Contabilizar status
                status_count[status] = status_count.get(status, 0) + 1
                
                # Dados da internação
                internacoes_lista.append({
                    'numero': intern.atendimento_id,
                    'cpf': cpf_formatado,
                    'paciente': pac.nome or 'Não informado',
                    'cid': intern.cid_principal or '-',
                    'diagnostico': (intern.diagnostico_inicial or intern.diagnostico or '-')[:50] + ('...' if (intern.diagnostico_inicial or intern.diagnostico or '') and len(intern.diagnostico_inicial or intern.diagnostico or '') > 50 else ''),
                    'leito': intern.leito or '-',
                    'data_internacao': intern.data_internacao.strftime('%d/%m/%Y') if intern.data_internacao else '-',
                    'data_alta': intern.data_alta.strftime('%d/%m/%Y') if intern.data_alta else 'Internado',
                    'dias_internacao': dias_internacao,
                    'status': status,
                    'badge_status': badge_status
                })
            
            # Calcular métricas
            media_diaria = round(total_internacoes / dias_periodo, 1) if dias_periodo > 0 else 0
            tempo_medio_permanencia = round(sum(tempo_permanencias) / len(tempo_permanencias), 1) if tempo_permanencias else 0
            taxa_mortalidade = f'{round((total_obitos / total_internacoes * 100), 1)}%' if total_internacoes > 0 else '0%'
            
            # Taxa de ocupação
            leitos = Leito.query.filter(Leito.status != 'Interditado').all()
            capacidade_total = sum(l.capacidade_maxima or 1 for l in leitos)
            ocupados_total = sum(l.ocupacao_atual or 0 for l in leitos)
            taxa_ocupacao_geral = round((ocupados_total / capacidade_total * 100), 1) if capacidade_total > 0 else 0
            
            # Dados de status
            dados_status = []
            for status, qtd in sorted(status_count.items(), key=lambda x: x[1], reverse=True):
                percentual = round((qtd / total_internacoes * 100), 1) if total_internacoes > 0 else 0
                badge_map = {'Alta': 'success', 'Óbito': 'danger', 'Transferência': 'warning', 'Internado': 'secondary'}
                dados_status.append({
                    'status': status,
                    'quantidade': qtd,
                    'percentual': percentual,
                    'badge_class': badge_map.get(status, 'secondary')
                })
            
            # Dados diários
            dados_por_dia = defaultdict(lambda: {'internacoes': 0, 'altas': 0, 'transferencias': 0, 'obitos': 0})
            
            for intern, pac, atend in internacoes:
                if intern.data_internacao:
                    data_key = intern.data_internacao.strftime('%d/%m/%Y')
                    dados_por_dia[data_key]['internacoes'] += 1
                
                if intern.data_alta:
                    data_alta_key = intern.data_alta.strftime('%d/%m/%Y')
                    if atend.status and 'obito' in atend.status.lower():
                        dados_por_dia[data_alta_key]['obitos'] += 1
                    elif atend.status and 'transferencia' in atend.status.lower():
                        dados_por_dia[data_alta_key]['transferencias'] += 1
                    else:
                        dados_por_dia[data_alta_key]['altas'] += 1
            
            # Gerar lista para todos os dias do período
            dados_diarios = []
            current_date = data_inicio.date()
            while current_date <= data_fim.date():
                data_str = current_date.strftime('%d/%m/%Y')
                dados_dia = dados_por_dia.get(data_str, {'internacoes': 0, 'altas': 0, 'transferencias': 0, 'obitos': 0})
                saldo = dados_dia['internacoes'] - (dados_dia['altas'] + dados_dia['transferencias'] + dados_dia['obitos'])
                
                dados_diarios.append({
                    'data': data_str,
                    'internacoes': dados_dia['internacoes'],
                    'altas': dados_dia['altas'],
                    'transferencias': dados_dia['transferencias'],
                    'obitos': dados_dia['obitos'],
                    'saldo': saldo
                })
                
                current_date += timedelta(days=1)
            
            # Atualizar contexto para série histórica de produtividade
            contexto.update({
                'internacoes_lista': internacoes_lista,
                'total_internacoes': total_internacoes,
                'media_diaria': media_diaria,
                'total_altas': total_altas,
                'tempo_medio_permanencia': tempo_medio_permanencia,
                'taxa_mortalidade': taxa_mortalidade,
                'taxa_ocupacao': f'{taxa_ocupacao_geral}%',
                'dados_status': dados_status,
                'dados_diarios': dados_diarios
            })
        
        elif tipo_relatorio == 'observacao' and modelo == 'analitico':
            # RELATÓRIO DE OBSERVAÇÃO - MODO ANALÍTICO
            from datetime import date as date_class
            from collections import defaultdict
            
            # Buscar atendimentos que passaram pela observação
            atendimentos_obs = db.session.query(
                Atendimento, Paciente
            ).join(
                Paciente, Atendimento.paciente_id == Paciente.id
            ).filter(
                Atendimento.horario_observacao.isnot(None),
                Atendimento.data_atendimento >= data_inicio.date(),
                Atendimento.data_atendimento <= data_fim.date()
            ).all()
            
            total_observacao = len(atendimentos_obs)
            dias_periodo = (data_fim.date() - data_inicio.date()).days + 1
            
            # Processar dados
            tempos_observacao = []
            condutas_count = {}
            pacientes_internados_lista = []
            total_internados = 0
            total_altas = 0
            genero_data = defaultdict(lambda: {'count': 0, 'idades': []})
            idade_total = []
            cid_count = {}
            bairro_count = defaultdict(lambda: {'count': 0, 'municipio': ''})
            
            for atend, pac in atendimentos_obs:
                # 1. Calcular tempo em observação
                if atend.horario_observacao and atend.hora_atendimento and atend.data_atendimento:
                    dt_entrada = datetime.combine(atend.data_atendimento, atend.hora_atendimento)
                    dt_obs = atend.horario_observacao
                    
                    if isinstance(dt_obs, datetime):
                        # Calcular tempo até alta ou internação
                        dt_saida = None
                        if atend.horario_alta:
                            dt_saida = atend.horario_alta
                        elif atend.horario_internacao:
                            dt_saida = atend.horario_internacao
                        
                        if dt_saida and isinstance(dt_saida, datetime):
                            tempo_horas = (dt_saida - dt_obs).total_seconds() / 3600
                            if tempo_horas >= 0:
                                tempos_observacao.append(tempo_horas)
                
                # 2. Analisar conduta final
                conduta = atend.conduta_final or atend.status or 'Não informado'
                condutas_count[conduta] = condutas_count.get(conduta, 0) + 1
                
                # Verificar se foi internado
                if atend.horario_internacao or (atend.status and 'internado' in atend.status.lower()):
                    total_internados += 1
                    
                    # Buscar dados da internação
                    internacao = Internacao.query.filter_by(atendimento_id=atend.id).first()
                    cid = internacao.cid_principal if internacao else '-'
                    
                    # Formatar CPF
                    cpf_formatado = pac.cpf or 'Não informado'
                    if cpf_formatado and cpf_formatado != 'Não informado' and len(cpf_formatado) == 11:
                        cpf_formatado = f'{cpf_formatado[:3]}.{cpf_formatado[3:6]}.{cpf_formatado[6:9]}-{cpf_formatado[9:]}'
                    
                    # Calcular tempo em observação antes de internar
                    tempo_obs_str = '-'
                    if atend.horario_observacao and atend.horario_internacao:
                        if isinstance(atend.horario_observacao, datetime) and isinstance(atend.horario_internacao, datetime):
                            tempo_obs_horas = (atend.horario_internacao - atend.horario_observacao).total_seconds() / 3600
                            if tempo_obs_horas >= 0:
                                horas = int(tempo_obs_horas)
                                minutos = int((tempo_obs_horas % 1) * 60)
                                tempo_obs_str = f'{horas}h{minutos:02d}min'
                    
                    pacientes_internados_lista.append({
                        'numero': atend.id,
                        'paciente': pac.nome or 'Não informado',
                        'cpf': cpf_formatado,
                        'data_obs': atend.horario_observacao.strftime('%d/%m/%Y %H:%M') if isinstance(atend.horario_observacao, datetime) else '-',
                        'data_internacao': atend.horario_internacao.strftime('%d/%m/%Y %H:%M') if isinstance(atend.horario_internacao, datetime) else '-',
                        'tempo_obs': tempo_obs_str,
                        'cid': cid
                    })
                
                # Contar altas
                if atend.horario_alta or (atend.status and 'alta' in atend.status.lower()):
                    total_altas += 1
                
                # 3. Perfil - Gênero e Idade
                if pac.data_nascimento:
                    hoje = date_class.today()
                    idade = hoje.year - pac.data_nascimento.year
                    if (hoje.month, hoje.day) < (pac.data_nascimento.month, pac.data_nascimento.day):
                        idade -= 1
                    idade_total.append(idade)
                    
                    genero = 'Masculino' if pac.sexo and pac.sexo.upper() in ['M', 'MASCULINO'] else 'Feminino' if pac.sexo and pac.sexo.upper() in ['F', 'FEMININO'] else 'Não informado'
                    genero_data[genero]['count'] += 1
                    genero_data[genero]['idades'].append(idade)
                
                # 4. CID mais comum
                if atend.dx:
                    # Tentar extrair CID do campo dx
                    cid_texto = atend.dx.split()[0] if atend.dx else 'Não informado'
                    cid_count[cid_texto] = cid_count.get(cid_texto, 0) + 1
                
                # 5. Bairro
                bairro = pac.bairro or 'Não informado'
                municipio = pac.municipio or 'Não informado'
                bairro_count[bairro]['count'] += 1
                bairro_count[bairro]['municipio'] = municipio
            
            # Calcular métricas
            tempo_medio_observacao = round(sum(tempos_observacao) / len(tempos_observacao), 1) if tempos_observacao else 0
            taxa_internacao = f'{round((total_internados / total_observacao * 100), 1)}%' if total_observacao > 0 else '0%'
            taxa_alta = f'{round((total_altas / total_observacao * 100), 1)}%' if total_observacao > 0 else '0%'
            idade_media = round(sum(idade_total) / len(idade_total), 1) if idade_total else 0
            media_diaria = round(total_observacao / dias_periodo, 1) if dias_periodo > 0 else 0
            
            # Processar condutas
            dados_condutas = []
            for conduta, qtd in sorted(condutas_count.items(), key=lambda x: x[1], reverse=True):
                percentual = round((qtd / total_observacao * 100), 1) if total_observacao > 0 else 0
                
                # Calcular tempo médio para cada conduta
                tempos_conduta = []
                for atend, pac in atendimentos_obs:
                    if (atend.conduta_final or atend.status or 'Não informado') == conduta:
                        if atend.horario_observacao and atend.hora_atendimento and atend.data_atendimento:
                            dt_obs = atend.horario_observacao
                            dt_saida = atend.horario_alta or atend.horario_internacao
                            if dt_saida and isinstance(dt_obs, datetime) and isinstance(dt_saida, datetime):
                                tempo_h = (dt_saida - dt_obs).total_seconds() / 3600
                                if tempo_h >= 0:
                                    tempos_conduta.append(tempo_h)
                
                tempo_medio_conduta = round(sum(tempos_conduta) / len(tempos_conduta), 1) if tempos_conduta else 0
                
                badge_map = {
                    'alta': 'success',
                    'internado': 'danger',
                    'internacao': 'danger',
                    'transferencia': 'warning',
                    'evasao': 'danger',
                    'aguardando': 'secondary'
                }
                badge_class = 'secondary'
                for key, value in badge_map.items():
                    if key in conduta.lower():
                        badge_class = value
                        break
                
                dados_condutas.append({
                    'conduta': conduta,
                    'quantidade': qtd,
                    'percentual': percentual,
                    'tempo_medio': f'{tempo_medio_conduta}h' if tempo_medio_conduta > 0 else '-',
                    'badge_class': badge_class
                })
            
            # Processar dados de gênero
            dados_genero = []
            for genero, data in genero_data.items():
                qtd = data['count']
                percentual = round((qtd / total_observacao * 100), 1) if total_observacao > 0 else 0
                idade_media_genero = round(sum(data['idades']) / len(data['idades']), 1) if data['idades'] else 0
                dados_genero.append({
                    'genero': genero,
                    'quantidade': qtd,
                    'percentual': percentual,
                    'idade_media': idade_media_genero
                })
            dados_genero.sort(key=lambda x: x['quantidade'], reverse=True)
            
            # Processar faixas de tempo
            faixas_tempo = {
                'Menos de 6h': (0, 6),
                '6h a 12h': (6, 12),
                '12h a 24h': (12, 24),
                '24h a 48h': (24, 48),
                'Mais de 48h': (48, float('inf'))
            }
            dados_tempo_permanencia = []
            for faixa_nome, (min_h, max_h) in faixas_tempo.items():
                qtd = sum(1 for t in tempos_observacao if min_h <= t < max_h)
                percentual = round((qtd / len(tempos_observacao) * 100), 1) if tempos_observacao else 0
                dados_tempo_permanencia.append({
                    'faixa': faixa_nome,
                    'quantidade': qtd,
                    'percentual': percentual
                })
            
            # Processar CIDs
            dados_cid = []
            for cid, qtd in sorted(cid_count.items(), key=lambda x: x[1], reverse=True)[:15]:
                percentual = round((qtd / total_observacao * 100), 1) if total_observacao > 0 else 0
                dados_cid.append({
                    'cid': cid,
                    'descricao': 'Diagnóstico',
                    'quantidade': qtd,
                    'percentual': percentual
                })
            
            # Processar bairros
            dados_bairro = []
            for bairro, data in sorted(bairro_count.items(), key=lambda x: x[1]['count'], reverse=True)[:20]:
                qtd = data['count']
                percentual = round((qtd / total_observacao * 100), 1) if total_observacao > 0 else 0
                dados_bairro.append({
                    'bairro': bairro,
                    'municipio': data['municipio'],
                    'quantidade': qtd,
                    'percentual': percentual
                })
            
            # Atualizar contexto
            contexto.update({
                'total_observacao': total_observacao,
                'tempo_medio_observacao': tempo_medio_observacao,
                'taxa_internacao': taxa_internacao,
                'taxa_alta': taxa_alta,
                'idade_media': idade_media,
                'media_diaria': media_diaria,
                'dados_condutas': dados_condutas,
                'pacientes_internados': pacientes_internados_lista,
                'dados_genero': dados_genero,
                'dados_tempo_permanencia': dados_tempo_permanencia,
                'dados_cid': dados_cid,
                'dados_bairro': dados_bairro
            })
        
        logging.info(f'Relatório gerado: {tipo_relatorio} ({modelo}) de {data_inicio_str} a {data_fim_str} por {user.nome}')
        
        return render_template(template_name, **contexto)
        
    except Exception as e:
        logging.error(f"Erro ao visualizar relatório: {str(e)}")
        logging.error(traceback.format_exc())
        flash('Erro ao visualizar relatório.', 'danger')
        return redirect(url_for('main.relatorios_admin'))

@bp.route('/api/pacientes/internados')
@login_required
def listar_pacientes_internados():
    try:
        current_user = get_current_user()
        if current_user.cargo.lower() not in ['medico', 'enfermeiro', 'multi']:
            return jsonify({'error': 'Acesso não autorizado'}), 403
        
        pacientes_list = []
        
        # Buscar pacientes com status "Internado" (case-insensitive) na tabela atendimentos
        # Inclui tanto pacientes ainda internados quanto aqueles com alta mas com prontuário para fechar
        # Exclui APENAS pacientes com dieta = '1' (prontuário fechado)
        # Inclui pacientes com dieta NULL ou dieta começando com 'PENDENTE:' (alta definida mas prontuário não fechado)
        internacoes = db.session.query(Internacao).join(
            Atendimento, Internacao.atendimento_id == Atendimento.id
        ).filter(
            Atendimento.status.ilike('%internado%'),
            db.or_(Internacao.dieta.is_(None), Internacao.dieta != '1')
        ).all()
        
        logging.info(f"Total de internações encontradas: {len(internacoes)}")
        
        for internacao in internacoes:
            logging.info(f"Processando internação {internacao.atendimento_id}: dieta={internacao.dieta}, data_alta={internacao.data_alta}")
            paciente = Paciente.query.get(internacao.paciente_id)
            atendimento = Atendimento.query.get(internacao.atendimento_id)
            
            if paciente and atendimento:
                # Determinar se é um paciente com prontuário para fechar
                prontuario_para_fechar = internacao.data_alta is not None and atendimento.status == 'Internado'
                
                logging.info(f"Adicionando à lista: {paciente.nome} - Status: {atendimento.status}, Dieta: {internacao.dieta}")
                
                pacientes_list.append({
                    'atendimento_id': internacao.atendimento_id,
                    'nome': paciente.nome,
                    'cpf': paciente.cpf,
                    'data_nascimento': paciente.data_nascimento.strftime('%Y-%m-%d') if paciente.data_nascimento else None,
                    'leito': internacao.leito,
                    'data_internacao': internacao.data_internacao.strftime('%Y-%m-%d %H:%M') if internacao.data_internacao else None,
                    'data_alta': internacao.data_alta.strftime('%Y-%m-%d %H:%M') if internacao.data_alta else None,
                    'diagnostico': internacao.diagnostico,
                    'diagnostico_inicial': internacao.diagnostico_inicial,
                    'cid_principal': internacao.cid_principal,
                    'carater_internacao': internacao.carater_internacao,
                    'tem_alta': internacao.data_alta is not None,
                    'prontuario_para_fechar': prontuario_para_fechar,
                    'status_atendimento': atendimento.status
                })

        # Adicionar também atendimentos com status "Internado" que não possuem registro em Internacao
        ids_com_internacao = {i.atendimento_id for i in internacoes}
        atendimentos_sem_internacao = Atendimento.query.filter(
            Atendimento.status.ilike('%internado%'),
            ~Atendimento.id.in_(ids_com_internacao)
        ).all()

        for at in atendimentos_sem_internacao:
            pac = Paciente.query.get(at.paciente_id)
            if not pac:
                continue
            pacientes_list.append({
                'atendimento_id': at.id,
                'nome': pac.nome,
                'cpf': pac.cpf,
                'data_nascimento': pac.data_nascimento.strftime('%Y-%m-%d') if pac.data_nascimento else None,
                'leito': '-',
                'data_internacao': (at.horario_internacao.strftime('%Y-%m-%d %H:%M') if getattr(at, 'horario_internacao', None) else None),
                'data_alta': None,
                'diagnostico': 'Não informado',
                'diagnostico_inicial': None,
                'cid_principal': None,
                'carater_internacao': None,
                'tem_alta': False,
                'prontuario_para_fechar': False,
                'status_atendimento': at.status
            })
        
        return jsonify({
            'success': True,
            'pacientes': pacientes_list
        })
        
    except Exception as e:
        logging.error(f"Erro ao listar pacientes internados: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'error': 'Erro interno do servidor'}), 500
    
# API para obter o nome do enfermeiro logado
@bp.route('/api/enfermeiro/nome')
@login_required
def get_enfermeiro_nome():
    try:
        current_user = get_current_user()
        if current_user.cargo.lower() not in ['enfermeiro', 'multi']:
            return jsonify({'error': 'Acesso não autorizado'}), 403
        
        return jsonify({
            'nome': current_user.nome,
            'cargo': current_user.cargo
        })
        
    except Exception as e:
        logging.error(f"Erro ao obter nome do enfermeiro: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'error': 'Erro interno do servidor'}), 500


# Dashboard do enfermeiro: contagens rápidas
@bp.route('/api/enfermeiro/dashboard', methods=['GET'])
@login_required
def api_enfermeiro_dashboard():
    try:
        current_user = get_current_user()
        if current_user.cargo.lower() not in ['enfermeiro', 'multi']:
            return jsonify({'success': False, 'message': 'Acesso não autorizado'}), 403

        # Determinar janelas de plantão (07-19 e 19-07) na data atual
        now = datetime.now()
        today = now.date()
        seven_am = time(7, 0, 0)
        seven_pm = time(19, 0, 0)

        if now.time() >= seven_pm:
            start_shift = datetime.combine(today, seven_pm)
            end_shift = datetime.combine(today + timedelta(days=1), seven_am)
        elif now.time() >= seven_am:
            start_shift = datetime.combine(today, seven_am)
            end_shift = datetime.combine(today, seven_pm)
        else:
            start_shift = datetime.combine(today - timedelta(days=1), seven_pm)
            end_shift = datetime.combine(today, seven_am)

        # Contagem de internados ativos (exclui observação)
        internados_count = Internacao.query.filter(
            Internacao.data_alta.is_(None),
            ~Internacao.carater_internacao.ilike('%observa%')
        ).count()

        # Contagem de pacientes em observação (com base no status do atendimento)
        observacao_count = Atendimento.query.filter(
            Atendimento.status == 'Em Observação'
        ).count()

        # Contagem de triagens realizadas pelo enfermeiro no plantão atual
        triagens_count = Atendimento.query.filter(
            Atendimento.enfermeiro_id == current_user.id,
            Atendimento.horario_triagem.isnot(None),
            Atendimento.horario_triagem >= start_shift,
            Atendimento.horario_triagem < end_shift
        ).count()

        return jsonify({
            'success': True,
            'nome': current_user.nome,
            'internados': internados_count,
            'observacao': observacao_count,
            'triagens_plantao': triagens_count,
            'plantao_inicio': start_shift.isoformat(),
            'plantao_fim': end_shift.isoformat()
        })

    except Exception as e:
        logging.error(f"Erro no dashboard do enfermeiro: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro interno do servidor'}), 500


# Dashboard do médico: KPIs e plantão
@bp.route('/api/medico/dashboard', methods=['GET'])
@login_required
def api_medico_dashboard():
    try:
        current_user = get_current_user()
        if current_user.cargo.strip().lower() != 'medico':
            return jsonify({'success': False, 'message': 'Acesso não autorizado'}), 403

        # Janela de plantão (07-19 e 19-07) com base no horário atual
        now = datetime.now()
        today = now.date()
        seven_am = time(7, 0, 0)
        seven_pm = time(19, 0, 0)

        if now.time() >= seven_pm:
            start_shift = datetime.combine(today, seven_pm)
            end_shift = datetime.combine(today + timedelta(days=1), seven_am)
        elif now.time() >= seven_am:
            start_shift = datetime.combine(today, seven_am)
            end_shift = datetime.combine(today, seven_pm)
        else:
            start_shift = datetime.combine(today - timedelta(days=1), seven_pm)
            end_shift = datetime.combine(today, seven_am)

        # Internados ativos (exclui observação)
        internados_count = Internacao.query.filter(
            Internacao.data_alta.is_(None),
            ~Internacao.carater_internacao.ilike('%observa%')
        ).count()

        # Em observação (com base no status do atendimento)
        observacao_count = Atendimento.query.filter(
            Atendimento.status == 'Em Observação'
        ).count()

        # Consultas realizadas no plantão atual
        # Contabiliza eventos explícitos de início de consulta registrados no FluxoPaciente
        consultas_count = db.session.query(FluxoPaciente)\
            .filter(
                FluxoPaciente.id_medico == current_user.id,
                FluxoPaciente.mudanca_status.ilike('%INICIO CONSULTA%'),
                FluxoPaciente.mudanca_hora >= start_shift,
                FluxoPaciente.mudanca_hora < end_shift
            ).count()

        # Observação aguardando conduta (lista_observacao sem médico_conduta e sem data_saida)
        try:
            obs_aguardando_conduta = ListaObservacao.query.filter(
                ListaObservacao.data_saida.is_(None),
                (ListaObservacao.medico_conduta.is_(None) | (ListaObservacao.medico_conduta == ''))
            ).count()
        except Exception:
            obs_aguardando_conduta = 0

        return jsonify({
            'success': True,
            'nome': current_user.nome,
            'internados': internados_count,
            'observacao': observacao_count,
            'consultas_plantao': consultas_count,
            'obs_aguardando_conduta': obs_aguardando_conduta,
            'plantao_inicio': start_shift.isoformat(),
            'plantao_fim': end_shift.isoformat()
        })

    except Exception as e:
        logging.error(f"Erro no dashboard do médico: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro interno do servidor'}), 500

# Rota para impressão das evoluções de fisioterapia
@bp.route('/impressao_fisio/<string:atendimento_id>')
@login_required
def impressao_fisio(atendimento_id):
    """Página de impressão das evoluções de fisioterapia"""
    try:
        # Buscar o atendimento
        atendimento = Atendimento.query.get_or_404(atendimento_id)
        
        # Buscar as evoluções de fisioterapia
        evolucoes = EvolucaoFisioterapia.query.filter_by(id_atendimento=atendimento_id)\
            .order_by(EvolucaoFisioterapia.data_evolucao.desc()).all()
        
        # Buscar dados do paciente
        paciente = atendimento.paciente
        
        # Buscar dados da internação
        internacao = Internacao.query.filter_by(atendimento_id=atendimento_id).first()
        
        # Preparar dados das evoluções com informações do fisioterapeuta
        evolucoes_dados = []
        for evolucao in evolucoes:
            funcionario = None
            if evolucao.funcionario_id:
                funcionario = Funcionario.query.get(evolucao.funcionario_id)
            
            evolucoes_dados.append({
                'id': evolucao.id,
                'data_evolucao': evolucao.data_evolucao,
                'evolucao_fisio': evolucao.evolucao_fisio,
                'fisioterapeuta_nome': funcionario.nome if funcionario else 'Não identificado',
                'fisioterapeuta_registro': funcionario.numero_profissional if funcionario else ''
            })
        
        # Função para obter data atual
        def moment():
            return datetime.now()
        
        return render_template('impressao_fisio.html',
                               paciente=paciente,
                               internacao=internacao,
                               atendimento=atendimento,
                               evolucoes=evolucoes_dados,
                               moment=moment)
                               
    except Exception as e:
        flash(f'Erro ao carregar dados para impressão: {str(e)}', 'error')
        return redirect(url_for('main.index'))

# Rota para impressão das evoluções de nutrição
@bp.route('/impressao_nutricao/<string:atendimento_id>')
@login_required
def impressao_nutricao(atendimento_id):
    """Página de impressão das evoluções de nutrição"""
    try:
        # Buscar o atendimento
        atendimento = Atendimento.query.get_or_404(atendimento_id)
        
        # Buscar as evoluções de nutrição
        evolucoes = EvolucaoNutricao.query.filter_by(id_atendimento=atendimento_id)\
            .order_by(EvolucaoNutricao.data_evolucao.desc()).all()
        
        # Buscar dados do paciente
        paciente = atendimento.paciente
        
        # Buscar dados da internação
        internacao = Internacao.query.filter_by(atendimento_id=atendimento_id).first()
        
        # Preparar dados das evoluções com informações do nutricionista
        evolucoes_dados = []
        for evolucao in evolucoes:
            funcionario = None
            if evolucao.funcionario_id:
                funcionario = Funcionario.query.get(evolucao.funcionario_id)
            
            evolucoes_dados.append({
                'id': evolucao.id,
                'data_evolucao': evolucao.data_evolucao,
                'evolucao_nutricao': evolucao.evolucao_nutricao,
                'nutricionista_nome': funcionario.nome if funcionario else 'Não identificado',
                'nutricionista_registro': funcionario.numero_profissional if funcionario else ''
            })
        
        # Função para obter data atual
        def moment():
            return datetime.now()
        
        return render_template('impressao_nutricao.html',
                               paciente=paciente,
                               internacao=internacao,
                               atendimento=atendimento,
                               evolucoes=evolucoes_dados,
                               moment=moment)
                               
    except Exception as e:
        flash(f'Erro ao carregar dados para impressão: {str(e)}', 'error')
        return redirect(url_for('main.index'))

# Rota para impressão de evolução individual de fisioterapia
@bp.route('/impressao_fisio_evolucao/<int:evolucao_id>')
@login_required
def impressao_fisio_evolucao(evolucao_id):
    """Página de impressão de uma evolução individual de fisioterapia"""
    try:
        # Buscar a evolução específica
        evolucao = EvolucaoFisioterapia.query.get_or_404(evolucao_id)
        
        # Buscar o atendimento
        atendimento = Atendimento.query.get_or_404(evolucao.id_atendimento)
        
        # Buscar dados do paciente
        paciente = atendimento.paciente
        
        # Buscar dados da internação
        internacao = Internacao.query.filter_by(atendimento_id=evolucao.id_atendimento).first()
        
        # Buscar dados do fisioterapeuta
        funcionario = None
        if evolucao.funcionario_id:
            funcionario = Funcionario.query.get(evolucao.funcionario_id)
        
        # Preparar dados da evolução
        evolucao_dados = {
            'id': evolucao.id,
            'data_evolucao': evolucao.data_evolucao,
            'evolucao_fisio': evolucao.evolucao_fisio,
            'fisioterapeuta_nome': funcionario.nome if funcionario else 'Não identificado',
            'fisioterapeuta_registro': funcionario.numero_profissional if funcionario else ''
        }
        
        # Função para obter data atual
        def moment():
            return datetime.now()
        
        return render_template('impressao_fisio_individual.html',
                               paciente=paciente,
                               internacao=internacao,
                               atendimento=atendimento,
                               evolucao=evolucao_dados,
                               moment=moment)
                               
    except Exception as e:
        flash(f'Erro ao carregar dados para impressão: {str(e)}', 'error')
        return redirect(url_for('main.index'))

# Rota para impressão de evolução individual de nutrição
@bp.route('/impressao_nutricao_evolucao/<int:evolucao_id>')
@login_required
def impressao_nutricao_evolucao(evolucao_id):
    """Página de impressão de uma evolução individual de nutrição"""
    try:
        # Buscar a evolução específica
        evolucao = EvolucaoNutricao.query.get_or_404(evolucao_id)
        
        # Buscar o atendimento
        atendimento = Atendimento.query.get_or_404(evolucao.id_atendimento)
        
        # Buscar dados do paciente
        paciente = atendimento.paciente
        
        # Buscar dados da internação
        internacao = Internacao.query.filter_by(atendimento_id=evolucao.id_atendimento).first()
        
        # Buscar dados do nutricionista
        funcionario = None
        if evolucao.funcionario_id:
            funcionario = Funcionario.query.get(evolucao.funcionario_id)
        
        # Preparar dados da evolução
        evolucao_dados = {
            'id': evolucao.id,
            'data_evolucao': evolucao.data_evolucao,
            'evolucao_nutricao': evolucao.evolucao_nutricao,
            'nutricionista_nome': funcionario.nome if funcionario else 'Não identificado',
            'nutricionista_registro': funcionario.numero_profissional if funcionario else ''
        }
        
        # Função para obter data atual
        def moment():
            return datetime.now()
        
        return render_template('impressao_nutricao_individual.html',
                               paciente=paciente,
                               internacao=internacao,
                               atendimento=atendimento,
                               evolucao=evolucao_dados,
                               moment=moment)
                               
    except Exception as e:
        flash(f'Erro ao carregar dados para impressão: {str(e)}', 'error')
        return redirect(url_for('main.index'))

# Rota principal
@bp.route('/')
def index():
    try:
        return render_template('index.html')
    except Exception as e:
        logging.error(f"Erro na rota index: {str(e)}")
        logging.error(traceback.format_exc())
        return f"Erro ao renderizar página: {str(e)}", 500

# Rota principal da clínica
@bp.route('/clinica')
@login_required
def clinica():
    try:
        current_user = get_current_user()
        if current_user.cargo.lower() not in ['medico', 'enfermeiro', 'multi']:
            flash('Acesso restrito a médicos, enfermeiros e profissionais multi.', 'danger')
            return redirect(url_for('main.index'))
        
        return render_template('clinica.html')
        
    except Exception as e:
        logging.error(f"Erro ao acessar página da clínica: {str(e)}")
        logging.error(traceback.format_exc())
        flash('Erro ao acessar a página. Por favor, tente novamente.', 'danger')
        return redirect(url_for('main.index'))

# Rota para listar pacientes internados
@bp.route('/clinica/pacientes-internados')
@login_required
def pacientes_internados():
    try:
        current_user = get_current_user()
        if current_user.cargo.lower() not in ['medico', 'enfermeiro', 'multi']:
            flash('Acesso restrito a médicos, enfermeiros e profissionais multi.', 'danger')
            return redirect(url_for('main.index'))
        
        return render_template('pacientes_internados.html')
        
    except Exception as e:
        logging.error(f"Erro ao acessar lista de pacientes internados: {str(e)}")
        logging.error(traceback.format_exc())
        flash('Erro ao acessar a lista de pacientes. Por favor, tente novamente.', 'danger')
        return redirect(url_for('main.clinica'))

# Rota de login com tratamento de erro mais robusto

@bp.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        try:
            # Log para diagnóstico
            logging.info("Recebendo requisição de login")
            
            # Obter dados do formulário
            cpf = request.form.get('cpf')
            senha = request.form.get('senha')
            
            if not cpf or not senha:
                logging.warning("Tentativa de login sem CPF ou senha")
                return jsonify({
                    'success': False,
                    'error': 'CPF e senha são obrigatórios.'
                }), 400
            
            # Log para diagnóstico
            logging.info(f"Tentativa de login para CPF: {cpf}")
            
            # Buscar funcionário pelo CPF
            funcionario = Funcionario.query.filter_by(cpf=cpf).first()
            
            if not funcionario:
                logging.warning(f"Login falhou: CPF {cpf} não encontrado")
                return jsonify({
                    'success': False,
                    'error': 'CPF ou senha inválidos.'
                })
            
            # Verificar a senha
            if not funcionario.check_password(senha):
                logging.warning(f"Login falhou: Senha incorreta para CPF {cpf}")
                return jsonify({
                    'success': False,
                    'error': 'CPF ou senha inválidos.'
                })
            
            # Login bem-sucedido
            login_user(funcionario)
            
            # Armazenar dados na sessão
            session['cargo'] = funcionario.cargo
            session['user_id'] = funcionario.id
            
            # Log para diagnóstico
            logging.info(f"Usuário {funcionario.nome} (ID: {funcionario.id}) com cargo {funcionario.cargo} logado com sucesso")
            
            return jsonify({
                'success': True,
                'cargo': funcionario.cargo,
                'funcionario_id': funcionario.id
            })
            
        except Exception as e:
            # Log detalhado do erro
            logging.error(f"Erro no processamento do login: {str(e)}")
            logging.error(traceback.format_exc())
            
            # Retornar resposta de erro
            return jsonify({
                'success': False,
                'error': 'Erro interno do servidor. Por favor, tente novamente.'
            }), 500
    else:
        return render_template('login.html')


# API para registrar SAE
@bp.route('/api/enfermagem/sae', methods=['POST'])
@login_required
def registrar_sae():
    try:
        # Verificar se o usuário é enfermeiro ou multi
        current_user = get_current_user()
        if current_user.cargo.lower() not in ['enfermeiro', 'multi']:
            return jsonify({
                'success': False,
                'message': 'Apenas enfermeiros podem registrar SAE'
            }), 403
        
        dados = request.get_json()
        
        # Validar dados obrigatórios
        campos_obrigatorios = [
            'paciente_id',
            'hipotese_diagnostica',
            'pa', 'fc', 'sat', 'dx', 'r', 't',
            'medicacao', 'alergias', 'antecedentes_pessoais',
            'sistema_neurologico', 'estado_geral', 'ventilacao',
            'diagnostico_de_enfermagem', 'pele',
            'sistema_gastrointerstinal', 'regulacao_vascular',
            'pulso', 'regulacao_abdominal', 'rha',
            'sistema_urinario', 'acesso_venoso', 'observacao'
        ]
        
        for campo in campos_obrigatorios:
            if campo not in dados:
                return jsonify({
                    'success': False,
                    'message': f'Campo obrigatório ausente: {campo}'
                }), 400
        
        # Criar nova SAE com timestamp atual
        from datetime import datetime, timezone, timedelta
        
        nova_sae = InternacaoSae(
            paciente_id=dados['paciente_id'],
            enfermeiro_id=current_user.id,
            data_registro=datetime.now(ZoneInfo("America/Sao_Paulo")),  # Timestamp único para cada SAE
            hipotese_diagnostica=dados['hipotese_diagnostica'],
            pa=dados['pa'],
            fc=dados['fc'],
            sat=dados['sat'],
            dx=dados['dx'],
            r=dados['r'],
            t=dados['t'],
            medicacao=dados['medicacao'],
            alergias=dados['alergias'],
            antecedentes_pessoais=dados['antecedentes_pessoais'],
            sistema_neurologico=dados['sistema_neurologico'],
            estado_geral=dados['estado_geral'],
            ventilacao=dados['ventilacao'],
            diagnostico_de_enfermagem=dados['diagnostico_de_enfermagem'],
            pele=dados['pele'],
            sistema_gastrointerstinal=dados['sistema_gastrointerstinal'],
            regulacao_vascular=dados['regulacao_vascular'],
            pulso=dados['pulso'],
            regulacao_abdominal=dados['regulacao_abdominal'],
            rha=dados['rha'],
            sistema_urinario=dados['sistema_urinario'],
            acesso_venoso=dados['acesso_venoso'],
            observacao=dados['observacao']
        )
        
        db.session.add(nova_sae)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': 'SAE registrada com sucesso',
            'id': nova_sae.id
        }), 201
        
    except Exception as e:
        db.session.rollback()
        logging.error(f'Erro ao registrar SAE: {str(e)}')
        logging.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': f'Erro interno do servidor: {str(e)}'
        }), 500

# API para listar evoluções de enfermagem por data (individual)
@bp.route('/api/enfermagem/evolucoes/lista/<string:atendimento_id>', methods=['GET'])
@login_required
def listar_evolucoes_enfermagem_por_data(atendimento_id):
    """
    Lista as evoluções de enfermagem de um atendimento em uma data específica.
    Retorna metadados para seleção e impressão individual.
    """
    try:
        current_user = get_current_user()
        if current_user.cargo.lower() not in ['enfermeiro', 'medico']:
            return jsonify({
                'success': False,
                'message': 'Acesso permitido apenas para enfermeiros e médicos'
            }), 403

        data_str = request.args.get('data')
        if not data_str:
            return jsonify({'success': False, 'message': 'Parâmetro "data" é obrigatório (YYYY-MM-DD)'}), 400

        try:
            data_selecionada = datetime.strptime(data_str, '%Y-%m-%d').date()
        except ValueError:
            return jsonify({'success': False, 'message': 'Formato de data inválido. Use YYYY-MM-DD.'}), 400

        internacao = Internacao.query.filter_by(atendimento_id=atendimento_id).first()
        if not internacao:
            return jsonify({'success': False, 'message': 'Internação não encontrada'}), 404

        inicio_dia = datetime.combine(data_selecionada, datetime.min.time())
        fim_dia = datetime.combine(data_selecionada, datetime.max.time())

        evolucoes = EvolucaoEnfermagem.query \
            .filter_by(atendimentos_clinica_id=internacao.id) \
            .filter(EvolucaoEnfermagem.data_evolucao >= inicio_dia) \
            .filter(EvolucaoEnfermagem.data_evolucao <= fim_dia) \
            .join(Funcionario, EvolucaoEnfermagem.funcionario_id == Funcionario.id) \
            .add_columns(
                EvolucaoEnfermagem.id,
                EvolucaoEnfermagem.data_evolucao,
                EvolucaoEnfermagem.texto,
                Funcionario.nome.label('enfermeiro_nome'),
                Funcionario.numero_profissional.label('enfermeiro_coren')
            ) \
            .order_by(EvolucaoEnfermagem.data_evolucao.asc()) \
            .all()

        lista = []
        for ev in evolucoes:
            preview_texto = ''
            if ev.texto:
                preview_texto = ev.texto.strip()
                if len(preview_texto) > 150:
                    preview_texto = preview_texto[:150] + '…'

            lista.append({
                'id': ev.id,
                'data_hora': ev.data_evolucao.strftime('%Y-%m-%d %H:%M') if ev.data_evolucao else None,
                'data_display': ev.data_evolucao.strftime('%d/%m/%Y %H:%M') if ev.data_evolucao else None,
                'enfermeiro_nome': ev.enfermeiro_nome,
                'enfermeiro_coren': ev.enfermeiro_coren,
                'preview': preview_texto
            })

        return jsonify({'success': True, 'evolucoes': lista}), 200

    except Exception as e:
        logging.error(f'Erro ao listar evoluções de enfermagem por data: {str(e)}')
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': f'Erro interno do servidor: {str(e)}'}), 500

# API para listar admissões de enfermagem de uma internação
@bp.route('/api/enfermagem/admissoes/<int:internacao_id>', methods=['GET'])
@login_required
def listar_admissoes_enfermagem_old(internacao_id):  # Renomeada para evitar conflito
    """
    Lista todas as admissões de enfermagem de uma internação.
    ---
    tags:
      - Enfermagem
    parameters:
      - name: internacao_id
        in: path
        type: integer
        required: true
        description: ID da internação
    responses:
      200:
        description: Lista de admissões
      404:
        description: Internação não encontrada
      500:
        description: Erro interno do servidor
    """
    try:
        # Verifica se a internação existe
        internacao = Internacao.query.get(internacao_id)
        if not internacao:
            return jsonify({
                'status': 'erro',
                'mensagem': 'Internação não encontrada'
            }), 404
        
        # Busca todas as admissões dessa internação
        admissoes = AdmissaoEnfermagem.query.filter_by(internacao_id=internacao_id).all()
        
        # Formata a resposta
        resultado = []
        for admissao in admissoes:
            enfermeiro = Funcionario.query.get(admissao.enfermeiro_id)
            resultado.append({
                'id': admissao.id,
                'data_hora': (admissao.data_hora - timedelta(hours=3)).strftime('%d/%m/%Y %H:%M'),
                'enfermeiro': enfermeiro.nome if enfermeiro else 'Não informado',
                'admissao_texto': admissao.admissao_texto
            })
        
        return jsonify({
            'status': 'sucesso',
            'admissoes': resultado
        }), 200
        
    except Exception as e:
        current_app.logger.error(f'Erro ao listar admissões de enfermagem: {str(e)}')
        return jsonify({
            'status': 'erro',
            'mensagem': 'Erro interno do servidor'
        }), 500

@bp.route('/api/medicamentos', methods=['GET'])
@login_required
def api_listar_medicamentos():
    try:
        # Verificar autenticação
        if 'user_id' not in session:
            return jsonify({'success': False, 'message': 'Usuário não autenticado'}), 401

        # Buscar medicamentos a partir dos aprazamentos
        aprazamentos = Aprazamento.query.all()

        result = []
        nomes_unicos = set()

        for apr in aprazamentos:
            chave = (apr.nome_medicamento.strip().lower(), (apr.descricao_uso or '').strip().lower())
            if chave not in nomes_unicos:
                nomes_unicos.add(chave)
                result.append({
                    'nome_medicamento': apr.nome_medicamento,
                    'descricao_uso': apr.descricao_uso or '',
                    'aprazamento': '',  
                    'enfermeiro_nome': ''  
                })

        return jsonify({
            'success': True,
            'medicamentos': result
        })

    except Exception as e:
        logging.error(f"Erro ao buscar medicamentos: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'error': 'Erro ao buscar medicamentos'
        }), 500


@bp.route('/api/prescricoes/<int:internacao_id>', methods=['GET'])
@login_required
def listar_prescricoes(internacao_id):
    """
    Lista todas as prescrições médicas de uma internação.
    Inclui aprazamentos da nova tabela Aprazamento.
    """
    try:
        internacao = Internacao.query.get(internacao_id)
        if not internacao:
            return jsonify({
                'success': False,
                'message': 'Internação não encontrada'
            }), 404

        prescricoes = PrescricaoClinica.query.filter_by(atendimentos_clinica_id=internacao_id).order_by(PrescricaoClinica.horario_prescricao.desc()).all()

        # ------- Função para corrigir o formato do aprazamento -------
        from datetime import date

        def corrigir_aprazamento(aprazamento_raw):
            if not aprazamento_raw:
                return ""
            aprazamento_str = str(aprazamento_raw).strip()

            padrao_completo = r'^\d{2}/\d{2}/\d{4}\s*:\s*\d{2}:\d{2}(?:\s*,\s*\d{2}:\d{2})*(?:\s*;\s*\d{2}/\d{2}/\d{4}\s*:\s*\d{2}:\d{2}(?:\s*,\s*\d{2}:\d{2})*)*$'
            padrao_apenas_horarios = r'^\d{2}:\d{2}(?:\s*,\s*\d{2}:\d{2})*$'

            if re.match(padrao_completo, aprazamento_str):
                return aprazamento_str

            if re.match(padrao_apenas_horarios, aprazamento_str):
                hoje = date.today()
                return f"{hoje.day:02d}/{hoje.month:02d}/{hoje.year}: {aprazamento_str}"

            # Corrigir formatos quebrados
            partes = aprazamento_str.split(':')
            if len(partes) > 2:
                data_match = re.match(r'^(\d{2}/\d{2}/\d{4})', aprazamento_str)
                if data_match:
                    data = data_match.group(1)
                    horarios_match = re.findall(r'(\d{2}:\d{2})(?:,\s*|$)', aprazamento_str)
                    if horarios_match:
                        return f"{data}: {', '.join(horarios_match)}"

            return aprazamento_str

        # --------------------------------------------------------------

        resultado = []
        for prescricao in prescricoes:
            medico = Funcionario.query.get(prescricao.medico_id)
            enfermeiro = Funcionario.query.get(prescricao.enfermeiro_id) if prescricao.enfermeiro_id else None

            # ----- PROCESSAR MEDICAMENTOS -----
            medicamentos = prescricao.medicamentos_json or []
            if isinstance(medicamentos, str):
                import json
                try:
                    medicamentos = json.loads(medicamentos)
                except:
                    medicamentos = []

            medicamentos_formatados = []
            for med in medicamentos:
                aprazamento_corrigido = corrigir_aprazamento(med.get('aprazamento'))
                
                # Obter aprazamentos da nova tabela para este medicamento
                aprazamentos_novos = []
                if med.get('nome_medicamento'):
                    aprazamentos_db = Aprazamento.query.filter_by(
                        prescricao_id=prescricao.id, 
                        nome_medicamento=med.get('nome_medicamento')
                    ).order_by(Aprazamento.data_hora_aprazamento).all()
                    
                    for apz in aprazamentos_db:
                        enfermeiro_resp = Funcionario.query.get(apz.enfermeiro_responsavel_id) if apz.enfermeiro_responsavel_id else None
                        
                        aprazamentos_novos.append({
                            'id': apz.id,
                            'data_hora': apz.data_hora_aprazamento.isoformat(),
                            'realizado': apz.realizado,
                            'enfermeiro_nome': enfermeiro_resp.nome if enfermeiro_resp else None,
                            'data_realizacao': apz.data_realizacao.isoformat() if apz.data_realizacao else None
                        })
                
                medicamentos_formatados.append({
                    'nome_medicamento': med.get('nome_medicamento'),
                    'descricao_uso': med.get('descricao_uso'),
                    'aprazamentos_novos': aprazamentos_novos  # Novos aprazamentos da tabela
                })

            # ----- BUSCAR APRAZAMENTOS GERAIS DA PRESCRIÇÃO -----
            aprazamentos_gerais = []
            aprazamentos_db = Aprazamento.query.filter_by(
                prescricao_id=prescricao.id
            ).order_by(Aprazamento.data_hora_aprazamento).all()
            
            for apz in aprazamentos_db:
                enfermeiro_resp = Funcionario.query.get(apz.enfermeiro_responsavel_id) if apz.enfermeiro_responsavel_id else None
                
                # Verificar se já existe na lista de medicamentos
                ja_incluido = False
                for med in medicamentos_formatados:
                    if med['nome_medicamento'] == apz.nome_medicamento:
                        ja_incluido = True
                        break
                
                # Se não foi incluído nos medicamentos específicos, adicionar aos gerais
                if not ja_incluido:
                    aprazamentos_gerais.append({
                        'id': apz.id,
                        'nome_medicamento': apz.nome_medicamento,
                        'descricao_uso': apz.descricao_uso,
                        'data_hora': apz.data_hora_aprazamento.isoformat(),
                        'realizado': apz.realizado,
                        'enfermeiro_nome': enfermeiro_resp.nome if enfermeiro_resp else None,
                        'data_realizacao': apz.data_realizacao.isoformat() if apz.data_realizacao else None
                    })

            # ----- MONTAR RESULTADO -----
            resultado.append({
                'id': prescricao.id,
                'data_prescricao': prescricao.horario_prescricao.isoformat() if prescricao.horario_prescricao else None,
                'medico_nome': medico.nome if medico else 'Não informado',
                'enfermeiro_nome': enfermeiro.nome if enfermeiro else None,
                'texto_dieta': prescricao.texto_dieta,
                'texto_procedimento_medico': prescricao.texto_procedimento_medico,
                'texto_procedimento_multi': prescricao.texto_procedimento_multi,
                'aprazamentos_gerais': aprazamentos_gerais,  # Novos aprazamentos da tabela
                'medicamentos': medicamentos_formatados
            })

        return jsonify({
            'success': True,
            'prescricoes': resultado
        }), 200
    except Exception as e:
        logging.error(f'Erro ao listar prescrições: {str(e)}')
        logging.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': 'Erro ao listar prescrições',
            'error': str(e)
        }), 500


# API para registrar ou atualizar prescrição
@bp.route('/api/prescricoes', methods=['POST'])
@login_required
def registrar_prescricao():
    """
    Registra uma nova prescrição médica de forma simplificada.
    - Salva os medicamentos no campo JSON da tabela prescricoes_clinica
    - Cria registros em FluxoDisp para cada medicamento (status pendente) para a Farmácia
    """
    try:
        current_user = get_current_user()
        dados = request.get_json()
        
        # Flexibilizar criação: aceitar atendimento_id OU atendimentos_clinica_id
        internacao = None
        internacao_id = dados.get('atendimentos_clinica_id')
        atendimento_id_payload = dados.get('atendimento_id')

        if internacao_id:
            internacao = Internacao.query.get(internacao_id)
            if not internacao:
                return jsonify({
                    'success': False,
                    'message': 'Internação não encontrada'
                }), 404
        else:
            # Criar/garantir uma internação vinculada ao atendimento, se vier atendimento_id
            if not atendimento_id_payload:
                return jsonify({
                    'success': False,
                    'message': 'Informe atendimento_id ou atendimentos_clinica_id'
                }), 400
            atendimento_ref = Atendimento.query.get(atendimento_id_payload)
            if not atendimento_ref:
                return jsonify({'success': False, 'message': 'Atendimento não encontrado'}), 404
            # Buscar internação existente para este atendimento
            internacao = Internacao.query.filter_by(atendimento_id=atendimento_ref.id).first()
            if not internacao:
                # Criar internação mínima para habilitar prescrições
                internacao = Internacao(
                    atendimento_id=atendimento_ref.id,
                    paciente_id=atendimento_ref.paciente_id,
                    medico_id=current_user.id,
                    data_internacao=datetime.now(ZoneInfo("America/Sao_Paulo"))
                )
                db.session.add(internacao)
                db.session.flush()
        
        # Montar lista simples de medicamentos (nome + descricao_uso)
        medicamentos_simples = []
        meds_payload = dados.get('medicamentos')
        if meds_payload is None and isinstance(dados.get('medicamentos_json'), list):
            meds_payload = dados.get('medicamentos_json')

        if isinstance(meds_payload, list):
            for med in meds_payload:
                nome = (med or {}).get('nome_medicamento')
                desc = (med or {}).get('descricao_uso')
                if nome:
                    medicamentos_simples.append({
                        'nome_medicamento': nome,
                        'descricao_uso': desc or ''
                    })

        nova_prescricao = PrescricaoClinica(
            atendimentos_clinica_id=internacao.id,
            medico_id=current_user.id,
            texto_dieta=dados.get('texto_dieta'),
            texto_procedimento_medico=dados.get('texto_procedimento_medico'),
            texto_procedimento_multi=dados.get('texto_procedimento_multi'),
            horario_prescricao=datetime.now(ZoneInfo("America/Sao_Paulo"))
        )

        if medicamentos_simples:
            nova_prescricao.medicamentos_json = medicamentos_simples

        db.session.add(nova_prescricao)
        db.session.flush()

        # Criar registros no FluxoDisp para a Farmácia (um por medicamento)
        try:
            for med in (medicamentos_simples or []):
                nome_medicamento = med.get('nome_medicamento')
                if not nome_medicamento:
                    continue
                novo_fluxo = FluxoDisp(
                    id_atendimento=internacao.atendimento_id,
                    id_medico=current_user.id,
                    id_responsavel=0,  # Definido pela Farmácia ao dispensar
                    id_prescricao=nova_prescricao.id,
                    hora=datetime.now(ZoneInfo("America/Sao_Paulo")).time(),
                    data=datetime.now(ZoneInfo("America/Sao_Paulo")).date(),
                    medicamento=nome_medicamento,
                    quantidade=0  # Sem quantidade definida neste fluxo simples
                )
                db.session.add(novo_fluxo)
        except Exception as e:
            logging.error(f"Erro ao criar FluxoDisp na criação da prescrição: {str(e)}")

        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': 'Prescrição registrada com sucesso',
            'id': nova_prescricao.id,
            'atendimentos_clinica_id': internacao.id
        }), 201

    except Exception as e:
        db.session.rollback()
        logging.error(f'Erro ao registrar prescrição: {str(e)}')
        logging.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': 'Erro interno do servidor'
        }), 500


# API para atualizar prescrição existente
@bp.route('/api/prescricoes/<int:prescricao_id>', methods=['PUT'])
@login_required
def atualizar_prescricao(prescricao_id):
    """
    Atualiza uma prescrição médica existente.
    Agora também atualiza os registros na tabela Aprazamento quando os medicamentos são alterados.
    """
    try:
        current_user = get_current_user()
        dados = request.get_json()
        
        prescricao = PrescricaoClinica.query.get(prescricao_id)
        if not prescricao:
            return jsonify({
                'success': False,
                'message': 'Prescrição não encontrada'
            }), 404
        
        if 'texto_dieta' in dados:
            prescricao.texto_dieta = dados['texto_dieta']
        if 'texto_procedimento_medico' in dados:
            prescricao.texto_procedimento_medico = dados['texto_procedimento_medico']
        if 'texto_procedimento_multi' in dados:
            prescricao.texto_procedimento_multi = dados['texto_procedimento_multi']
        
        if 'medicamentos' in dados:
            # Obter lista anterior de medicamentos para comparar
            medicamentos_anteriores = prescricao.medicamentos_json
            
            # Atualizar medicamentos na prescrição
            prescricao.medicamentos_json = dados['medicamentos']
            
            # Processar aprazamentos atualizados
            if isinstance(dados['medicamentos'], list):
                # Manter um registro dos medicamentos que ainda existem
                nomes_medicamentos_atuais = set()
                
                for medicamento in dados['medicamentos']:
                    nome_medicamento = medicamento.get('nome_medicamento', '')
                    nomes_medicamentos_atuais.add(nome_medicamento)
                    
                    # Verificar se existe informação de aprazamento
                    if medicamento.get('aprazamento'):
                        # Se for um medicamento novo ou se o aprazamento foi alterado, atualizar
                        medicamento_anterior = None
                        if isinstance(medicamentos_anteriores, list):
                            for med_ant in medicamentos_anteriores:
                                if med_ant.get('nome_medicamento') == nome_medicamento:
                                    medicamento_anterior = med_ant
                                    break
                        
                        aprazamento_alterado = (
                            not medicamento_anterior or 
                            medicamento_anterior.get('aprazamento') != medicamento.get('aprazamento')
                        )
                        
                        if aprazamento_alterado:
                            # Remover aprazamentos anteriores deste medicamento
                            Aprazamento.query.filter_by(
                                prescricao_id=prescricao_id, 
                                nome_medicamento=nome_medicamento
                            ).delete()
                            
                            try:
                                # Processar o texto de aprazamento para criar registros individuais
                                aprazamento_texto = medicamento.get('aprazamento')
                                
                                # Formato esperado: "DD/MM/YYYY: HH:MM, HH:MM; DD/MM/YYYY: HH:MM, HH:MM"
                                dias = aprazamento_texto.split(';')
                                
                                for dia in dias:
                                    dia = dia.strip()
                                    if not dia:
                                        continue
                                        
                                    partes = dia.split(':')
                                    if len(partes) < 2:
                                        continue
                                        
                                    data_str = partes[0].strip()
                                    horarios_str = ':'.join(partes[1:]).strip()  # Reconectar caso haja múltiplos ":" após o primeiro
                                    
                                    horarios = [h.strip() for h in horarios_str.split(',')]
                                    
                                    try:
                                        # Converter data no formato DD/MM/YYYY para objeto Date
                                        dia, mes, ano = map(int, data_str.split('/'))
                                        
                                        for horario in horarios:
                                            if not horario:
                                                continue
                                                
                                            # Converter horário no formato HH:MM para objeto Time
                                            hora, minuto = map(int, horario.split(':'))
                                            
                                            # Criar o objeto DateTime combinando a data e o horário
                                            data_hora = datetime(ano, mes, dia, hora, minuto)
                                            
                                            # Criar registro de Aprazamento
                                            novo_aprazamento = Aprazamento(
                                                prescricao_id=prescricao_id,
                                                nome_medicamento=nome_medicamento,
                                                descricao_uso=medicamento.get('descricao_uso', ''),
                                                data_hora_aprazamento=data_hora,
                                                enfermeiro_responsavel_id=dados.get('enfermeiro_id')
                                            )
                                            
                                            db.session.add(novo_aprazamento)
                                            
                                    except Exception as e:
                                        logging.error(f"Erro ao processar aprazamento para a data {data_str}: {str(e)}")
                                        continue
                            except Exception as e:
                                logging.error(f"Erro ao processar aprazamento do medicamento {nome_medicamento}: {str(e)}")
                                continue
                
                # Remover aprazamentos de medicamentos que não existem mais
                if isinstance(medicamentos_anteriores, list):
                    for med_ant in medicamentos_anteriores:
                        nome_ant = med_ant.get('nome_medicamento', '')
                        if nome_ant and nome_ant not in nomes_medicamentos_atuais:
                            # Este medicamento foi removido, então devemos remover seus aprazamentos
                            Aprazamento.query.filter_by(
                                prescricao_id=prescricao_id, 
                                nome_medicamento=nome_ant
                            ).delete()

        # Atualizar registros no FluxoDisp quando medicamentos são alterados
        if 'medicamentos' in dados and isinstance(dados['medicamentos'], list):
            # Remover registros FluxoDisp dos medicamentos que não existem mais
            if isinstance(medicamentos_anteriores, list):
                nomes_medicamentos_anteriores = set(med_ant.get('nome_medicamento', '') for med_ant in medicamentos_anteriores)
                nomes_medicamentos_atuais = set(med.get('nome_medicamento', '') for med in dados['medicamentos'])

                # Remover FluxoDisp dos medicamentos removidos
                for nome_ant in nomes_medicamentos_anteriores:
                    if nome_ant and nome_ant not in nomes_medicamentos_atuais:
                        FluxoDisp.query.filter_by(
                            id_prescricao=prescricao_id,
                            medicamento=nome_ant
                        ).delete()

            # Criar ou atualizar registros FluxoDisp para medicamentos atuais
            for medicamento in dados['medicamentos']:
                try:
                    nome_medicamento = medicamento.get('nome_medicamento', '')
                    quantidade = medicamento.get('quantidade', 0)

                    # Verificar se já existe registro FluxoDisp para este medicamento nesta prescrição
                    fluxo_existente = FluxoDisp.query.filter_by(
                        id_prescricao=prescricao_id,
                        medicamento=nome_medicamento
                    ).first()

                    if fluxo_existente:
                        # Atualizar registro existente
                        fluxo_existente.quantidade = quantidade
                        fluxo_existente.hora = datetime.now(ZoneInfo("America/Sao_Paulo")).time()
                        fluxo_existente.data = datetime.now(ZoneInfo("America/Sao_Paulo")).date()
                    else:
                        # Criar novo registro FluxoDisp
                        internacao = prescricao.internacao
                        novo_fluxo = FluxoDisp(
                            id_atendimento=internacao.atendimento_id,
                            id_medico=current_user.id,
                            id_responsavel=0,  # Será definido quando dispensado
                            id_prescricao=prescricao_id,
                            hora=datetime.now(ZoneInfo("America/Sao_Paulo")).time(),
                            data=datetime.now(ZoneInfo("America/Sao_Paulo")).date(),
                            medicamento=nome_medicamento,
                            quantidade=quantidade
                        )
                        db.session.add(novo_fluxo)

                except Exception as e:
                    logging.error(f"Erro ao atualizar FluxoDisp para medicamento {medicamento.get('nome_medicamento')}: {str(e)}")
                    continue

        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': 'Prescrição atualizada com sucesso'
        }), 200

    except Exception as e:
        db.session.rollback()
        logging.error(f'Erro ao atualizar prescrição: {str(e)}')
        logging.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': 'Erro interno do servidor'
        }), 500


@bp.route('/clinica/evolucao-paciente-enfermeiro/<string:atendimento_id>')
@login_required
def evolucao_paciente_enfermeiro(atendimento_id):
    try:
        # Verificar se o usuário é enfermeiro ou multi
        current_user = get_current_user()
        if current_user.cargo.lower() not in ['enfermeiro', 'multi']:
            flash('Acesso restrito a enfermeiros.', 'danger')
            return redirect(url_for('main.index'))
        
        # Buscar dados do paciente e internação
        internacao = Internacao.query.filter_by(atendimento_id=atendimento_id).first()
        if not internacao:
            flash('Internação não encontrada.', 'danger')
            return redirect(url_for('main.pacientes_internados'))
        
        paciente = Paciente.query.get(internacao.paciente_id)
        if not paciente:
            flash('Paciente não encontrado.', 'danger')
            return redirect(url_for('main.pacientes_internados'))
        
        # Buscar o atendimento para ter acesso às alergias
        atendimento = Atendimento.query.get(atendimento_id)
        
        return render_template('clinica_evolucao_paciente_enfermeiro.html', 
                            paciente=paciente, 
                            internacao=internacao,
                            atendimento=atendimento)
        
    except Exception as e:
        logging.error(f"Erro ao acessar evolução do paciente (enfermeiro): {str(e)}")
        logging.error(traceback.format_exc())
        flash('Erro ao acessar a evolução do paciente. Por favor, tente novamente.', 'danger')
        return redirect(url_for('main.pacientes_internados'))

@bp.route('/clinica/evolucao-paciente-multi/<string:atendimento_id>')
@login_required
def evolucao_paciente_multi(atendimento_id):
    try:
        # Verificar se o usuário é multi
        current_user = get_current_user()
        if current_user.cargo.lower() != 'multi':
            flash('Acesso restrito a profissionais multi.', 'danger')
            return redirect(url_for('main.index'))
        
        # Buscar dados do paciente e internação
        internacao = Internacao.query.filter_by(atendimento_id=atendimento_id).first()
        if not internacao:
            flash('Internação não encontrada.', 'danger')
            return redirect(url_for('main.pacientes_internados'))
        
        paciente = Paciente.query.get(internacao.paciente_id)
        if not paciente:
            flash('Paciente não encontrado.', 'danger')
            return redirect(url_for('main.pacientes_internados'))
        
        return render_template('clinica_evolucao_paciente_multi.html', 
                            paciente=paciente, 
                            internacao=internacao)
        
    except Exception as e:
        logging.error(f"Erro ao acessar evolução do paciente (multi): {str(e)}")
        logging.error(traceback.format_exc())
        flash('Erro ao acessar a evolução do paciente. Por favor, tente novamente.', 'danger')
        return redirect(url_for('main.pacientes_internados'))

@bp.route('/clinica/evolucao-paciente-medico/<string:atendimento_id>')
@login_required
def evolucao_paciente_medico(atendimento_id):
    try:
        # Verificar se o usuário é médico
        current_user = get_current_user()
        if current_user.cargo.lower() != 'medico':
            flash('Acesso restrito a médicos.', 'danger')
            return redirect(url_for('main.index'))
        
        # Buscar dados do paciente e internação
        internacao = Internacao.query.filter_by(atendimento_id=atendimento_id).first()
        if not internacao:
            flash('Internação não encontrada.', 'danger')
            return redirect(url_for('main.pacientes_internados'))
        
        paciente = Paciente.query.get(internacao.paciente_id)
        if not paciente:
            flash('Paciente não encontrado.', 'danger')
            return redirect(url_for('main.pacientes_internados'))
        
        # Procurar se este paciente é responsável por algum RN e, se sim, montar info para link rápido
        rn_ativo_info = None
        try:
            rns_responsavel = PacienteRN.query.filter_by(responsavel_id=paciente.id).all()
            for rn_rel in rns_responsavel:
                paciente_rn = Paciente.query.get(rn_rel.paciente_id)
                if not paciente_rn:
                    continue
                # Preferir uma internação ativa (sem alta). Se não houver, usar a mais recente
                internacao_rn_ativa = (
                    Internacao.query
                    .filter_by(paciente_id=paciente_rn.id, data_alta=None)
                    .order_by(Internacao.id.desc())
                    .first()
                )
                internacao_rn = internacao_rn_ativa or (
                    Internacao.query
                    .filter_by(paciente_id=paciente_rn.id)
                    .order_by(Internacao.id.desc())
                    .first()
                )
                if internacao_rn:
                    rn_ativo_info = {
                        'nome': paciente_rn.nome,
                        'atendimento_id': internacao_rn.atendimento_id,
                    }
                    break
        except Exception as e:
            logging.error(f"Erro ao buscar RN vinculado ao responsável {paciente.id}: {str(e)}")
            logging.error(traceback.format_exc())
        
        return render_template('clinica_evolucao_paciente_medico.html', 
                            paciente=paciente, 
                            internacao=internacao,
                            rn_ativo_info=rn_ativo_info)
        
    except Exception as e:
        logging.error(f"Erro ao acessar evolução do paciente (médico): {str(e)}")
        logging.error(traceback.format_exc())
        flash('Erro ao acessar a evolução do paciente. Por favor, tente novamente.', 'danger')
        return redirect(url_for('main.pacientes_internados'))

## Evoluçao:

@bp.route('/api/evolucoes/<int:internacao_id>', methods=['GET'])
def get_evolucoes(internacao_id):
    try:
        # Buscar todas as evoluções dessa internação
        evolucoes = EvolucaoAtendimentoClinica.query \
            .filter_by(atendimentos_clinica_id=internacao_id) \
            .join(Funcionario, EvolucaoAtendimentoClinica.funcionario_id == Funcionario.id) \
            .add_columns(
                EvolucaoAtendimentoClinica.id,
                EvolucaoAtendimentoClinica.data_evolucao,
                EvolucaoAtendimentoClinica.evolucao,
                Funcionario.nome.label('nome_medico')
            ) \
            .order_by(EvolucaoAtendimentoClinica.data_evolucao.desc()) \
            .all()

        evolucoes_list = []
        for ev in evolucoes:
            evolucao_dict = {
                'id': ev.id,
                'data_evolucao': ev.data_evolucao.isoformat() if ev.data_evolucao else '',
                'evolucao': ev.evolucao or '',
                'nome_medico': ev.nome_medico or ''
            }
            evolucoes_list.append(evolucao_dict)

        return jsonify({
            'success': True,
            'evolucoes': evolucoes_list
        })

    except Exception as e:
        print('Erro ao buscar evoluções:', str(e))
        return jsonify({
            'success': False,
            'message': 'Erro ao buscar evoluções',
            'error': str(e)
        }), 500
    

@bp.route('/api/evolucoes/registrar', methods=['POST'])
def registrar_evolucao():
    try:
        data = request.get_json()

        # Aceitar tanto internacao_id quanto atendimentos_clinica_id
        atendimentos_clinica_id = data.get('atendimentos_clinica_id')
        if not atendimentos_clinica_id:
            atendimentos_clinica_id = data.get('internacao_id')
        
        funcionario_id = data.get('funcionario_id')
        evolucao_texto = data.get('evolucao')

        if not atendimentos_clinica_id or not funcionario_id or not evolucao_texto:
            return jsonify({
                'success': False,
                'message': 'Campos obrigatórios não foram preenchidos.'
            }), 400

        # Logging para debug
        print(f"Registrando evolução - atendimentos_clinica_id: {atendimentos_clinica_id}, funcionario_id: {funcionario_id}")
        
        nova_evolucao = EvolucaoAtendimentoClinica(
            atendimentos_clinica_id=atendimentos_clinica_id,
            funcionario_id=funcionario_id,
            data_evolucao=datetime.now(ZoneInfo("America/Sao_Paulo")),
            evolucao=evolucao_texto
        )

        db.session.add(nova_evolucao)
        db.session.commit()

        return jsonify({
            'success': True,
            'message': 'Evolução registrada com sucesso.'
        })

    except Exception as e:
        print('Erro ao registrar evolução:', str(e))
        db.session.rollback()
        return jsonify({
            'success': False,
            'message': 'Erro ao registrar evolução.',
            'error': str(e)
        }), 500


    
@bp.route('/api/enfermagem/evolucao', methods=['GET'])
def listar_evolucoes_enfermagem():
    evolucoes = EvolucaoEnfermagem.query.all()
    resultado = []
    for evolucao in evolucoes:
        resultado.append({
            'id': evolucao.id,
            'atendimentos_clinica_id': evolucao.atendimentos_clinica_id,
            'funcionario_id': evolucao.funcionario_id,
            'data_evolucao': evolucao.data_evolucao.isoformat(),
            'texto': evolucao.texto
        })
    return jsonify(resultado), 200

# Adicionar rota POST para registrar evolução de enfermagem
@bp.route('/api/enfermagem/evolucao', methods=['POST'])
def registrar_evolucao_enfermagem():
    data = request.get_json()
    if not data:
        return jsonify({'erro': 'Dados não fornecidos'}), 400

    try:
        nova_evolucao = EvolucaoEnfermagem(
            atendimentos_clinica_id=data['atendimentos_clinica_id'],
            funcionario_id=data['funcionario_id'],
            texto=data['texto'],
            data_evolucao=datetime.now(ZoneInfo("America/Sao_Paulo"))
        )
        db.session.add(nova_evolucao)
        db.session.commit()
        
        return jsonify({
            'mensagem': 'Evolução de enfermagem registrada com sucesso!',
            'evolucao_id': nova_evolucao.id
        }), 201
        
    except Exception as e:
        db.session.rollback()
        logging.error(f'Erro ao registrar evolução de enfermagem: {str(e)}')
        logging.error(traceback.format_exc())
        return jsonify({'erro': f'Erro ao registrar evolução: {str(e)}'}), 500


@bp.route('/api/enfermagem/admissao', methods=['POST'])
def salvar_admissao_enfermagem():
    """
    Registra uma nova admissão de enfermagem usando a tabela específica AdmissaoEnfermagem.
    """
    try:
        # Verificar se o usuário é enfermeiro ou multi
        current_user = get_current_user()
        if current_user.cargo.lower() not in ['enfermeiro', 'multi']:
            return jsonify({
                'success': False,
                'message': 'Apenas enfermeiros podem registrar admissões'
            }), 403

        dados = request.get_json()
        
        if not dados:
            return jsonify({
                'success': False,
                'message': 'Dados não fornecidos'
            }), 400

        internacao_id = dados.get('internacao_id')
        admissao_texto = dados.get('admissao_texto')

        if not internacao_id or not admissao_texto:
            return jsonify({
                'success': False,
                'message': 'ID da internação e texto da admissão são obrigatórios'
            }), 400

        # Verificar se a internação existe
        internacao = Internacao.query.get(internacao_id)
        if not internacao:
            return jsonify({
                'success': False,
                'message': 'Internação não encontrada'
            }), 404

        # Criar nova admissão de enfermagem
        nova_admissao = AdmissaoEnfermagem(
            internacao_id=internacao_id,
            enfermeiro_id=current_user.id,
            admissao_texto=admissao_texto,
            data_hora=datetime.now(ZoneInfo("America/Sao_Paulo"))
        )
        
        db.session.add(nova_admissao)

        # Também atualizar o campo legado na tabela Internacao (para compatibilidade)
        internacao.admissao_enfermagem = admissao_texto

        db.session.commit()

        return jsonify({
            'success': True,
            'message': 'Admissão de enfermagem registrada com sucesso',
            'admissao_id': nova_admissao.id
        }), 201

    except Exception as e:
        db.session.rollback()
        logging.error(f'Erro ao salvar admissão de enfermagem: {str(e)}')
        logging.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': f'Erro ao registrar admissão: {str(e)}'
        }), 500

# API para buscar admissões de enfermagem de uma internação
@bp.route('/api/enfermagem/admissao/<int:internacao_id>', methods=['GET'])
@login_required
def buscar_admissao_enfermagem(internacao_id):
    """
    Busca a admissão de enfermagem mais recente de uma internação.
    """
    try:
        # Verificar se o usuário é médico, enfermeiro ou multi
        current_user = get_current_user()
        if current_user.cargo.lower() not in ['medico', 'enfermeiro', 'multi']:
            return jsonify({
                'success': False,
                'message': 'Acesso permitido apenas para médicos e enfermeiros'
            }), 403

        # Verificar se a internação existe
        internacao = Internacao.query.get(internacao_id)
        if not internacao:
            return jsonify({
                'success': False,
                'message': 'Internação não encontrada'
            }), 404

        # Buscar a admissão mais recente desta internação
        admissao = AdmissaoEnfermagem.query.filter_by(
            internacao_id=internacao_id
        ).order_by(
            AdmissaoEnfermagem.data_hora.desc()
        ).first()

        if admissao:
            # Buscar dados do enfermeiro
            enfermeiro = Funcionario.query.get(admissao.enfermeiro_id) if admissao.enfermeiro_id else None
            
            return jsonify({
                'success': True,
                'admissao': {
                    'id': admissao.id,
                    'admissao_texto': admissao.admissao_texto,
                    'data_hora': (admissao.data_hora - timedelta(hours=3)).strftime('%d/%m/%Y %H:%M'),
                    'enfermeiro_nome': enfermeiro.nome if enfermeiro else 'Não informado'
                }
            })
        else:
            # Se não há admissão na tabela específica, verificar campo legado
            if internacao.admissao_enfermagem:
                return jsonify({
                    'success': True,
                    'admissao': {
                        'id': None,
                        'admissao_texto': internacao.admissao_enfermagem,
                        'data_hora': 'Data não registrada',
                        'enfermeiro_nome': 'Não informado'
                    }
                })
            else:
                return jsonify({
                    'success': False,
                    'message': 'Nenhuma admissão de enfermagem encontrada'
                }), 404

    except Exception as e:
        logging.error(f'Erro ao buscar admissão de enfermagem: {str(e)}')
        logging.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': f'Erro ao buscar admissão: {str(e)}'
        }), 500

# API para listar todas as admissões de enfermagem de uma internação
@bp.route('/api/enfermagem/admissoes/<int:internacao_id>', methods=['GET'])
@login_required
def listar_admissoes_enfermagem(internacao_id):
    """
    Lista todas as admissões de enfermagem de uma internação.
    """
    try:
        # Verificar se o usuário é médico, enfermeiro ou multi
        current_user = get_current_user()
        if current_user.cargo.lower() not in ['medico', 'enfermeiro', 'multi']:
            return jsonify({
                'success': False,
                'message': 'Acesso permitido apenas para médicos e enfermeiros'
            }), 403

        # Verificar se a internação existe
        internacao = Internacao.query.get(internacao_id)
        if not internacao:
            return jsonify({
                'success': False,
                'message': 'Internação não encontrada'
            }), 404
        
        # Buscar todas as admissões dessa internação
        admissoes = AdmissaoEnfermagem.query.filter_by(
            internacao_id=internacao_id
        ).order_by(
            AdmissaoEnfermagem.data_hora.desc()
        ).all()
        
        # Formatar a resposta
        resultado = []
        for admissao in admissoes:
            enfermeiro = Funcionario.query.get(admissao.enfermeiro_id) if admissao.enfermeiro_id else None
            resultado.append({
                'id': admissao.id,
                'data_hora': (admissao.data_hora - timedelta(hours=3)).strftime('%d/%m/%Y %H:%M'),
                'enfermeiro_nome': enfermeiro.nome if enfermeiro else 'Não informado',
                'admissao_texto': admissao.admissao_texto
            })
        
        return jsonify({
            'success': True,
            'admissoes': resultado
        })
        
    except Exception as e:
        logging.error(f'Erro ao listar admissões de enfermagem: {str(e)}')
        logging.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': f'Erro ao listar admissões: {str(e)}'
        }), 500

# API para atualizar uma admissão de enfermagem
@bp.route('/api/enfermagem/admissao/<int:admissao_id>', methods=['PUT'])
@login_required
def atualizar_admissao_enfermagem(admissao_id):
    """
    Atualiza uma admissão de enfermagem existente.
    """
    try:
        # Verificar se o usuário é enfermeiro ou multi
        current_user = get_current_user()
        if current_user.cargo.lower() not in ['enfermeiro', 'multi']:
            return jsonify({
                'success': False,
                'message': 'Apenas enfermeiros podem atualizar admissões'
            }), 403

        dados = request.get_json()
        
        if not dados:
            return jsonify({
                'success': False,
                'message': 'Dados não fornecidos'
            }), 400

        # Buscar a admissão
        admissao = AdmissaoEnfermagem.query.get(admissao_id)
        if not admissao:
            return jsonify({
                'success': False,
                'message': 'Admissão não encontrada'
            }), 404

        # Atualizar o texto da admissão
        if 'admissao_texto' in dados:
            admissao.admissao_texto = dados['admissao_texto']
            
            # Também atualizar o campo legado na Internacao
            internacao = Internacao.query.get(admissao.internacao_id)
            if internacao:
                internacao.admissao_enfermagem = dados['admissao_texto']

        db.session.commit()

        return jsonify({
            'success': True,
            'message': 'Admissão atualizada com sucesso'
        })

    except Exception as e:
        db.session.rollback()
        logging.error(f'Erro ao atualizar admissão de enfermagem: {str(e)}')
        logging.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': f'Erro ao atualizar admissão: {str(e)}'
        }), 500

@bp.route('/api/enfermagem/atualizar/<int:id>', methods=['PUT'])
def atualizar_evolucao_enfermagem(id):
    dados = request.get_json()
    try:
        evolucao = EvolucaoEnfermagem.query.get_or_404(id)
        evolucao.texto = dados.get('texto', evolucao.texto)
        evolucao.data_evolucao = datetime.now(ZoneInfo("America/Sao_Paulo"))
        db.session.commit()
        return jsonify({'mensagem': 'Evolução de enfermagem atualizada com sucesso.'}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'erro': str(e)}), 400


# Buscar SAE por paciente_id
@bp.route('/api/enfermagem/sae/<int:paciente_id>', methods=['GET'])
def obter_sae_por_paciente(paciente_id):
    try:
        # Verificar se o usuário é médico ou enfermeiro
        current_user = get_current_user()
        if not current_user:
            return jsonify({
                'success': False,
                'message': 'Usuário não autenticado'
            }), 401
            
        if current_user.cargo.lower() not in ['enfermeiro', 'medico', 'multi']:
            return jsonify({
                'success': False,
                'message': 'Acesso permitido apenas para enfermeiros e médicos'
            }), 403
        
        # Buscar todas as SAEs do paciente, ordenadas por data
        saes = InternacaoSae.query.filter_by(paciente_id=paciente_id).order_by(InternacaoSae.data_registro.desc(), InternacaoSae.id.desc()).all()
        
        if not saes:
            return jsonify({'success': False, 'error': 'SAE não registrada'}), 404
        
        # Separar SAEs de hoje e anteriores (como nas evoluções)
        hoje = datetime.now(ZoneInfo("America/Sao_Paulo")).date()
        sae_hoje = None
        sae_antiga = None
        
        for sae in saes:
            data_sae = sae.data_registro.date()
            if data_sae == hoje and not sae_hoje:
                sae_hoje = sae
            elif not sae_antiga:
                sae_antiga = sae
        
        # Priorizar SAE de hoje, se existir
        sae = sae_hoje if sae_hoje else sae_antiga
        
        # Formatar a resposta
        return jsonify({
            'success': True,
            'sae': {
                'id': sae.id,
                'paciente_id': sae.paciente_id,
                'enfermeiro_id': sae.enfermeiro_id,
                'hipotese_diagnostica': sae.hipotese_diagnostica,
                'pa': sae.pa,
                'fc': sae.fc,
                'sat': sae.sat,
                'dx': sae.dx,
                'r': sae.r,
                't': sae.t,
                'medicacao': sae.medicacao,
                'alergias': sae.alergias,
                'antecedentes_pessoais': sae.antecedentes_pessoais,
                'sistema_neurologico': sae.sistema_neurologico,
                'estado_geral': sae.estado_geral,
                'ventilacao': sae.ventilacao,
                'diagnostico_de_enfermagem': sae.diagnostico_de_enfermagem,
                'pele': sae.pele,
                'sistema_gastrointerstinal': sae.sistema_gastrointerstinal,
                'regulacao_vascular': sae.regulacao_vascular,
                'pulso': sae.pulso,
                'regulacao_abdominal': sae.regulacao_abdominal,
                'rha': sae.rha,
                'sistema_urinario': sae.sistema_urinario,
                'acesso_venoso': sae.acesso_venoso,
                'observacao': sae.observacao,
                'data_registro': (sae.data_registro - timedelta(hours=3)).strftime('%Y-%m-%d %H:%M:%S') if sae.data_registro else None,
                'eh_hoje': True if sae_hoje and sae.id == sae_hoje.id else False
            }
        })
    except Exception as e:
        logging.error(f'Erro ao buscar SAE: {str(e)}')
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'error': f'Erro ao buscar SAE: {str(e)}'}), 500

# Nova rota para listar o histórico completo de SAE do paciente
@bp.route('/api/enfermagem/sae/historico/<int:paciente_id>', methods=['GET'])
def obter_historico_sae_paciente(paciente_id):
    try:
        # Verificar se o usuário é médico, enfermeiro ou multi
        current_user = get_current_user()
        if current_user.cargo.lower() not in ['enfermeiro', 'medico', 'multi']:
            return jsonify({
                'success': False,
                'message': 'Acesso permitido apenas para enfermeiros e médicos'
            }), 403
            
        # Buscar todas as SAEs do paciente, ordenadas por data
        saes = InternacaoSae.query.filter_by(paciente_id=paciente_id).order_by(InternacaoSae.data_registro.desc(), InternacaoSae.id.desc()).all()
        
        if not saes:
            return jsonify({'success': False, 'error': 'SAE não registrada'}), 404
        
        # Formatar a resposta para incluir todos os registros SAE
        resultado = []
        for sae in saes:
            resultado.append({
                'id': sae.id,
                'paciente_id': sae.paciente_id,
                'enfermeiro_id': sae.enfermeiro_id,
                'hipotese_diagnostica': sae.hipotese_diagnostica,
                'pa': sae.pa,
                'fc': sae.fc,
                'sat': sae.sat,
                'dx': sae.dx,
                'r': sae.r,
                't': sae.t,
                'medicacao': sae.medicacao,
                'alergias': sae.alergias,
                'antecedentes_pessoais': sae.antecedentes_pessoais,
                'sistema_neurologico': sae.sistema_neurologico,
                'estado_geral': sae.estado_geral,
                'ventilacao': sae.ventilacao,
                'diagnostico_de_enfermagem': sae.diagnostico_de_enfermagem,
                'pele': sae.pele,
                'sistema_gastrointerstinal': sae.sistema_gastrointerstinal,
                'regulacao_vascular': sae.regulacao_vascular,
                'pulso': sae.pulso,
                'regulacao_abdominal': sae.regulacao_abdominal,
                'rha': sae.rha,
                'sistema_urinario': sae.sistema_urinario,
                'acesso_venoso': sae.acesso_venoso,
                'observacao': sae.observacao,
                'data_registro': (sae.data_registro - timedelta(hours=3)).strftime('%Y-%m-%d %H:%M:%S') if sae.data_registro else None
            })
        
        return jsonify({
            'success': True,
            'sae': resultado
        })
    except Exception as e:
        logging.error(f'Erro ao buscar histórico SAE: {str(e)}')
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'error': f'Erro ao buscar histórico SAE: {str(e)}'}), 500

# Atualizar SAE existente
@bp.route('/api/enfermagem/sae/<int:id>', methods=['PUT'])
def atualizar_sae(id):
    data = request.get_json()
    sae = InternacaoSae.query.get(id)

    if not sae:
        return jsonify({'success': False, 'error': 'Registro não encontrado'}), 404

    try:
        for campo, valor in data.items():
            if hasattr(sae, campo):
                setattr(sae, campo, valor)
        
        db.session.commit()
        return jsonify({'success': True, 'message': 'SAE atualizado com sucesso'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 400


# Listar todos os registros (opcional)
@bp.route('/api/enfermagem/sae', methods=['GET'])
def listar_saes():
    saes = InternacaoSae.query.all()
    resultado = []
    for sae in saes:
        resultado.append({
            'id': sae.id,
            'paciente_id': sae.paciente_id,
            'enfermeiro_id': sae.enfermeiro_id,
            'data_registro': (sae.data_registro - timedelta(hours=3)).strftime('%Y-%m-%d %H:%M:%S') if sae.data_registro else None
        })
    return jsonify({'success': True, 'saes': resultado})

# Rota para buscar evoluções de enfermagem por ID de internação
@bp.route('/api/enfermagem/evolucao/<int:internacao_id>', methods=['GET'])
def buscar_evolucoes_enfermagem_por_internacao(internacao_id):
    """
    Busca todas as evoluções de enfermagem para uma internação específica.
    """
    try:
        evolucoes = EvolucaoEnfermagem.query\
            .filter_by(atendimentos_clinica_id=internacao_id)\
            .join(Funcionario, EvolucaoEnfermagem.funcionario_id == Funcionario.id)\
            .add_columns(
                EvolucaoEnfermagem.id,
                EvolucaoEnfermagem.atendimentos_clinica_id,
                EvolucaoEnfermagem.data_evolucao,
                EvolucaoEnfermagem.texto,
                Funcionario.nome.label('enfermeiro_nome'),
                Funcionario.numero_profissional.label('enfermeiro_coren')
            )\
            .order_by(EvolucaoEnfermagem.data_evolucao.desc())\
            .all()
        
        resultado = []
        for ev in evolucoes:
            resultado.append({
                'id': ev.id,
                'atendimentos_clinica_id': ev.atendimentos_clinica_id,
                'data_evolucao': (ev.data_evolucao - timedelta(hours=3)).isoformat() if ev.data_evolucao else None,
                'texto': ev.texto,
                'enfermeiro_nome': ev.enfermeiro_nome,
                'enfermeiro_coren': ev.enfermeiro_coren
            })
        
        return jsonify(resultado), 200
    except Exception as e:
        logging.error(f'Erro ao buscar evoluções de enfermagem: {str(e)}')
        logging.error(traceback.format_exc())
        return jsonify({'erro': f'Erro ao buscar evoluções: {str(e)}'}), 500

# API para listar prescrições de enfermagem por ID de internação
@bp.route('/api/enfermagem/prescricao/<int:internacao_id>', methods=['GET'])
@login_required
def buscar_prescricoes_enfermagem_por_internacao(internacao_id):
    """
    Busca todas as prescrições de enfermagem para uma internação específica.
    """
    try:
        # Verificar se o usuário é enfermeiro, médico ou multi
        current_user = get_current_user()
        if not current_user or current_user.cargo.lower() not in ['enfermeiro', 'medico', 'multi']:
            return jsonify({
                'success': False,
                'message': 'Acesso negado'
            }), 403

        # Verificar se a internação existe
        internacao = Internacao.query.get(internacao_id)
        if not internacao:
            return jsonify({
                'success': False,
                'message': 'Internação não encontrada'
            }), 404

        # Buscar todas as prescrições de enfermagem da internação
        prescricoes = db.session.query(
            PrescricaoEnfermagem.id,
            PrescricaoEnfermagem.atendimentos_clinica_id,
            PrescricaoEnfermagem.data_prescricao,
            PrescricaoEnfermagem.texto,
            Funcionario.nome.label("enfermeiro_nome"),
            Funcionario.numero_profissional.label("enfermeiro_coren")
        )\
        .outerjoin(Funcionario, PrescricaoEnfermagem.funcionario_id == Funcionario.id)\
        .filter(PrescricaoEnfermagem.atendimentos_clinica_id == internacao_id)\
        .order_by(PrescricaoEnfermagem.data_prescricao.desc())\
        .all()

        resultado = []
        for p in prescricoes:
            # Converter para timezone brasileiro
            data_prescricao_br = None
            if p.data_prescricao:
                data_prescricao_br = p.data_prescricao.isoformat()

            resultado.append({
                'id': p.id,
                'atendimentos_clinica_id': p.atendimentos_clinica_id,
                'data_prescricao': data_prescricao_br,
                'texto': p.texto,
                'enfermeiro_nome': p.enfermeiro_nome or 'Não informado',
                'enfermeiro_coren': p.enfermeiro_coren
            })

        return jsonify({
            'success': True,
            'prescricoes': resultado,
            'total': len(resultado)
        }), 200

    except Exception as e:
        logging.error(f'Erro ao buscar prescrições de enfermagem: {str(e)}')
        logging.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': f'Erro ao buscar prescrições: {str(e)}'
        }), 500

# API para registrar nova prescrição de enfermagem
@bp.route('/api/enfermagem/prescricao', methods=['POST'])
@login_required
def registrar_prescricao_enfermagem():
    """
    Registra uma nova prescrição de enfermagem
    """
    try:
        data = request.get_json()
        if not data:
            return jsonify({
                'success': False,
                'message': 'Dados não fornecidos'
            }), 400

        # Validações básicas
        if not data.get('atendimentos_clinica_id'):
            return jsonify({
                'success': False,
                'message': 'ID da internação é obrigatório'
            }), 400
            
        if not data.get('funcionario_id'):
            return jsonify({
                'success': False,
                'message': 'ID do funcionário é obrigatório'
            }), 400
            
        if not data.get('texto') or not data.get('texto').strip():
            return jsonify({
                'success': False,
                'message': 'Texto da prescrição é obrigatório'
            }), 400

        # Verificar se o usuário atual é enfermeiro ou multi
        current_user = get_current_user()
        if not current_user or current_user.cargo.lower() not in ['enfermeiro', 'multi']:
            return jsonify({
                'success': False,
                'message': 'Apenas enfermeiros podem criar prescrições de enfermagem'
            }), 403

        # Verificar se a internação existe
        internacao = Internacao.query.get(data['atendimentos_clinica_id'])
        if not internacao:
            return jsonify({
                'success': False,
                'message': 'Internação não encontrada'
            }), 404

        # Criar nova prescrição
        nova_prescricao = PrescricaoEnfermagem(
            atendimentos_clinica_id=data['atendimentos_clinica_id'],
            funcionario_id=data['funcionario_id'],
            texto=data['texto'].strip(),
            data_prescricao=datetime.now(ZoneInfo("America/Sao_Paulo"))
        )
        
        db.session.add(nova_prescricao)
        db.session.commit()
        
        logging.info(f'Prescrição de enfermagem registrada com sucesso - ID: {nova_prescricao.id}, Enfermeiro: {current_user.nome}')
        
        return jsonify({
            'success': True,
            'message': 'Prescrição de enfermagem registrada com sucesso!',
            'prescricao_id': nova_prescricao.id
        }), 201
        
    except Exception as e:
        db.session.rollback()
        logging.error(f'Erro ao registrar prescrição de enfermagem: {str(e)}')
        logging.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': f'Erro ao registrar prescrição: {str(e)}'
        }), 500

# API para atualizar prescrição de enfermagem
@bp.route('/api/enfermagem/prescricao/<int:id>', methods=['PUT'])
def atualizar_prescricao_enfermagem(id):
    dados = request.get_json()
    try:
        prescricao = PrescricaoEnfermagem.query.get_or_404(id)
        prescricao.texto = dados.get('texto', prescricao.texto)
        prescricao.data_prescricao = datetime.now(ZoneInfo("America/Sao_Paulo"))
        db.session.commit()
        return jsonify({'mensagem': 'Prescrição de enfermagem atualizada com sucesso.'}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'erro': str(e)}), 400

# API para listar todas as prescrições de enfermagem (opcional)
@bp.route('/api/enfermagem/prescricao', methods=['GET'])
def listar_prescricoes_enfermagem():
    try:
        prescricoes = db.session.query(
            PrescricaoEnfermagem.id,
            PrescricaoEnfermagem.atendimentos_clinica_id,
            PrescricaoEnfermagem.data_prescricao,
            PrescricaoEnfermagem.texto,
            Funcionario.nome.label("enfermeiro_nome"),
            Funcionario.numero_profissional.label("enfermeiro_coren")
        ).select_from(PrescricaoEnfermagem)\
        .join(Funcionario, PrescricaoEnfermagem.funcionario_id == Funcionario.id)\
        .order_by(PrescricaoEnfermagem.data_prescricao.asc())\
        .all()

        resultado = [{
            'id': presc.id,
            'atendimentos_clinica_id': presc.atendimentos_clinica_id,
            'data_prescricao': presc.data_prescricao.isoformat() if presc.data_prescricao else None,
            'texto': presc.texto,
            'enfermeiro_nome': presc.enfermeiro_nome,
            'enfermeiro_coren': presc.enfermeiro_coren
        } for presc in prescricoes]

        return jsonify({
            'success': True,
            'prescricoes': resultado
        }), 200

    except Exception as e:
        logging.error(f'Erro ao buscar prescrições de enfermagem para médico: {str(e)}')
        logging.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': f'Erro ao buscar prescrições: {str(e)}'
        }), 500

# API para buscar detalhes de uma internação específica
@bp.route('/api/internacao/<string:internacao_id>', methods=['GET'])
@login_required
def buscar_internacao(internacao_id):
    """
    Busca os detalhes de uma internação específica.
    """
    try:
        logging.info(f'Buscando internação com atendimento_id: {internacao_id}')
        
        # Verificar se o usuário é médico, enfermeiro ou multi
        current_user = get_current_user()
        if current_user.cargo.lower() not in ['medico', 'enfermeiro', 'multi']:
            return jsonify({
                'success': False,
                'message': 'Acesso permitido apenas para médicos e enfermeiros'
            }), 403
        
        # Buscar a internação
        internacao = Internacao.query.filter_by(atendimento_id=internacao_id).first()
        if not internacao:
            logging.warning(f'Internação não encontrada para atendimento_id: {internacao_id}')
            return jsonify({
                'success': False,
                'message': 'Internação não encontrada',
                'internacao': None
            }), 200
        
        logging.info(f'Internação encontrada: ID={internacao.id}, paciente_id={internacao.paciente_id}')
        
        # Buscar dados do paciente
        paciente = Paciente.query.get(internacao.paciente_id)
        if not paciente:
            logging.error(f'INCONSISTÊNCIA: Internação {internacao.id} tem paciente_id {internacao.paciente_id} que não existe na tabela pacientes')
            return jsonify({
                'success': False,
                'message': f'Dados do paciente inconsistentes para esta internação (paciente_id: {internacao.paciente_id})'
            }), 500
        
        logging.info(f'Paciente encontrado: ID={paciente.id}, nome={paciente.nome}')
        
        # Buscar dados do atendimento para obter anamnese_exame_fisico
        atendimento = Atendimento.query.filter_by(id=internacao_id).first()
        
        # Formatar resposta
        resultado = {
            'id': internacao.id,
            'atendimento_id': internacao.atendimento_id,
            'paciente_id': internacao.paciente_id,
            'nome_paciente': paciente.nome if paciente else 'Não informado',
            'leito': internacao.leito,
            'data_internacao': internacao.data_internacao.isoformat() if internacao.data_internacao else None,
            'data_alta': internacao.data_alta.isoformat() if internacao.data_alta else None,
            'diagnostico': internacao.diagnostico,
            'diagnostico_inicial': internacao.diagnostico_inicial,
            'cid_principal': internacao.cid_principal,
            'hda': internacao.hda,  # Campo HDA que estava faltando
            'historico_internacao': internacao.historico_internacao,
            'conduta': internacao.conduta,
            'cuidados_gerais': internacao.cuidados_gerais,
            'antibiotico': internacao.antibiotico,  # Novo campo adicionado
            'carater_internacao': internacao.carater_internacao,  # Campo que estava faltando
            # Campo da tabela Internacao (correto)
            'anamnese_exame_fisico': internacao.folha_anamnese,  # Mapear folha_anamnese para o nome esperado
            # CAMPOS EXTRAS
            'justificativa_internacao_sinais_e_sintomas': internacao.justificativa_internacao_sinais_e_sintomas,
            'justificativa_internacao_condicoes': internacao.justificativa_internacao_condicoes,
            'justificativa_internacao_principais_resultados_diagnostico': internacao.justificativa_internacao_principais_resultados_diagnostico,
            'cid_10_secundario': internacao.cid_10_secundario,
            'cid_10_causas_associadas': internacao.cid_10_causas_associadas,
            'descr_procedimento_solicitado': internacao.descr_procedimento_solicitado,
            'codigo_procedimento': internacao.codigo_procedimento,
            'acidente_de_trabalho': internacao.acidente_de_trabalho
        }
        
        logging.info(f'Resposta preparada para atendimento_id {internacao_id}: paciente_id={resultado["paciente_id"]}')
        
        return jsonify({
            'success': True,
            'internacao': resultado
        })
        
    except Exception as e:
        logging.error(f'Erro ao buscar internação: {str(e)}')
        logging.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': f'Erro ao buscar internação: {str(e)}'
        }), 500

# API para registrar alta de paciente
@bp.route('/api/internacao/<string:internacao_id>/alta', methods=['POST'])
@login_required
def registrar_alta_paciente(internacao_id):
    """
    Registra a alta de um paciente internado, atribuindo a data de alta.
    Monta o histórico da internação com: hda + folha_anamnese + conduta + exames_laboratoriais.
    """
    try:
        current_user = get_current_user()
        if current_user.cargo.lower() != 'medico':
            return jsonify({
                'success': False,
                'message': 'Apenas médicos podem registrar alta.'
            }), 403

        internacao = Internacao.query.filter_by(atendimento_id=internacao_id).first()
        if not internacao:
            return jsonify({
                'success': False,
                'message': 'Internação não encontrada.'
            }), 404

        if internacao.data_alta:
            return jsonify({
                'success': False,
                'message': 'Paciente já recebeu alta anteriormente.'
            }), 400

        dados = request.get_json()
        if not dados:
            return jsonify({
                'success': False,
                'message': 'Nenhum dado enviado para registrar alta.'
            }), 400

        # Montar histórico da internação de forma organizada com cabeçalhos
        # Formato: seções bem definidas para melhor organização
        secoes_historico = []
        
        # HDA - História da Doença Atual
        if internacao.hda and internacao.hda.strip():
            secoes_historico.append(f"# HDA (História da Doença Atual):\n{internacao.hda.strip()}")
        
        # Anamnese
        if internacao.folha_anamnese and internacao.folha_anamnese.strip():
            secoes_historico.append(f"# ANAMNESE:\n{internacao.folha_anamnese.strip()}")
        
        # Conduta
        if internacao.conduta and internacao.conduta.strip():
            secoes_historico.append(f"# CONDUTA:\n{internacao.conduta.strip()}")
        
        # Exames Laboratoriais
        if internacao.exames_laboratoriais and internacao.exames_laboratoriais.strip():
            secoes_historico.append(f"# EXAMES LABORATORIAIS:\n{internacao.exames_laboratoriais.strip()}")
        
        # Juntar todas as seções com dupla quebra de linha para separação visual
        historico_final = '\n\n'.join(secoes_historico) if secoes_historico else 'Histórico da internação não disponível.'

        # Atualizar campos da internação com dados da alta
        internacao.historico_internacao = dados.get('historico_internacao', historico_final)  # Usar o enviado ou o montado automaticamente
        internacao.diagnostico = dados.get('diagnostico') or internacao.diagnostico
        internacao.cuidados_gerais = dados.get('cuidados_gerais') or internacao.cuidados_gerais
        internacao.data_alta = datetime.now(ZoneInfo("America/Sao_Paulo"))

        # IMPORTANTE: NÃO mudar o status do atendimento aqui!
        # O status continua 'Internado' até que o prontuário seja fechado
        # Armazenar temporariamente o status_final desejado no campo dieta (prefixado com PENDENTE:)
        try:
            atendimento = Atendimento.query.get(internacao_id)
        except Exception:
            atendimento = None

        if atendimento is not None:
            status_raw = (dados.get('status_final') or dados.get('status') or '').strip().lower()
            status_map = {
                'alta': 'Alta',
                'obito': 'Óbito',
                'óbito': 'Óbito',
                'transferencia': 'Transferido',
                'transferência': 'Transferido',
                'transferido': 'Transferido',
                'evasao': 'Evasão',
                'evasão': 'Evasão',
                'alta a pedido': 'Alta a pedido',
                'a pedido': 'Alta a pedido',
            }
            finais_permitidos = {'Alta', 'Óbito', 'Transferido', 'Evasão', 'Alta a pedido'}
            desejado = status_map.get(status_raw) if status_raw else 'Alta'
            if desejado in finais_permitidos:
                # Armazenar o status desejado no campo dieta temporariamente (será usado ao fechar prontuário)
                internacao.dieta = f'PENDENTE:{desejado}'
                logging.info(f"Status pendente armazenado: {desejado} (será aplicado ao fechar prontuário)")

        # Atualizar ocupação do leito
        if internacao.leito:
            leito = Leito.query.filter_by(nome=internacao.leito).first()
            if leito:
                if leito.ocupacao_atual > 0:
                    leito.ocupacao_atual -= 1
                if leito.status == 'Ocupado' and leito.ocupacao_atual < leito.capacidade_maxima:
                    leito.status = 'Disponível'
                db.session.add(leito)

        db.session.commit()

        return jsonify({
            'success': True,
            'message': 'Alta registrada com sucesso.'
        }), 200

    except Exception as e:
        db.session.rollback()
        logging.error(f'Erro ao registrar alta: {str(e)}')
        logging.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': 'Erro interno do servidor ao registrar alta.'
        }), 500

# Rota para a página de histórico de internações
@bp.route('/clinica/historico-internacoes')
@login_required
def historico_internacoes():
    """
    Página que exibe o histórico de pacientes que receberam alta.
    """
    try:
        current_user = get_current_user()
        if current_user.cargo.lower() not in ['medico', 'enfermeiro', 'multi']:
            flash('Acesso restrito a médicos, enfermeiros e profissionais multi.', 'danger')
            return redirect(url_for('main.index'))
        
        return render_template('historico_internacoes.html')
        
    except Exception as e:
        logging.error(f"Erro ao acessar histórico de internações: {str(e)}")
        logging.error(traceback.format_exc())
        flash('Erro ao acessar o histórico. Por favor, tente novamente.', 'danger')
        return redirect(url_for('main.clinica'))

# API para listar histórico de internações (com alta)
@bp.route('/api/internacoes/historico')
@login_required
def listar_historico_internacoes():
    """
    Lista todas as internações que já receberam alta.
    """
    try:
        current_user = get_current_user()
        if current_user.cargo.lower() not in ['medico', 'enfermeiro', 'multi']:
            return jsonify({'error': 'Acesso não autorizado'}), 403
        
        # Verificar se há filtro por mês
        filtro_mes = request.args.get('mes')
        
        # Buscar internações que possuem data de alta
        query = Internacao.query.filter(Internacao.data_alta != None)
        
        # Aplicar filtro por mês se fornecido
        if filtro_mes:
            # Formato esperado: YYYY-MM
            try:
                ano, mes = filtro_mes.split('-')
                ano = int(ano)
                mes = int(mes)
                # Filtra pelo mês e ano da alta
                query = query.filter(
                    db.extract('year', Internacao.data_alta) == ano,
                    db.extract('month', Internacao.data_alta) == mes
                )
            except (ValueError, AttributeError):
                # Se formato inválido, ignora o filtro
                pass
        
        # Ordenar por data de alta (mais recente primeiro)
        internacoes = query.order_by(Internacao.data_alta.desc()).all()
        
        # Listar internações com detalhes do paciente
        resultado = []
        for internacao in internacoes:
            paciente = Paciente.query.get(internacao.paciente_id)
            if paciente:
                resultado.append({
                    'id': internacao.id,
                    'atendimento_id': internacao.atendimento_id,
                    'nome_paciente': paciente.nome,
                    'cpf': paciente.cpf,
                    'data_nascimento': paciente.data_nascimento.strftime('%Y-%m-%d') if paciente.data_nascimento else None,
                    'leito': internacao.leito,
                    'data_internacao': internacao.data_internacao.strftime('%Y-%m-%d %H:%M') if internacao.data_internacao else None,
                    'data_alta': internacao.data_alta.strftime('%Y-%m-%d %H:%M') if internacao.data_alta else None,
                    'diagnostico': internacao.diagnostico,
                })
        
        return jsonify({
            'success': True,
            'internacoes': resultado
        })
        
    except Exception as e:
        logging.error(f"Erro ao listar histórico de internações: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'error': 'Erro interno do servidor'}), 500

# Rota para relatório do paciente (somente leitura)
@bp.route('/clinica/relatorio-paciente/<int:internacao_id>')
@login_required
def relatorio_paciente(internacao_id):
    try:
        # Verificar autenticação
        current_user = get_current_user()
        
        # Buscar dados do paciente e da internação
        internacao = Internacao.query.get(internacao_id)
        
        if not internacao:
            flash('Internação não encontrada.', 'danger')
            return redirect(url_for('main.pacientes_internados'))
        
        paciente = Paciente.query.get(internacao.paciente_id)
        
        if not paciente:
            flash('Paciente não encontrado.', 'danger')
            return redirect(url_for('main.pacientes_internados'))
        
        # Buscar evoluções médicas
        evolucoes = EvolucaoAtendimentoClinica.query.filter_by(
            atendimentos_clinica_id=internacao_id
        ).order_by(
            EvolucaoAtendimentoClinica.data_evolucao.desc()
        ).all()
        
        # Formatar dados das evoluções
        evolucoes_formatadas = []
        for e in evolucoes:
            medico_nome = e.funcionario.nome if hasattr(e, 'funcionario') and e.funcionario else 'Desconhecido'
            evolucoes_formatadas.append({
                'id': e.id,
                'data_evolucao': e.data_evolucao.strftime('%d/%m/%Y %H:%M') if hasattr(e, 'data_evolucao') else '',
                'nome_medico': medico_nome,
                'evolucao': e.evolucao if hasattr(e, 'evolucao') else ''
            })
        
        # Buscar prescrições médicas
        prescricoes = PrescricaoClinica.query.filter_by(
            atendimentos_clinica_id=internacao_id
        ).order_by(
            PrescricaoClinica.horario_prescricao.desc()
        ).all()
        
        # Formatar dados das prescrições
        prescricoes_formatadas = []
        for p in prescricoes:
            medico_nome = p.funcionario.nome if hasattr(p, 'funcionario') and p.funcionario else 'Desconhecido'
            
            # Buscar medicamentos da prescrição
            medicamentos = []
            if hasattr(p, 'medicamentos_json') and p.medicamentos_json:
                import json
                try:
                    medicamentos_data = json.loads(p.medicamentos_json)
                    medicamentos = medicamentos_data
                except json.JSONDecodeError:
                    logging.error(f"Erro ao decodificar JSON de medicamentos para prescrição ID {p.id}")
            
            prescricoes_formatadas.append({
                'id': p.id,
                'data_prescricao': p.horario_prescricao.strftime('%d/%m/%Y %H:%M') if hasattr(p, 'horario_prescricao') else '',
                'medico_nome': medico_nome,
                'texto_dieta': p.texto_dieta,
                'texto_procedimento_medico': p.texto_procedimento_medico,
                'texto_procedimento_multi': p.texto_procedimento_multi,
                'medicamentos': medicamentos
            })
        
        # Buscar evoluções de enfermagem
        evolucoes_enfermagem = []
        try:
            evolucoes_enf = EvolucaoEnfermagem.query.filter_by(
                atendimentos_clinica_id=internacao_id
            ).order_by(
                EvolucaoEnfermagem.data_evolucao.desc()
            ).all()
            
            for e in evolucoes_enf:
                enfermeiro = Funcionario.query.get(e.funcionario_id) if e.funcionario_id else None
                evolucoes_enfermagem.append({
                    'data_evolucao': e.data_evolucao.strftime('%d/%m/%Y %H:%M') if hasattr(e, 'data_evolucao') else '',
                    'enfermeiro_nome': enfermeiro.nome if enfermeiro else 'Não informado',
                    'texto': e.texto
                })
        except Exception as e:
            logging.error(f"Erro ao buscar evoluções de enfermagem: {str(e)}")
        
        # Buscar prescrições de enfermagem
        prescricoes_enfermagem = []
        try:
            prescricoes_enf = PrescricaoEnfermagem.query.filter_by(
                atendimentos_clinica_id=internacao_id
            ).order_by(
                PrescricaoEnfermagem.data_prescricao.desc()
            ).all()
            
            for p in prescricoes_enf:
                enfermeiro = Funcionario.query.get(p.funcionario_id) if p.funcionario_id else None
                prescricoes_enfermagem.append({
                    'data_prescricao': p.data_prescricao.strftime('%d/%m/%Y %H:%M') if hasattr(p, 'data_prescricao') else '',
                    'enfermeiro_nome': enfermeiro.nome if enfermeiro else 'Não informado',
                    'texto': p.texto
                })
        except Exception as e:
            logging.error(f"Erro ao buscar prescrições de enfermagem: {str(e)}")
        
        # Obter data e hora atual do Brasil
        now = datetime.now(ZoneInfo("America/Sao_Paulo"))
        
        return render_template('clinica_relatorio_paciente.html', 
                              internacao=internacao, 
                              paciente=paciente, 
                              evolucoes=evolucoes_formatadas,
                              prescricoes=prescricoes_formatadas,
                              evolucoes_enfermagem=evolucoes_enfermagem,
                              prescricoes_enfermagem=prescricoes_enfermagem,
                              now=lambda: datetime.now(ZoneInfo("America/Sao_Paulo")))
    
    except Exception as e:
        logging.error(f"Erro ao gerar relatório do paciente: {str(e)}")
        logging.error(traceback.format_exc())
        flash('Erro ao gerar relatório. Por favor, tente novamente.', 'danger')
        return redirect(url_for('main.pacientes_internados'))
    


@bp.route('/api/internacao/atualizar-hda', methods=['POST'])
@login_required
def atualizar_hda():
    """
    Atualiza a História da Doença Atual (HDA) de uma internação.
    """
    try:
        # Verificar se o usuário é médico (ignorar variações de capitalização e espaços)
        current_user = get_current_user()
        if current_user.cargo and current_user.cargo.strip().lower() != 'medico':
            return jsonify({
                'success': False,
                'message': 'Apenas médicos podem atualizar o HDA'
            }), 403
        
        dados = request.get_json()
        
        if not dados or 'atendimentos_clinica_id' not in dados or 'hda' not in dados:
            return jsonify({
                'success': False,
                'message': 'Dados obrigatórios não fornecidos'
            }), 400
            
        # Log para debug
        print(f"Recebendo atualização de HDA - ID: {dados['atendimentos_clinica_id']}, HDA: {dados['hda'][:50]}...")
        
        # Buscar a internação
        internacao = Internacao.query.filter_by(id=dados['atendimentos_clinica_id']).first()
        if not internacao:
            # Tentar buscar por atendimento_id se não encontrar por id
            internacao = Internacao.query.filter_by(atendimento_id=dados['atendimentos_clinica_id']).first()
            if not internacao:
                return jsonify({
                    'success': False,
                    'message': 'Internação não encontrada'
                }), 404
        
        # Atualizar o HDA
        internacao.hda = dados['hda']
        db.session.commit()
        
        # Log para confirmar atualização
        print(f"HDA atualizado com sucesso para internação {internacao.id}")
        
        return jsonify({
            'success': True,
            'message': 'HDA atualizado com sucesso'
        })
        
    except Exception as e:
        db.session.rollback()
        logging.error(f'Erro ao atualizar HDA: {str(e)}')
        logging.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': 'Erro ao atualizar HDA',
            'error': str(e)
        }), 500

# API para registrar aprazamento
@bp.route('/api/aprazamentos', methods=['POST'])
@login_required
def registrar_aprazamento():
    """
    Registra um novo aprazamento para uma prescrição médica usando o novo modelo Aprazamento.
    """
    try:
        current_user = get_current_user()
        if current_user.cargo.lower() not in ['medico', 'enfermeiro', 'multi']:
            return jsonify({
                'success': False,
                'message': 'Acesso permitido apenas para médicos e enfermeiros'
            }), 403

        dados = request.get_json()
        
        if not dados or 'prescricao_id' not in dados or 'nome_medicamento' not in dados or 'data_hora_aprazamento' not in dados:
            return jsonify({
                'success': False,
                'message': 'Dados obrigatórios não fornecidos (prescricao_id, nome_medicamento, data_hora_aprazamento)'
            }), 400

        prescricao = PrescricaoClinica.query.get(dados['prescricao_id'])
        if not prescricao:
            return jsonify({
                'success': False,
                'message': 'Prescrição não encontrada'
            }), 404

        # Converter string para datetime, se for string
        data_hora = dados['data_hora_aprazamento']
        if isinstance(data_hora, str):
            from datetime import datetime
            try:
                # Formato esperado: DD/MM/YYYY HH:MM ou YYYY-MM-DD HH:MM:SS
                if '/' in data_hora:
                    data_hora = datetime.strptime(data_hora, '%d/%m/%Y %H:%M')
                else:
                    data_hora = datetime.strptime(data_hora, '%Y-%m-%d %H:%M:%S')
            except ValueError:
                return jsonify({
                    'success': False,
                    'message': 'Formato de data/hora inválido'
                }), 400

        # Criar novo registro de aprazamento
        novo_aprazamento = Aprazamento(
            prescricao_id=dados['prescricao_id'],
            nome_medicamento=dados['nome_medicamento'],
            descricao_uso=dados.get('descricao_uso', ''),
            data_hora_aprazamento=data_hora,
            enfermeiro_responsavel_id=dados.get('enfermeiro_responsavel_id', current_user.id)
        )
        
        db.session.add(novo_aprazamento)
        db.session.commit()

        return jsonify({
            'success': True,
            'message': 'Aprazamento registrado com sucesso',
            'id': novo_aprazamento.id
        })

    except Exception as e:
        db.session.rollback()
        logging.error(f'Erro ao registrar aprazamento: {str(e)}')
        logging.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': 'Erro ao registrar aprazamento',
            'error': str(e)
        }), 500

# API para listar aprazamentos
@bp.route('/api/aprazamentos/<int:prescricao_id>', methods=['GET'])
@login_required
def listar_aprazamentos(prescricao_id):
    """
    Lista os aprazamentos de uma prescrição específica usando o novo modelo Aprazamento.
    """
    try:
        current_user = get_current_user()
        if current_user.cargo.lower() not in ['medico', 'enfermeiro', 'multi']:
            return jsonify({
                'success': False,
                'message': 'Acesso permitido apenas para médicos e enfermeiros'
            }), 403

        prescricao = PrescricaoClinica.query.get(prescricao_id)
        if not prescricao:
            return jsonify({
                'success': False,
                'message': 'Prescrição não encontrada'
            }), 404

        # Buscar aprazamentos da tabela
        aprazamentos = Aprazamento.query.filter_by(prescricao_id=prescricao_id).order_by(Aprazamento.data_hora_aprazamento).all()
        
        aprazamentos_formatados = []
        for apr in aprazamentos:
            enfermeiro = Funcionario.query.get(apr.enfermeiro_responsavel_id) if apr.enfermeiro_responsavel_id else None
            
            aprazamentos_formatados.append({
                'id': apr.id,
                'nome_medicamento': apr.nome_medicamento,
                'descricao_uso': apr.descricao_uso,
                'data_hora_aprazamento': apr.data_hora_aprazamento.strftime('%d/%m/%Y %H:%M'),
                'realizado': apr.realizado,
                'enfermeiro_responsavel': enfermeiro.nome if enfermeiro else None,
                'enfermeiro_responsavel_id': apr.enfermeiro_responsavel_id,
                'data_realizacao': apr.data_realizacao.strftime('%d/%m/%Y %H:%M') if apr.data_realizacao else None
            })

        return jsonify({
            'success': True,
            'aprazamentos': aprazamentos_formatados,
            'prescricao_id': prescricao_id
        })

    except Exception as e:
        logging.error(f'Erro ao listar aprazamentos: {str(e)}')
        logging.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': 'Erro ao listar aprazamentos',
            'error': str(e)
        }), 500

# API para atualizar aprazamento
@bp.route('/api/aprazamentos/<int:aprazamento_id>', methods=['PUT'])
@login_required
def atualizar_aprazamento(aprazamento_id):
    """
    Atualiza um aprazamento específico usando o novo modelo Aprazamento.
    """
    try:
        current_user = get_current_user()
        if current_user.cargo.lower() not in ['medico', 'enfermeiro', 'multi']:
            return jsonify({
                'success': False,
                'message': 'Acesso permitido apenas para médicos e enfermeiros'
            }), 403

        dados = request.get_json()
        if not dados:
            return jsonify({
                'success': False,
                'message': 'Dados não fornecidos'
            }), 400

        aprazamento = Aprazamento.query.get(aprazamento_id)
        if not aprazamento:
            return jsonify({
                'success': False,
                'message': 'Aprazamento não encontrado'
            }), 404

        # Atualizar campos do aprazamento
        if 'nome_medicamento' in dados:
            aprazamento.nome_medicamento = dados['nome_medicamento']
            
        if 'descricao_uso' in dados:
            aprazamento.descricao_uso = dados['descricao_uso']
            
        if 'data_hora_aprazamento' in dados:
            data_hora = dados['data_hora_aprazamento']
            if isinstance(data_hora, str):
                from datetime import datetime
                try:
                    # Formato esperado: DD/MM/YYYY HH:MM ou YYYY-MM-DD HH:MM:SS
                    if '/' in data_hora:
                        data_hora = datetime.strptime(data_hora, '%d/%m/%Y %H:%M')
                    else:
                        data_hora = datetime.strptime(data_hora, '%Y-%m-%d %H:%M:%S')
                except ValueError:
                    return jsonify({
                        'success': False,
                        'message': 'Formato de data/hora inválido'
                    }), 400
            aprazamento.data_hora_aprazamento = data_hora
            
        if 'realizado' in dados:
            aprazamento.realizado = dados['realizado']
            # Se o medicamento foi marcado como administrado, registrar a data
            if dados['realizado']:
                from datetime import datetime
                aprazamento.data_realizacao = datetime.now(ZoneInfo("America/Sao_Paulo"))
            else:
                aprazamento.data_realizacao = None
                
        if 'enfermeiro_responsavel_id' in dados:
            aprazamento.enfermeiro_responsavel_id = dados['enfermeiro_responsavel_id']
            
        db.session.commit()

        return jsonify({
            'success': True,
            'message': 'Aprazamento atualizado com sucesso'
        })

    except Exception as e:
        db.session.rollback()
        logging.error(f'Erro ao atualizar aprazamento: {str(e)}')
        logging.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': 'Erro ao atualizar aprazamento',
            'error': str(e)
        }), 500

# API para excluir aprazamento
@bp.route('/api/aprazamentos/<int:aprazamento_id>', methods=['DELETE'])
@login_required
def excluir_aprazamento(aprazamento_id):
    """
    Remove um aprazamento específico usando o novo modelo Aprazamento.
    """
    try:
        current_user = get_current_user()
        if current_user.cargo.lower() not in ['medico', 'enfermeiro', 'multi']:
            return jsonify({
                'success': False,
                'message': 'Acesso permitido apenas para médicos e enfermeiros'
            }), 403

        aprazamento = Aprazamento.query.get(aprazamento_id)
        if not aprazamento:
            return jsonify({
                'success': False,
                'message': 'Aprazamento não encontrado'
            }), 404

        # Excluir o aprazamento
        db.session.delete(aprazamento)
        db.session.commit()

        return jsonify({
            'success': True,
            'message': 'Aprazamento removido com sucesso'
        })

    except Exception as e:
        db.session.rollback()
        logging.error(f'Erro ao excluir aprazamento: {str(e)}')
        logging.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': 'Erro ao excluir aprazamento',
            'error': str(e)
        }), 500

@bp.route('/api/prescricoes/aprazamento', methods=['POST'])
@login_required
def registrar_aprazamento_prescricao():
    """
    Registra novos aprazamentos para uma prescrição médica seguindo a nova modelagem.
    """
    try:
        current_user = get_current_user()
        if current_user.cargo.lower() not in ['medico', 'enfermeiro', 'multi']:
            return jsonify({'success': False, 'message': 'Acesso permitido apenas para médicos, enfermeiros e profissionais multi'}), 403

        dados = request.get_json()

        if not dados or 'prescricao_id' not in dados or 'aprazamentos' not in dados:
            return jsonify({'success': False, 'message': 'Dados obrigatórios não fornecidos'}), 400

        prescricao = PrescricaoClinica.query.get(dados['prescricao_id'])
        if not prescricao:
            return jsonify({'success': False, 'message': 'Prescrição não encontrada'}), 404

        # Recuperar o medicamento específico usando o índice, se fornecido
        medicamento_index = dados.get('medicamento_index')
        medicamento_nome = None
        
        # Obter os medicamentos como lista JSON
        # Ensure medicamentos_json is parsed as a list
        if isinstance(prescricao.medicamentos_json, str):
            try:
                medicamentos_lista = json.loads(prescricao.medicamentos_json)
                if not isinstance(medicamentos_lista, list):
                    medicamentos_lista = []
            except json.JSONDecodeError:
                medicamentos_lista = []
        else:
            medicamentos_lista = prescricao.medicamentos_json if isinstance(prescricao.medicamentos_json, list) else []
        
        if medicamento_index is not None and medicamentos_lista and len(medicamentos_lista) > medicamento_index:
            medicamento = medicamentos_lista[medicamento_index]
            medicamento_nome = medicamento.get('nome_medicamento')
            
            # Deletar aprazamentos existentes apenas para este medicamento
            if medicamento_nome:
                Aprazamento.query.filter_by(
                    prescricao_id=prescricao.id,
                    nome_medicamento=medicamento_nome
                ).delete()
                db.session.commit()
        
        # ID do enfermeiro responsável
        enfermeiro_id = dados.get('enfermeiro_id') or current_user.id

        aprazamentos_recebidos = dados['aprazamentos']  # Lista de objetos com detalhes de aprazamento
        aprazamentos_validos = 0
        erros_aprazamento = []

        for apr in aprazamentos_recebidos:
            nome_medicamento = apr.get('nome_medicamento')
            descricao_uso = apr.get('descricao_uso', '')
            data_hora_str = apr.get('data_hora_aprazamento')

            if not nome_medicamento or not data_hora_str:
                continue  # Pular registros incompletos
                
            try:
                # Verificar formato da data/hora (deve ser DD/MM/YYYY HH:MM)
                if not re.match(r'^\d{2}/\d{2}/\d{4} \d{2}:\d{2}$', data_hora_str):
                    erros_aprazamento.append(f"Formato inválido: {data_hora_str}")
                    continue
                    
                # Converte a string para datetime
                data_hora = datetime.strptime(data_hora_str, '%d/%m/%Y %H:%M')
                
                novo_aprazamento = Aprazamento(
                    prescricao_id=prescricao.id,
                    nome_medicamento=nome_medicamento,
                    descricao_uso=descricao_uso,
                    data_hora_aprazamento=data_hora,
                    enfermeiro_responsavel_id=enfermeiro_id
                )
                db.session.add(novo_aprazamento)
                aprazamentos_validos += 1
            except ValueError as e:
                # Registrar erro específico de formatação de data/hora
                erros_aprazamento.append(f"Erro na data/hora '{data_hora_str}': {str(e)}")
                logging.warning(f"Erro ao converter data/hora: '{data_hora_str}' - {str(e)}")
                continue
            except Exception as e:
                # Registrar outros erros
                erros_aprazamento.append(f"Erro desconhecido: {str(e)}")
                logging.error(f"Erro desconhecido ao processar aprazamento: {str(e)}")
                continue

        # Se nenhum aprazamento válido foi processado
        if aprazamentos_validos == 0:
            if erros_aprazamento:
                return jsonify({
                    'success': False, 
                    'message': f'Nenhum aprazamento válido processado. Erros: {"; ".join(erros_aprazamento[:5])}'
                }), 400
            else:
                return jsonify({
                    'success': False,
                    'message': 'Nenhum aprazamento válido encontrado nos dados enviados'
                }), 400

        db.session.commit()

        return jsonify({
            'success': True, 
            'message': f'Aprazamentos registrados com sucesso ({aprazamentos_validos} horários)'
        })

    except Exception as e:
        db.session.rollback()
        logging.error(f'Erro ao registrar aprazamentos: {str(e)}')
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro ao registrar aprazamentos', 'error': str(e)}), 500

@bp.route('/api/prescricoes/aprazamento/<int:prescricao_id>', methods=['GET'])
@login_required
def listar_aprazamento_prescricao(prescricao_id):
    """
    Lista os aprazamentos da nova tabela relacionados à prescrição.
    """
    try:
        current_user = get_current_user()
        if current_user.cargo.lower() not in ['medico', 'enfermeiro', 'multi']:
            return jsonify({'success': False, 'message': 'Acesso permitido apenas para médicos, enfermeiros e profissionais multi'}), 403

        prescricao = PrescricaoClinica.query.get(prescricao_id)
        if not prescricao:
            return jsonify({'success': False, 'message': 'Prescrição não encontrada'}), 404

        aprazamentos_formatados = []
        for apr in prescricao.aprazamentos:
            enfermeiro_nome = apr.enfermeiro_responsavel.nome if apr.enfermeiro_responsavel else None
            aprazamentos_formatados.append({
                'id': apr.id,
                'nome_medicamento': apr.nome_medicamento,
                'descricao_uso': apr.descricao_uso,
                'data_hora_aprazamento': apr.data_hora_aprazamento.strftime('%d/%m/%Y %H:%M'),
                'realizado': apr.realizado,
                'enfermeiro_responsavel': enfermeiro_nome,
                'data_realizacao': apr.data_realizacao.strftime('%d/%m/%Y %H:%M') if apr.data_realizacao else None
            })

        return jsonify({'success': True, 'aprazamentos': aprazamentos_formatados})

    except Exception as e:
        logging.error(f'Erro ao listar aprazamentos: {str(e)}')
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro ao listar aprazamentos', 'error': str(e)}), 500


# Rota para logout
@bp.route('/logout')
def logout():
    try:
        # Remover dados da sessão
        session.pop('user_id', None)
        session.pop('cargo', None)
        # Fazer logout do flask-login
        logout_user()
        # Redirecionar para a página inicial
        return redirect(url_for('main.index'))
    except Exception as e:
        logging.error(f"Erro ao fazer logout: {str(e)}")
        logging.error(traceback.format_exc())
        flash('Erro ao fazer logout. Por favor, tente novamente.', 'danger')
        return redirect(url_for('main.index'))

@bp.route('/api/aprazamentos/<int:aprazamento_id>/realizar', methods=['PUT'])
@login_required
def marcar_aprazamento_realizado(aprazamento_id):
    """
    Marca um aprazamento como realizado.
    """
    try:
        current_user = get_current_user()
        if current_user.cargo.lower() not in ['medico', 'enfermeiro', 'multi']:
            return jsonify({'success': False, 'message': 'Acesso permitido apenas para médicos, enfermeiros e profissionais multi'}), 403

        aprazamento = Aprazamento.query.get(aprazamento_id)
        if not aprazamento:
            return jsonify({'success': False, 'message': 'Aprazamento não encontrado'}), 404

        if aprazamento.realizado:
            return jsonify({'success': False, 'message': 'Este aprazamento já foi realizado'}), 400

        aprazamento.realizado = True
        aprazamento.data_realizacao = datetime.now(ZoneInfo("America/Sao_Paulo"))
        aprazamento.enfermeiro_responsavel_id = current_user.id

        db.session.commit()

        return jsonify({
            'success': True,
            'message': 'Aprazamento marcado como realizado',
            'data_realizacao': aprazamento.data_realizacao.strftime('%d/%m/%Y %H:%M')
        })

    except Exception as e:
        db.session.rollback()
        logging.error(f'Erro ao marcar aprazamento como realizado: {str(e)}')
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro ao marcar aprazamento como realizado', 'error': str(e)}), 500


@bp.route('/api/aprazamentos/realizados', methods=['GET'])
@login_required
def listar_aprazamentos_realizados():
    """
    Lista todos os aprazamentos realizados usando o novo modelo Aprazamento.
    """
    try:
        current_user = get_current_user()
        if current_user.cargo.lower() not in ['medico', 'enfermeiro', 'admin']:
            return jsonify({
                'success': False,
                'message': 'Acesso permitido apenas para médicos, enfermeiros e administradores'
            }), 403

        # Buscar aprazamentos realizados
        aprazamentos = Aprazamento.query.filter(
            Aprazamento.realizado == True
        ).order_by(
            Aprazamento.data_realizacao.desc()
        ).all()

        # Formatar resposta
        aprazamentos_formatados = []
        for aprazamento in aprazamentos:
            aprazamentos_formatados.append({
                'id': aprazamento.id,
                'prescricao_id': aprazamento.prescricao_id,
                'nome_medicamento': aprazamento.nome_medicamento,
                'descricao_uso': aprazamento.descricao_uso,
                'data_hora_aprazamento': aprazamento.data_hora_aprazamento.strftime('%Y-%m-%d %H:%M:%S'),
                'realizado': aprazamento.realizado,
                'enfermeiro_responsavel_id': aprazamento.enfermeiro_responsavel_id,
                'data_realizacao': aprazamento.data_realizacao.strftime('%Y-%m-%d %H:%M:%S') if aprazamento.data_realizacao else None
            })

        return jsonify({
            'success': True,
            'aprazamentos': aprazamentos_formatados
        })

    except Exception as e:
        logging.error(f'Erro ao listar aprazamentos realizados: {str(e)}')
        logging.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': 'Erro ao listar aprazamentos realizados',
            'error': str(e)
        }), 500

@bp.route('/api/aprazamentos/ativos', methods=['GET'])
@login_required
def listar_aprazamentos_ativos():
    """
    Lista todos os aprazamentos ativos usando o novo modelo Aprazamento.
    """
    try:
        current_user = get_current_user()
        if current_user.cargo.lower() not in ['medico', 'enfermeiro', 'admin']:
            return jsonify({
                'success': False,
                'message': 'Acesso permitido apenas para médicos, enfermeiros e administradores'
            }), 403

        # Buscar aprazamentos ativos
        aprazamentos = Aprazamento.query.filter(
            Aprazamento.realizado == False,
            Aprazamento.data_hora_aprazamento >= datetime.now(ZoneInfo("America/Sao_Paulo"))
        ).order_by(
            Aprazamento.data_hora_aprazamento
        ).all()

        # Formatar resposta
        aprazamentos_formatados = []
        for aprazamento in aprazamentos:
            aprazamentos_formatados.append({
                'id': aprazamento.id,
                'prescricao_id': aprazamento.prescricao_id,
                'nome_medicamento': aprazamento.nome_medicamento,
                'descricao_uso': aprazamento.descricao_uso,
                'data_hora_aprazamento': aprazamento.data_hora_aprazamento.strftime('%Y-%m-%d %H:%M:%S'),
                'realizado': aprazamento.realizado,
                'enfermeiro_responsavel_id': aprazamento.enfermeiro_responsavel_id,
                'data_realizacao': aprazamento.data_realizacao.strftime('%Y-%m-%d %H:%M:%S') if aprazamento.data_realizacao else None
            })

        return jsonify({
            'success': True,
            'aprazamentos': aprazamentos_formatados
        })

    except Exception as e:
        logging.error(f'Erro ao listar aprazamentos ativos: {str(e)}')
        logging.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': 'Erro ao listar aprazamentos ativos',
            'error': str(e)
        }), 500

@bp.route('/api/aprazamentos/prescricao/<int:prescricao_id>', methods=['GET'])
@login_required
def listar_aprazamentos_por_prescricao(prescricao_id):
    """
    Lista todos os aprazamentos de uma prescrição específica usando o novo modelo Aprazamento.
    """
    try:
        current_user = get_current_user()
        if current_user.cargo.lower() not in ['medico', 'enfermeiro', 'admin']:
            return jsonify({
                'success': False,
                'message': 'Acesso permitido apenas para médicos, enfermeiros e administradores'
            }), 403

        # Verificar se a prescrição existe
        prescricao = PrescricaoClinica.query.get(prescricao_id)
        if not prescricao:
            return jsonify({
                'success': False,
                'message': 'Prescrição não encontrada'
            }), 404

        # Buscar aprazamentos da prescrição
        aprazamentos = Aprazamento.query.filter(
            Aprazamento.prescricao_id == prescricao_id
        ).order_by(
            Aprazamento.data_hora_aprazamento
        ).all()

        # Formatar resposta
        aprazamentos_formatados = []
        for aprazamento in aprazamentos:
            aprazamentos_formatados.append({
                'id': aprazamento.id,
                'prescricao_id': aprazamento.prescricao_id,
                'nome_medicamento': aprazamento.nome_medicamento,
                'descricao_uso': aprazamento.descricao_uso,
                'data_hora_aprazamento': aprazamento.data_hora_aprazamento.strftime('%Y-%m-%d %H:%M:%S'),
                'realizado': aprazamento.realizado,
                'enfermeiro_responsavel_id': aprazamento.enfermeiro_responsavel_id,
                'data_realizacao': aprazamento.data_realizacao.strftime('%Y-%m-%d %H:%M:%S') if aprazamento.data_realizacao else None
            })

        return jsonify({
            'success': True,
            'aprazamentos': aprazamentos_formatados
        })

    except Exception as e:
        logging.error(f'Erro ao listar aprazamentos por prescrição: {str(e)}')
        logging.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': 'Erro ao listar aprazamentos',
            'error': str(e)
        }), 500

@bp.route('/api/aprazamentos/data/<data>', methods=['GET'])
@login_required
def listar_aprazamentos_por_data(data):
    """
    Lista todos os aprazamentos de uma data específica usando o novo modelo Aprazamento.
    """
    try:
        current_user = get_current_user()
        if current_user.cargo.lower() not in ['medico', 'enfermeiro', 'admin']:
            return jsonify({
                'success': False,
                'message': 'Acesso permitido apenas para médicos, enfermeiros e administradores'
            }), 403

        # Converter a data para datetime
        try:
            data_inicio = datetime.strptime(data, '%Y-%m-%d')
            data_fim = data_inicio + timedelta(days=1)
        except ValueError:
            return jsonify({
                'success': False,
                'message': 'Formato de data inválido. Use YYYY-MM-DD'
            }), 400

        # Buscar aprazamentos da data
        aprazamentos = Aprazamento.query.filter(
            Aprazamento.data_hora_aprazamento >= data_inicio,
            Aprazamento.data_hora_aprazamento < data_fim
        ).order_by(
            Aprazamento.data_hora_aprazamento
        ).all()

        # Formatar resposta
        aprazamentos_formatados = []
        for aprazamento in aprazamentos:
            aprazamentos_formatados.append({
                'id': aprazamento.id,
                'prescricao_id': aprazamento.prescricao_id,
                'nome_medicamento': aprazamento.nome_medicamento,
                'descricao_uso': aprazamento.descricao_uso,
                'data_hora_aprazamento': aprazamento.data_hora_aprazamento.strftime('%Y-%m-%d %H:%M:%S'),
                'realizado': aprazamento.realizado,
                'enfermeiro_responsavel_id': aprazamento.enfermeiro_responsavel_id,
                'data_realizacao': aprazamento.data_realizacao.strftime('%Y-%m-%d %H:%M:%S') if aprazamento.data_realizacao else None
            })

        return jsonify({
            'success': True,
            'aprazamentos': aprazamentos_formatados
        })

    except Exception as e:
        logging.error(f'Erro ao listar aprazamentos por data: {str(e)}')
        logging.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': 'Erro ao listar aprazamentos',
            'message': 'Erro ao listar aprazamentos pendentes',
            'error': str(e)
        }), 500

@bp.route('/api/aprazamentos/enfermeiro/<int:enfermeiro_id>', methods=['GET'])
@login_required
def listar_aprazamentos_por_enfermeiro(enfermeiro_id):
    """
    Lista todos os aprazamentos realizados por um enfermeiro específico usando o novo modelo Aprazamento.
    """
    try:
        current_user = get_current_user()
        if current_user.cargo.lower() not in ['medico', 'enfermeiro', 'admin']:
            return jsonify({
                'success': False,
                'message': 'Acesso permitido apenas para médicos, enfermeiros e administradores'
            }), 403

        # Buscar aprazamentos do enfermeiro
        aprazamentos = Aprazamento.query.filter(
            Aprazamento.enfermeiro_responsavel_id == enfermeiro_id
        ).order_by(
            Aprazamento.data_hora_aprazamento
        ).all()

        # Formatar resposta
        aprazamentos_formatados = []
        for aprazamento in aprazamentos:
            aprazamentos_formatados.append({
                'id': aprazamento.id,
                'prescricao_id': aprazamento.prescricao_id,
                'nome_medicamento': aprazamento.nome_medicamento,
                'descricao_uso': aprazamento.descricao_uso,
                'data_hora_aprazamento': aprazamento.data_hora_aprazamento.strftime('%Y-%m-%d %H:%M:%S'),
                'realizado': aprazamento.realizado,
                'enfermeiro_responsavel_id': aprazamento.enfermeiro_responsavel_id,
                'data_realizacao': aprazamento.data_realizacao.strftime('%Y-%m-%d %H:%M:%S') if aprazamento.data_realizacao else None
            })

        return jsonify({
            'success': True,
            'aprazamentos': aprazamentos_formatados
        })

    except Exception as e:
        logging.error(f'Erro ao listar aprazamentos por enfermeiro: {str(e)}')
        logging.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': 'Erro ao listar aprazamentos',
            'error': str(e)
        }), 500

@bp.route('/api/aprazamentos/pendentes', methods=['GET'])
@login_required
def listar_aprazamentos_pendentes():
    """
    Lista todos os aprazamentos pendentes usando o novo modelo Aprazamento.
    """
    try:
        current_user = get_current_user()
        if current_user.cargo.lower() not in ['medico', 'enfermeiro', 'admin']:
            return jsonify({
                'success': False,
                'message': 'Acesso permitido apenas para médicos, enfermeiros e administradores'
            }), 403

        # Buscar aprazamentos pendentes
        aprazamentos = Aprazamento.query.filter(
            Aprazamento.realizado == False,
            Aprazamento.data_hora_aprazamento >= datetime.now(ZoneInfo("America/Sao_Paulo"))
        ).order_by(
            Aprazamento.data_hora_aprazamento
        ).all()

        # Formatar resposta
        aprazamentos_formatados = []
        for aprazamento in aprazamentos:
            aprazamentos_formatados.append({
                'id': aprazamento.id,
                'prescricao_id': aprazamento.prescricao_id,
                'nome_medicamento': aprazamento.nome_medicamento,
                'descricao_uso': aprazamento.descricao_uso,
                'data_hora_aprazamento': aprazamento.data_hora_aprazamento.strftime('%Y-%m-%d %H:%M:%S'),
                'realizado': aprazamento.realizado,
                'enfermeiro_responsavel_id': aprazamento.enfermeiro_responsavel_id,
                'data_realizacao': aprazamento.data_realizacao.strftime('%Y-%m-%d %H:%M:%S') if aprazamento.data_realizacao else None
            })

        return jsonify({
            'success': True,
            'aprazamentos': aprazamentos_formatados
        })

    except Exception as e:
        logging.error(f'Erro ao listar aprazamentos pendentes: {str(e)}')
        logging.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': 'Erro ao listar aprazamentos pendentes',
            'error': str(e)
        }), 500

@bp.route('/api/aprazamentos/paciente/<int:paciente_id>', methods=['GET'])
@login_required
def listar_aprazamentos_por_paciente(paciente_id):
    """
    Lista todos os aprazamentos de um paciente específico usando o novo modelo Aprazamento.
    """
    try:
        current_user = get_current_user()
        if current_user.cargo.lower() not in ['medico', 'enfermeiro', 'admin']:
            return jsonify({
                'success': False,
                'message': 'Acesso permitido apenas para médicos, enfermeiros e administradores'
            }), 403

        # Buscar todas as prescrições do paciente
        prescricoes = PrescricaoClinica.query.join(
            Internacao, PrescricaoClinica.atendimentos_clinica_id == Internacao.id
        ).filter(
            Internacao.paciente_id == paciente_id
        ).all()

        # Coletar todos os aprazamentos das prescrições
        aprazamentos_formatados = []
        for prescricao in prescricoes:
            aprazamentos = Aprazamento.query.filter_by(
                prescricao_id=prescricao.id
            ).order_by(
                Aprazamento.data_hora_aprazamento
            ).all()

            for aprazamento in aprazamentos:
                enfermeiro = Funcionario.query.get(aprazamento.enfermeiro_responsavel_id) if aprazamento.enfermeiro_responsavel_id else None
                
                aprazamentos_formatados.append({
                    'id': aprazamento.id,
                    'prescricao_id': aprazamento.prescricao_id,
                    'nome_medicamento': aprazamento.nome_medicamento,
                    'descricao_uso': aprazamento.descricao_uso,
                    'data_hora_aprazamento': aprazamento.data_hora_aprazamento.strftime('%Y-%m-%d %H:%M:%S'),
                    'realizado': aprazamento.realizado,
                    'enfermeiro_responsavel': enfermeiro.nome if enfermeiro else None,
                    'enfermeiro_responsavel_id': aprazamento.enfermeiro_responsavel_id,
                    'data_realizacao': aprazamento.data_realizacao.strftime('%Y-%m-%d %H:%M:%S') if aprazamento.data_realizacao else None
                })

        return jsonify({
            'success': True,
            'aprazamentos': aprazamentos_formatados
        })

    except Exception as e:
        logging.error(f'Erro ao listar aprazamentos por paciente: {str(e)}')
        logging.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': 'Erro ao listar aprazamentos',
            'error': str(e)
        }), 500

@bp.route('/api/prescricoes/aprazamento_horarios/<int:prescricao_id>/<int:medicamento_index>', methods=['GET'])
@login_required
def buscar_horarios_aprazamento(prescricao_id, medicamento_index):
    """
    Busca os horários de aprazamento para um medicamento específico de uma prescrição.
    """
    try:
        # Verificar se o usuário é médico ou enfermeiro
        current_user = get_current_user()
        if current_user.cargo.lower() not in ['medico', 'enfermeiro', 'multi']:
            return jsonify({
                'success': False,
                'message': 'Acesso permitido apenas para médicos e enfermeiros'
            }), 403

        # Buscar a prescrição
        prescricao = PrescricaoClinica.query.get(prescricao_id)
        if not prescricao:
            return jsonify({
                'success': False,
                'message': 'Prescrição não encontrada'
            }), 404

        # Obter o nome do medicamento da lista de medicamentos da prescrição
        medicamentos = prescricao.medicamentos_json
        if not medicamentos or medicamento_index >= len(medicamentos):
            return jsonify({
                'success': False,
                'message': 'Medicamento não encontrado'
            }), 404

        nome_medicamento = medicamentos[medicamento_index].get('nome_medicamento')
        if not nome_medicamento:
            return jsonify({
                'success': False,
                'message': 'Nome do medicamento não encontrado'
            }), 404

        # Buscar aprazamentos para este medicamento
        aprazamentos = Aprazamento.query.filter_by(
            prescricao_id=prescricao_id,
            nome_medicamento=nome_medicamento
        ).order_by(Aprazamento.data_hora_aprazamento).all()

        # Formatar horários
        horarios = []
        for apr in aprazamentos:
            horario = apr.data_hora_aprazamento.strftime('%d/%m/%Y %H:%M')
            status = "Realizado" if apr.realizado else "Pendente"
            enfermeiro = Funcionario.query.get(apr.enfermeiro_responsavel_id) if apr.enfermeiro_responsavel_id else None
            
            horarios.append({
                'horario': horario,
                'status': status,
                'enfermeiro': enfermeiro.nome if enfermeiro else None,
                'data_realizacao': apr.data_realizacao.strftime('%d/%m/%Y %H:%M') if apr.data_realizacao else None
            })

        return jsonify({
            'success': True,
            'horarios': horarios
        })

    except Exception as e:
        logging.error(f'Erro ao buscar horários de aprazamento: {str(e)}')
        logging.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': 'Erro ao buscar horários de aprazamento',
            'error': str(e)
        }), 500

@bp.route('/api/aprazamentos/atendimento/<string:atendimento_id>/medicamento/<path:nome_medicamento>', methods=['GET'])
@login_required
def buscar_aprazamentos_por_medicamento(atendimento_id, nome_medicamento):
    """
    Busca aprazamentos filtrados por atendimento e medicamento.
    """
    try:
        # Verificar se o usuário é médico ou enfermeiro
        current_user = get_current_user()
        if current_user.cargo.lower() not in ['medico', 'enfermeiro', 'multi']:
            return jsonify({
                'success': False,
                'message': 'Acesso permitido apenas para médicos e enfermeiros'
            }), 403

        # Buscar a internação pelo atendimento_id
        internacao = Internacao.query.filter_by(atendimento_id=atendimento_id).first()
        if not internacao:
            return jsonify({
                'success': False,
                'message': 'Internação não encontrada'
            }), 404

        # Buscar prescrições da internação
        prescricoes = PrescricaoClinica.query.filter_by(atendimentos_clinica_id=internacao.id).all()
        prescricao_ids = [p.id for p in prescricoes]

        # Buscar aprazamentos
        aprazamentos = Aprazamento.query\
            .filter(Aprazamento.prescricao_id.in_(prescricao_ids))\
            .filter(Aprazamento.nome_medicamento == nome_medicamento)\
            .order_by(Aprazamento.data_hora_aprazamento)\
            .all()

        # Formatar resposta
        aprazamentos_formatados = []
        for apr in aprazamentos:
            enfermeiro = None
            if apr.enfermeiro_responsavel_id:
                enfermeiro = Funcionario.query.get(apr.enfermeiro_responsavel_id)

            # Verificar se está atrasado
            agora = datetime.now(ZoneInfo("America/Sao_Paulo"))
            data_hora_apz = apr.data_hora_aprazamento.replace(tzinfo=ZoneInfo("America/Sao_Paulo"))
            atrasado = not apr.realizado and data_hora_apz < agora

            aprazamentos_formatados.append({
                'id': apr.id,
                'data_hora_aprazamento': apr.data_hora_aprazamento.strftime('%Y-%m-%d %H:%M'),
                'realizado': apr.realizado,
                'atrasado': atrasado,
                'enfermeiro_responsavel': enfermeiro.nome if enfermeiro else None,
                'data_realizacao': apr.data_realizacao.strftime('%Y-%m-%d %H:%M') if apr.data_realizacao else None
            })

        return jsonify({
            'success': True,
            'aprazamentos': aprazamentos_formatados
        })

    except Exception as e:
        logging.error(f'Erro ao buscar aprazamentos: {str(e)}')
        logging.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': 'Erro ao buscar aprazamentos',
            'error': str(e)
        }), 500

# API para atualizar dados da internação
@bp.route('/api/internacao/<int:internacao_id>/atualizar', methods=['PUT'])
@login_required
def atualizar_internacao(internacao_id):
    """
    Atualiza dados específicos de uma internação.
    """
    try:
        # Verificar se o usuário é médico
        current_user = get_current_user()
        if current_user.cargo.lower() != 'medico':
            return jsonify({
                'success': False,
                'message': 'Apenas médicos podem atualizar dados da internação'
            }), 403
        
        # Buscar a internação
        internacao = Internacao.query.get(internacao_id)
        if not internacao:
            return jsonify({
                'success': False,
                'message': 'Internação não encontrada'
            }), 404
        
        dados = request.get_json()
        
        # Atualizar campos permitidos
        if 'exames_laboratoriais' in dados:
            internacao.exames_laboratoriais = dados['exames_laboratoriais']
        
        if 'folha_anamnese' in dados:
            internacao.folha_anamnese = dados['folha_anamnese']
        
        if 'conduta' in dados:
            internacao.conduta = dados['conduta']
        
        if 'diagnostico_inicial' in dados:
            internacao.diagnostico_inicial = dados['diagnostico_inicial']
        
        if 'cid_principal' in dados:
            internacao.cid_principal = dados['cid_principal']
        
        if 'antibiotico' in dados:
            internacao.antibiotico = dados['antibiotico']
        
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': 'Dados da internação atualizados com sucesso'
        })
        
    except Exception as e:
        db.session.rollback()
        logging.error(f'Erro ao atualizar dados da internação: {str(e)}')
        logging.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': 'Erro ao atualizar dados da internação',
            'error': str(e)
        }), 500


@bp.route('/clinica/impressoes/<int:id>')
@login_required
def lobby_impressoes(id):
    atendimento = Atendimento.query.filter_by(id=str(id)).first()
    internacao = Internacao.query.filter_by(atendimento_id=str(id)).first()

    if not atendimento or not internacao:
        return abort(404, description="Atendimento ou Internação não encontrados.")

    return render_template(
        'clinica_impressoes.html',
        id=id,
        atendimento=atendimento,
        internacao=internacao
    )




# ---------------------------
# REGISTRAR NOVO RECEITUÁRIO
# ---------------------------
@bp.route('/api/receituarios', methods=['POST'])
@login_required
def criar_receituario():
    try:
        dados = request.get_json()

        atendimento_id = dados.get('atendimento_id')
        medico_id = dados.get('medico_id')
        tipo_receita = dados.get('tipo_receita')
        conteudo_receita = dados.get('conteudo_receita')

        # Validações básicas
        if not atendimento_id or not medico_id or not tipo_receita or not conteudo_receita:
            return jsonify({'success': False, 'message': 'Todos os campos são obrigatórios.'}), 400

        if tipo_receita not in ['normal', 'especial']:
            return jsonify({'success': False, 'message': 'Tipo de receita inválido.'}), 400

        novo_receituario = ReceituarioClinica(
            atendimento_id=atendimento_id,
            medico_id=medico_id,
            tipo_receita=tipo_receita,
            conteudo_receita=conteudo_receita
        )
        db.session.add(novo_receituario)
        db.session.commit()

        return jsonify({'success': True, 'message': 'Receituário criado com sucesso.', 'id': novo_receituario.id})

    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': f'Erro ao criar receituário: {str(e)}'}), 500

# --------------------------------------
# CONSULTAR RECEITUÁRIOS POR ATENDIMENTO
# --------------------------------------
@bp.route('/api/receituarios/<string:atendimento_id>', methods=['GET'])
@login_required
def listar_receituarios(atendimento_id):
    try:
        receituarios = ReceituarioClinica.query.filter_by(atendimento_id=atendimento_id).all()

        resultado = []
        for r in receituarios:
            resultado.append({
                'id': r.id,
                'atendimento_id': r.atendimento_id,
                'medico_id': r.medico_id,
                'tipo_receita': r.tipo_receita,
                'conteudo_receita': r.conteudo_receita,
                'data_receita': formatar_datetime_br_completo(r.data_receita)
            })

        return jsonify({'success': True, 'receituarios': resultado})

    except Exception as e:
        return jsonify({'success': False, 'message': f'Erro ao buscar receituários: {str(e)}'}), 500
    
@bp.route('/clinica/receituario/<int:receituario_id>/gerar_pdf')
def gerar_receita_pdf(receituario_id):
    try:
        receituario = ReceituarioClinica.query.get(receituario_id)
        atendimento = receituario.atendimento
        paciente = atendimento.paciente
        medico = receituario.medico

        contexto = {
            'paciente_nome': paciente.nome,
            'data': datetime.now(ZoneInfo("America/Sao_Paulo")).strftime('%d/%m/%Y'),
            'endereco': paciente.endereco or "Não informado",
            'medicamentos': receituario.conteudo_receita,
            'medico': medico.nome
        }

        # Caminho do HTML (modelo exportado do Word)
        base_dir = os.path.dirname(os.path.abspath(__file__))
        template_path = os.path.join(base_dir, 'templates_docx', 'modelo_receita.htm')

        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template não encontrado em: {template_path}")

        # Corrigido: usar encoding latin1 para arquivos exportados do Word
        with open(template_path, 'r', encoding='latin1') as f:
            html_content = f.read()

        html_renderizado = render_template_string(html_content, **contexto)

        temp_dir = os.path.join(base_dir, 'temp')
        if not os.path.exists(temp_dir):
            os.makedirs(temp_dir)

        output_pdf = os.path.join(temp_dir, f"receita_{receituario_id}.pdf")

        # Configuração do caminho do wkhtmltopdf
        config = pdfkit.configuration(wkhtmltopdf=r'C:\Program Files\wkhtmltopdf\bin\wkhtmltopdf.exe')

        # Gerar o PDF
        pdfkit.from_string(html_renderizado, output_pdf, configuration=config)

        return send_file(output_pdf, as_attachment=True)

    except Exception as e:
        return f"<h2>Erro ao gerar receita:</h2><pre>{str(e)}</pre>", 500
    

# -------------------------------------------
# REGISTRAR NOVO ATESTADO
# -------------------------------------------a
@bp.route('/api/atestados', methods=['POST'])
@login_required
def criar_atestado():
    try:
        dados = request.get_json()

        atendimento_id = dados.get('atendimento_id')
        medico_id = dados.get('medico_id')
        conteudo_atestado = dados.get('conteudo_atestado')
        dias_afastamento = dados.get('dias_afastamento')  # Pode ser None

        # Validação básica
        if not atendimento_id or not medico_id or not conteudo_atestado:
            return jsonify({'success': False, 'message': 'Campos obrigatórios não preenchidos.'}), 400

        novo_atestado = AtestadoClinica(
            atendimento_id=atendimento_id,
            medico_id=medico_id,
            conteudo_atestado=conteudo_atestado,
            dias_afastamento=dias_afastamento
        )

        db.session.add(novo_atestado)
        db.session.commit()

        return jsonify({'success': True, 'message': 'Atestado criado com sucesso.', 'id': novo_atestado.id})

    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': f'Erro ao criar atestado: {str(e)}'}), 500

# ----------------------------------------------------
# CONSULTAR ATESTADOS POR ATENDIMENTO
# ----------------------------------------------------
@bp.route('/api/atestados/<string:atendimento_id>', methods=['GET'])
@login_required
def listar_atestados(atendimento_id):
    try:
        atestados = AtestadoClinica.query.filter_by(atendimento_id=atendimento_id).all()

        resultado = []
        for a in atestados:
            resultado.append({
                'id': a.id,
                'atendimento_id': a.atendimento_id,
                'medico_id': a.medico_id,
                'conteudo_atestado': a.conteudo_atestado,
                'dias_afastamento': a.dias_afastamento,
                'data_atestado': a.data_atestado.strftime('%Y-%m-%d %H:%M:%S')
            })

        return jsonify({'success': True, 'atestados': resultado})

    except Exception as e:
        return jsonify({'success': False, 'message': f'Erro ao buscar atestados: {str(e)}'}), 500

# -----------------------------------------
# REGISTRAR NOVO RN
# -----------------------------------------
@bp.route('/api/pacientes_rn', methods=['POST'])
@login_required
def criar_rn():
    try:
        dados = request.get_json()

        paciente_id = dados.get('paciente_id')  # RN
        responsavel_id = dados.get('responsavel_id')  # Mãe ou responsável

        data_nascimento = dados.get('data_nascimento')
        tipo_parto = dados.get('tipo_parto')
        idade_gestacional = dados.get('idade_gestacional')
        peso_ao_nascer = dados.get('peso_ao_nascer')
        observacoes = dados.get('observacoes')

        # Validação básica
        if not paciente_id or not responsavel_id:
            return jsonify({'success': False, 'message': 'Paciente RN e responsável são obrigatórios.'}), 400

        # Verifica se já existe esse vínculo
        existente = PacienteRN.query.filter_by(paciente_id=paciente_id).first()
        if existente:
            return jsonify({'success': False, 'message': 'Este paciente já está registrado como RN.'}), 400

        novo_rn = PacienteRN(
            paciente_id=paciente_id,
            responsavel_id=responsavel_id,
            data_nascimento=datetime.strptime(data_nascimento, '%Y-%m-%d') if data_nascimento else None,
            tipo_parto=tipo_parto,
            idade_gestacional=idade_gestacional,
            peso_ao_nascer=peso_ao_nascer,
            observacoes=observacoes
        )

        db.session.add(novo_rn)
        db.session.commit()

        return jsonify({'success': True, 'message': 'RN registrado com sucesso.', 'id': novo_rn.id})

    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': f'Erro ao criar RN: {str(e)}'}), 500

# -----------------------------------------
# LISTAR RNs DE UM RESPONSÁVEL
# -----------------------------------------
@bp.route('/api/pacientes_rn/responsavel/<int:responsavel_id>', methods=['GET'])
@login_required
def listar_rns_responsavel(responsavel_id):
    try:
        rns = PacienteRN.query.filter_by(responsavel_id=responsavel_id).all()

        resultado = []
        for rn in rns:
            paciente = Paciente.query.get(rn.paciente_id)
            resultado.append({
                'rn_id': rn.id,
                'paciente_id': rn.paciente_id,
                'nome': paciente.nome if paciente else 'Desconhecido',
                'data_nascimento': rn.data_nascimento.strftime('%Y-%m-%d') if rn.data_nascimento else None,
                'tipo_parto': rn.tipo_parto,
                'idade_gestacional': rn.idade_gestacional,
                'peso_ao_nascer': float(rn.peso_ao_nascer) if rn.peso_ao_nascer else None,
                'observacoes': rn.observacoes
            })

        return jsonify({'success': True, 'rns': resultado})

    except Exception as e:
        return jsonify({'success': False, 'message': f'Erro ao buscar RNs: {str(e)}'}), 500

# -----------------------------------------
# CONSULTAR DADOS DE UM RN ESPECÍFICO
# -----------------------------------------
@bp.route('/api/pacientes_rn/<int:rn_id>', methods=['GET'])
@login_required
def consultar_rn(rn_id):
    try:
        rn = PacienteRN.query.get(rn_id)
        if not rn:
            return jsonify({'success': False, 'message': 'RN não encontrado.'}), 404

        paciente = Paciente.query.get(rn.paciente_id)
        responsavel = Paciente.query.get(rn.responsavel_id)

        resultado = {
            'rn_id': rn.id,
            'paciente_id': rn.paciente_id,
            'nome': paciente.nome if paciente else 'Desconhecido',
            'responsavel_id': rn.responsavel_id,
            'nome_responsavel': responsavel.nome if responsavel else 'Desconhecido',
            'data_nascimento': rn.data_nascimento.strftime('%Y-%m-%d') if rn.data_nascimento else None,
            'tipo_parto': rn.tipo_parto,
            'idade_gestacional': rn.idade_gestacional,
            'peso_ao_nascer': float(rn.peso_ao_nascer) if rn.peso_ao_nascer else None,
            'observacoes': rn.observacoes
        }

        return jsonify({'success': True, 'rn': resultado})

    except Exception as e:
        return jsonify({'success': False, 'message': f'Erro ao consultar RN: {str(e)}'}), 500

@bp.route('/api/fichas-referencia/atendimento/<string:atendimento_id>', methods=['GET'])
@bp.route('/api/pacientes', methods=['POST'])
@login_required
def criar_paciente():
    try:
        dados = request.get_json()
        logging.info(f'[criar_paciente] Dados recebidos: {dados}')
        
        if not dados:
            return jsonify({'success': False, 'message': 'Dados não fornecidos'}), 400

        # Validações básicas
        if not dados.get('nome'):
            return jsonify({'success': False, 'message': 'Nome é obrigatório'}), 400
            
        if not dados.get('data_nascimento'):
            return jsonify({'success': False, 'message': 'Data de nascimento é obrigatória'}), 400
            
        if not dados.get('sexo'):
            return jsonify({'success': False, 'message': 'Sexo é obrigatório'}), 400

        # Verificar se CPF já existe (se fornecido)
        cpf_fornecido = dados.get('cpf', '').strip()
        if cpf_fornecido:
            # Limpar CPF
            cpf_limpo = re.sub(r'\D', '', cpf_fornecido)
            # Aceitar CPFs com 11 dígitos ou gerar temporário para pacientes não identificados
            if len(cpf_limpo) == 11:
                # Verificar se já existe
                paciente_existente = Paciente.query.filter_by(cpf=cpf_limpo).first()
                if paciente_existente:
                    return jsonify({
                        'success': False,
                        'message': 'CPF já cadastrado no sistema'
                    }), 400
                dados['cpf'] = cpf_limpo
            elif len(cpf_limpo) > 0:
                # CPF fornecido mas não tem 11 dígitos - aceitar como está para casos especiais
                dados['cpf'] = cpf_limpo
                logging.warning(f'[criar_paciente] CPF com tamanho não padrão aceito: {cpf_limpo}')
            else:
                # CPF vazio - gerar temporário
                import random
                data = datetime.now(ZoneInfo("America/Sao_Paulo")).strftime('%y%m%d')
                random_digits = str(random.randint(0, 9999)).zfill(4)
                dados['cpf'] = f'RN{data}{random_digits}'[:14]
                logging.info(f'[criar_paciente] CPF gerado: {dados["cpf"]}')
        else:
            # Gerar CPF temporário se não fornecido
            import random
            data = datetime.now(ZoneInfo("America/Sao_Paulo")).strftime('%y%m%d')
            random_digits = str(random.randint(0, 9999)).zfill(4)
            dados['cpf'] = f'RN{data}{random_digits}'[:14]
            logging.info(f'[criar_paciente] CPF gerado: {dados["cpf"]}')

        # Verificar se Cartão SUS já existe (se fornecido)
        cartao_sus = dados.get('cartao_sus', '').strip()
        if cartao_sus:
            cartao_sus_limpo = re.sub(r'\D', '', cartao_sus)
            if cartao_sus_limpo:
                paciente_existente_sus = Paciente.query.filter_by(cartao_sus=cartao_sus_limpo).first()
                if paciente_existente_sus:
                    return jsonify({
                        'success': False, 
                        'message': 'Cartão SUS já cadastrado no sistema'
                    }), 400

        # Converter data de nascimento
        try:
            data_nascimento = datetime.strptime(dados['data_nascimento'], '%Y-%m-%d').date()
            logging.info(f'[criar_paciente] Data de nascimento convertida: {data_nascimento}')
        except (ValueError, KeyError) as e:
            logging.error(f'[criar_paciente] Erro ao converter data: {str(e)}')
            return jsonify({
                'success': False,
                'message': 'Data de nascimento inválida'
            }), 400

        # Criar paciente
        try:
            paciente = Paciente()
            paciente.nome = dados['nome'].strip()
            paciente.cpf = dados['cpf']
            paciente.data_nascimento = data_nascimento
            paciente.sexo = dados['sexo']
            
            # Campos opcionais
            paciente.nome_social = dados.get('nome_social', '').strip() or None
            paciente.filiacao = dados.get('filiacao', '').strip() or None
            paciente.telefone = dados.get('telefone', '').strip() or None
            paciente.endereco = dados.get('endereco', '').strip() or None
            paciente.municipio = dados.get('municipio', '').strip() or None
            paciente.bairro = dados.get('bairro', '').strip() or None
            paciente.cor = dados.get('cor', '').strip() or 'Não informada'

            # Prioridade administrativa
            paciente.prioridade = dados.get('prioridade', False)
            if paciente.prioridade:
                paciente.desc_prioridade = dados.get('desc_prioridade', '').strip() or None
            else:
                paciente.desc_prioridade = None

            # Auto-detectar idosos (60+) e marcar como prioridade
            from datetime import date
            idade = (date.today() - data_nascimento).days // 365
            if idade >= 60 and not paciente.prioridade:
                paciente.prioridade = True
                paciente.desc_prioridade = f'Idoso - {idade} anos'

            # Alergias
            paciente.alergias = dados.get('alergias', '').strip() or None

            # Cartão SUS
            if cartao_sus:
                paciente.cartao_sus = re.sub(r'\D', '', cartao_sus)
            
            # Definir se é identificado baseado na presença de CPF real
            paciente.identificado = not dados['cpf'].startswith('RN')
            
            # Para pacientes não identificados
            if not paciente.identificado:
                paciente.descricao_nao_identificado = dados.get('descricao_nao_identificado', '').strip()

            logging.info(f'[criar_paciente] Objeto paciente criado: {paciente.nome}, {paciente.cpf}, {paciente.data_nascimento}')

            db.session.add(paciente)
            logging.info('[criar_paciente] Paciente adicionado à sessão')
            
            db.session.commit()
            logging.info(f'[criar_paciente] Commit realizado com sucesso. ID do paciente: {paciente.id}')

            return jsonify({
                'success': True,
                'message': 'Paciente criado com sucesso',
                'paciente_id': paciente.id,
                'paciente': {
                    'id': paciente.id,
                    'nome': paciente.nome,
                    'cpf': paciente.cpf,
                    'cartao_sus': paciente.cartao_sus,
                    'data_nascimento': paciente.data_nascimento.strftime('%Y-%m-%d'),
                    'sexo': paciente.sexo,
                    'identificado': paciente.identificado
                }
            })

        except Exception as e:
            db.session.rollback()
            logging.error(f'[criar_paciente] Erro ao criar objeto paciente: {str(e)}')
            logging.error(traceback.format_exc())
            return jsonify({
                'success': False,
                'message': 'Erro interno ao salvar paciente no banco de dados'
            }), 500

    except Exception as e:
        logging.error(f'[criar_paciente] Erro geral: {str(e)}')
        logging.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': 'Erro interno do servidor. Tente novamente.'
        }), 500


@bp.route('/api/pacientes/<int:paciente_id>', methods=['GET'])
@login_required
def obter_paciente(paciente_id):
    """
    Obter dados de um paciente específico
    """
    try:
        paciente = Paciente.query.get(paciente_id)
        if not paciente:
            return jsonify({'success': False, 'message': 'Paciente não encontrado'}), 404

        # Calcular idade
        idade = None
        if paciente.data_nascimento:
            hoje = datetime.now().date()
            idade = hoje.year - paciente.data_nascimento.year - ((hoje.month, hoje.day) < (paciente.data_nascimento.month, paciente.data_nascimento.day))

        paciente_data = {
            'id': paciente.id,
            'nome': paciente.nome,
            'cpf': paciente.cpf,
            'data_nascimento': paciente.data_nascimento.strftime('%Y-%m-%d') if paciente.data_nascimento else None,
            'sexo': paciente.sexo,
            'cor': paciente.cor,
            'nome_social': paciente.nome_social,
            'filiacao': paciente.filiacao,
            'cartao_sus': paciente.cartao_sus,
            'telefone': paciente.telefone,
            'endereco': paciente.endereco,
            'bairro': paciente.bairro,
            'municipio': paciente.municipio,
            'identificado': paciente.identificado,
            'idade': idade
        }

        return jsonify({
            'success': True,
            'paciente': paciente_data
        })

    except Exception as e:
        logging.error(f'Erro ao obter paciente {paciente_id}: {str(e)}')
        return jsonify({'success': False, 'message': 'Erro interno do servidor'}), 500


@bp.route('/api/criar_atendimento_rn', methods=['POST'])
def criar_atendimento_rn():
    try:
        dados = request.get_json()

        paciente_id = dados.get('paciente_id')
        funcionario_id = dados.get('profissional_id')
        medico_id = dados.get('medico_id')
        enfermeiro_id = dados.get('enfermeiro_id')

        if not paciente_id or not funcionario_id:
            return jsonify({
                'success': False,
                'message': 'Campos obrigatórios ausentes: paciente_id e profissional_id'
            }), 400

        # Gerar ID que começa com 999
        def gerar_id_rn():
            while True:
                sufixo = str(random.randint(0, 99999)).zfill(5)
                novo_id = f"999{sufixo}"
                if not Atendimento.query.get(novo_id):
                    return novo_id

        novo_atendimento = Atendimento(
            id=gerar_id_rn(),
            paciente_id=paciente_id,
            funcionario_id=funcionario_id,
            medico_id=medico_id,
            enfermeiro_id=enfermeiro_id,
            data_atendimento=datetime.now(ZoneInfo("America/Sao_Paulo")).date(),
            hora_atendimento=datetime.now(ZoneInfo("America/Sao_Paulo")).time(),
            status='internado'  # fixo para RN
        )

        db.session.add(novo_atendimento)
        db.session.commit()

        return jsonify({
            'success': True,
            'message': 'Atendimento RN criado com sucesso',
            'atendimento_id': novo_atendimento.id
        })

    except Exception as e:
        print('[ERRO criar_atendimento_rn]:', e)
        return jsonify({'success': False, 'message': 'Erro ao criar atendimento do RN'}), 500
    
@bp.route('/api/internacao/paciente/<int:paciente_id>/ativo', methods=['GET'])
def verificar_internacao_ativa(paciente_id):
    try:
        internacao = Internacao.query.filter_by(paciente_id=paciente_id, status='ativo').first()
        if internacao:
            return jsonify({
                'success': True,
                'internacao_id': internacao.id
            })
        else:
            return jsonify({
                'success': False,
                'message': 'Nenhuma internação ativa encontrada.'
            }), 404
    except Exception as e:
        print('[ERRO verificar_internacao_ativa]:', e)
        return jsonify({'success': False, 'message': 'Erro interno'}), 500

@bp.route('/api/internacao/<string:internacao_id>', methods=['PUT'])
@login_required
def editar_internacao(internacao_id):
    """
    Permite que médicos editem os dados de uma internação existente.
    """
    try:
        # Verificar se o usuário é médico
        current_user = get_current_user()
        if current_user.cargo.lower() != 'medico':
            return jsonify({
                'success': False,
                'message': 'Apenas médicos podem editar dados de internação.'
            }), 403

        # Buscar a internação pelo atendimento_id
        internacao = Internacao.query.filter_by(atendimento_id=internacao_id).first()
        if not internacao:
            return jsonify({
                'success': False,
                'message': 'Internação não encontrada.'
            }), 404

        # Obter dados enviados no JSON
        dados = request.get_json()

        # Atualizar os campos da internação
        campos_editaveis = [
            'diagnostico', 'diagnostico_inicial', 'cid_principal', 'historico_internacao',
            'relatorio_alta', 'conduta', 'cuidados_gerais', 'antibiotico',
            'carater_internacao', 'folha_anamnese',  # Campo correto para anamnese
            'justificativa_internacao_sinais_e_sintomas',
            'justificativa_internacao_condicoes',
            'justificativa_internacao_principais_resultados_diagnostico',
            'cid_10_secundario', 'cid_10_causas_associadas',
            'descr_procedimento_solicitado', 'codigo_procedimento',
            'acidente_de_trabalho'
            # Removidos: 'leito', 'data_internacao', 'data_alta' - não devem ser editáveis
        ]

        # Atualizar campos da internação
        for campo in campos_editaveis:
            if campo in dados:
                setattr(internacao, campo, dados[campo])

        # Se o campo anamnese_exame_fisico foi enviado, mapear para folha_anamnese
        if 'anamnese_exame_fisico' in dados:
            internacao.folha_anamnese = dados['anamnese_exame_fisico']

        db.session.commit()

        # Recarregar os dados atualizados para retornar na resposta
        db.session.refresh(internacao)

        return jsonify({
            'success': True,
            'message': 'Internação atualizada com sucesso.',
            'internacao': {
                'id': internacao.id,
                'diagnostico': internacao.diagnostico,
                'diagnostico_inicial': internacao.diagnostico_inicial,
                'cid_principal': internacao.cid_principal,
                'carater_internacao': internacao.carater_internacao,
                'antibiotico': internacao.antibiotico,
                'anamnese_exame_fisico': internacao.folha_anamnese,
                'conduta': internacao.conduta,
                'justificativa_internacao_sinais_e_sintomas': internacao.justificativa_internacao_sinais_e_sintomas,
                'justificativa_internacao_condicoes': internacao.justificativa_internacao_condicoes,
                'justificativa_internacao_principais_resultados_diagnostico': internacao.justificativa_internacao_principais_resultados_diagnostico,
                'cid_10_secundario': internacao.cid_10_secundario,
                'cid_10_causas_associadas': internacao.cid_10_causas_associadas,
                'descr_procedimento_solicitado': internacao.descr_procedimento_solicitado,
                'codigo_procedimento': internacao.codigo_procedimento,
                'acidente_de_trabalho': internacao.acidente_de_trabalho
            }
        })

    except Exception as e:
        logging.error(f'Erro ao editar internação: {str(e)}')
        logging.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': f'Erro ao editar internação: {str(e)}'
        }), 500


@bp.route('/clinica/historico-internacao/<string:atendimento_id>', methods=['GET'])
@login_required
def historico_internacao(atendimento_id):
    try:
        # Verificar permissão de acesso
        current_user = get_current_user()
        if current_user.cargo.lower() not in ['medico', 'enfermeiro', 'multi']:
            flash('Acesso restrito a médicos, enfermeiros e profissionais multi.', 'danger')
            return redirect(url_for('main.index'))
        
        atendimento = Atendimento.query.get_or_404(atendimento_id)
        paciente = Paciente.query.get_or_404(atendimento.paciente_id)
        medico = Funcionario.query.get(atendimento.medico_id)

        # Buscar a internação vinculada ao atendimento
        internacao = Internacao.query.filter_by(atendimento_id=atendimento_id).first()

        if not internacao:
            flash('Nenhuma internação encontrada para este atendimento.', 'warning')
            return redirect(url_for('main.historico_internacoes'))

        return render_template(
            'clinica_evolucao_historico.html',
            internacao=internacao,
            atendimento=atendimento,
            paciente=paciente,
            medico=medico
        )
    except Exception as e:
        logging.error(f"Erro ao acessar histórico da internação {atendimento_id}: {str(e)}")
        logging.error(traceback.format_exc())
        flash('Erro ao carregar histórico da internação.', 'danger')
        return redirect(url_for('main.historico_internacoes'))

@bp.route('/api/internacao/<string:atendimento_id>', methods=['GET'])
@login_required
def get_internacao_por_id(atendimento_id):
    # Buscar a internação pelo atendimento_id
    internacao = Internacao.query.filter_by(atendimento_id=atendimento_id).first()
    if not internacao:
        return jsonify({'success': False, 'message': 'Internação não encontrada', 'internacao': None}), 200

    atendimento = Atendimento.query.get(internacao.atendimento_id)
    paciente = atendimento.paciente if atendimento else None
    medico = Funcionario.query.get(atendimento.medico_id) if atendimento else None

    return jsonify({
        'success': True,
        'internacao': {
            'nome_paciente': paciente.nome,
            'cpf': paciente.cpf,
            'data_internacao': internacao.data_internacao.isoformat() if internacao.data_internacao else None,
            'data_alta': internacao.data_alta.isoformat() if internacao.data_alta else None,
            'diagnostico': internacao.diagnostico,
            'hda': internacao.hda,
            'anamnese_exame_fisico': internacao.folha_anamnese,
            'cid_principal': internacao.cid_principal,
            'historico_internacao': internacao.historico_internacao,
            'conduta': internacao.conduta,
            'cuidados_gerais': internacao.cuidados_gerais,
            'medico': medico.nome if medico else None
        }
    })

@bp.route('/api/internacao/<int:internacao_id>/hda', methods=['GET'])
@login_required
def obter_hda(internacao_id):
    try:
        internacao = Internacao.query.get_or_404(internacao_id)
        
        return jsonify({
            'success': True,
            'hda': internacao.hda
        })
    
    except Exception as e:
        print(f"Erro ao obter HDA: {str(e)}")
        return jsonify({
            'success': False,
            'error': 'Erro interno do servidor'
        }), 500

@bp.route('/imprimir/aih/<string:atendimento_id>')
@login_required
def imprimir_aih(atendimento_id):
    try:
        # Busca o atendimento
        atendimento = Atendimento.query.filter_by(id=atendimento_id).first()
        if not atendimento:
            return abort(404, description="Atendimento não encontrado")

        paciente = atendimento.paciente
        internacao = Internacao.query.filter_by(atendimento_id=atendimento_id).first()
        if not internacao:
            return abort(404, description="Internação não encontrada")

        # Verifica permissão: permitir médico e enfermeiro imprimir (enfermeiro não assina)
        current_user = get_current_user()
        if not current_user or current_user.cargo.lower() not in ['medico', 'enfermeiro', 'multi']:
            return abort(403, description="Acesso não autorizado para imprimir AIH.")
        
        medico = current_user if current_user.cargo.lower() == 'medico' else None

        # Carrega template HTML
        caminho_template = os.path.join(current_app.root_path, 'static', 'impressos', 'AIH.html')
        
        try:
            with open(caminho_template, 'r', encoding='utf-8') as f:
                template_content = f.read()
        except Exception as e:
            logging.error(f"Erro ao carregar template HTML: {str(e)}")
            return abort(500, description=f"Erro ao carregar template: {str(e)}")

        # Contexto com as variáveis para substituição
        try:
            contexto = {
                'paciente_nome': paciente.nome or '',
                'id_atendimento': atendimento.id,
                'paciente_cartao_sus': paciente.cartao_sus or '',
                'paciente_data_nascimento': paciente.data_nascimento.strftime('%d/%m/%Y') if paciente.data_nascimento else '',
                'paciente_sexo': paciente.sexo or '',
                'paciente_nome_filiacao': paciente.filiacao or '',
                'paciente_telefone': paciente.telefone or '',
                'paciente_cor': paciente.cor or '',
                'paciente_endereco': paciente.endereco or '',
                'municipio_residencia': paciente.municipio or '',
                'sinais_e_sintomas': internacao.justificativa_internacao_sinais_e_sintomas or '',
                'justificativa_de_internacao': internacao.justificativa_internacao_condicoes or '',
                'exames': internacao.justificativa_internacao_principais_resultados_diagnostico or '',
                'diagnostico_inicial': internacao.diagnostico_inicial or '',
                'cid_principal': internacao.cid_principal or '',
                'cid_secundario': internacao.cid_10_secundario or '',
                'cid_causas_associadas': internacao.cid_10_causas_associadas or '',
                'descricao_procedimento': internacao.descr_procedimento_solicitado or '',
                'codigo_procedimento': internacao.codigo_procedimento or '',
                'leito': internacao.leito or '',
                'carater_de_internacao': internacao.carater_internacao or '',
                'funcionario_nome': (medico.nome if medico and getattr(medico, 'nome', None) else '') or '',
                'medico_cpf': (medico.cpf if medico and getattr(medico, 'cpf', None) else '') or '',
                'data_atual': datetime.now(ZoneInfo("America/Sao_Paulo")).strftime('%d/%m/%Y')
            }
        except Exception as e:
            logging.error(f"Erro ao montar contexto: {str(e)}")
            return abort(500, description=f"Erro ao montar contexto: {str(e)}")

        # Renderiza e retorna o template HTML
        try:
            html_content = render_template_string(template_content, **contexto)
            return html_content
        except Exception as e:
            logging.error(f"Erro ao renderizar template: {str(e)}")
            return abort(500, description=f"Erro ao renderizar template: {str(e)}")

    except Exception as e:
        logging.error(f"Erro geral no endpoint imprimir_aih: {str(e)}")
        logging.error(traceback.format_exc())
        return abort(500, description=f"Erro interno: {str(e)}")

@bp.route('/api/leitos', methods=['GET'])
def listar_leitos():
    # Filtros opcionais via query string
    setor = request.args.get('setor')
    tipo = request.args.get('tipo')
    status = request.args.get('status')

    query = Leito.query

    if setor:
        query = query.filter(Leito.setor.ilike(f"%{setor}%"))
    if tipo:
        query = query.filter(Leito.tipo.ilike(f"%{tipo}%"))
    if status:
        query = query.filter(Leito.status.ilike(f"%{status}%"))

    leitos = query.order_by(Leito.nome).all()

    return jsonify([
        {
            'id': leito.id,
            'nome': leito.nome,
            'tipo': leito.tipo,
            'setor': leito.setor,
            'capacidade_maxima': leito.capacidade_maxima,
            'status': leito.status,
            'observacoes': leito.observacoes
        }
        for leito in leitos
    ])

@bp.route('/api/leitos/opcoes', methods=['GET'])
@login_required
def listar_opcoes_leitos():
    leitos = Leito.query.order_by(Leito.nome).all()
    return jsonify([
        {'nome': leito.nome, 'status': leito.status}
        for leito in leitos
    ])

@bp.route('/clinica/gestao-leitos')
@login_required
def gestao_leitos():
    leitos = Leito.query.order_by(Leito.nome).all()
    return render_template('gestao_leitos.html', leitos=leitos)

@bp.route('/api/internacoes/leito/<string:leito_nome>')
@login_required
def listar_internacoes_por_leito(leito_nome):
    """
    Lista todas as internações ativas (sem data de alta) de um leito específico.
    """
    try:
        current_user = get_current_user()
        if current_user.cargo.lower() not in ['medico', 'enfermeiro', 'multi']:
            return jsonify({'error': 'Acesso não autorizado'}), 403
        
        logging.info(f"Buscando pacientes do leito: {leito_nome}")
        
        # Buscar internações ativas no leito especificado
        internacoes = Internacao.query.filter_by(
            leito=leito_nome,
            data_alta=None
        ).all()
        
        logging.info(f"Encontradas {len(internacoes)} internações ativas no leito {leito_nome}")
        
        pacientes_list = []
        
        for internacao in internacoes:
            paciente = Paciente.query.get(internacao.paciente_id)
            if paciente:
                pacientes_list.append({
                    'id': paciente.id,
                    'paciente_id': paciente.id,
                    'nome': paciente.nome,
                    'cpf': paciente.cpf,
                    'data_nascimento': paciente.data_nascimento.strftime('%Y-%m-%d') if paciente.data_nascimento else None,
                    'data_internacao': internacao.data_internacao.strftime('%d/%m/%Y %H:%M') if internacao.data_internacao else None,
                    'diagnostico': internacao.diagnostico,
                    'atendimento_id': internacao.atendimento_id
                })
        
        logging.info(f"Retornando {len(pacientes_list)} pacientes do leito {leito_nome}")
        return jsonify(pacientes_list)
        
    except Exception as e:
        logging.error(f"Erro ao listar pacientes do leito {leito_nome}: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'error': 'Erro interno do servidor'}), 500

@bp.route('/api/realocar-paciente', methods=['POST'])
@login_required
def realocar_paciente():
    """
    Realoca um paciente de um leito para outro.
    """
    try:
        # Verificar se o usuário é médico ou enfermeiro
        current_user = get_current_user()
        if current_user.cargo.lower() not in ['medico', 'enfermeiro', 'multi']:
            return jsonify({
                'success': False,
                'message': 'Apenas médicos e enfermeiros podem realocar pacientes'
            }), 403
        
        dados = request.get_json()
        atendimento_id = dados.get('atendimento_id')
        leito_destino = dados.get('leito_destino')
        
        if not atendimento_id or not leito_destino:
            return jsonify({
                'success': False,
                'message': 'Atendimento ID e leito destino são obrigatórios'
            }), 400
        
        # Buscar a internação pelo atendimento_id
        internacao = Internacao.query.filter_by(
            atendimento_id=atendimento_id,
            data_alta=None
        ).first()
        
        if not internacao:
            return jsonify({
                'success': False,
                'message': 'Internação ativa não encontrada para este atendimento'
            }), 404
        
        # Verificar se o leito destino existe
        leito = Leito.query.filter_by(nome=leito_destino).first()
        if not leito:
            return jsonify({
                'success': False,
                'message': f'Leito {leito_destino} não encontrado'
            }), 404
        
        # Verificar capacidade do leito destino
        internacoes_destino = Internacao.query.filter_by(
            leito=leito_destino,
            data_alta=None
        ).count()
        
        if internacoes_destino >= leito.capacidade_maxima:
            return jsonify({
                'success': False,
                'message': f'Leito {leito_destino} está com capacidade máxima'
            }), 400
        
        # Salvar leito anterior para log
        leito_anterior = internacao.leito
        
        # Atualizar o leito da internação
        internacao.leito = leito_destino
        
        # Registrar a mudança no histórico
        if internacao.historico_internacao:
            internacao.historico_internacao += f"\n\n[{datetime.now(ZoneInfo('America/Sao_Paulo')).strftime('%d/%m/%Y %H:%M')}] Paciente realocado do leito {leito_anterior} para o leito {leito_destino} por {current_user.nome}"
        else:
            internacao.historico_internacao = f"[{datetime.now(ZoneInfo('America/Sao_Paulo')).strftime('%d/%m/%Y %H:%M')}] Paciente realocado do leito {leito_anterior} para o leito {leito_destino} por {current_user.nome}"
                
        db.session.commit()
        
        # Buscar dados do paciente para o log
        paciente = Paciente.query.get(internacao.paciente_id)
        
        logging.info(f"Paciente {paciente.nome} (ID: {paciente.id}) realocado do leito {leito_anterior} para {leito_destino} por {current_user.nome}")
        
        return jsonify({
            'success': True,
            'message': 'Paciente realocado com sucesso',
            'leito_anterior': leito_anterior,
            'leito_novo': leito_destino
        })
        
    except Exception as e:
        db.session.rollback()
        logging.error(f'Erro ao realocar paciente: {str(e)}')
        logging.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': f'Erro ao realocar paciente: {str(e)}'
        }), 500

# POST /api/enfermeiro/mudar-senha
@bp.route('/api/enfermeiro/mudar-senha', methods=['POST'])
@login_required
def mudar_senha_enfermeiro():
    """
    Permite que o enfermeiro altere a própria senha.
    """
    try:
        # Usar o sistema de autenticação personalizado
        current_user = get_current_user()
        if not current_user:
            return jsonify({'success': False, 'message': 'Usuário não autenticado.'}), 401
        
        if current_user.cargo.lower() != 'enfermeiro':
            return jsonify({'success': False, 'message': 'Usuário não é enfermeiro.'}), 403

        dados = request.get_json()
        if not dados:
            return jsonify({'success': False, 'message': 'Dados não fornecidos.'}), 400

        senha_atual = dados.get('senha_atual')
        nova_senha = dados.get('nova_senha')

        if not senha_atual or not nova_senha:
            return jsonify({'success': False, 'message': 'Campos obrigatórios não preenchidos.'}), 400

        # Verificar se a senha atual confere
        if not current_user.check_password(senha_atual):
            return jsonify({'success': False, 'message': 'Senha atual incorreta.'}), 400

        # Atualizar a senha
        current_user.set_password(nova_senha)
        db.session.commit()

        return jsonify({'success': True, 'message': 'Senha alterada com sucesso.'})

    except Exception as e:
        db.session.rollback()
        logging.error(f"Erro ao alterar senha do enfermeiro: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': 'Erro ao alterar a senha.',
            'error': str(e)
        }), 500
    
@bp.route('/recepcionista')
@login_required
def painel_recepcionista():
    try:
        # Verificar se o usuário é recepcionista
        current_user = get_current_user()
        if current_user.cargo.lower() != 'recepcionista':
            flash('Acesso restrito a recepcionistas.', 'danger')
            return redirect(url_for('main.index'))
        
        # Renderizar o template do painel do recepcionista
        return render_template('recepcionista.html')
        
    except Exception as e:
        logging.error(f"Erro ao acessar painel do recepcionista: {str(e)}")
        logging.error(traceback.format_exc())
        flash('Erro ao acessar o painel. Por favor, tente novamente.', 'danger')
        return redirect(url_for('main.index'))

# GET /api/recepcionista/nome
@bp.route('/api/recepcionista/nome', methods=['GET'])
@login_required
def get_nome_recepcionista():
    """
    Retorna o nome do recepcionista logado.
    """
    try:
        # Usar o sistema de autenticação personalizado
        current_user = get_current_user()
        if not current_user:
            return jsonify({'success': False, 'message': 'Usuário não autenticado.'}), 401
        
        if current_user.cargo.lower() != 'recepcionista':
            return jsonify({'success': False, 'message': 'Usuário não é recepcionista.'}), 403

        return jsonify({
            'success': True,
            'nome': current_user.nome
        })

    except Exception as e:
        logging.error(f"Erro ao buscar nome do recepcionista: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': 'Erro ao buscar nome do recepcionista.',
            'error': str(e)
        }), 500
    


@bp.route('/api/pacientes/buscar/nome', methods=['GET'])
@login_required
def buscar_paciente_por_nome():
    try:
        valor = request.args.get('valor', '').strip()
        
        if not valor:
            return jsonify({'success': False, 'message': 'Nome não fornecido'})
        
        if len(valor) < 3:
            return jsonify({'success': False, 'message': 'Digite pelo menos 3 caracteres'})
        
        # Buscar pacientes cujo nome contenha o valor (case insensitive)
        pacientes = Paciente.query.filter(
            Paciente.nome.ilike(f'%{valor}%')
        ).order_by(Paciente.nome).limit(20).all()
        
        # Converter para dicionário
        pacientes_data = []
        for paciente in pacientes:
            pacientes_data.append({
                'id': paciente.id,
                'nome': paciente.nome,
                'cpf': paciente.cpf,
                'data_nascimento': paciente.data_nascimento.strftime('%Y-%m-%d') if paciente.data_nascimento else None,
                'sexo': paciente.sexo,
                'telefone': paciente.telefone,
                'cartao_sus': paciente.cartao_sus,
                'endereco': paciente.endereco,
                'filiacao': paciente.filiacao
            })
        
        return jsonify({
            'success': True,
            'pacientes': pacientes_data,
            'total': len(pacientes_data)
        })
        
    except Exception as e:
        logging.error(f"Erro ao buscar paciente por nome: {str(e)}")
        return jsonify({'success': False, 'message': 'Erro interno do servidor'})


# Endpoint para buscar paciente por nome da mãe
@bp.route('/api/pacientes/buscar/filiacao', methods=['GET'])
@login_required
def buscar_paciente_por_filiacao():
    try:
        valor = request.args.get('valor', '').strip()

        if not valor:
            return jsonify({'success': False, 'message': 'Nome da mãe não fornecido'})

        if len(valor) < 3:
            return jsonify({'success': False, 'message': 'Digite pelo menos 3 caracteres'})

        # Buscar pacientes cujo nome da mãe contenha o valor (case insensitive)
        pacientes = Paciente.query.filter(
            Paciente.filiacao.ilike(f'%{valor}%')
        ).order_by(Paciente.nome).limit(20).all()

        # Converter para dicionário
        pacientes_data = []
        for paciente in pacientes:
            pacientes_data.append({
                'id': paciente.id,
                'nome': paciente.nome,
                'cpf': paciente.cpf,
                'data_nascimento': paciente.data_nascimento.strftime('%Y-%m-%d') if paciente.data_nascimento else None,
                'sexo': paciente.sexo,
                'telefone': paciente.telefone,
                'cartao_sus': paciente.cartao_sus,
                'endereco': paciente.endereco,
                'filiacao': paciente.filiacao
            })

        return jsonify({
            'success': True,
            'pacientes': pacientes_data,
            'total': len(pacientes_data)
        })

    except Exception as e:
        logging.error(f"Erro ao buscar paciente por filiação: {str(e)}")
        return jsonify({'success': False, 'message': 'Erro interno do servidor'})


# Endpoint para buscar paciente por cartão SUS
@bp.route('/api/pacientes/buscar/cartao-sus', methods=['GET'])
@login_required
def buscar_paciente_por_cartao_sus():
    try:
        valor = request.args.get('valor', '').strip()

        if not valor:
            return jsonify({'success': False, 'message': 'Cartão SUS não fornecido'})

        if len(valor) < 3:
            return jsonify({'success': False, 'message': 'Digite pelo menos 3 caracteres'})

        # Buscar pacientes cujo cartão SUS contenha o valor
        pacientes = Paciente.query.filter(
            Paciente.cartao_sus.ilike(f'%{valor}%')
        ).order_by(Paciente.nome).limit(20).all()

        # Converter para dicionário
        pacientes_data = []
        for paciente in pacientes:
            pacientes_data.append({
                'id': paciente.id,
                'nome': paciente.nome,
                'cpf': paciente.cpf,
                'data_nascimento': paciente.data_nascimento.strftime('%Y-%m-%d') if paciente.data_nascimento else None,
                'sexo': paciente.sexo,
                'telefone': paciente.telefone,
                'cartao_sus': paciente.cartao_sus,
                'endereco': paciente.endereco,
                'filiacao': paciente.filiacao
            })

        return jsonify({
            'success': True,
            'pacientes': pacientes_data,
            'total': len(pacientes_data)
        })

    except Exception as e:
        logging.error(f"Erro ao buscar paciente por cartão SUS: {str(e)}")
        return jsonify({'success': False, 'message': 'Erro interno do servidor'})


# Endpoint para buscar paciente por CPF
@bp.route('/api/pacientes/buscar/cpf', methods=['GET'])
@login_required
def buscar_paciente_por_cpf():
    try:
        valor = request.args.get('valor', '').strip()
        
        if not valor:
            return jsonify({'success': False, 'message': 'CPF não fornecido'})
        
        # Remover formatação do CPF
        cpf_limpo = re.sub(r'\D', '', valor)
        
        if len(cpf_limpo) != 11:
            return jsonify({'success': False, 'message': 'CPF deve ter 11 dígitos'})
        
        # Buscar paciente pelo CPF comparando ignorando máscara no banco também
        paciente = Paciente.query.filter(
            func.replace(func.replace(Paciente.cpf, '.', ''), '-', '') == cpf_limpo
        ).first()
        
        pacientes_data = []
        if paciente:
            pacientes_data.append({
                'id': paciente.id,
                'nome': paciente.nome,
                'cpf': paciente.cpf,
                'data_nascimento': paciente.data_nascimento.strftime('%Y-%m-%d') if paciente.data_nascimento else None,
                'sexo': paciente.sexo,
                'telefone': paciente.telefone,
                'cartao_sus': paciente.cartao_sus,
                'endereco': paciente.endereco,
                'filiacao': paciente.filiacao
            })
        
        return jsonify({
            'success': True,
            'pacientes': pacientes_data,
            'total': len(pacientes_data)
        })
        
    except Exception as e:
        logging.error(f"Erro ao buscar paciente por CPF: {str(e)}")
        return jsonify({'success': False, 'message': 'Erro interno do servidor'})
 

# Buscar SAE por atendimento_id (novo)
@bp.route('/api/enfermagem/sae/atendimento/<string:atendimento_id>', methods=['GET'])
def obter_sae_por_atendimento(atendimento_id):
    try:
        # Verificar se o usuário é médico ou enfermeiro
        current_user = get_current_user()
        if not current_user:
            return jsonify({
                'success': False,
                'message': 'Usuário não autenticado'
            }), 401
            
        if current_user.cargo.lower() not in ['enfermeiro', 'medico']:
            return jsonify({
                'success': False,
                'message': 'Acesso permitido apenas para enfermeiros e médicos'
            }), 403
        
        # Buscar a internação pelo atendimento_id
        internacao = Internacao.query.filter_by(atendimento_id=atendimento_id).first()
        if not internacao:
            return jsonify({'success': False, 'error': 'Internação não encontrada'}), 404
        
        # Buscar todas as SAEs do paciente, ordenadas por data
        saes = InternacaoSae.query.filter_by(paciente_id=internacao.paciente_id).order_by(InternacaoSae.data_registro.desc(), InternacaoSae.id.desc()).all()
        
        if not saes:
            return jsonify({'success': False, 'error': 'SAE não registrada'}), 404
        
        # Separar SAEs de hoje e anteriores (como nas evoluções)
        hoje = datetime.now(ZoneInfo("America/Sao_Paulo")).date()
        sae_hoje = None
        sae_antiga = None
        
        for sae in saes:
            data_sae = sae.data_registro.date()
            if data_sae == hoje and not sae_hoje:
                sae_hoje = sae
            elif not sae_antiga:
                sae_antiga = sae
        
        # Priorizar SAE de hoje, se existir
        sae = sae_hoje if sae_hoje else sae_antiga
        
        # Formatar a resposta
        return jsonify({
            'success': True,
            'sae': {
                'id': sae.id,
                'paciente_id': sae.paciente_id,
                'enfermeiro_id': sae.enfermeiro_id,
                'hipotese_diagnostica': sae.hipotese_diagnostica,
                'pa': sae.pa,
                'fc': sae.fc,
                'sat': sae.sat,
                'dx': sae.dx,
                'r': sae.r,
                't': sae.t,
                'medicacao': sae.medicacao,
                'alergias': sae.alergias,
                'antecedentes_pessoais': sae.antecedentes_pessoais,
                'sistema_neurologico': sae.sistema_neurologico,
                'estado_geral': sae.estado_geral,
                'ventilacao': sae.ventilacao,
                'diagnostico_de_enfermagem': sae.diagnostico_de_enfermagem,
                'pele': sae.pele,
                'sistema_gastrointerstinal': sae.sistema_gastrointerstinal,
                'regulacao_vascular': sae.regulacao_vascular,
                'pulso': sae.pulso,
                'regulacao_abdominal': sae.regulacao_abdominal,
                'rha': sae.rha,
                'sistema_urinario': sae.sistema_urinario,
                'acesso_venoso': sae.acesso_venoso,
                'observacao': sae.observacao,
                'data_registro': (sae.data_registro - timedelta(hours=3)).strftime('%Y-%m-%d %H:%M:%S') if sae.data_registro else None,
                'eh_hoje': True if sae_hoje and sae.id == sae_hoje.id else False
            }
        })
    except Exception as e:
        logging.error(f'Erro ao buscar SAE por atendimento: {str(e)}')
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'error': f'Erro ao buscar SAE: {str(e)}'}), 500

# Buscar SAE por paciente_id (original)
def obter_sae_por_paciente(paciente_id):
    try:
        current_user = get_current_user()
        if not current_user:
            return jsonify({'success': False, 'message': 'Usuário não autenticado'}), 401
        
        if current_user.cargo.lower() not in ['enfermeiro', 'medico']:
            return jsonify({'success': False, 'message': 'Acesso permitido apenas para enfermeiros e médicos'}), 403

        # Recuperar todas as SAEs do paciente, ordenadas por data mais recente
        saes = InternacaoSae.query.filter_by(paciente_id=paciente_id).order_by(InternacaoSae.data_registro.desc(), InternacaoSae.id.desc()).all()

        if not saes:
            return jsonify({'success': False, 'message': 'Nenhuma SAE registrada para este paciente.'}), 404

        # Selecionar SAE de hoje, se existir; senão, a mais recente anterior
        hoje = datetime.now(ZoneInfo("America/Sao_Paulo")).date()
        sae_hoje = next((s for s in saes if s.data_registro.date() == hoje), None)
        sae = sae_hoje if sae_hoje else saes[0]

        resposta_sae = {
            'id': sae.id,
            'paciente_id': sae.paciente_id,
            'enfermeiro_id': sae.enfermeiro_id,
            'hipotese_diagnostica': sae.hipotese_diagnostica,
            'pa': sae.pa,
            'fc': sae.fc,
            'fr': sae.fr,
            'temperatura': sae.temperatura,
            'saturacao': sae.saturacao,
            'glasgow': sae.glasgow,
            'data_registro': sae.data_registro.isoformat()
        }

        return jsonify({'success': True, 'sae': resposta_sae}), 200

    except Exception as e:
        logging.error(f"Erro ao obter SAE do paciente {paciente_id}: {str(e)}")
        return jsonify({'success': False, 'message': 'Erro interno do servidor'}), 500


@bp.route('/api/pacientes/buscar', methods=['GET', 'POST'])
@login_required
def buscar_paciente():
    """
    Busca paciente por CPF (exato) ou nome (parcial, case-insensitive).
    Permissão para médicos, enfermeiros e recepcionistas.
    """
    try:
        import re
        
        current_user = get_current_user()
        if not current_user or current_user.cargo.lower() not in ['medico', 'enfermeiro', 'recepcionista']:
            return jsonify({'success': False, 'message': 'Acesso não autorizado.'}), 403

        # Verificar se é GET ou POST
        if request.method == 'POST':
            dados = request.get_json()
            if not dados:
                return jsonify({'success': False, 'message': 'Dados não fornecidos.'}), 400
            cpf = dados.get('cpf', '').strip()
            nome = dados.get('nome', '').strip()
        else:  # GET
            filtro = request.args.get('filtro', '').strip()
            if not filtro:
                return jsonify({'success': False, 'message': 'Parâmetro filtro é obrigatório.'}), 400

            # Tentar identificar se é CPF ou nome baseado no formato
            cpf_limpo = re.sub(r'\D', '', filtro)
            if len(cpf_limpo) == 11:
                cpf = cpf_limpo
                nome = ''
            else:
                cpf = ''
                nome = filtro

        if not cpf and not nome:
            return jsonify({'success': False, 'message': 'Informe o CPF ou nome do paciente.'}), 400

        query = Paciente.query

        pacientes_encontrados: list[Paciente] = []
        if cpf:
            cpf_limpo = re.sub(r'\D', '', cpf)  # Remove pontos e traços
            # Comparar ignorando máscara no banco também
            p = query.filter(
                func.replace(func.replace(Paciente.cpf, '.', ''), '-', '') == cpf_limpo
            ).first()
            if p:
                pacientes_encontrados = [p]
        else:
            # Usar a mesma lógica avançada de busca por nome
            import unicodedata
            from difflib import SequenceMatcher
            
            # Função para normalizar texto (remover acentos e converter para minúsculas)
            def normalizar_texto(texto):
                if not texto:
                    return ''
                # Converter para minúsculas
                texto = texto.lower()
                # Remover acentos
                texto = unicodedata.normalize('NFD', texto)
                texto = ''.join(char for char in texto if unicodedata.category(char) != 'Mn')
                # Remover caracteres especiais, manter apenas letras, números e espaços
                texto = re.sub(r'[^a-z0-9\s]', '', texto)
                return texto
            
            # Função para calcular similaridade entre textos
            def calcular_similaridade(texto1, texto2):
                return SequenceMatcher(None, texto1, texto2).ratio()
            
            # Normalizar o valor de busca
            nome_normalizado = normalizar_texto(nome)
            
            # Buscar todos os pacientes e pontuar similaridade
            todos_pacientes = Paciente.query.all()
            pacientes_com_score: list[tuple[Paciente, float]] = []

            for p in todos_pacientes:
                nome_paciente_normalizado = normalizar_texto(p.nome)
                
                # Verificar se há correspondência exata (substring)
                if nome_normalizado in nome_paciente_normalizado:
                    pacientes_com_score.append((p, 1.0))  # Score máximo para correspondência exata
                else:
                    # Verificar similaridade com cada palavra do nome
                    palavras_nome = nome_paciente_normalizado.split()
                    max_similaridade = 0
                    
                    for palavra in palavras_nome:
                        # Similaridade da palavra completa
                        similaridade = calcular_similaridade(nome_normalizado, palavra)
                        max_similaridade = max(max_similaridade, similaridade)
                        
                        # Verificar se o valor de busca está contido na palavra
                        if nome_normalizado in palavra and len(nome_normalizado) >= 3:
                            max_similaridade = max(max_similaridade, 0.8)
                    
                    # Similaridade com o nome completo
                    similaridade_completa = calcular_similaridade(nome_normalizado, nome_paciente_normalizado)
                    max_similaridade = max(max_similaridade, similaridade_completa)
                    
                    # Apenas incluir se a similaridade for >= 0.6 (60%)
                    if max_similaridade >= 0.6:
                        pacientes_com_score.append((p, max_similaridade))
            # Ordenar por score (maior primeiro) e pegar os top 20
            if pacientes_com_score:
                pacientes_com_score.sort(key=lambda x: -x[1])
                pacientes_encontrados = [p for p, _ in pacientes_com_score[:20]]

        if not pacientes_encontrados:
            return jsonify({'success': True, 'pacientes': [], 'total': 0, 'message': 'Nenhum paciente encontrado.'}), 200

        # Mapear lista completa de pacientes encontrados
        pacientes_data = []
        for paciente in pacientes_encontrados:
            pacientes_data.append({
                'id': paciente.id,
                'nome': paciente.nome,
                'cpf': paciente.cpf,
                'data_nascimento': paciente.data_nascimento.strftime('%Y-%m-%d') if paciente.data_nascimento else None,
                'sexo': paciente.sexo,
                'filiacao': paciente.filiacao,
                'telefone': paciente.telefone,
                'endereco': paciente.endereco,
                'bairro': paciente.bairro,
                'municipio': paciente.municipio,
                'cartao_sus': getattr(paciente, 'cartao_sus', None),
                'cor': getattr(paciente, 'cor', None)
            })

        return jsonify({'success': True, 'pacientes': pacientes_data, 'total': len(pacientes_data)}), 200

    except Exception as e:
        logging.error(f"Erro ao buscar paciente: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro interno do servidor.'}), 500

@bp.route('/api/internacao/<string:internacao_id>/alta', methods=['GET'])
@login_required
def buscar_informacoes_alta(internacao_id):
    """
    Busca as informações de alta de uma internação específica.
    """
    try:
        # Verificar se o usuário é médico, enfermeiro ou multi
        current_user = get_current_user()
        if current_user.cargo.lower() not in ['medico', 'enfermeiro', 'multi']:
            return jsonify({
                'success': False,
                'message': 'Acesso permitido apenas para médicos, enfermeiros e profissionais multi'
            }), 403

        # Buscar a internação
        internacao = Internacao.query.filter_by(atendimento_id=internacao_id).first()
        if not internacao:
            return jsonify({
                'success': False,
                'message': 'Internação não encontrada'
            }), 404

        # Buscar dados do paciente
        paciente = Paciente.query.get(internacao.paciente_id)
        if not paciente:
            return jsonify({
                'success': False,
                'message': 'Paciente não encontrado'
            }), 404

        # Montar histórico automaticamente se ainda não houver data de alta
        historico_internacao = internacao.historico_internacao
        if not internacao.data_alta:
            secoes_historico = []

            # HDA - História da Doença Atual
            if internacao.hda and internacao.hda.strip():
                secoes_historico.append(f"# HDA (História da Doença Atual):\n{internacao.hda.strip()}")

            # Evolução médica (última)
            ultima_evolucao = (
                EvolucaoAtendimentoClinica.query
                .filter_by(atendimentos_clinica_id=internacao.id)
                .order_by(EvolucaoAtendimentoClinica.data_evolucao.desc())
                .first()
            )
            if ultima_evolucao and ultima_evolucao.evolucao and ultima_evolucao.evolucao.strip():
                secoes_historico.append(f"# EVOLUÇÃO MÉDICA:\n{ultima_evolucao.evolucao.strip()}")

            # Anamnese
            if internacao.folha_anamnese and internacao.folha_anamnese.strip():
                secoes_historico.append(f"# ANAMNESE:\n{internacao.folha_anamnese.strip()}")

            # Conduta
            if internacao.conduta and internacao.conduta.strip():
                secoes_historico.append(f"# CONDUTA:\n{internacao.conduta.strip()}")

            # Exames Laboratoriais
            if internacao.exames_laboratoriais and internacao.exames_laboratoriais.strip():
                secoes_historico.append(f"# EXAMES LABORATORIAIS:\n{internacao.exames_laboratoriais.strip()}")

            historico_internacao = '\n\n'.join(secoes_historico) if secoes_historico else 'Histórico será montado automaticamente'

        # Formatar resposta
        resultado = {
            'success': True,
            'data_alta': internacao.data_alta.strftime('%d/%m/%Y %H:%M') if internacao.data_alta else None,
            'diagnostico': internacao.diagnostico or 'Não informado',
            'historico_internacao': historico_internacao or 'Não informado',
            'cuidados_gerais': internacao.cuidados_gerais or 'Não informado',
            'nome_paciente': paciente.nome or 'Não informado',
            'cpf_paciente': paciente.cpf or 'Não informado'
        }

        return jsonify(resultado)

    except Exception as e:
        logging.error(f'Erro ao buscar informações de alta para impressão: {str(e)}')
        logging.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': f'Erro ao buscar informações de alta: {str(e)}'
        }), 500



@bp.route('/api/imprimir-prescricao/<int:prescricao_id>', methods=['GET'])
def imprimir_prescricao(prescricao_id):
    prescricao = PrescricaoClinica.query.get(prescricao_id)
    if not prescricao:
        return abort(404, 'Prescrição não encontrada.')

    internacao = prescricao.internacao
    if not internacao:
        return abort(404, 'Internação não vinculada à prescrição.')

    atendimento = Atendimento.query.get(internacao.atendimento_id)
    if not atendimento:
        return abort(404, 'Atendimento não encontrado.')

    paciente = atendimento.paciente
    if not paciente:
        return abort(404, 'Paciente não encontrado.')

    # Abrir o modelo
    doc = Document('app/static/impressos/prescricao.docx')

    def substituir_variaveis(doc, substituicoes):
        substituicoes_realizadas = 0
        
        # Processar parágrafos
        for p in doc.paragraphs:
            for chave, valor in substituicoes.items():
                if chave in p.text:
                    p.text = p.text.replace(chave, str(valor))
                    substituicoes_realizadas += 1
                    
        # Processar tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for chave, valor in substituicoes.items():
                            if chave in p.text:
                                p.text = p.text.replace(chave, str(valor))
                                substituicoes_realizadas += 1
        
        logging.info(f"Realizadas {substituicoes_realizadas} substituições")
        return substituicoes_realizadas

    # Preencher os campos principais
    substituicoes = {
        '{{ paciente_nome }}': paciente.nome or '',
        '{{ id_atendimento}}': atendimento.id,
        '{{paciente_cor }}': paciente.cor or '',
        '{{ paciente_data_nascimento }}': paciente.data_nascimento.strftime('%d/%m/%Y') if paciente.data_nascimento else '',
        '{{paciente_sexo }}': paciente.sexo or '',
        '{{ paciente_cartao_sus }}': paciente.cartao_sus or '',
        '{{ paciente_telefone}}': paciente.telefone or '',
        '{{internamento_leito}}': internacao.leito or '',
        '{{internamento_data}}': internacao.data_internacao.strftime('%d/%m/%Y %H:%M') if internacao.data_internacao else '',
        '{{ paciente_nome_filiacao}}': paciente.filiacao or '',
        '{{ paciente_endereço }}': paciente.endereco or '',
        '{{ municipio_residencia }}': paciente.municipio or '',
        '{{atendimentos_alergia}}': paciente.alergias or '',
        '{{ prescricao_dieta }}': prescricao.texto_dieta or '',
    }

    substituir_variaveis(doc, substituicoes)

    # Substituir medicamentos da lista JSON
    medicamentos = prescricao.medicamentos_json or []
    data_prescricao = prescricao.horario_prescricao.strftime('%d/%m/%Y %H:%M') if prescricao.horario_prescricao else ''

    for med in medicamentos:
        doc_texto = doc.get_paragraphs()
        for p in doc.paragraphs:
            if '{{ prescricao_medicamento }}' in p.text:
                p.text = p.text.replace('{{ prescricao_medicamento }}', med.get('nome_medicamento', ''))
                break
        for p in doc.paragraphs:
            if '{{ prescricao_descricao }}' in p.text:
                p.text = p.text.replace('{{ prescricao_descricao }}', med.get('descricao_uso', ''))
                break
        for p in doc.paragraphs:
            if '{{ prescricao_datahorario }}' in p.text:
                p.text = p.text.replace('{{ prescricao_datahorario }}', data_prescricao)
                break

    # Salvar temporariamente em memória
    output = BytesIO()
    doc.save(output)
    output.seek(0)

    # Enviar para download
    nome_arquivo = f"prescricao_{prescricao.id}_{paciente.nome.replace(' ', '_')}.docx"
    return send_file(output, as_attachment=True, download_name=nome_arquivo, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

@bp.route('/imprimir/prescricao/<string:atendimento_id>', methods=['GET'])
def imprimir_prescricao_por_atendimento(atendimento_id):
    """
    Gera a prescrição mais recente de um atendimento específico
    """
    try:
        logging.info(f"Iniciando geração de prescrição para atendimento: {atendimento_id}")
        
        # Buscar o atendimento
        atendimento = Atendimento.query.get(atendimento_id)
        if not atendimento:
            return abort(404, 'Atendimento não encontrado.')

        # Buscar a internação relacionada ao atendimento
        internacao = Internacao.query.filter_by(atendimento_id=atendimento_id).first()
        if not internacao:
            return abort(404, 'Internação não encontrada para este atendimento.')

        # Buscar a prescrição mais recente desta internação
        prescricao = PrescricaoClinica.query.filter_by(
            atendimentos_clinica_id=internacao.id
        ).order_by(PrescricaoClinica.horario_prescricao.desc()).first()
        
        if not prescricao:
            return abort(404, 'Nenhuma prescrição encontrada para este atendimento.')

        paciente = atendimento.paciente
        if not paciente:
            return abort(404, 'Paciente não encontrado.')

        # Abrir o modelo
        doc = Document('app/static/impressos/prescricao.docx')

        def substituir_variaveis(doc, substituicoes):
            substituicoes_realizadas = 0
            
            # Processar parágrafos
            for p in doc.paragraphs:
                for chave, valor in substituicoes.items():
                    if chave in p.text:
                        p.text = p.text.replace(chave, str(valor))
                        substituicoes_realizadas += 1
                        
            # Processar tabelas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for chave, valor in substituicoes.items():
                                if chave in p.text:
                                    p.text = p.text.replace(chave, str(valor))
                                    substituicoes_realizadas += 1
            
            logging.info(f"Realizadas {substituicoes_realizadas} substituições")
            return substituicoes_realizadas

        # Preencher os campos principais
        substituicoes = {
            '{{ paciente_nome }}': paciente.nome or '',
            '{{ id_atendimento}}': str(atendimento.id),
            '{{paciente_cor }}': paciente.cor or '',
            '{{ paciente_data_nascimento }}': paciente.data_nascimento.strftime('%d/%m/%Y') if paciente.data_nascimento else '',
            '{{paciente_sexo }}': paciente.sexo or '',
            '{{ paciente_cartao_sus }}': paciente.cartao_sus or '',
            '{{ paciente_telefone}}': paciente.telefone or '',
            '{{internamento_leito}}': internacao.leito or '',
            '{{internamento_data}}': internacao.data_internacao.strftime('%d/%m/%Y %H:%M') if internacao.data_internacao else '',
            '{{ paciente_nome_filiacao}}': paciente.filiacao or '',
            '{{ paciente_endereço }}': paciente.endereco or '',
            '{{ municipio_residencia }}': paciente.municipio or '',
            '{{atendimentos_alergia}}': paciente.alergias or '',
            '{{ prescricao_dieta }}': prescricao.texto_dieta or '',
        }

        substituir_variaveis(doc, substituicoes)

        # Substituir medicamentos da lista JSON
        medicamentos = prescricao.medicamentos_json or []
        data_prescricao = prescricao.horario_prescricao.strftime('%d/%m/%Y %H:%M') if prescricao.horario_prescricao else ''

        logging.info(f"Processando {len(medicamentos)} medicamentos")

        # Processar medicamentos
        if medicamentos:
            primeiro_med = medicamentos[0]
            substituicoes_medicamentos = {
                '{{ prescricao_medicamento }}': primeiro_med.get('nome_medicamento', ''),
                '{{ prescricao_descricao }}': primeiro_med.get('descricao_uso', ''),
                '{{ prescricao_datahorario }}': data_prescricao
            }
            
            logging.info(f"Medicamento: {primeiro_med.get('nome_medicamento', 'N/A')}")
            substituir_variaveis(doc, substituicoes_medicamentos)
        else:
            # Limpar placeholders se não houver medicamentos
            substituicoes_vazias = {
                '{{ prescricao_medicamento }}': '',
                '{{ prescricao_descricao }}': '',
                '{{ prescricao_datahorario }}': data_prescricao
            }
            substituir_variaveis(doc, substituicoes_vazias)

        # Salvar temporariamente em memória
        output = BytesIO()
        doc.save(output)
        output.seek(0)

        # Enviar para download
        nome_arquivo = f"prescricao_{atendimento_id}_{paciente.nome.replace(' ', '_').replace('/', '')}.docx"
        return send_file(output, as_attachment=True, download_name=nome_arquivo, 
                        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    except Exception as e:
        logging.error(f"Erro ao gerar prescrição para atendimento {atendimento_id}: {str(e)}")
        return abort(500, f'Erro interno do servidor: {str(e)}')

@bp.route('/api/imprimir-prescricao-html/<int:prescricao_id>', methods=['GET'])
@login_required
def imprimir_prescricao_html(prescricao_id):
    """
    Gera uma página HTML para impressão de prescrição médica
    Formatada para papel A4 horizontal com design profissional
    """
    try:
        # Buscar a prescrição
        prescricao = PrescricaoClinica.query.get(prescricao_id)
        if not prescricao:
            abort(404, 'Prescrição não encontrada.')

        # Buscar a internação relacionada
        internacao = prescricao.internacao
        if not internacao:
            abort(404, 'Internação não vinculada à prescrição.')

        # Buscar o atendimento
        atendimento = Atendimento.query.get(internacao.atendimento_id)
        if not atendimento:
            abort(404, 'Atendimento não encontrado.')

        # Buscar o paciente
        paciente = atendimento.paciente
        if not paciente:
            abort(404, 'Paciente não encontrado.')

        # Buscar o médico que prescreveu
        medico = Funcionario.query.get(prescricao.medico_id) if prescricao.medico_id else None

        # Buscar aprazamentos relacionados a esta prescrição
        aprazamentos = Aprazamento.query.filter_by(prescricao_id=prescricao_id).order_by(
            Aprazamento.data_hora_aprazamento
        ).all()

        # Organizar medicamentos com seus respectivos aprazamentos
        medicamentos_com_aprazamentos = []
        medicamentos_json = prescricao.medicamentos_json or []
        
        for medicamento in medicamentos_json:
            nome_med = medicamento.get('nome_medicamento', '')
            aprazamentos_med = [a for a in aprazamentos if a.nome_medicamento == nome_med]
            medicamentos_com_aprazamentos.append({
                'medicamento': medicamento,
                'aprazamentos': aprazamentos_med
            })

        # Renderizar template HTML para impressão
        return render_template('impressao_prescricao.html',
                             prescricao=prescricao,
                             internacao=internacao,
                             atendimento=atendimento,
                             paciente=paciente,
                             medico=medico,
                             medicamentos_com_aprazamentos=medicamentos_com_aprazamentos)

    except Exception as e:
        logging.error(f"Erro ao gerar impressão de prescrição: {str(e)}")
        abort(500, f'Erro interno: {str(e)}')

@bp.route('/api/imprimir-evolucao-html/<int:evolucao_id>', methods=['GET'])
@login_required
def imprimir_evolucao_html(evolucao_id):
    """
    Gera uma página HTML para impressão de evolução médica
    Formatada para papel A4 com design profissional
    """
    try:
        # Buscar a evolução
        evolucao = EvolucaoAtendimentoClinica.query.get(evolucao_id)
        if not evolucao:
            abort(404, 'Evolução não encontrada.')

        # Buscar a internação relacionada
        internacao = Internacao.query.get(evolucao.atendimentos_clinica_id)
        if not internacao:
            abort(404, 'Internação não vinculada à evolução.')

        # Buscar o atendimento
        atendimento = Atendimento.query.get(internacao.atendimento_id)
        if not atendimento:
            abort(404, 'Atendimento não encontrado.')

        # Buscar o paciente
        paciente = atendimento.paciente
        if not paciente:
            abort(404, 'Paciente não encontrado.')

        # Buscar o médico que fez a evolução
        medico = Funcionario.query.get(evolucao.funcionario_id) if evolucao.funcionario_id else None

        # Renderizar template HTML para impressão
        return render_template('impressao_evolucao.html',
                             evolucao=evolucao,
                             internacao=internacao,
                             atendimento=atendimento,
                             paciente=paciente,
                             medico=medico)

    except Exception as e:
        logging.error(f"Erro ao gerar impressão de evolução: {str(e)}")
        abort(500, f'Erro interno: {str(e)}')


# Impressão de evolução individual de enfermagem
@bp.route('/api/imprimir-evolucao-enfermagem-html/<int:evolucao_id>', methods=['GET'])
@login_required
def imprimir_evolucao_enfermagem_html(evolucao_id):
    """
    Gera uma página HTML para impressão de uma evolução de enfermagem individual
    """
    try:
        # Buscar a evolução de enfermagem
        evolucao = EvolucaoEnfermagem.query.get(evolucao_id)
        if not evolucao:
            abort(404, 'Evolução de enfermagem não encontrada.')

        # Buscar a internação relacionada
        internacao = Internacao.query.get(evolucao.atendimentos_clinica_id)
        if not internacao:
            abort(404, 'Internação não vinculada à evolução de enfermagem.')

        # Buscar o atendimento
        atendimento = Atendimento.query.get(internacao.atendimento_id)
        if not atendimento:
            abort(404, 'Atendimento não encontrado.')

        # Buscar o paciente
        paciente = atendimento.paciente
        if not paciente:
            abort(404, 'Paciente não encontrado.')

        # Buscar o profissional de enfermagem que fez a evolução
        enfermeiro = Funcionario.query.get(evolucao.funcionario_id) if evolucao.funcionario_id else None

        # Renderizar template HTML para impressão
        return render_template('impressao_evolucao_enfermagem.html',
                               evolucao=evolucao,
                               internacao=internacao,
                               atendimento=atendimento,
                               paciente=paciente,
                               enfermeiro=enfermeiro)

    except Exception as e:
        logging.error(f"Erro ao gerar impressão de evolução de enfermagem: {str(e)}")
        abort(500, f'Erro interno: {str(e)}')

@bp.route('/api/ultima-evolucao-id/<string:atendimento_id>', methods=['GET'])
@login_required
def buscar_ultima_evolucao_id(atendimento_id):
    """
    Busca o ID da última evolução registrada para um atendimento específico
    """
    try:
        # Buscar a internação pelo atendimento_id
        internacao = Internacao.query.filter_by(atendimento_id=atendimento_id).first()
        if not internacao:
            return jsonify({'success': False, 'message': 'Internação não encontrada'})

        # Buscar a evolução mais recente desta internação
        evolucao = EvolucaoAtendimentoClinica.query.filter_by(
            atendimentos_clinica_id=internacao.id
        ).order_by(EvolucaoAtendimentoClinica.data_evolucao.desc()).first()
        
        if not evolucao:
            return jsonify({'success': False, 'message': 'Nenhuma evolução encontrada'})

        return jsonify({
            'success': True,
            'evolucao_id': evolucao.id,
            'data_evolucao': evolucao.data_evolucao.isoformat() if evolucao.data_evolucao else None
        })

    except Exception as e:
        logging.error(f"Erro ao buscar última evolução para atendimento {atendimento_id}: {str(e)}")
        return jsonify({'success': False, 'message': f'Erro interno: {str(e)}'})


@bp.route('/api/imprimir-observacao-html/<int:observacao_id>', methods=['GET'])
@login_required
def imprimir_observacao_html(observacao_id):
    """
    Gera uma página HTML para impressão de evolução médica
    Formatada para papel A4 com design profissional
    """
    try:
        # Buscar a evolução
        evolucao = EvolucaoAtendimentoClinica.query.get(observacao_id)
        if not evolucao:
            abort(404, 'Evolução não encontrada.')

        # Buscar a internação relacionada
        internacao = Internacao.query.get(evolucao.atendimentos_clinica_id)
        if not internacao:
            abort(404, 'Internação não vinculada à evolução.')

        # Buscar o atendimento
        atendimento = Atendimento.query.get(internacao.atendimento_id)
        if not atendimento:
            abort(404, 'Atendimento não encontrado.')

        # Buscar o paciente
        paciente = atendimento.paciente
        if not paciente:
            abort(404, 'Paciente não encontrado.')

        # Buscar o médico que fez a evolução
        medico = Funcionario.query.get(evolucao.funcionario_id) if evolucao.funcionario_id else None

        # Renderizar template HTML para impressão
        return render_template('impressao_observacao.html',
                             evolucao=evolucao,
                             internacao=internacao,
                             atendimento=atendimento,
                             paciente=paciente,
                             medico=medico)

    except Exception as e:
        logging.error(f"Erro ao gerar impressão de evolução: {str(e)}")
        abort(500, f'Erro interno: {str(e)}')


@bp.route('/imprimir/sae/<string:atendimento_id>')
@login_required
def imprimir_sae(atendimento_id):
    """
    Gera uma página HTML para impressão da SAE (Sistematização da Assistência de Enfermagem)
    """
    try:
        from datetime import datetime
        
        # Buscar o atendimento
        atendimento = Atendimento.query.filter_by(id=atendimento_id).first()
        if not atendimento:
            abort(404, 'Atendimento não encontrado.')

        # Buscar a internação relacionada ao atendimento
        internacao = Internacao.query.filter_by(atendimento_id=atendimento_id).first()
        if not internacao:
            abort(404, 'Internação não encontrada para este atendimento.')

        # Buscar o paciente
        paciente = atendimento.paciente
        if not paciente:
            abort(404, 'Paciente não encontrado.')

        # Buscar a SAE mais recente para este paciente
        sae = InternacaoSae.query.filter_by(
            paciente_id=internacao.paciente_id
        ).order_by(
            InternacaoSae.data_registro.desc()
        ).first()
        
        if not sae:
            flash('Nenhuma SAE encontrada para este paciente.', 'warning')
            return redirect(url_for('main.pacientes_internados'))
        
        return render_template('impressao_sae_enfermagem.html',
                             paciente=paciente,
                             internacao=internacao,
                             atendimento=atendimento,
                             sae=sae,
                             atendimento_id=atendimento_id)

    except Exception as e:
        logging.error(f"Erro ao gerar impressão de SAE: {str(e)}")
        abort(500, f'Erro interno: {str(e)}')

@bp.route('/imprimir/ficha_admissao/<string:atendimento_id>')
@login_required
def imprimir_ficha_admissao(atendimento_id):
    """
    Gera uma página HTML para impressão da Ficha de Admissão
    Formatada para papel A4 horizontal com design profissional
    """
    try:
        from datetime import datetime
        
        # Buscar o atendimento
        atendimento = Atendimento.query.filter_by(id=atendimento_id).first()
        if not atendimento:
            abort(404, 'Atendimento não encontrado.')

        # Buscar a internação relacionada ao atendimento
        internacao = Internacao.query.filter_by(atendimento_id=atendimento_id).first()
        if not internacao:
            abort(404, 'Internação não encontrada para este atendimento.')

        # Buscar o paciente
        paciente = atendimento.paciente
        if not paciente:
            abort(404, 'Paciente não encontrado.')

        # Buscar o médico responsável
        medico = None
        if internacao.medico_id:
            medico = Funcionario.query.get(internacao.medico_id)
        elif atendimento.medico_id:
            medico = Funcionario.query.get(atendimento.medico_id)

        # Buscar enfermeiro responsável
        enfermeiro = None
        if internacao.enfermeiro_id:
            enfermeiro = Funcionario.query.get(internacao.enfermeiro_id)
        elif atendimento.enfermeiro_id:
            enfermeiro = Funcionario.query.get(atendimento.enfermeiro_id)

        # Data atual para o rodapé
        data_impressao = datetime.now(ZoneInfo("America/Sao_Paulo"))

        # Calcular idade do paciente
        idade = None
        if paciente.data_nascimento:
            hoje = datetime.now(ZoneInfo("America/Sao_Paulo")).date()
            idade = hoje.year - paciente.data_nascimento.year
            if hoje.month < paciente.data_nascimento.month or (hoje.month == paciente.data_nascimento.month and hoje.day < paciente.data_nascimento.day):
                idade -= 1

        # Renderizar template HTML para impressão
        return render_template('imprimir_ficha_admissao.html',
                             internacao=internacao,
                             atendimento=atendimento,
                             paciente=paciente,
                             medico=medico,
                             enfermeiro=enfermeiro,
                             idade=idade,
                             data_impressao=data_impressao,
                             hospital_nome='Hospital Sistema HSOP')

    except Exception as e:
        logging.error(f"Erro ao gerar impressão de ficha de admissão: {str(e)}")
        abort(500, f'Erro interno: {str(e)}')

@bp.route('/imprimir/ficha_atendimento/<string:atendimento_id>')
@login_required
def imprimir_ficha_atendimento(atendimento_id):
    """
    Gera uma página HTML para impressão da Ficha de Atendimento Médico
    Formatada para papel A4 frente e verso com design profissional
    """
    try:
        from datetime import datetime
        from zoneinfo import ZoneInfo
        
        # Buscar o atendimento
        atendimento = Atendimento.query.filter_by(id=atendimento_id).first()
        if not atendimento:
            abort(404, 'Atendimento não encontrado.')

        # Buscar o paciente
        paciente = atendimento.paciente
        if not paciente:
            abort(404, 'Paciente não encontrado.')

        # Buscar o médico responsável
        medico = None
        if atendimento.medico_id:
            medico = Funcionario.query.get(atendimento.medico_id)

        # Buscar enfermeiro responsável
        enfermeiro = None
        if atendimento.enfermeiro_id:
            enfermeiro = Funcionario.query.get(atendimento.enfermeiro_id)

        # Buscar a internação relacionada (se houver)
        internacao = Internacao.query.filter_by(atendimento_id=atendimento_id).first()

        # Helpers para formatação com conversão para horário do Brasil (UTC-3)
        def fmt_datetime(dt):
            return formatar_datetime_br(dt, '%d/%m/%Y %H:%M') if dt else None

        # Buscar prescrições de emergência (porta) pelo atendimento_id
        prescricoes_processadas = []
        prescricoes = PrescricaoEmergencia.query.filter_by(
            atendimento_id=atendimento_id
        ).order_by(PrescricaoEmergencia.horario_prescricao.desc()).all()

        for prescricao in prescricoes:
            prescricoes_processadas.append({
                'id': prescricao.id,
                'horario_prescricao_iso': prescricao.horario_prescricao.isoformat() if prescricao.horario_prescricao else None,
                'horario_prescricao_pretty': fmt_datetime(prescricao.horario_prescricao),
                'texto_dieta': prescricao.texto_dieta,
                'texto_procedimento_medico': prescricao.texto_procedimento_medico,
                'texto_procedimento_multi': prescricao.texto_procedimento_multi,
                'medicamentos_json': prescricao.medicamentos,  # Nova tabela usa campo 'medicamentos'
                'medicamentos': prescricao.medicamentos  # Disponibiliza também como 'medicamentos'
            })

        # Data atual para o rodapé
        data_impressao = datetime.now(ZoneInfo("America/Sao_Paulo"))

        # Calcular idade do paciente
        idade = None
        if paciente.data_nascimento:
            hoje = datetime.now(ZoneInfo("America/Sao_Paulo")).date()
            idade = hoje.year - paciente.data_nascimento.year
            if hoje.month < paciente.data_nascimento.month or (hoje.month == paciente.data_nascimento.month and hoje.day < paciente.data_nascimento.day):
                idade -= 1

        # Renderizar template HTML para impressão
        return render_template('ficha_atendimento_medico.html',
                             atendimento=atendimento,
                             paciente=paciente,
                             medico=medico,
                             enfermeiro=enfermeiro,
                             internacao=internacao,
                             prescricoes=prescricoes_processadas,
                             idade=idade,
                             current_date=data_impressao,
                             # Dados do paciente
                             paciente_cartao_sus=paciente.cartao_sus or '',
                             paciente_nome=paciente.nome or '',
                             paciente_nome_social=paciente.nome_social or '',
                             paciente_filiacao=paciente.filiacao or '',
                             paciente_data_nascimento=paciente.data_nascimento.strftime('%d/%m/%Y') if paciente.data_nascimento else '',
                             paciente_idade=str(idade) if idade else '',
                             paciente_sexo=paciente.sexo or '',
                             paciente_cpf=paciente.cpf or '',
                             paciente_endereco=paciente.endereco or '',
                             paciente_bairro=paciente.bairro or '',
                             paciente_municipio=paciente.municipio or '',
                             paciente_telefone=paciente.telefone or '',
                             paciente_cor=paciente.cor or '',
                             # Dados do atendimento
                             atendimento_id=atendimento.id,
                             atendimento_data=atendimento.data_atendimento.strftime('%d/%m/%Y') if atendimento.data_atendimento else '',
                             atendimento_hora=atendimento.hora_atendimento.strftime('%H:%M') if atendimento.hora_atendimento else '',
                             atendimento_status=atendimento.status or '',
                             atendimento_classificacao_risco=atendimento.classificacao_risco or '',
                             atendimento_pressao=atendimento.pressao or '',
                             atendimento_pulso=atendimento.pulso or '',
                             atendimento_sp02=atendimento.sp02 or '',
                             atendimento_temp=atendimento.temp or '',
                             atendimento_peso=atendimento.peso or '',
                             atendimento_altura=atendimento.altura or '',
                             atendimento_dx=atendimento.dx or '',
                             atendimento_fr=atendimento.fr or '',
                             atendimento_alergias=paciente.alergias or '',
                             atendimento_horario_triagem=fmt_datetime(atendimento.horario_triagem) or 'Não realizado',
                             atendimento_triagem=atendimento.triagem or '',
                             atendimento_horario_consulta_medica=fmt_datetime(atendimento.horario_consulta_medica) or 'Não realizado',
                             atendimento_anamnese_exame_fisico=atendimento.anamnese_exame_fisico or '',
                             atendimento_reavaliacao=atendimento.reavaliacao or '',
                             atendimento_exames=atendimento.exames or '',
                             atendimento_conduta_final=atendimento.conduta_final or '',
                             atendimento_observacao=atendimento.observacao or '',
                             # Dados do médico
                             medico_nome=medico.nome if medico else '',
                             medico_crm=medico.numero_profissional if medico else '')

    except Exception as e:
        logging.error(f"Erro ao gerar impressão de ficha de atendimento: {str(e)}")
        abort(500, f'Erro interno: {str(e)}')

@bp.route('/imprimir/identificacao_paciente/<string:atendimento_id>')
@login_required
def imprimir_identificacao_paciente(atendimento_id):
    """
    Gera uma página HTML para impressão de identificação do paciente
    Formatada para uso em pulseiras, adesivos ou etiquetas de identificação
    Focada nos dados do paciente, leito, diagnóstico e informações para enfermeiros e técnicos
    """
    try:
        from datetime import datetime
        
        # Buscar o atendimento
        atendimento = Atendimento.query.filter_by(id=atendimento_id).first()
        if not atendimento:
            abort(404, 'Atendimento não encontrado.')

        # Buscar a internação relacionada ao atendimento
        internacao = Internacao.query.filter_by(atendimento_id=atendimento_id).first()
        if not internacao:
            abort(404, 'Internação não encontrada para este atendimento.')

        # Buscar o paciente
        paciente = atendimento.paciente
        if not paciente:
            abort(404, 'Paciente não encontrado.')

        # Buscar o médico responsável
        medico = None
        if internacao.medico_id:
            medico = Funcionario.query.get(internacao.medico_id)
        elif atendimento.medico_id:
            medico = Funcionario.query.get(atendimento.medico_id)

        # Buscar enfermeiro responsável
        enfermeiro = None
        if internacao.enfermeiro_id:
            enfermeiro = Funcionario.query.get(internacao.enfermeiro_id)
        elif atendimento.enfermeiro_id:
            enfermeiro = Funcionario.query.get(atendimento.enfermeiro_id)

        # Calcular idade do paciente
        idade = None
        if paciente.data_nascimento:
            hoje = datetime.now(ZoneInfo("America/Sao_Paulo")).date()
            idade = hoje.year - paciente.data_nascimento.year
            if hoje.month < paciente.data_nascimento.month or (hoje.month == paciente.data_nascimento.month and hoje.day < paciente.data_nascimento.day):
                idade -= 1

        # Data atual para o rodapé
        data_impressao = datetime.now(ZoneInfo("America/Sao_Paulo"))

        # Buscar prescrições médicas ativas para alertas
        prescricoes = PrescricaoClinica.query.filter_by(atendimentos_clinica_id=internacao.id).order_by(
            PrescricaoClinica.horario_prescricao.desc()
        ).limit(3).all()

        # Renderizar template HTML para impressão
        return render_template('imprimir_identificacao_paciente.html',
                             internacao=internacao,
                             atendimento=atendimento,
                             paciente=paciente,
                             medico=medico,
                             enfermeiro=enfermeiro,
                             idade=idade,
                             data_impressao=data_impressao,
                             prescricoes=prescricoes,
                             hospital_nome='Hospital Sistema HSOP')

    except Exception as e:
        logging.error(f"Erro ao gerar impressão de identificação do paciente: {str(e)}")
        abort(500, f'Erro interno: {str(e)}')

# API para verificar duplicatas de pacientes
@bp.route('/api/pacientes/verificar-duplicatas', methods=['POST'])
@login_required
def verificar_duplicatas_pacientes():
    """
    Verifica se já existe um paciente com o mesmo CPF ou Cartão SUS.
    """
    try:
        current_user = get_current_user()
        if current_user.cargo.lower() not in ['medico', 'enfermeiro', 'multi']:
            return jsonify({
                'success': False,
                'message': 'Acesso não autorizado'
            }), 403

        dados = request.get_json()
        
        if not dados:
            return jsonify({
                'success': False,
                'message': 'Dados não fornecidos'
            }), 400

        cpf = dados.get('cpf', '').strip()
        cartao_sus = dados.get('cartao_sus', '').strip()

        if not cpf and not cartao_sus:
            return jsonify({
                'success': True,
                'duplicata_encontrada': False,
                'message': 'Nenhum documento fornecido para verificação'
            })

        paciente_encontrado = None
        campo_duplicado = None

        # Verificar duplicata por CPF
        if cpf:
            paciente_encontrado = Paciente.query.filter_by(cpf=cpf).first()
            if paciente_encontrado:
                campo_duplicado = 'cpf'

        # Se não encontrou por CPF, verificar por Cartão SUS
        if not paciente_encontrado and cartao_sus:
            paciente_encontrado = Paciente.query.filter_by(cartao_sus=cartao_sus).first()
            if paciente_encontrado:
                campo_duplicado = 'cartao_sus'

        if paciente_encontrado:
            return jsonify({
                'success': True,
                'duplicata_encontrada': True,
                'campo_duplicado': campo_duplicado,
                'paciente_existente': {
                    'id': paciente_encontrado.id,
                    'nome': paciente_encontrado.nome,
                    'cpf': paciente_encontrado.cpf,
                    'cartao_sus': paciente_encontrado.cartao_sus,
                    'data_nascimento': paciente_encontrado.data_nascimento.isoformat() if paciente_encontrado.data_nascimento else None,
                    'sexo': paciente_encontrado.sexo
                }
            })

        return jsonify({
            'success': True,
            'duplicata_encontrada': False,
            'message': 'Nenhuma duplicata encontrada'
        })

    except Exception as e:
        logging.error(f'Erro ao verificar duplicatas: {str(e)}')
        logging.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': f'Erro interno do servidor: {str(e)}'
        }), 500


@bp.route('/api/atendimentos/cadastrar', methods=['POST'])
@login_required
def cadastrar_atendimento():
    """
    Cria um novo atendimento automaticamente com dados básicos.
    Recebe: paciente_id, data_atendimento, hora_atendimento
    """
    try:
        current_user = get_current_user()
        if not current_user or current_user.cargo.lower() != 'recepcionista':
            return jsonify({'success': False, 'message': 'Acesso não autorizado. Apenas recepcionistas podem criar atendimentos.'}), 403

        dados = request.get_json()

        if not dados:
            return jsonify({'success': False, 'message': 'Dados não fornecidos'}), 400

        paciente_id = dados.get('paciente_id')
        data_atendimento = dados.get('data_atendimento')
        hora_atendimento = dados.get('hora_atendimento')

        # Validações básicas
        if not paciente_id:
            return jsonify({'success': False, 'message': 'ID do paciente é obrigatório'}), 400

        if not data_atendimento:
            return jsonify({'success': False, 'message': 'Data do atendimento é obrigatória'}), 400

        if not hora_atendimento:
            return jsonify({'success': False, 'message': 'Hora do atendimento é obrigatória'}), 400

        # Verificar se o paciente existe
        paciente = Paciente.query.get(paciente_id)
        if not paciente:
            return jsonify({'success': False, 'message': 'Paciente não encontrado'}), 404

        # Gerar ID único no formato YYMMNNNN (ano, mês e 4 dígitos aleatórios)
        import random
        from datetime import datetime as dt

        while True:
            prefixo = dt.now().strftime('%y%m')
            sufixo = str(random.randint(0, 9999)).zfill(4)
            novo_id = f"{prefixo}{sufixo}"
            if not Atendimento.query.get(novo_id):
                break

        # Criar o atendimento
        novo_atendimento = Atendimento(
            id=novo_id,
            paciente_id=paciente_id,
            funcionario_id=current_user.id,  # Recepcionista que criou
            data_atendimento=datetime.strptime(data_atendimento, '%Y-%m-%d').date(),
            hora_atendimento=datetime.strptime(hora_atendimento, '%H:%M').time(),
            status='Aguardando Triagem'  # Status inicial para novos atendimentos
        )

        db.session.add(novo_atendimento)
        db.session.commit()

        logging.info(f'Atendimento criado: {novo_id} para paciente {paciente_id} por recepcionista {current_user.id}')

        return jsonify({
            'success': True,
            'message': 'Atendimento criado com sucesso',
            'atendimento': {
                'id': novo_atendimento.id,
                'paciente_id': novo_atendimento.paciente_id,
                'funcionario_id': novo_atendimento.funcionario_id,
                'data_atendimento': novo_atendimento.data_atendimento.isoformat(),
                'hora_atendimento': novo_atendimento.hora_atendimento.strftime('%H:%M')
            }
        }), 201

    except Exception as e:
        db.session.rollback()
        logging.error(f'Erro ao criar atendimento: {str(e)}')
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': f'Erro interno do servidor: {str(e)}'        }), 500


@bp.route('/ficha-de-atendimento')
@login_required
def ficha_de_atendimento():
    """
    Página de ficha de atendimento para recepcionista.
    Recebe paciente_id como parâmetro GET.
    """
    try:
        current_user = get_current_user()
        if not current_user or current_user.cargo.lower() != 'recepcionista':
            flash('Acesso não autorizado. Apenas recepcionistas podem acessar esta página.', 'danger')
            return redirect(url_for('main.index'))

        paciente_id = request.args.get('paciente_id')

        if not paciente_id:
            flash('ID do paciente não fornecido.', 'danger')
            return redirect(url_for('main.recepcionista'))

        # Buscar dados do paciente
        paciente = Paciente.query.get(paciente_id)
        if not paciente:
            flash('Paciente não encontrado.', 'danger')
            return redirect(url_for('main.recepcionista'))

        # Calcular idade do paciente
        idade = None
        if paciente.data_nascimento:
            hoje = datetime.now().date()
            idade = hoje.year - paciente.data_nascimento.year - ((hoje.month, hoje.day) < (paciente.data_nascimento.month, paciente.data_nascimento.day))

        return render_template('ficha_de_atendimento_recepçao.html',
                             paciente=paciente,
                             idade=idade,
                             current_user=current_user)

    except Exception as e:
        logging.error(f'Erro ao carregar ficha de atendimento: {str(e)}')
        logging.error(traceback.format_exc())
        flash('Erro ao carregar ficha de atendimento. Por favor, tente novamente.', 'danger')
        return redirect(url_for('main.recepcionista'))


        return render_template('clinica_evolucao_paciente_enfermeiro.html',
                            paciente=paciente,
                            internacao=internacao)
        
    except Exception as e:
        logging.error(f"Erro ao acessar evolução do paciente (enfermeiro): {str(e)}")
        logging.error(traceback.format_exc())
        flash('Erro ao acessar a evolução do paciente. Por favor, tente novamente.', 'danger')
        return redirect(url_for('main.pacientes_internados'))

@bp.route('/clinica/impressoes-enfermagem/<string:atendimento_id>')
@login_required
def impressoes_enfermagem(atendimento_id):
    try:
        # Buscar dados do paciente e internação
        internacao = Internacao.query.filter_by(atendimento_id=atendimento_id).first()
        if not internacao:
            flash('Internação não encontrada.', 'danger')
            return redirect(url_for('main.pacientes_internados'))
        
        paciente = Paciente.query.get(internacao.paciente_id)
        if not paciente:
            flash('Paciente não encontrado.', 'danger')
            return redirect(url_for('main.pacientes_internados'))
        
        
        return render_template('impressoes_enfermagem.html', 
                            paciente=paciente, 
                            internacao=internacao,
                            atendimento_id=atendimento_id)
        
    except Exception as e:
        logging.error(f"Erro ao acessar impressões de enfermagem: {str(e)}")
        logging.error(traceback.format_exc())
        flash('Erro ao acessar as impressões de enfermagem. Por favor, tente novamente.', 'danger')
        return redirect(url_for('main.pacientes_internados'))

@bp.route('/imprimir/admissao_enfermagem/<string:atendimento_id>')
@login_required
def imprimir_admissao_enfermagem(atendimento_id):
    """
    Imprime a admissão de enfermagem do paciente
    """
    try:
        # Buscar o atendimento
        atendimento = Atendimento.query.get(atendimento_id)
        if not atendimento:
            flash('Atendimento não encontrado.', 'danger')
            return redirect(url_for('main.pacientes_internados'))

        # Buscar a internação
        internacao = Internacao.query.filter_by(atendimento_id=atendimento_id).first()
        if not internacao:
            flash('Internação não encontrada.', 'danger')
            return redirect(url_for('main.pacientes_internados'))
        
        # Buscar o paciente
        paciente = Paciente.query.get(internacao.paciente_id)
        if not paciente:
            flash('Paciente não encontrado.', 'danger')
            return redirect(url_for('main.pacientes_internados'))
        
        # Buscar a admissão de enfermagem mais recente
        admissao = AdmissaoEnfermagem.query.filter_by(
            internacao_id=internacao.id
        ).order_by(
            AdmissaoEnfermagem.data_hora.desc()
        ).first()
        
        # Buscar dados do enfermeiro responsável pela admissão
        enfermeiro = None
        if admissao and admissao.enfermeiro_id:
            enfermeiro = Funcionario.query.get(admissao.enfermeiro_id)
        
        # Se não houver admissão na tabela específica, verificar campo legado
        texto_admissao = None
        data_admissao = None
        
        if admissao:
            texto_admissao = admissao.admissao_texto
            data_admissao = admissao.data_hora
        elif internacao.admissao_enfermagem:
            texto_admissao = internacao.admissao_enfermagem
            data_admissao = internacao.data_internacao
        
        if not texto_admissao:
            flash('Nenhuma admissão de enfermagem encontrada para este paciente.', 'warning')
            return redirect(url_for('main.pacientes_internados'))
        
        return render_template('impressao_admissao_enfermagem.html',
                             paciente=paciente,
                             internacao=internacao,
                             atendimento=atendimento,
                             admissao=admissao,
                             texto_admissao=texto_admissao,
                             data_admissao=data_admissao,
                             enfermeiro=enfermeiro,
                             atendimento_id=atendimento_id)
        
    except Exception as e:
        logging.error(f"Erro ao imprimir admissão de enfermagem: {str(e)}")
        logging.error(traceback.format_exc())
        flash('Erro ao gerar impressão da admissão de enfermagem. Por favor, tente novamente.', 'danger')
        return redirect(url_for('main.pacientes_internados'))

@bp.route('/imprimir/sae_enfermagem/<string:atendimento_id>')
@login_required
def imprimir_sae_enfermagem(atendimento_id):
    """
    Imprime a SAE (Sistematização da Assistência de Enfermagem) do paciente
    """
    try:
        # Buscar o atendimento
        atendimento = Atendimento.query.get(atendimento_id)
        if not atendimento:
            flash('Atendimento não encontrado.', 'danger')
            return redirect(url_for('main.pacientes_internados'))

        # Buscar a internação
        internacao = Internacao.query.filter_by(atendimento_id=atendimento_id).first()
        if not internacao:
            flash('Internação não encontrada.', 'danger')
            return redirect(url_for('main.pacientes_internados'))
        
        # Buscar o paciente
        paciente = Paciente.query.get(internacao.paciente_id)
        if not paciente:
            flash('Paciente não encontrado.', 'danger')
            return redirect(url_for('main.pacientes_internados'))
        
        # Buscar a SAE mais recente deste paciente
        sae = InternacaoSae.query.filter_by(
            paciente_id=internacao.paciente_id
        ).order_by(
            InternacaoSae.data_registro.desc()
        ).first()
        
        if not sae:
            flash('Nenhuma SAE encontrada para este paciente.', 'warning')
            return redirect(url_for('main.pacientes_internados'))
        
        return render_template('impressao_sae_enfermagem.html',
                             paciente=paciente,
                             internacao=internacao,
                             atendimento=atendimento,
                             sae=sae,
                             atendimento_id=atendimento_id)
        
    except Exception as e:
        logging.error(f"Erro ao imprimir SAE de enfermagem: {str(e)}")
        logging.error(traceback.format_exc())
        flash('Erro ao gerar impressão da SAE. Por favor, tente novamente.', 'danger')
        return redirect(url_for('main.pacientes_internados'))

# API para listar datas disponíveis das evoluções de enfermagem
@bp.route('/api/enfermagem/evolucoes/datas/<string:atendimento_id>', methods=['GET'])
@login_required
def listar_datas_evolucoes_enfermagem(atendimento_id):
    """
    Lista as datas disponíveis das evoluções de enfermagem para um atendimento específico.
    """
    try:
        # Verificar se o usuário é enfermeiro ou médico
        current_user = get_current_user()
        if current_user.cargo.lower() not in ['enfermeiro', 'medico']:
            return jsonify({
                'success': False,
                'message': 'Acesso permitido apenas para enfermeiros e médicos'
            }), 403

        # Buscar a internação
        internacao = Internacao.query.filter_by(atendimento_id=atendimento_id).first()
        if not internacao:
            return jsonify({
                'success': False,
                'message': 'Internação não encontrada'
            }), 404

        # Buscar as datas das evoluções
        evolucoes = EvolucaoEnfermagem.query\
            .filter_by(atendimentos_clinica_id=internacao.id)\
            .order_by(EvolucaoEnfermagem.data_evolucao.desc())\
            .all()

        # Agrupar por data
        datas_disponiveis = {}
        for evolucao in evolucoes:
            if evolucao.data_evolucao:
                data_str = evolucao.data_evolucao.strftime('%Y-%m-%d')
                data_display = evolucao.data_evolucao.strftime('%d/%m/%Y')
                
                if data_str not in datas_disponiveis:
                    datas_disponiveis[data_str] = {
                        'data': data_str,
                        'data_display': data_display,
                        'quantidade': 0
                    }
                datas_disponiveis[data_str]['quantidade'] += 1

        # Converter para lista ordenada
        datas_lista = list(datas_disponiveis.values())
        datas_lista.sort(key=lambda x: x['data'], reverse=True)

        return jsonify({
            'success': True,
            'datas': datas_lista
        }), 200

    except Exception as e:
        logging.error(f'Erro ao listar datas das evoluções de enfermagem: {str(e)}')
        logging.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': f'Erro interno do servidor: {str(e)}'
        }), 500

# Rota para imprimir evoluções de enfermagem de uma data específica
@bp.route('/imprimir/evolucoes_enfermagem/<string:atendimento_id>')
@login_required
def imprimir_evolucoes_enfermagem(atendimento_id):
    """
    Imprime as evoluções de enfermagem de uma data específica
    """
    try:
        # Pegar a data da query string
        data_selecionada_str = request.args.get('data')
        if not data_selecionada_str:
            flash('Data não especificada para impressão.', 'warning')
            return redirect(url_for('main.impressoes_enfermagem', atendimento_id=atendimento_id))

        # Converter a data
        try:
            data_selecionada = datetime.strptime(data_selecionada_str, '%Y-%m-%d').date()
        except ValueError:
            flash('Formato de data inválido.', 'danger')
            return redirect(url_for('main.impressoes_enfermagem', atendimento_id=atendimento_id))

        # Buscar o atendimento
        atendimento = Atendimento.query.get(atendimento_id)
        if not atendimento:
            flash('Atendimento não encontrado.', 'danger')
            return redirect(url_for('main.pacientes_internados'))

        # Buscar a internação
        internacao = Internacao.query.filter_by(atendimento_id=atendimento_id).first()
        if not internacao:
            flash('Internação não encontrada.', 'danger')
            return redirect(url_for('main.pacientes_internados'))
        
        # Buscar o paciente
        paciente = Paciente.query.get(internacao.paciente_id)
        if not paciente:
            flash('Paciente não encontrado.', 'danger')
            return redirect(url_for('main.pacientes_internados'))
        
        # Buscar as evoluções da data selecionada
        inicio_dia = datetime.combine(data_selecionada, datetime.min.time())
        fim_dia = datetime.combine(data_selecionada, datetime.max.time())
        
        evolucoes = EvolucaoEnfermagem.query\
            .filter_by(atendimentos_clinica_id=internacao.id)\
            .filter(EvolucaoEnfermagem.data_evolucao >= inicio_dia)\
            .filter(EvolucaoEnfermagem.data_evolucao <= fim_dia)\
            .join(Funcionario, EvolucaoEnfermagem.funcionario_id == Funcionario.id)\
            .add_columns(
                EvolucaoEnfermagem.id,
                EvolucaoEnfermagem.data_evolucao,
                EvolucaoEnfermagem.texto,
                Funcionario.nome.label('enfermeiro_nome'),
                Funcionario.numero_profissional.label('enfermeiro_coren')
            )\
            .order_by(EvolucaoEnfermagem.data_evolucao.asc())\
            .all()
        
        if not evolucoes:
            flash(f'Nenhuma evolução de enfermagem encontrada para o dia {data_selecionada.strftime("%d/%m/%Y")}.', 'warning')
            return redirect(url_for('main.impressoes_enfermagem', atendimento_id=atendimento_id))
        
        # Preparar dados das evoluções
        evolucoes_dados = []
        for evolucao in evolucoes:
            evolucoes_dados.append({
                'id': evolucao.id,
                'data_evolucao': evolucao.data_evolucao,
                'texto': evolucao.texto,
                'enfermeiro_nome': evolucao.enfermeiro_nome,
                'enfermeiro_coren': evolucao.enfermeiro_coren
            })
        
        return render_template('impressao_evolucoes_enfermagem.html',
                             paciente=paciente,
                             internacao=internacao,
                             atendimento=atendimento,
                             evolucoes=evolucoes_dados,
                             data_selecionada=data_selecionada,
                             data_geracao=datetime.now(ZoneInfo("America/Sao_Paulo")),
                             atendimento_id=atendimento_id)
        
    except Exception as e:
        logging.error(f"Erro ao imprimir evoluções de enfermagem: {str(e)}")
        logging.error(traceback.format_exc())
        flash('Erro ao gerar impressão das evoluções de enfermagem. Por favor, tente novamente.', 'danger')
        return redirect(url_for('main.impressoes_enfermagem', atendimento_id=atendimento_id))

# API para listar datas disponíveis das prescrições de enfermagem
@bp.route('/api/enfermagem/prescricoes/datas/<string:atendimento_id>', methods=['GET'])
@login_required
def listar_datas_prescricoes_enfermagem(atendimento_id):
    """
    Lista as datas disponíveis das prescrições de enfermagem para um atendimento específico.
    """
    try:
        # Verificar se o usuário é enfermeiro ou médico
        current_user = get_current_user()
        if current_user.cargo.lower() not in ['enfermeiro', 'medico']:
            return jsonify({
                'success': False,
                'message': 'Acesso permitido apenas para enfermeiros e médicos'
            }), 403

        # Buscar a internação
        internacao = Internacao.query.filter_by(atendimento_id=atendimento_id).first()
        if not internacao:
            return jsonify({
                'success': False,
                'message': 'Internação não encontrada'
            }), 404

        # Buscar as datas das prescrições
        prescricoes = db.session.query(
            PrescricaoEnfermagem.data_prescricao
        ).filter(PrescricaoEnfermagem.atendimentos_clinica_id == internacao.id)\
            .filter(PrescricaoEnfermagem.data_prescricao.isnot(None))\
            .all()

        # Agrupar por data
        datas_disponiveis = {}
        for prescricao in prescricoes:
            if prescricao.data_prescricao:
                data_str = prescricao.data_prescricao.strftime('%Y-%m-%d')
                data_display = prescricao.data_prescricao.strftime('%d/%m/%Y')
                
                if data_str not in datas_disponiveis:
                    datas_disponiveis[data_str] = {
                        'data': data_str,
                        'data_display': data_display,
                        'quantidade': 0
                    }
                datas_disponiveis[data_str]['quantidade'] += 1

        # Converter para lista ordenada
        datas_lista = list(datas_disponiveis.values())
        datas_lista.sort(key=lambda x: x['data'], reverse=True)

        return jsonify({
            'success': True,
            'datas': datas_lista
        }), 200

    except Exception as e:
        logging.error(f'Erro ao listar datas das prescrições de enfermagem: {str(e)}')
        logging.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': f'Erro interno do servidor: {str(e)}'
        }), 500

# Rota para imprimir prescrições de enfermagem de uma data específica
@bp.route('/imprimir/prescricoes_enfermagem/<string:atendimento_id>')
@login_required
def imprimir_prescricoes_enfermagem(atendimento_id):
    """
    Imprime as prescrições de enfermagem de uma data específica
    """
    try:
        # Pegar a data da query string
        data_selecionada_str = request.args.get('data')
        if not data_selecionada_str:
            flash('Data não especificada para impressão.', 'warning')
            return redirect(url_for('main.impressoes_enfermagem', atendimento_id=atendimento_id))

        # Converter a data
        try:
            data_selecionada = datetime.strptime(data_selecionada_str, '%Y-%m-%d').date()
        except ValueError:
            flash('Formato de data inválido.', 'danger')
            return redirect(url_for('main.impressoes_enfermagem', atendimento_id=atendimento_id))

        # Buscar o atendimento
        atendimento = Atendimento.query.get(atendimento_id)
        if not atendimento:
            flash('Atendimento não encontrado.', 'danger')
            return redirect(url_for('main.pacientes_internados'))

        # Buscar a internação
        internacao = Internacao.query.filter_by(atendimento_id=atendimento_id).first()
        if not internacao:
            flash('Internação não encontrada.', 'danger')
            return redirect(url_for('main.pacientes_internados'))
        
        # Buscar o paciente
        paciente = Paciente.query.get(internacao.paciente_id)
        if not paciente:
            flash('Paciente não encontrado.', 'danger')
            return redirect(url_for('main.pacientes_internados'))
        
        # Buscar as prescrições da data selecionada
        inicio_dia = datetime.combine(data_selecionada, datetime.min.time())
        fim_dia = datetime.combine(data_selecionada, datetime.max.time())
        
        prescricoes = db.session.query(
            PrescricaoEnfermagem.id,
            PrescricaoEnfermagem.data_prescricao,
            PrescricaoEnfermagem.texto,
            Funcionario.nome.label("enfermeiro_nome"),
            Funcionario.numero_profissional.label("enfermeiro_coren")
        ).select_from(PrescricaoEnfermagem)\
            .join(Funcionario, PrescricaoEnfermagem.funcionario_id == Funcionario.id)\
            .filter(PrescricaoEnfermagem.atendimentos_clinica_id == internacao.id)\
            .filter(PrescricaoEnfermagem.data_prescricao >= inicio_dia)\
            .filter(PrescricaoEnfermagem.data_prescricao <= fim_dia)\
            .order_by(PrescricaoEnfermagem.data_prescricao.asc())\
            .all()
        
        if not prescricoes:
            flash(f'Nenhuma prescrição de enfermagem encontrada para o dia {data_selecionada.strftime("%d/%m/%Y")}.', 'warning')
            return redirect(url_for('main.impressoes_enfermagem', atendimento_id=atendimento_id))
        
        # Preparar dados das prescrições
        prescricoes_dados = []
        for prescricao in prescricoes:
            prescricoes_dados.append({
                'id': prescricao.id,
                'data_prescricao': prescricao.data_prescricao,
                'texto': prescricao.texto,
                'enfermeiro_nome': prescricao.enfermeiro_nome,
                'enfermeiro_coren': prescricao.enfermeiro_coren
            })
        
        return render_template('impressao_prescricoes_enfermagem.html',
                             paciente=paciente,
                             internacao=internacao,
                             atendimento=atendimento,
                             prescricoes=prescricoes_dados,
                             data_selecionada=data_selecionada,
                             data_geracao=datetime.now(ZoneInfo("America/Sao_Paulo")),
                             atendimento_id=atendimento_id)
        
    except Exception as e:
        logging.error(f"Erro ao imprimir prescrições de enfermagem: {str(e)}")
        logging.error(traceback.format_exc())
        flash('Erro ao gerar impressão das prescrições de enfermagem. Por favor, tente novamente.', 'danger')
        return redirect(url_for('main.impressoes_enfermagem', atendimento_id=atendimento_id))

@bp.route('/api/paciente/<int:paciente_id>/historico-internamentos', methods=['GET'])
@login_required
def obter_historico_internamentos_paciente(paciente_id):
    """
    Busca o histórico de internações de um paciente específico.
    Exclui a internação atual se fornecida via query parameter.
    """
    try:
        current_user = get_current_user()
        if current_user.cargo.lower() not in ['medico', 'enfermeiro', 'multi']:
            return jsonify({
                'success': False,
                'message': 'Acesso permitido apenas para médicos, enfermeiros e profissionais multi'
            }), 403
        
        # Verificar se o paciente existe
        paciente = Paciente.query.get(paciente_id)
        if not paciente:
            return jsonify({
                'success': False,
                'message': 'Paciente não encontrado'
            }), 404
        
        # Obter atendimento atual para excluir da listagem
        atendimento_atual = request.args.get('atendimento_atual')
        
        # Buscar todas as internações do paciente
        query = Internacao.query.filter_by(paciente_id=paciente_id)
        
        # Excluir internação atual se fornecido
        if atendimento_atual:
            query = query.filter(Internacao.atendimento_id != atendimento_atual)
        
        # Ordenar por data de internação (mais recente primeiro)
        internacoes = query.order_by(Internacao.data_internacao.desc()).all()
        
        # Formatar resposta
        internamentos = []
        for internacao in internacoes:
            # Buscar dados do médico responsável
            medico = None
            if internacao.medico_id:
                medico = Funcionario.query.get(internacao.medico_id)
            
            internamentos.append({
                'id': internacao.id,
                'atendimento_id': internacao.atendimento_id,
                'data_internacao': internacao.data_internacao.isoformat() if internacao.data_internacao else None,
                'data_alta': internacao.data_alta.isoformat() if internacao.data_alta else None,
                'diagnostico': internacao.diagnostico,
                'diagnostico_inicial': internacao.diagnostico_inicial,
                'leito': internacao.leito,
                'medico_nome': medico.nome if medico else None,
                'cid_principal': internacao.cid_principal,
                'carater_internacao': internacao.carater_internacao
            })
        
        return jsonify({
            'success': True,
            'internamentos': internamentos
        })
        
    except Exception as e:
        logging.error(f"Erro ao obter histórico de internamentos do paciente {paciente_id}: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'error': 'Erro interno do servidor'
        }), 500

@bp.route('/api/prescricoes/<int:prescricao_id>/base', methods=['GET'])
@login_required
def buscar_prescricao_base(prescricao_id):
    """
    Buscar dados de uma prescrição específica para usar como base para nova prescrição
    """
    try:
        prescricao = PrescricaoClinica.query.get(prescricao_id)
        
        if not prescricao:
            return jsonify({
                'success': False,
                'error': 'Prescrição não encontrada'
            }), 404
        
        # Verificar se o usuário tem acesso à prescrição
        if prescricao.internacao.paciente_id != prescricao.internacao.paciente_id:
            return jsonify({
                'success': False,
                'error': 'Acesso negado'
            }), 403
        
        # Preparar dados da prescrição para ser base
        prescricao_base = {
            'id': prescricao.id,
            'texto_dieta': prescricao.texto_dieta,
            'texto_procedimento_medico': prescricao.texto_procedimento_medico,
            'texto_procedimento_multi': prescricao.texto_procedimento_multi,
            'medicamentos': prescricao.medicamentos_json if prescricao.medicamentos_json else []
        }
        
        return jsonify({
            'success': True,
            'prescricao_base': prescricao_base
        })

    except Exception as e:
        current_app.logger.error(f"Erro ao buscar prescrição base: {str(e)}")
        return jsonify({
            'success': False,
            'error': 'Erro interno do servidor'
        }), 500

# API para fechar prontuário (define dieta = '1')
@bp.route('/api/internacao/<string:internacao_id>/fechar-prontuario', methods=['POST'])
@login_required
def fechar_prontuario(internacao_id):
    """
    Fecha o prontuário de um paciente que já teve alta, definindo dieta = '1'
    e atualizando o status final do atendimento (por padrão 'Alta').
    Isso remove o paciente da listagem de pacientes internados.
    """
    try:
        current_user = get_current_user()
        if current_user.cargo.lower() not in ['medico', 'enfermeiro', 'multi']:
            return jsonify({
                'success': False,
                'message': 'Acesso não autorizado. Apenas médicos e enfermeiros podem fechar prontuários.'
            }), 403
        
        # Buscar a internação
        internacao = Internacao.query.filter_by(atendimento_id=internacao_id).first()
        
        if not internacao:
            return jsonify({
                'success': False,
                'message': 'Internação não encontrada.'
            }), 404
        
        # Buscar o atendimento relacionado
        atendimento = Atendimento.query.get(internacao_id)
        
        if not atendimento:
            return jsonify({
                'success': False,
                'message': 'Atendimento não encontrado.'
            }), 404
        
        # Verificar se o paciente já teve alta (necessário para exibição do botão no front)
        if not internacao.data_alta:
            return jsonify({
                'success': False,
                'message': 'Só é possível fechar prontuário de pacientes que já tiveram alta.'
            }), 400
        
        # Extrair o status final pendente armazenado no campo dieta
        status_final = 'Alta'  # Padrão
        if internacao.dieta and internacao.dieta.startswith('PENDENTE:'):
            # Extrair o status após 'PENDENTE:'
            status_final = internacao.dieta.split(':', 1)[1].strip()
            logging.info(f"Status final extraído do campo dieta: {status_final}")
        else:
            # Se não houver status pendente, verificar se foi enviado no payload
            dados = request.get_json(silent=True) or {}
            status_raw = (dados.get('status_final') or dados.get('status') or '').strip().lower()
            if status_raw:
                status_map = {
                    'alta': 'Alta',
                    'obito': 'Óbito',
                    'óbito': 'Óbito',
                    'transferencia': 'Transferido',
                    'transferência': 'Transferido',
                    'transferido': 'Transferido',
                    'evasao': 'Evasão',
                    'evasão': 'Evasão',
                    'alta a pedido': 'Alta a pedido',
                    'a pedido': 'Alta a pedido',
                }
                status_final = status_map.get(status_raw, 'Alta')
        
        # AGORA SIM: Aplicar o status final ao atendimento
        atendimento.status = status_final
        logging.info(f"Status do atendimento alterado para: {status_final}")
        
        # Fechar o prontuário definindo dieta = '1'
        internacao.dieta = '1'
        
        # Commit das alterações
        db.session.commit()
        
        # Log da ação
        logging.info(
            f"Prontuário fechado pelo {current_user.cargo} {current_user.nome} (ID: {current_user.id}) "
            f"para internação {internacao_id} (dieta='1')"
        )
        
        return jsonify({
            'success': True,
            'message': 'Prontuário fechado com sucesso. Paciente removido da lista de internados.'
        })
        
    except Exception as e:
        db.session.rollback()
        logging.error(f"Erro ao fechar prontuário: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': 'Erro interno do servidor ao fechar prontuário.'
        }), 500


# ======= NOVOS ENDPOINTS PARA DIAGNÓSTICO E ANAMNESE =======

@bp.route('/api/internacao/<string:atendimento_id>/diagnostico', methods=['PUT'])
@login_required
def atualizar_diagnostico_internacao(atendimento_id):
    """
    Atualiza dados de diagnóstico e classificação de uma internação
    """
    try:
        # Verificar se o usuário é médico
        current_user = get_current_user()
        if current_user.cargo.lower() != 'medico':
            return jsonify({
                'success': False,
                'message': 'Apenas médicos podem atualizar diagnósticos'
            }), 403

        # Buscar a internação pelo atendimento_id
        internacao = Internacao.query.filter_by(atendimento_id=atendimento_id).first()
        if not internacao:
            return jsonify({
                'success': False,
                'message': 'Internação não encontrada'
            }), 404

        # Obter dados do JSON
        dados = request.get_json()
        if not dados:
            return jsonify({
                'success': False,
                'message': 'Dados não fornecidos'
            }), 400

        # Campos permitidos para atualização de diagnóstico
        campos_diagnostico = [
            'diagnostico_inicial',
            'cid_principal', 
            'diagnostico',
            'cid_10_secundario',
            'cid_10_causas_associadas',
            'justificativa_internacao_sinais_e_sintomas',
            'justificativa_internacao_condicoes',
            'justificativa_internacao_principais_resultados_diagnostico'
        ]

        # Atualizar campos de diagnóstico
        for campo in campos_diagnostico:
            if campo in dados:
                setattr(internacao, campo, dados[campo])

        # Salvar alterações
        db.session.commit()

        # Log da ação
        logging.info(f"Diagnóstico atualizado pelo médico {current_user.nome} (ID: {current_user.id}) para internação {atendimento_id}")

        return jsonify({
            'success': True,
            'message': 'Diagnóstico atualizado com sucesso'
        })

    except Exception as e:
        db.session.rollback()
        logging.error(f"Erro ao atualizar diagnóstico da internação {atendimento_id}: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': 'Erro interno do servidor',
            'error': str(e)
        }), 500


@bp.route('/api/internacao/<string:atendimento_id>/anamnese-conduta', methods=['PUT'])
@login_required
def atualizar_anamnese_conduta_internacao(atendimento_id):
    """
    Atualiza dados de anamnese e conduta médica de uma internação
    """
    try:
        # Verificar se o usuário é médico
        current_user = get_current_user()
        if current_user.cargo.lower() != 'medico':
            return jsonify({
                'success': False,
                'message': 'Apenas médicos podem atualizar anamnese e conduta'
            }), 403

        # Buscar a internação pelo atendimento_id
        internacao = Internacao.query.filter_by(atendimento_id=atendimento_id).first()
        if not internacao:
            return jsonify({
                'success': False,
                'message': 'Internação não encontrada'
            }), 404

        # Obter dados do JSON
        dados = request.get_json()
        if not dados:
            return jsonify({
                'success': False,
                'message': 'Dados não fornecidos'
            }), 400

        # Atualizar campos específicos com mapeamento correto
        if 'anamnese_exame_fisico' in dados:
            internacao.folha_anamnese = dados['anamnese_exame_fisico']
            
        if 'hda' in dados:
            internacao.hda = dados['hda']
            
        if 'conduta' in dados:
            internacao.conduta = dados['conduta']
            
        if 'exames_laboratoriais' in dados:
            internacao.exames_laboratoriais = dados['exames_laboratoriais']
            
        if 'descr_procedimento_solicitado' in dados:
            internacao.descr_procedimento_solicitado = dados['descr_procedimento_solicitado']
            
        if 'codigo_procedimento' in dados:
            internacao.codigo_procedimento = dados['codigo_procedimento']

        # Salvar alterações
        db.session.commit()

        # Log da ação
        logging.info(f"Anamnese e conduta atualizadas pelo médico {current_user.nome} (ID: {current_user.id}) para internação {atendimento_id}")

        return jsonify({
            'success': True,
            'message': 'Anamnese e conduta atualizadas com sucesso'
        })

    except Exception as e:
        db.session.rollback()
        logging.error(f"Erro ao atualizar anamnese e conduta da internação {atendimento_id}: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': 'Erro interno do servidor',
            'error': str(e)
        }), 500

@bp.route('/api/internacao/<string:atendimento_id>/informacoes-alta', methods=['GET'])
@login_required
def buscar_informacoes_alta_para_impressao(atendimento_id):
    """
    API específica para buscar informações de alta formatadas para impressão.
    Usada pela página de pacientes internados para imprimir informações de alta.
    """
    try:
        # Verificar se o usuário é médico ou enfermeiro
        current_user = get_current_user()
        if current_user.cargo.lower() not in ['medico', 'enfermeiro', 'multi']:
            return jsonify({
                'success': False,
                'message': 'Acesso permitido apenas para médicos e enfermeiros'
            }), 403
        
        # Buscar a internação
        internacao = Internacao.query.filter_by(atendimento_id=atendimento_id).first()
        if not internacao:
            return jsonify({
                'success': False,
                'message': 'Internação não encontrada'
            }), 404
        
        # Buscar dados do paciente
        paciente = Paciente.query.get(internacao.paciente_id)
        
        # Se ainda não teve alta, montar histórico automaticamente
        historico_internacao = internacao.historico_internacao
        if not internacao.data_alta:
            # Montar histórico automaticamente usando: HDA + Anamnese + Conduta + Exames
            partes_historico = [
                internacao.hda or '',
                internacao.folha_anamnese or '',
                internacao.conduta or '',
                internacao.exames_laboratoriais or ''
            ]
            
            # Filtrar partes vazias e juntar com uma linha simples de separação
            historico_automatico = '\n'.join([parte.strip() for parte in partes_historico if parte.strip()])
            historico_internacao = historico_automatico or 'Histórico será montado automaticamente'
        if not paciente:
            return jsonify({
                'success': False,
                'message': 'Paciente não encontrado'
            }), 404
        
        # Formatar resposta com as informações de alta no formato esperado pelo JavaScript
        resultado = {
            'success': True,
            'informacoes_alta': {
                'data_alta': internacao.data_alta.isoformat() if internacao.data_alta else None,
            'diagnostico': internacao.diagnostico or 'Não informado',
            'historico_internacao': historico_internacao or 'Não informado', 
                'conduta': internacao.conduta or 'Não informado',
                'medicacao': getattr(internacao, 'medicacao_alta', None) or internacao.medicacao or 'Não informado',
                'cuidados_gerais': internacao.cuidados_gerais or 'Não informado'
            },
            'paciente': {
                'nome': paciente.nome or 'Não informado',
                'cpf': paciente.cpf or 'Não informado'
            }
        }
        
        return jsonify(resultado)
        
    except Exception as e:
        logging.error(f'Erro ao buscar informações de alta para impressão: {str(e)}')
        logging.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': f'Erro ao buscar informações de alta: {str(e)}'
        }), 500

@bp.route('/clinica/receituario/<int:receita_id>/imprimir_html')
def imprimir_receita(receita_id):
    # Localiza o receituário
    receita = ReceituarioClinica.query.get_or_404(receita_id)

    # Se for receita normal, usar template de impressão comum
    try:
        tipo = (receita.tipo_receita or '').lower().strip()
    except Exception:
        tipo = ''
    if tipo == 'normal':
        from datetime import datetime
        from markupsafe import Markup
        paciente = receita.atendimento.paciente
        data_hoje = datetime.now().strftime('%d/%m/%Y')
        conteudo_receita = receita.conteudo_receita or ''
        return render_template(
            'impressao_receita_comum.html',
            paciente_nome=paciente.nome,
            paciente_endereco=paciente.endereco,
            paciente_filiacao=paciente.filiacao,
            paciente_cartao_sus=paciente.cartao_sus,
            paciente_data_nascimento=paciente.data_nascimento.strftime('%d/%m/%Y') if paciente.data_nascimento else None,
            receituario_receita=Markup(conteudo_receita),
            data_hoje=data_hoje
        )

    # Dados do paciente
    paciente = receita.atendimento.paciente

    # Dados do médico
    medico = receita.medico

    # Data atual para impressão
    from datetime import datetime
    data_hoje = datetime.now().strftime('%d/%m/%Y')

    return render_template(
        'impressao_receita.html',
        paciente_nome=paciente.nome,
        paciente_endereco=paciente.endereco,
        paciente_telefone=paciente.telefone,
        funcionario_nome=medico.nome,
        funcionario_nprofissional=medico.numero_profissional,
        receituario_receita=receita.conteudo_receita,
        data_hoje=data_hoje
    )

@bp.route('/clinica/receituario/<int:receita_id>/imprimir_html_comum')
def imprimir_receita_comum(receita_id):
    import re
    from markupsafe import Markup
    
    # Localiza o receituário
    receita = ReceituarioClinica.query.get_or_404(receita_id)

    # Dados do paciente
    paciente = receita.atendimento.paciente

    # Data atual para impressão
    from datetime import datetime
    data_hoje = datetime.now().strftime('%d/%m/%Y')

    # Processar conteúdo da receita para melhor formatação preservando todo o conteúdo
    conteudo_receita = receita.conteudo_receita or ''
    
    # Se não há conteúdo, retorna vazio
    if not conteudo_receita.strip():
        conteudo_receita = ''
    else:
        # Limpar espaços extras entre tags, mas preservar todo o conteúdo
        conteudo_receita = re.sub(r'>\s+<', '><', conteudo_receita)
        
        # Garantir que quebras de linha sejam preservadas em elementos <p> vazios
        conteudo_receita = re.sub(r'<p>\s*</p>', '<p>&nbsp;</p>', conteudo_receita)
        
        # Converter quebras de linha simples em <br> apenas se não estiverem dentro de tags
        if not re.search(r'<[^>]+>', conteudo_receita):
            # Se o conteúdo não tem tags HTML, converter quebras de linha em <br>
            conteudo_receita = conteudo_receita.replace('\n', '<br>')
        
        # Normalizar espaços múltiplos consecutivos, mas preservar formatação
        conteudo_receita = re.sub(r'\s{3,}', '  ', conteudo_receita)

    # Usar template padronizado de receita comum
    return render_template(
        'impressao_receita_comum.html',
        paciente_nome=paciente.nome,
        paciente_endereco=paciente.endereco,
        paciente_filiacao=paciente.filiacao,
        paciente_cartao_sus=paciente.cartao_sus,
        paciente_data_nascimento=paciente.data_nascimento.strftime('%d/%m/%Y') if paciente.data_nascimento else None,
        receituario_receita=Markup(conteudo_receita),
        data_hoje=data_hoje
    )

@bp.route('/clinica/atestado/<int:atestado_id>/imprimir')
@login_required
def impressao_atestado(atestado_id):
    """
    Gera página de impressão para atestado médico
    """
    try:
        # Buscar o atestado com informações relacionadas
        atestado = AtestadoClinica.query.options(
            joinedload(AtestadoClinica.atendimento).joinedload(Atendimento.paciente),
            joinedload(AtestadoClinica.medico)
        ).filter_by(id=atestado_id).first()
        
        if not atestado:
            flash('Atestado não encontrado.', 'error')
            return redirect(url_for('main.dashboard'))
        
        # Buscar dados do paciente
        paciente = atestado.atendimento.paciente
        medico = atestado.medico
        
        # Calcular idade do paciente
        idade = None
        if paciente.data_nascimento:
            hoje = date.today()
            idade = hoje.year - paciente.data_nascimento.year
            if hoje.month < paciente.data_nascimento.month or \
               (hoje.month == paciente.data_nascimento.month and hoje.day < paciente.data_nascimento.day):
                idade -= 1
        
        # Formatar data do atestado em America/Sao_Paulo (UTC-3)
        data_atestado_formatada = formatar_data_br(atestado.data_atestado)
        
        # Dados para o template
        contexto = {
            'atestado': atestado,
            'paciente': paciente,
            'medico': medico,
            'idade': idade,
            'data_atestado': data_atestado_formatada,
            'conteudo_atestado': atestado.conteudo_atestado,
            'dias_afastamento': atestado.dias_afastamento or 'Não especificado'
        }
        
        return render_template('impressao_atestado.html', **contexto)
        
    except Exception as e:
        current_app.logger.error(f'Erro ao gerar impressão do atestado {atestado_id}: {str(e)}')
        flash('Erro interno do servidor ao gerar impressão do atestado.', 'error')
        return redirect(url_for('main.dashboard'))

@bp.route('/clinica/ficha-referencia/<int:ficha_id>/imprimir')
@login_required
def impressao_ficha_referencia(ficha_id):
    """
    Gera página de impressão para ficha de referência
    """
    try:
        # Buscar a ficha de referência com informações relacionadas
        ficha = FichaReferencia.query.options(
            joinedload(FichaReferencia.atendimento).joinedload(Atendimento.paciente),
            joinedload(FichaReferencia.medico)
        ).filter_by(id=ficha_id).first()
        
        if not ficha:
            flash('Ficha de referência não encontrada.', 'error')
            return redirect(url_for('main.dashboard'))
        
        # Buscar dados do paciente e médico
        paciente = ficha.atendimento.paciente
        medico = ficha.medico
        
        # Calcular idade do paciente
        idade = None
        if paciente.data_nascimento:
            hoje = date.today()
            idade = hoje.year - paciente.data_nascimento.year
            if hoje.month < paciente.data_nascimento.month or \
               (hoje.month == paciente.data_nascimento.month and hoje.day < paciente.data_nascimento.day):
                idade -= 1
        
        # Formatar data da ficha
        data_ficha_formatada = f"{ficha.data.strftime('%d/%m/%Y')} às {ficha.hora.strftime('%H:%M')}"
        
        # Dados para o template
        contexto = {
            'ficha': ficha,
            'paciente': paciente,
            'medico': medico,
            'idade': idade,
            'data_ficha': data_ficha_formatada,
            'atendimento_id': ficha.atendimento_id,
            'paciente_nome': paciente.nome,
            'paciente_sexo': paciente.sexo,
            'paciente_data': paciente.data_nascimento.strftime('%d/%m/%Y') if paciente.data_nascimento else '',
            'paciente_endereco': paciente.endereco or '',
            'texto_referencia': ficha.texto_referencia or '',
            'tipo_atendimento': ficha.encaminhamento_atendimento or '',
            'referencia_unidade': ficha.unidade_referencia or '',
            'referencia_atendimento': ficha.procedimento or '',
            'referencia_data': data_ficha_formatada
        }
        
        return render_template('impressao_referencia.html', **contexto)
        
    except Exception as e:
        current_app.logger.error(f'Erro ao gerar impressão da ficha de referência {ficha_id}: {str(e)}')
        flash('Erro interno do servidor ao gerar impressão da ficha de referência.', 'error')
        return redirect(url_for('main.dashboard'))

@bp.route('/api/fichas-referencia', methods=['POST'])
@login_required
def criar_ficha_referencia():
    dados = request.get_json()
    
    try:
        # Obter dados do usuário atual
        current_user = get_current_user()
        if current_user.cargo.lower() != 'medico':
            return jsonify({'success': False, 'error': 'Acesso negado. Apenas médicos podem criar fichas de referência.'}), 403
        
        # Buscar internação pelo atendimento_id se fornecido
        internacao = None
        if 'atendimento_id' in dados:
            internacao = Internacao.query.filter_by(atendimento_id=dados['atendimento_id']).first()
        
        # Obter data e hora atuais
        from datetime import datetime, timezone, timedelta
        agora = datetime.now(ZoneInfo("America/Sao_Paulo"))  # Timezone do Brasil
        
        nova_ficha = FichaReferencia(
            atendimento_id=dados['atendimento_id'],
            medico_id=current_user.id,  # Usar o ID do médico logado
            internacao_id=internacao.id if internacao else None,  # Preencher automaticamente se encontrado
            texto_referencia=dados.get('texto_referencia'),
            encaminhamento_atendimento=dados.get('encaminhamento_atendimento'),
            procedimento=dados.get('procedimento'),
            unidade_referencia=dados.get('unidade_referencia'),
            data=agora.date(),  # Data atual automaticamente
            hora=agora.time()   # Hora atual automaticamente
        )
        db.session.add(nova_ficha)
        db.session.commit()

        return jsonify({'success': True, 'id': nova_ficha.id}), 201

    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 400

@bp.route('/api/fichas-referencia', methods=['GET'])
def listar_fichas_referencia():
    fichas = FichaReferencia.query.all()
    resultado = []
    for ficha in fichas:
        resultado.append({
            'id': ficha.id,
            'atendimento_id': ficha.atendimento_id,
            'medico_id': ficha.medico_id,
            'internacao_id': ficha.internacao_id,
            'texto_referencia': ficha.texto_referencia,
            'encaminhamento_atendimento': ficha.encaminhamento_atendimento,
            'procedimento': ficha.procedimento,
            'unidade_referencia': ficha.unidade_referencia,
            'data': ficha.data.strftime('%Y-%m-%d'),
            'hora': ficha.hora.strftime('%H:%M')
        })
    return jsonify(resultado), 200

@bp.route('/api/fichas-referencia/lista/<string:atendimento_id>', methods=['GET'])
@login_required
def api_listar_fichas_por_atendimento(atendimento_id):
    """
    Lista as fichas de referência de um atendimento específico (endpoint específico)
    """
    try:
        fichas = FichaReferencia.query.filter_by(atendimento_id=atendimento_id).order_by(FichaReferencia.data.desc(), FichaReferencia.hora.desc()).all()
        
        resultado = []
        for ficha in fichas:
            # Buscar dados do médico
            medico = Funcionario.query.get(ficha.medico_id)
            
            resultado.append({
                'id': ficha.id,
                'atendimento_id': ficha.atendimento_id,
                'medico_id': ficha.medico_id,
                'medico_nome': medico.nome if medico else 'Médico não encontrado',
                'internacao_id': ficha.internacao_id,
                'texto_referencia': ficha.texto_referencia,
                'encaminhamento_atendimento': ficha.encaminhamento_atendimento,
                'procedimento': ficha.procedimento,
                'unidade_referencia': ficha.unidade_referencia,
                'data': ficha.data.strftime('%Y-%m-%d'),
                'hora': ficha.hora.strftime('%H:%M'),
                'data_criacao': f"{ficha.data.strftime('%d/%m/%Y')} às {ficha.hora.strftime('%H:%M')}"
            })
        
        return jsonify({'success': True, 'fichas': resultado}), 200
        
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@bp.route('/api/fichas-referencia/<int:id>', methods=['GET'])
@login_required
def buscar_ficha_por_id(id):
    """
    Busca uma ficha específica por ID
    """
    try:
        ficha = FichaReferencia.query.get(id)
        if not ficha:
            return jsonify({'success': False, 'message': 'Ficha não encontrada'}), 404

        # Buscar dados do médico
        medico = Funcionario.query.get(ficha.medico_id)

        return jsonify({
            'id': ficha.id,
            'atendimento_id': ficha.atendimento_id,
            'medico_id': ficha.medico_id,
            'medico_nome': medico.nome if medico else 'Médico não encontrado',
            'internacao_id': ficha.internacao_id,
            'texto_referencia': ficha.texto_referencia,
            'encaminhamento_atendimento': ficha.encaminhamento_atendimento,
            'procedimento': ficha.procedimento,
            'unidade_referencia': ficha.unidade_referencia,
            'data': ficha.data.strftime('%Y-%m-%d'),
            'hora': ficha.hora.strftime('%H:%M'),
            'data_criacao': f"{ficha.data.strftime('%d/%m/%Y')} às {ficha.hora.strftime('%H:%M')}"
        }), 200
        
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@bp.route('/api/fichas-referencia/ativas', methods=['GET'])
@login_required
def listar_fichas_referencia_ativas():
    """
    Lista fichas de referência cujos atendimentos associados estão ativos
    (status "Internado" ou "Em Observação").
    """
    try:
        # Buscar fichas com seus atendimentos
        query = (
            db.session.query(FichaReferencia, Atendimento)
            .join(Atendimento, FichaReferencia.atendimento_id == Atendimento.id)
            .order_by(FichaReferencia.data.desc(), FichaReferencia.hora.desc())
        )

        fichas_data = []
        for ficha, atendimento in query.all():
            status_norm = normalize_status(atendimento.status or '')
            if status_norm in ['INTERNADO', 'EM OBSERVACAO']:
                medico = Funcionario.query.get(ficha.medico_id)
                paciente = Paciente.query.get(atendimento.paciente_id)

                fichas_data.append({
                    'id': ficha.id,
                    'atendimento_id': ficha.atendimento_id,
                    'paciente_id': atendimento.paciente_id,
                    'paciente_nome': (paciente.nome if paciente else None),
                    'medico_id': ficha.medico_id,
                    'medico_nome': (medico.nome if medico else None),
                    'texto_referencia': ficha.texto_referencia,
                    'encaminhamento_atendimento': ficha.encaminhamento_atendimento,
                    'procedimento': ficha.procedimento,
                    'unidade_referencia': ficha.unidade_referencia,
                    'data': ficha.data.strftime('%Y-%m-%d'),
                    'hora': ficha.hora.strftime('%H:%M'),
                    'status_atendimento': atendimento.status,
                })

        return jsonify({'success': True, 'total': len(fichas_data), 'fichas': fichas_data}), 200

    except Exception as e:
        current_app.logger.error(f'Erro ao listar fichas de referência ativas: {str(e)}')
        return jsonify({'success': False, 'message': 'Erro interno do servidor'}), 500

@bp.route('/api/fichas-referencia/<int:id>', methods=['PUT'])
def atualizar_ficha_referencia(id):
    dados = request.get_json()
    ficha = FichaReferencia.query.get(id)

    if not ficha:
        return jsonify({'success': False, 'message': 'Ficha não encontrada'}), 404

    try:
        ficha.texto_referencia = dados.get('texto_referencia', ficha.texto_referencia)
        ficha.encaminhamento_atendimento = dados.get('encaminhamento_atendimento', ficha.encaminhamento_atendimento)
        ficha.procedimento = dados.get('procedimento', ficha.procedimento)
        ficha.unidade_referencia = dados.get('unidade_referencia', ficha.unidade_referencia)
        
        if 'data' in dados:
            ficha.data = datetime.strptime(dados['data'], "%Y-%m-%d").date()
        if 'hora' in dados:
            ficha.hora = datetime.strptime(dados['hora'], "%H:%M").time()

        db.session.commit()
        return jsonify({'success': True}), 200

    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 400
    
@bp.route('/api/listar_prescricoes_enf', methods=['GET'])
def listar_prescricoes_padrao():
    """
    Lista as prescrições de enfermagem padrão com funcionalidades de busca
    """
    try:
        # Parâmetros de filtro
        id_param = request.args.get('id')
        termo = request.args.get('termo', '').strip()
        categoria = request.args.get('categoria', '').strip()
        titulo = request.args.get('titulo', '').strip()
        nic = request.args.get('nic', '').strip()
        apenas_padrao = request.args.get('apenas_padrao', 'true').lower() == 'true'
        
        # Se solicitar uma prescrição específica
        if id_param:
            prescricao = PrescricaoEnfermagemTemplate.query.get(id_param)
            if not prescricao:
                return jsonify({'success': False, 'message': 'Prescrição não encontrada'}), 404
            return jsonify({
                'success': True,
                'template': {
                    'id': prescricao.id,
                    'padrao': prescricao.padrao,
                    'titulo': prescricao.titulo,
                    'texto_prescricao': prescricao.texto_prescricao,
                    'nic': prescricao.nic,
                    'nic_tipo': prescricao.nic_tipo
                }
            })
        
        # Query base
        query = PrescricaoEnfermagemTemplate.query
        
        # Filtrar apenas prescrições padrão se solicitado
        if apenas_padrao:
            query = query.filter_by(padrao=True)
        
        # Filtrar por termo de busca (título, texto ou tipo NIC)
        if termo:
            query = query.filter(
                db.or_(
                    PrescricaoEnfermagemTemplate.titulo.ilike(f'%{termo}%'),
                    PrescricaoEnfermagemTemplate.texto_prescricao.ilike(f'%{termo}%'),
                    PrescricaoEnfermagemTemplate.nic_tipo.ilike(f'%{termo}%')
                )
            )
        
        # Filtrar por categoria NIC
        if categoria:
            query = query.filter(PrescricaoEnfermagemTemplate.nic_tipo.ilike(f'%{categoria}%'))
        
        # Filtrar por título específico
        if titulo:
            query = query.filter(PrescricaoEnfermagemTemplate.titulo.ilike(f'%{titulo}%'))
        
        # Filtrar por código NIC específico
        if nic:
            query = query.filter(PrescricaoEnfermagemTemplate.nic == nic)
        
        # Ordenar por título
        prescricoes = query.order_by(PrescricaoEnfermagemTemplate.titulo).all()
        
        # Serializar os dados
        prescricoes_data = []
        for prescricao in prescricoes:
            prescricoes_data.append({
                'id': prescricao.id,
                'padrao': prescricao.padrao,
                'titulo': prescricao.titulo,
                'texto_prescricao': prescricao.texto_prescricao,
                'nic': prescricao.nic,
                'nic_tipo': prescricao.nic_tipo
            })
        
        return jsonify({
            'success': True,
            'prescricoes': prescricoes_data,
            'total': len(prescricoes_data)
        })
        
    except Exception as e:
        print(f"Erro ao listar prescrições de enfermagem: {str(e)}")
        return jsonify({
            'success': False,
            'error': 'Erro interno do servidor',
            'message': str(e)
        }), 500

@bp.route('/evolucoes_fisioterapia/atendimento/<string:id_atendimento>', methods=['GET'])
def listar_evolucoes_por_atendimento(id_atendimento):
    from sqlalchemy.orm import joinedload
    
    evolucoes = EvolucaoFisioterapia.query.filter_by(id_atendimento=id_atendimento).all()
    resultado = []
    
    for e in evolucoes:
        # Buscar dados do funcionário
        funcionario = None
        if e.funcionario_id:
            funcionario = Funcionario.query.get(e.funcionario_id)
        
        resultado.append({
            'id': e.id,
            'id_atendimento': e.id_atendimento,
            'funcionario_id': e.funcionario_id,
            'fisioterapeuta_nome': funcionario.nome if funcionario else None,
            'fisioterapeuta_registro': funcionario.numero_profissional if funcionario else None,
            'data_evolucao': e.data_evolucao.strftime('%Y-%m-%d %H:%M:%S'),
            'evolucao_fisio': e.evolucao_fisio
        })
    return jsonify(resultado)

@bp.route('/evolucoes_fisioterapia/atendimento/<string:id_atendimento>', methods=['POST'])
def criar_evolucao_para_atendimento(id_atendimento):
    dados = request.get_json()

    nova_evolucao = EvolucaoFisioterapia(
        id_atendimento=id_atendimento,
        funcionario_id=dados['funcionario_id'],
        data_evolucao=datetime.strptime(dados['data_evolucao'], '%Y-%m-%d %H:%M:%S'),
        evolucao_fisio=dados['evolucao_fisio']
    )

    db.session.add(nova_evolucao)
    db.session.commit()

    return jsonify({'mensagem': 'Evolução criada com sucesso.', 'id': nova_evolucao.id}), 201

@bp.route('/evolucoes_fisioterapia/<int:id>', methods=['PUT'])
def atualizar_evolucao(id):
    e = EvolucaoFisioterapia.query.get_or_404(id)
    dados = request.get_json()

    e.funcionario_id = dados.get('funcionario_id', e.funcionario_id)
    if 'data_evolucao' in dados:
        e.data_evolucao = datetime.strptime(dados['data_evolucao'], '%Y-%m-%d %H:%M:%S')
    e.evolucao_fisio = dados.get('evolucao_fisio', e.evolucao_fisio)

    db.session.commit()

    return jsonify({'mensagem': 'Evolução atualizada com sucesso.'})


@bp.route('/evolucoes_nutricao/atendimento/<string:id_atendimento>', methods=['GET'])
def listar_evolucoes_nutricao(id_atendimento):
    evolucoes = EvolucaoNutricao.query.filter_by(id_atendimento=id_atendimento).all()
    resultado = []
    
    for e in evolucoes:
        # Buscar dados do funcionário
        funcionario = None
        if e.funcionario_id:
            funcionario = Funcionario.query.get(e.funcionario_id)
        
        resultado.append({
            'id': e.id,
            'id_atendimento': e.id_atendimento,
            'funcionario_id': e.funcionario_id,
            'nutricionista_nome': funcionario.nome if funcionario else None,
            'nutricionista_registro': funcionario.numero_profissional if funcionario else None,
            'data_evolucao': e.data_evolucao.strftime('%Y-%m-%d %H:%M:%S'),
            'evolucao_nutricao': e.evolucao_nutricao
        })
    return jsonify(resultado)

@bp.route('/evolucoes_nutricao/atendimento/<string:id_atendimento>', methods=['POST'])
def criar_evolucao_nutricao(id_atendimento):
    dados = request.get_json()

    nova_evolucao = EvolucaoNutricao(
        id_atendimento=id_atendimento,
        funcionario_id=dados['funcionario_id'],
        data_evolucao=datetime.strptime(dados['data_evolucao'], '%Y-%m-%d %H:%M:%S'),
        evolucao_nutricao=dados['evolucao_nutricao']
    )

    db.session.add(nova_evolucao)
    db.session.commit()

    return jsonify({'mensagem': 'Evolução nutricional criada com sucesso.', 'id': nova_evolucao.id}), 201

@bp.route('/evolucoes_nutricao/<int:id>', methods=['GET'])
def buscar_evolucao_nutricao(id):
    e = EvolucaoNutricao.query.get_or_404(id)
    return jsonify({
        'id': e.id,
        'id_atendimento': e.id_atendimento,
        'funcionario_id': e.funcionario_id,
        'data_evolucao': e.data_evolucao.strftime('%Y-%m-%d %H:%M:%S'),
        'evolucao_nutricao': e.evolucao_nutricao
    })

@bp.route('/evolucoes_assistentesocial/atendimento/<string:id_atendimento>', methods=['GET'])
def listar_evolucoes_assistentesocial(id_atendimento):
    evolucoes = EvolucaoAssistenteSocial.query.filter_by(id_atendimento=id_atendimento).all()
    resultado = []
    
    for e in evolucoes:
        # Buscar dados do funcionário
        funcionario = None
        if e.funcionario_id:
            funcionario = Funcionario.query.get(e.funcionario_id)
        
        resultado.append({
            'id': e.id,
            'id_atendimento': e.id_atendimento,
            'funcionario_id': e.funcionario_id,
            'assistente_social_nome': funcionario.nome if funcionario else None,
            'assistente_social_registro': funcionario.numero_profissional if funcionario else None,
            'data_evolucao': e.data_evolucao.strftime('%Y-%m-%d %H:%M:%S'),
            'evolucao_assistente_social': e.evolucao_assistentesocial
        })
    return jsonify(resultado)

@bp.route('/evolucoes_assistentesocial/atendimento/<string:id_atendimento>', methods=['POST'])
def criar_evolucao_assistentesocial(id_atendimento):
    dados = request.get_json()

    nova_evolucao = EvolucaoAssistenteSocial(
        id_atendimento=id_atendimento,
        funcionario_id=dados['funcionario_id'],
        data_evolucao=datetime.strptime(dados['data_evolucao'], '%Y-%m-%d %H:%M:%S'),
        evolucao_assistentesocial=dados['evolucao_assistente_social']
    )

    db.session.add(nova_evolucao)
    db.session.commit()

    return jsonify({'mensagem': 'Evolução de assistência social criada com sucesso.', 'id': nova_evolucao.id}), 201

@bp.route('/evolucoes_assistentesocial/<int:id>', methods=['GET'])
def buscar_evolucao_assistentesocial(id):
    e = EvolucaoAssistenteSocial.query.get_or_404(id)
    return jsonify({
        'id': e.id,
        'id_atendimento': e.id_atendimento,
        'funcionario_id': e.funcionario_id,
        'data_evolucao': e.data_evolucao.strftime('%Y-%m-%d %H:%M:%S'),
        'evolucao_assistentesocial': e.evolucao_assistentesocial
    })

@bp.route('/impressao_assistente_social/<string:atendimento_id>')
@login_required
def impressao_assistente_social(atendimento_id):
    """Página de impressão de evoluções de assistência social de um atendimento"""
    try:
        # Verificar se o atendimento existe
        atendimento = Atendimento.query.get_or_404(atendimento_id)
        
        # Buscar dados do paciente
        paciente = atendimento.paciente
        
        # Buscar dados da internação
        internacao = Internacao.query.filter_by(atendimento_id=atendimento_id).first()
        
        # Buscar evoluções de assistência social
        evolucoes_query = EvolucaoAssistenteSocial.query.filter_by(id_atendimento=atendimento_id)
        evolucoes_raw = evolucoes_query.order_by(EvolucaoAssistenteSocial.data_evolucao.desc()).all()
        
        # Formatar evoluções com dados do profissional
        evolucoes = []
        for e in evolucoes_raw:
            funcionario = None
            if e.funcionario_id:
                funcionario = Funcionario.query.get(e.funcionario_id)
            
            evolucoes.append({
                'id': e.id,
                'data_evolucao': e.data_evolucao,
                'evolucao_assistente_social': e.evolucao_assistentesocial,
                'assistente_social_nome': funcionario.nome if funcionario else 'Não identificado',
                'assistente_social_registro': funcionario.numero_profissional if funcionario else ''
            })
        
        # Função para obter data atual
        def moment():
            return datetime.now()
        
        return render_template('impressao_assistente_social.html',
                               paciente=paciente,
                               internacao=internacao,
                               atendimento=atendimento,
                               evolucoes=evolucoes,
                               moment=moment)
                               
    except Exception as e:
        flash(f'Erro ao carregar dados para impressão: {str(e)}', 'error')
        return redirect(url_for('main.index'))

@bp.route('/impressao_assistente_social_evolucao/<int:evolucao_id>')
@login_required
def impressao_assistente_social_evolucao(evolucao_id):
    """Página de impressão de uma evolução individual de assistência social"""
    try:
        # Buscar a evolução específica
        evolucao = EvolucaoAssistenteSocial.query.get_or_404(evolucao_id)
        
        # Buscar o atendimento
        atendimento = Atendimento.query.get_or_404(evolucao.id_atendimento)
        
        # Buscar dados do paciente
        paciente = atendimento.paciente
        
        # Buscar dados da internação
        internacao = Internacao.query.filter_by(atendimento_id=evolucao.id_atendimento).first()
        
        # Buscar dados do assistente social
        funcionario = None
        if evolucao.funcionario_id:
            funcionario = Funcionario.query.get(evolucao.funcionario_id)
        
        # Preparar dados da evolução
        evolucao_dados = {
            'id': evolucao.id,
            'data_evolucao': evolucao.data_evolucao,
            'evolucao_assistente_social': evolucao.evolucao_assistentesocial,
            'assistente_social_nome': funcionario.nome if funcionario else 'Não identificado',
            'assistente_social_registro': funcionario.numero_profissional if funcionario else ''
        }
        
        # Função para obter data atual
        def moment():
            return datetime.now()
        
        return render_template('impressao_assistente_social_individual.html',
                               paciente=paciente,
                               internacao=internacao,
                               atendimento=atendimento,
                               evolucao=evolucao_dados,
                               moment=moment)
                               
    except Exception as e:
        flash(f'Erro ao carregar dados para impressão: {str(e)}', 'error')
        return redirect(url_for('main.index'))


# ===================== FARMÁCIA =====================
@bp.route('/farmacia')
@login_required
def painel_farmacia():
    try:
        usuario_atual = get_current_user()
        if not usuario_atual:
            flash('Sessão expirada. Por favor, faça login novamente.', 'warning')
            return redirect(url_for('main.index'))

        if usuario_atual.cargo.strip().lower() != 'farmacia':
            flash('Acesso restrito à Farmácia.', 'danger')
            return redirect(url_for('main.index'))

        return render_template('farmacia.html')

    except Exception as e:
        logging.error(f"Erro ao acessar painel da farmácia: {str(e)}")
        logging.error(traceback.format_exc())
        flash('Erro ao acessar o painel. Por favor, tente novamente.', 'danger')
        return redirect(url_for('main.index'))


# GET /api/farmacia/nome
@bp.route('/api/farmacia/nome', methods=['GET'])
@login_required
def get_nome_farmacia():
    try:
        usuario_atual = get_current_user()
        if not usuario_atual:
            return jsonify({'success': False, 'message': 'Usuário não autenticado.'}), 401

        if usuario_atual.cargo.strip().lower() != 'farmacia':
            return jsonify({'success': False, 'message': 'Usuário não é da farmácia.'}), 403

        return jsonify({'success': True, 'nome': usuario_atual.nome})

    except Exception as e:
        logging.error(f"Erro ao buscar nome do usuário da farmácia: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro ao buscar nome do usuário da farmácia.', 'error': str(e)}), 500


# POST /api/farmacia/mudar-senha
@bp.route('/api/farmacia/mudar-senha', methods=['POST'])
@login_required
def mudar_senha_farmacia():
    try:
        usuario_atual = get_current_user()
        if not usuario_atual:
            return jsonify({'success': False, 'message': 'Usuário não autenticado.'}), 401

        if usuario_atual.cargo.strip().lower() != 'farmacia':
            return jsonify({'success': False, 'message': 'Usuário não é da farmácia.'}), 403

        dados = request.get_json()
        if not dados:
            return jsonify({'success': False, 'message': 'Dados não fornecidos.'}), 400

        senha_atual = dados.get('senha_atual')
        nova_senha = dados.get('nova_senha')

        if not senha_atual or not nova_senha:
            return jsonify({'success': False, 'message': 'Campos obrigatórios não preenchidos.'}), 400

        if not usuario_atual.check_password(senha_atual):
            return jsonify({'success': False, 'message': 'Senha atual incorreta.'}), 400

        usuario_atual.set_password(nova_senha)
        db.session.commit()

        return jsonify({'success': True, 'message': 'Senha alterada com sucesso.'})

    except Exception as e:
        db.session.rollback()
        logging.error(f"Erro ao alterar senha do usuário da farmácia: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro ao alterar a senha.', 'error': str(e)}), 500

@bp.route('/farmacia/estoque')
@login_required
def farmacia_estoque():
    try:
        usuario_atual = get_current_user()
        if not usuario_atual:
            flash('Sessão expirada. Por favor, faça login novamente.', 'warning')
            return redirect(url_for('main.index'))

        if usuario_atual.cargo.strip().lower() != 'farmacia':
            flash('Acesso restrito à Farmácia.', 'danger')
            return redirect(url_for('main.index'))

        return render_template('farmacia_estoque.html')

    except Exception as e:
        logging.error(f"Erro ao acessar estoque da farmácia: {str(e)}")
        logging.error(traceback.format_exc())
        flash('Erro ao acessar o estoque. Por favor, tente novamente.', 'danger')
        return redirect(url_for('main.index'))

@bp.route('/farmacia/dispensacoes')
@login_required
def farmacia_dispensacoes():
    try:
        usuario_atual = get_current_user()
        if not usuario_atual:
            flash('Sessão expirada. Por favor, faça login novamente.', 'warning')
            return redirect(url_for('main.index'))

        if usuario_atual.cargo.strip().lower() != 'farmacia':
            flash('Acesso restrito à Farmácia.', 'danger')
            return redirect(url_for('main.index'))

        return render_template('farmacia_disp.html')

    except Exception as e:
        logging.error(f"Erro ao acessar dispensações: {str(e)}")
        logging.error(traceback.format_exc())
        flash('Erro ao acessar dispensações. Por favor, tente novamente.', 'danger')
        return redirect(url_for('main.index'))

### Dispensações - Atendimentos e Prescrições

@bp.route('/api/dispensacoes/atendimentos-aguardando', methods=['GET'])
@login_required
def buscar_atendimentos_aguardando_medicacao():
    """
    Endpoint para buscar atendimentos aguardando medicação
    """
    try:
        usuario_atual = get_current_user()
        if not usuario_atual:
            return jsonify({'success': False, 'message': 'Sessão expirada. Faça login novamente.'}), 401

        if usuario_atual.cargo.strip().lower() != 'farmacia':
            return jsonify({'success': False, 'message': 'Acesso restrito à Farmácia.'}), 403

        from app.models import Atendimento, Paciente, Funcionario

        # Buscar atendimentos com status aguardando medicação
        atendimentos = Atendimento.query.filter_by(status='Aguardando Medicação').all()

        atendimentos_data = []
        for atendimento in atendimentos:
            # Verificar se tem prescrição médica
            if not atendimento.prescricao_medica:
                continue

            try:
                # Parse da prescrição JSON
                prescricoes = json.loads(atendimento.prescricao_medica)
            except (json.JSONDecodeError, TypeError):
                continue

            if not prescricoes:
                continue

            # Buscar dados do paciente
            paciente = Paciente.query.get(atendimento.paciente_id)

            # Buscar dados do médico
            medico = None
            if atendimento.medico_id:
                medico = Funcionario.query.get(atendimento.medico_id)

            atendimento_info = {
                'id': atendimento.id,
                'data_atendimento': atendimento.data_atendimento.strftime('%d/%m/%Y') if atendimento.data_atendimento else None,
                'hora_atendimento': atendimento.hora_atendimento.strftime('%H:%M') if atendimento.hora_atendimento else None,
                'paciente': {
                    'id': paciente.id if paciente else None,
                    'nome': paciente.nome if paciente else 'Paciente não encontrado',
                    'data_nascimento': paciente.data_nascimento.strftime('%d/%m/%Y') if paciente and paciente.data_nascimento else None
                },
                'medico': {
                    'id': medico.id if medico else None,
                    'nome': medico.nome if medico else 'Médico não informado'
                },
                'prescricoes': prescricoes,
                'observacao': atendimento.observacao,
                'dx': atendimento.dx
            }

            atendimentos_data.append(atendimento_info)

        return jsonify({
            'success': True,
            'data': atendimentos_data,
            'total': len(atendimentos_data)
        }), 200

    except Exception as e:
        logging.error(f"Erro ao buscar atendimentos aguardando medicação: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro interno do servidor.', 'error': str(e)}), 500

@bp.route('/api/dispensacoes/mapeamento-medicamentos', methods=['POST'])
@login_required
def mapear_medicamentos():
    """
    Endpoint para mapear nomes de medicamentos da prescrição para classes do estoque
    """
    try:
        usuario_atual = get_current_user()
        if not usuario_atual:
            return jsonify({'success': False, 'message': 'Sessão expirada. Faça login novamente.'}), 401

        if usuario_atual.cargo.strip().lower() != 'farmacia':
            return jsonify({'success': False, 'message': 'Acesso restrito à Farmácia.'}), 403

        dados = request.get_json()

        if not dados or 'prescricoes' not in dados:
            return jsonify({'success': False, 'message': 'Prescrições não fornecidas.'}), 400

        from app.models import MedicacaoClasse, MedicacaoItem

        prescricoes = dados['prescricoes']
        mapeamento = []

        for prescricao in prescricoes:
            nome_medicamento = prescricao.get('nome_medicamento', '').strip()

            if not nome_medicamento:
                continue

            # Buscar classe de medicação que contenha o nome
            # Usar LIKE para busca aproximada
            from sqlalchemy import text
            classe = MedicacaoClasse.query.filter(
                text("LOWER(nome) LIKE LOWER(:nome)")
            ).params(nome=f"%{nome_medicamento}%").first()

            if classe:
                # Buscar itens disponíveis na farmácia satélite
                itens_satelite = MedicacaoItem.query.filter_by(
                    id_med_classe=classe.id,
                    local='Farmácia Satélite 1'
                ).filter(MedicacaoItem.quantidade > 0).all()

                mapeamento.append({
                    'prescricao_original': prescricao,
                    'medicamento_encontrado': {
                        'id': classe.id,
                        'nome': classe.nome,
                        'apresentacao': classe.apresentacao,
                        'codigo': classe.codigo,
                        'unidade': classe.unidade
                    },
                    'itens_disponiveis': [{
                        'id': item.id,
                        'lote': item.lote,
                        'validade': item.validade.strftime('%d/%m/%Y'),
                        'quantidade': item.quantidade
                    } for item in itens_satelite],
                    'total_disponivel': sum(item.quantidade for item in itens_satelite)
                })
            else:
                mapeamento.append({
                    'prescricao_original': prescricao,
                    'medicamento_encontrado': None,
                    'erro': 'Medicamento não encontrado no estoque'
                })

        return jsonify({
            'success': True,
            'mapeamento': mapeamento
        }), 200

    except Exception as e:
        logging.error(f"Erro ao mapear medicamentos: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro interno do servidor.', 'error': str(e)}), 500

@bp.route('/api/fluxo-disp/aguardando', methods=['GET'])
@login_required
def listar_medicamentos_aguardando():
    """
    Endpoint para listar medicamentos aguardando dispensação via FluxoDisp
    """
    try:
        usuario_atual = get_current_user()
        if not usuario_atual:
            return jsonify({'success': False, 'message': 'Sessão expirada. Faça login novamente.'}), 401

        # Verificar se é farmacêutico ou admin
        if usuario_atual.cargo.strip().lower() not in ['farmacia', 'admin']:
            return jsonify({'success': False, 'message': 'Acesso restrito à Farmácia.'}), 403

        # Buscar medicamentos aguardando dispensação (id_responsavel = 0)
        medicamentos_aguardando = FluxoDisp.query.filter(
            FluxoDisp.id_responsavel == 0
        ).order_by(FluxoDisp.data.desc(), FluxoDisp.hora.desc()).all()

        resultado = []
        for fluxo in medicamentos_aguardando:
            # Buscar informações do atendimento e paciente
            atendimento = Atendimento.query.get(fluxo.id_atendimento)
            if not atendimento:
                continue

            paciente = Paciente.query.get(atendimento.paciente_id)
            medico = Funcionario.query.get(fluxo.id_medico)

            resultado.append({
                'id': fluxo.id,
                'atendimento_id': fluxo.id_atendimento,
                'paciente_nome': paciente.nome if paciente else 'N/A',
                'paciente_cartao_sus': paciente.cartao_sus if paciente else 'N/A',
                'medico_nome': medico.nome if medico else 'N/A',
                'medicamento': fluxo.medicamento,
                'quantidade': fluxo.quantidade,
                'data': fluxo.data.strftime('%d/%m/%Y') if fluxo.data else None,
                'hora': fluxo.hora.strftime('%H:%M') if fluxo.hora else None,
                'prescricao_id': fluxo.id_prescricao
            })

        return jsonify({
            'success': True,
            'medicamentos': resultado
        }), 200

    except Exception as e:
        logging.error(f"Erro ao listar medicamentos aguardando dispensação: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro interno do servidor.', 'error': str(e)}), 500


@bp.route('/api/fluxo-disp/dispensar/<int:fluxo_id>', methods=['PUT'])
@login_required
def dispensar_medicamento_fluxo(fluxo_id):
    """
    Endpoint para dispensar medicamento através do FluxoDisp
    Define o id_responsavel quando o farmacêutico realiza a dispensação
    """
    try:
        usuario_atual = get_current_user()
        if not usuario_atual:
            return jsonify({'success': False, 'message': 'Sessão expirada. Faça login novamente.'}), 401

        # Verificar se é farmacêutico ou admin
        if usuario_atual.cargo.strip().lower() not in ['farmacia', 'admin']:
            return jsonify({'success': False, 'message': 'Acesso restrito à Farmácia.'}), 403

        # Buscar o registro do FluxoDisp
        fluxo = FluxoDisp.query.get(fluxo_id)
        if not fluxo:
            return jsonify({'success': False, 'message': 'Registro de fluxo não encontrado.'}), 404

        # Verificar se já foi dispensado
        if fluxo.id_responsavel != 0:
            return jsonify({'success': False, 'message': 'Medicamento já foi dispensado.'}), 400

        # Definir o responsável pela dispensação
        fluxo.id_responsavel = usuario_atual.id

        # Log da dispensação
        logging.info(f"DISPENSAÇÃO FLUXO - Medicamento: {fluxo.medicamento} - Qtd: {fluxo.quantidade} - Farmacêutico: {usuario_atual.nome}")

        db.session.commit()

        return jsonify({
            'success': True,
            'message': f'Medicamento "{fluxo.medicamento}" dispensado com sucesso!',
            'fluxo_id': fluxo_id
        }), 200

    except Exception as e:
        db.session.rollback()
        logging.error(f"Erro ao dispensar medicamento via fluxo: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro interno do servidor.', 'error': str(e)}), 500


@bp.route('/api/dispensacoes/processar', methods=['POST'])
@login_required
def processar_dispensacao():
    """
    Endpoint para processar a dispensação de medicamentos
    """
    try:
        usuario_atual = get_current_user()
        if not usuario_atual:
            return jsonify({'success': False, 'message': 'Sessão expirada. Faça login novamente.'}), 401

        if usuario_atual.cargo.strip().lower() != 'farmacia':
            return jsonify({'success': False, 'message': 'Acesso restrito à Farmácia.'}), 403

        dados = request.get_json()

        campos_obrigatorios = ['atendimento_id', 'dispensacoes']
        for campo in campos_obrigatorios:
            if campo not in dados:
                return jsonify({'success': False, 'message': f'Campo {campo} obrigatório.'}), 400

        from app.models import Atendimento, MedicacaoItem

        # Verificar se o atendimento existe e está aguardando medicação
        atendimento = Atendimento.query.get(dados['atendimento_id'])
        if not atendimento:
            return jsonify({'success': False, 'message': 'Atendimento não encontrado.'}), 404

        if atendimento.status != 'Aguardando Medicação':
            return jsonify({'success': False, 'message': 'Atendimento não está aguardando medicação.'}), 400

        dispensacoes_processadas = []

        # Processar cada dispensação
        for dispensacao in dados['dispensacoes']:
            item_id = dispensacao.get('item_id')
            quantidade = dispensacao.get('quantidade')

            if not item_id or not quantidade:
                continue

            # Buscar o item
            item = MedicacaoItem.query.get(item_id)
            if not item:
                continue

            # Verificar se tem quantidade suficiente
            if item.quantidade < quantidade:
                return jsonify({
                    'success': False,
                    'message': f'Quantidade insuficiente no lote {item.lote}. Disponível: {item.quantidade}'
                }), 400

            # Reduzir quantidade do item
            item.quantidade -= quantidade

            # Se quantidade zerar, remover o item
            if item.quantidade == 0:
                db.session.delete(item)

            dispensacoes_processadas.append({
                'medicamento': item.medicacao_classe.nome,
                'lote': item.lote,
                'quantidade': quantidade,
                'local': item.local
            })

        # Atualizar status do atendimento
        atendimento.status = 'Medicamentos Dispensados'
        atendimento.horario_medicacao = datetime.now()

        # Registrar medicamentos utilizados
        if dispensacoes_processadas:
            atendimento.medicacao_utilizada = json.dumps(dispensacoes_processadas)

        # Log da dispensação
        logging.info(f"DISPENSAÇÃO REALIZADA - Atendimento: {atendimento.id} - Farmacêutico: {usuario_atual.nome}")
        for disp in dispensacoes_processadas:
            logging.info(f"  • {disp['medicamento']} - Lote: {disp['lote']} - Qtd: {disp['quantidade']}")

        db.session.commit()

        return jsonify({
            'success': True,
            'message': f'Dispensação realizada com sucesso! {len(dispensacoes_processadas)} medicamento(s) dispensado(s).',
            'dispensacoes': dispensacoes_processadas
        }), 200

    except Exception as e:
        db.session.rollback()
        logging.error(f"Erro ao processar dispensação: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro interno do servidor.', 'error': str(e)}), 500


### Dispensações - Prescrições da tabela fluxo_disp

@bp.route('/api/dispensacoes/fluxo-pendentes', methods=['GET'])
@login_required
def buscar_prescricoes_fluxo_pendentes():
    """
    Endpoint para buscar prescrições pendentes da tabela fluxo_disp (id_responsavel=0)
    """
    try:
        usuario_atual = get_current_user()
        if not usuario_atual:
            return jsonify({'success': False, 'message': 'Sessão expirada. Faça login novamente.'}), 401

        if usuario_atual.cargo.strip().lower() != 'farmacia':
            return jsonify({'success': False, 'message': 'Acesso restrito à Farmácia.'}), 403

        from app.models import FluxoDisp, Atendimento, Paciente, Funcionario

        # Buscar prescrições pendentes (id_responsavel=0)
        prescricoes_pendentes = FluxoDisp.query.filter_by(id_responsavel=0).all()

        prescricoes_data = []
        for prescricao in prescricoes_pendentes:
            # Buscar informações relacionadas
            # Tentar buscar como string primeiro, depois como int se necessário
            atendimento = None
            try:
                atendimento = Atendimento.query.get(str(prescricao.id_atendimento))
            except:
                try:
                    atendimento = Atendimento.query.get(prescricao.id_atendimento)
                except:
                    atendimento = None

            paciente = Paciente.query.get(atendimento.paciente_id) if atendimento else None
            medico = Funcionario.query.get(prescricao.id_medico) if prescricao.id_medico else None

            if not atendimento or not paciente:
                continue

            prescricao_info = {
                'id': prescricao.id,
                'id_atendimento': prescricao.id_atendimento,
                'id_prescricao': prescricao.id_prescricao,
                'medicamento': prescricao.medicamento,
                'quantidade': prescricao.quantidade,
                'status': prescricao.status,
                'data': prescricao.data.isoformat() if prescricao.data else None,
                'hora': prescricao.hora.isoformat() if prescricao.hora else None,
                'paciente': {
                    'id': paciente.id,
                    'nome': paciente.nome,
                    'cpf': paciente.cpf,
                    'telefone': paciente.telefone
                },
                'atendimento': {
                    'id': atendimento.id,
                    'data_atendimento': atendimento.data_atendimento.isoformat(),
                    'hora_atendimento': atendimento.hora_atendimento.isoformat() if atendimento.hora_atendimento else None,
                    'prescricao_medica': atendimento.prescricao_medica,
                    'dx': atendimento.dx
                },
                'medico': {
                    'id': medico.id if medico else None,
                    'nome': medico.nome if medico else None
                }
            }

            prescricoes_data.append(prescricao_info)

        return jsonify({
            'success': True,
            'data': prescricoes_data,
            'total': len(prescricoes_data)
        }), 200

    except Exception as e:
        logging.error(f"Erro ao buscar prescrições pendentes do fluxo: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro interno do servidor.', 'error': str(e)}), 500


@bp.route('/api/dispensacoes/fluxo-processar', methods=['POST'])
@login_required
def processar_dispensacao_fluxo():
    """
    Endpoint para processar dispensação de prescrições da tabela fluxo_disp
    """
    try:
        usuario_atual = get_current_user()
        if not usuario_atual:
            return jsonify({'success': False, 'message': 'Sessão expirada. Faça login novamente.'}), 401

        if usuario_atual.cargo.strip().lower() != 'farmacia':
            return jsonify({'success': False, 'message': 'Acesso restrito à Farmácia.'}), 403

        dados = request.get_json()

        campos_obrigatorios = ['prescricao_id', 'quantidade_confirmada', 'medicamento_id', 'status']
        for campo in campos_obrigatorios:
            if campo not in dados:
                return jsonify({'success': False, 'message': f'Campo {campo} obrigatório.'}), 400

        status_permitidos = ['Dispensado', 'Parcial', 'Cancelado']
        if dados['status'] not in status_permitidos:
            return jsonify({'success': False, 'message': 'Status inválido. Deve ser: Dispensado, Parcial ou Cancelado.'}), 400

        from app.models import FluxoDisp, MedicacaoClasse, MedicacaoItem

        # Buscar a prescrição no fluxo
        prescricao = FluxoDisp.query.get(dados['prescricao_id'])
        if not prescricao:
            return jsonify({'success': False, 'message': 'Prescrição não encontrada.'}), 404

        if prescricao.id_responsavel != 0:
            return jsonify({'success': False, 'message': 'Esta prescrição já foi processada.'}), 400

        # Buscar a classe de medicação
        medicamento_classe = MedicacaoClasse.query.get(dados['medicamento_id'])
        if not medicamento_classe:
            return jsonify({'success': False, 'message': 'Medicamento não encontrado.'}), 404

        status_dispensacao = dados['status']
        quantidade_solicitada = dados['quantidade_confirmada']

        # Se for cancelado, não reduz estoque
        if status_dispensacao == 'Cancelado':
            dispensacoes_processadas = []
        else:
            # Verificar se há quantidade suficiente no estoque (apenas para Dispensado e Parcial)
            setor = dados.get('setor')
            itens_query = MedicacaoItem.query.filter_by(id_med_classe=medicamento_classe.id)\
                .filter(MedicacaoItem.quantidade > 0)
            if setor and str(setor).strip() != '':
                itens_query = itens_query.filter(MedicacaoItem.local.ilike(f"%{setor}%"))
            itens_disponiveis = itens_query.order_by(MedicacaoItem.validade.asc()).all()

            quantidade_total_disponivel = sum(item.quantidade for item in itens_disponiveis)

            if quantidade_total_disponivel < quantidade_solicitada:
                return jsonify({
                    'success': False,
                    'message': f'Quantidade insuficiente em estoque. Disponível: {quantidade_total_disponivel}, Necessário: {quantidade_solicitada}'
                }), 400

            # Processar dispensação - reduzir estoque
            quantidade_restante = quantidade_solicitada
            dispensacoes_processadas = []

            for item in itens_disponiveis:
                if quantidade_restante <= 0:
                    break

                quantidade_deste_lote = min(quantidade_restante, item.quantidade)
                item.quantidade -= quantidade_deste_lote
                quantidade_restante -= quantidade_deste_lote

                dispensacoes_processadas.append({
                    'lote': item.lote,
                    'quantidade': quantidade_deste_lote,
                    'local': item.local
                })

                # Se o lote ficou vazio, remover
                if item.quantidade == 0:
                    db.session.delete(item)

        # Atualizar prescrição no fluxo com o responsável e status
        prescricao.id_responsavel = usuario_atual.id
        prescricao.status = status_dispensacao

        # Log da dispensação baseado no status
        if status_dispensacao == 'Cancelado':
            logging.info(f"DISPENSAÇÃO FLUXO CANCELADA - Prescrição: {prescricao.id} - Farmacêutico: {usuario_atual.nome}")
            logging.info(f"  • {medicamento_classe.nome} - Status: Cancelado")
        else:
            logging.info(f"DISPENSAÇÃO FLUXO REALIZADA - Prescrição: {prescricao.id} - Status: {status_dispensacao} - Farmacêutico: {usuario_atual.nome}")
            for disp in dispensacoes_processadas:
                logging.info(f"  • {medicamento_classe.nome} - Lote: {disp['lote']} - Qtd: {disp['quantidade']} - Local: {disp['local']}")

        db.session.commit()

        # Mensagem baseada no status
        if status_dispensacao == 'Cancelado':
            mensagem = f'Prescrição cancelada com sucesso!'
        elif status_dispensacao == 'Parcial':
            mensagem = f'Dispensação parcial realizada com sucesso! {quantidade_solicitada} unidade(s) dispensada(s).'
        else:
            mensagem = f'Dispensação realizada com sucesso! {quantidade_solicitada} unidade(s) dispensada(s).'

        return jsonify({
            'success': True,
            'message': mensagem,
            'dispensacoes': dispensacoes_processadas,
            'medicamento': medicamento_classe.nome,
            'status': status_dispensacao
        }), 200

    except Exception as e:
        db.session.rollback()
        logging.error(f"Erro ao processar dispensação do fluxo: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro interno do servidor.', 'error': str(e)}), 500


@bp.route('/api/dispensacoes/fluxo-processar-multiplos', methods=['POST'])
@login_required
def processar_dispensacao_fluxo_multiplos():
    """
    Processa múltiplos itens de uma mesma prescrição (tabela fluxo_disp) em uma única operação.
    Espera JSON: {
      prescricao_id: int,
      setor?: str,
      itens: [ { medicamento_id, quantidade_confirmada, status, observacoes? }, ... ]
    }
    """
    try:
        usuario_atual = get_current_user()
        if not usuario_atual:
            return jsonify({'success': False, 'message': 'Sessão expirada. Faça login novamente.'}), 401

        if usuario_atual.cargo.strip().lower() != 'farmacia':
            return jsonify({'success': False, 'message': 'Acesso restrito à Farmácia.'}), 403

        dados = request.get_json() or {}
        prescricao_id = dados.get('prescricao_id')
        itens_payload = dados.get('itens') or []
        setor = dados.get('setor')

        if not prescricao_id or not isinstance(itens_payload, list) or len(itens_payload) == 0:
            return jsonify({'success': False, 'message': 'Parâmetros inválidos: prescricao_id e itens são obrigatórios.'}), 400

        from app.models import FluxoDisp, MedicacaoClasse, MedicacaoItem

        prescricao = FluxoDisp.query.get(prescricao_id)
        if not prescricao:
            return jsonify({'success': False, 'message': 'Prescrição não encontrada.'}), 404

        if prescricao.id_responsavel != 0:
            return jsonify({'success': False, 'message': 'Esta prescrição já foi processada.'}), 400

        resultados = []
        houve_dispensado = False
        houve_parcial = False
        houve_cancelado = True  # assume cancelado até achar algo diferente

        for item in itens_payload:
            try:
                med_id = item.get('medicamento_id')
                qtd = item.get('quantidade_confirmada')
                status = item.get('status')
                observacoes = item.get('observacoes') or ''

                if med_id is None or qtd is None or status not in ['Dispensado', 'Parcial', 'Cancelado']:
                    resultados.append({'success': False, 'medicamento_id': med_id, 'message': 'Item inválido.'})
                    continue

                medicamento_classe = MedicacaoClasse.query.get(med_id)
                if not medicamento_classe:
                    resultados.append({'success': False, 'medicamento_id': med_id, 'message': 'Medicamento não encontrado.'})
                    continue

                dispensacoes_processadas = []

                if status == 'Cancelado':
                    resultados.append({
                        'success': True,
                        'medicamento': medicamento_classe.nome,
                        'status': status,
                        'quantidade': 0,
                        'dispensacoes': []
                    })
                else:
                    itens_query = MedicacaoItem.query.filter_by(id_med_classe=medicamento_classe.id) \
                        .filter(MedicacaoItem.quantidade > 0)
                    if setor and str(setor).strip() != '':
                        itens_query = itens_query.filter(MedicacaoItem.local.ilike(f"%{setor}%"))
                    itens_disponiveis = itens_query.order_by(MedicacaoItem.validade.asc()).all()

                    quantidade_total_disponivel = sum(itm.quantidade for itm in itens_disponiveis)

                    if quantidade_total_disponivel < int(qtd):
                        resultados.append({
                            'success': False,
                            'medicamento': medicamento_classe.nome,
                            'message': f'Quantidade insuficiente em estoque. Disponível: {quantidade_total_disponivel}, Necessário: {qtd}'
                        })
                        continue

                    quantidade_restante = int(qtd)
                    for itm in itens_disponiveis:
                        if quantidade_restante <= 0:
                            break
                        quantidade_deste_lote = min(quantidade_restante, itm.quantidade)
                        itm.quantidade -= quantidade_deste_lote
                        quantidade_restante -= quantidade_deste_lote
                        dispensacoes_processadas.append({
                            'lote': itm.lote,
                            'quantidade': quantidade_deste_lote,
                            'local': itm.local
                        })
                        if itm.quantidade == 0:
                            db.session.delete(itm)

                    resultados.append({
                        'success': True,
                        'medicamento': medicamento_classe.nome,
                        'status': status,
                        'quantidade': int(qtd),
                        'dispensacoes': dispensacoes_processadas,
                        'observacoes': observacoes
                    })

                    if status == 'Parcial':
                        houve_parcial = True
                        houve_cancelado = False
                    elif status == 'Dispensado':
                        houve_dispensado = True
                        houve_cancelado = False

            except Exception as ie:
                logging.error(f"Erro ao processar item em lote: {str(ie)}")
                logging.error(traceback.format_exc())
                resultados.append({'success': False, 'message': 'Erro ao processar item.', 'error': str(ie)})

        # Determinar status geral
        if houve_cancelado and not houve_dispensado and not houve_parcial:
            status_geral = 'Cancelado'
        elif houve_parcial or (houve_dispensado and any((r.get('status') != 'Dispensado') for r in resultados if r.get('success'))):
            status_geral = 'Parcial'
        else:
            status_geral = 'Dispensado'

        # Atualiza prescrição e responsável
        prescricao.id_responsavel = usuario_atual.id
        prescricao.status = status_geral

        # Logs
        logging.info(f"DISPENSAÇÃO FLUXO (LOTE) - Prescrição: {prescricao.id} - Status: {status_geral} - Farmacêutico: {usuario_atual.nome}")
        for r in resultados:
            if r.get('success') and r.get('dispensacoes'):
                for disp in r['dispensacoes']:
                    logging.info(f"  • {r.get('medicamento')} - Lote: {disp['lote']} - Qtd: {disp['quantidade']} - Local: {disp['local']}")

        db.session.commit()

        total_ok = len([r for r in resultados if r.get('success')])
        total_fail = len([r for r in resultados if not r.get('success')])
        mensagem = f"Processados {total_ok} item(ns); falhas em {total_fail}."

        return jsonify({
            'success': True,
            'message': mensagem,
            'status': status_geral,
            'resultados': resultados
        }), 200

    except Exception as e:
        db.session.rollback()
        logging.error(f"Erro ao processar dispensação do fluxo (lote): {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro interno do servidor.', 'error': str(e)}), 500

@bp.route('/api/medicamentos/buscar', methods=['GET'])
@login_required
def buscar_medicamentos():
    """
    Endpoint para buscar medicamentos disponíveis no estoque
    """
    try:
        from app.models import MedicacaoClasse, MedicacaoItem

        termo_busca = request.args.get('q', '').strip()
        medicamento_id = request.args.get('medicamento_id', type=int)
        setor = request.args.get('setor', type=str)

        # Se foi passado um ID específico de medicamento
        if medicamento_id:
            medicamento = MedicacaoClasse.query.get(medicamento_id)
            if not medicamento:
                return jsonify({'success': False, 'message': 'Medicamento não encontrado.'}), 404

            # Buscar itens disponíveis deste medicamento
            itens_query = MedicacaoItem.query.filter_by(id_med_classe=medicamento_id)\
                .filter(MedicacaoItem.quantidade > 0)
            if setor and str(setor).strip() != '':
                itens_query = itens_query.filter(MedicacaoItem.local.ilike(f"%{setor}%"))
            itens = itens_query.order_by(MedicacaoItem.validade.asc()).all()

            quantidade_total = sum(item.quantidade for item in itens)

            return jsonify({
                'success': True,
                'medicamento': {
                    'id': medicamento.id,
                    'nome': medicamento.nome,
                    'apresentacao': medicamento.apresentacao,
                    'classe': medicamento.classe,
                    'unidade': medicamento.unidade,
                    'quantidade_total': quantidade_total,
                    'itens': [{
                        'id': item.id,
                        'lote': item.lote,
                        'validade': item.validade.isoformat() if item.validade else None,
                        'local': item.local,
                        'quantidade': item.quantidade
                    } for item in itens]
                }
            }), 200

        # Busca geral de medicamentos
        query = MedicacaoClasse.query

        if termo_busca:
            query = query.filter(MedicacaoClasse.nome.ilike(f'%{termo_busca}%'))

        medicamentos = query.all()

        medicamentos_data = []
        for med in medicamentos:
            # Calcular quantidade total disponível
            if setor and str(setor).strip() != '':
                quantidade_total = sum(item.quantidade for item in med.itens if item.quantidade > 0 and item.local and setor.lower() in item.local.lower())
            else:
                quantidade_total = sum(item.quantidade for item in med.itens if item.quantidade > 0)

            if quantidade_total > 0:  # Só mostrar medicamentos com estoque
                medicamentos_data.append({
                    'id': med.id,
                    'nome': med.nome,
                    'apresentacao': med.apresentacao,
                    'classe': med.classe,
                    'unidade': med.unidade,
                    'quantidade_total': quantidade_total
                })

        return jsonify({
            'success': True,
            'medicamentos': medicamentos_data,
            'total': len(medicamentos_data)
        }), 200

    except Exception as e:
        logging.error(f"Erro ao buscar medicamentos: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro interno do servidor.', 'error': str(e)}), 500


@bp.route('/api/medicamentos/disponiveis', methods=['GET'])
@login_required
def buscar_medicamentos_disponiveis():
    """
    Endpoint para buscar medicamentos disponíveis em estoque (com MedicacaoItem)
    """
    try:
        from app.models import MedicacaoClasse, MedicacaoItem

        # Buscar medicamentos que têm itens em estoque
        medicamentos_com_estoque = MedicacaoClasse.query\
            .join(MedicacaoItem)\
            .filter(MedicacaoItem.quantidade > 0)\
            .distinct()\
            .all()

        medicamentos_data = []
        for med in medicamentos_com_estoque:
            # Calcular quantidade total disponível
            quantidade_total = sum(item.quantidade for item in med.itens if item.quantidade > 0)

            if quantidade_total > 0:
                medicamentos_data.append({
                    'id': med.id,
                    'nome': med.nome,
                    'apresentacao': med.apresentacao,
                    'classe': med.classe,
                    'unidade': med.unidade,
                    'quantidade_total': quantidade_total
                })

        # Ordenar por nome
        medicamentos_data.sort(key=lambda x: x['nome'])

        return jsonify({
            'success': True,
            'medicamentos': medicamentos_data,
            'total': len(medicamentos_data)
        }), 200

    except Exception as e:
        logging.error(f"Erro ao buscar medicamentos disponíveis: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro interno do servidor.', 'error': str(e)}), 500
# Listagem de setores do estoque (locais)
@bp.route('/api/estoque/setores', methods=['GET'])
@login_required
def listar_setores_estoque():
    try:
        usuario_atual = get_current_user()
        if not usuario_atual:
            return jsonify({'success': False, 'message': 'Sessão expirada. Faça login novamente.'}), 401

        # Opcional: restringe a Farmácia
        if usuario_atual.cargo.strip().lower() != 'farmacia':
            return jsonify({'success': False, 'message': 'Acesso restrito à Farmácia.'}), 403

        from app.models import MedicacaoItem

        setores_rows = db.session.query(MedicacaoItem.local) \
            .filter(MedicacaoItem.local.isnot(None)) \
            .filter(MedicacaoItem.local != '') \
            .filter(MedicacaoItem.quantidade > 0) \
            .distinct() \
            .all()

        setores = sorted([row[0] for row in setores_rows if row and row[0]])

        return jsonify({'success': True, 'setores': setores}), 200
    except Exception as e:
        logging.error(f"Erro ao listar setores do estoque: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro interno do servidor.', 'error': str(e)}), 500


### Medicamentos - Cadastro de Classes e Itens

@bp.route('/api/medicamentos/classes', methods=['POST'])
@login_required
def criar_medicacao_classe():
    """
    Endpoint para criar uma nova classe de medicação
    """
    try:
        usuario_atual = get_current_user()
        if not usuario_atual:
            return jsonify({'success': False, 'message': 'Sessão expirada. Faça login novamente.'}), 401

        if usuario_atual.cargo.strip().lower() != 'farmacia':
            return jsonify({'success': False, 'message': 'Acesso restrito à Farmácia.'}), 403

        dados = request.get_json()

        # Validações
        if not dados:
            return jsonify({'success': False, 'message': 'Dados não fornecidos.'}), 400

        campos_obrigatorios = ['nome', 'apresentacao', 'classe', 'codigo', 'unidade']
        for campo in campos_obrigatorios:
            if campo not in dados or not dados[campo] or str(dados[campo]).strip() == '':
                return jsonify({'success': False, 'message': f'Campo {campo} é obrigatório.'}), 400

        # Verificar se código já existe
        from app.models import MedicacaoClasse
        classe_existente = MedicacaoClasse.query.filter_by(codigo=dados['codigo'].strip()).first()
        if classe_existente:
            return jsonify({'success': False, 'message': 'Código já cadastrado para outra medicação.'}), 400

        # Criar nova classe
        nova_classe = MedicacaoClasse(
            nome=dados['nome'].strip(),
            apresentacao=dados['apresentacao'].strip(),
            classe=dados['classe'].strip(),
            codigo=dados['codigo'].strip(),
            unidade=dados['unidade'].strip()
        )

        db.session.add(nova_classe)
        db.session.commit()

        return jsonify({
            'success': True,
            'message': 'Classe de medicação cadastrada com sucesso!',
            'data': {
                'id': nova_classe.id,
                'nome': nova_classe.nome,
                'apresentacao': nova_classe.apresentacao,
                'classe': nova_classe.classe,
                'codigo': nova_classe.codigo,
                'unidade': nova_classe.unidade
            }
        }), 201

    except Exception as e:
        db.session.rollback()
        logging.error(f"Erro ao criar classe de medicação: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro interno do servidor.', 'error': str(e)}), 500

@bp.route('/api/medicamentos/classes', methods=['GET'])
@login_required
def listar_medicacao_classes():
    """
    Endpoint para listar todas as classes de medicação
    """
    try:
        usuario_atual = get_current_user()
        if not usuario_atual:
            return jsonify({'success': False, 'message': 'Sessão expirada. Faça login novamente.'}), 401

        if usuario_atual.cargo.strip().lower() != 'farmacia':
            return jsonify({'success': False, 'message': 'Acesso restrito à Farmácia.'}), 403

        from app.models import MedicacaoClasse
        classes = MedicacaoClasse.query.order_by(MedicacaoClasse.nome).all()

        classes_data = []
        for classe in classes:
            classes_data.append({
                'id': classe.id,
                'nome': classe.nome,
                'apresentacao': classe.apresentacao,
                'classe': classe.classe,
                'codigo': classe.codigo,
                'unidade': classe.unidade
            })

        return jsonify({
            'success': True,
            'data': classes_data,
            'total': len(classes_data)
        }), 200

    except Exception as e:
        logging.error(f"Erro ao listar classes de medicação: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro interno do servidor.', 'error': str(e)}), 500


@bp.route('/api/medicamentos/classes/<int:classe_id>', methods=['PUT'])
@login_required
def atualizar_medicacao_classe(classe_id):
    """
    Endpoint para atualizar uma classe de medicação existente
    """
    try:
        usuario_atual = get_current_user()
        if not usuario_atual:
            return jsonify({'success': False, 'message': 'Sessão expirada. Faça login novamente.'}), 401

        if usuario_atual.cargo.strip().lower() != 'farmacia':
            return jsonify({'success': False, 'message': 'Acesso restrito à Farmácia.'}), 403

        from app.models import MedicacaoClasse

        classe = MedicacaoClasse.query.get(classe_id)
        if not classe:
            return jsonify({'success': False, 'message': 'Classe de medicação não encontrada.'}), 404

        dados = request.get_json() or {}

        # Atualiza somente campos permitidos
        campos_permitidos = ['nome', 'apresentacao', 'classe', 'codigo', 'unidade']

        # Se for alterar código, verificar duplicidade
        if 'codigo' in dados and dados['codigo'] and dados['codigo'].strip() != classe.codigo:
            existente = MedicacaoClasse.query.filter_by(codigo=dados['codigo'].strip()).first()
            if existente and existente.id != classe.id:
                return jsonify({'success': False, 'message': 'Código já cadastrado para outra medicação.'}), 400

        for campo in campos_permitidos:
            if campo in dados and dados[campo] is not None and str(dados[campo]).strip() != '':
                setattr(classe, campo, str(dados[campo]).strip())

        db.session.commit()

        return jsonify({
            'success': True,
            'message': 'Classe de medicação atualizada com sucesso!',
            'data': {
                'id': classe.id,
                'nome': classe.nome,
                'apresentacao': classe.apresentacao,
                'classe': classe.classe,
                'codigo': classe.codigo,
                'unidade': classe.unidade
            }
        }), 200

    except Exception as e:
        db.session.rollback()
        logging.error(f"Erro ao atualizar classe de medicação: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro interno do servidor.', 'error': str(e)}), 500
@bp.route('/api/medicamentos/itens', methods=['POST'])
@login_required
def criar_medicacao_item():
    """
    Endpoint para criar um novo item de medicação (lote)
    """
    try:
        usuario_atual = get_current_user()
        if not usuario_atual:
            return jsonify({'success': False, 'message': 'Sessão expirada. Faça login novamente.'}), 401

        if usuario_atual.cargo.strip().lower() != 'farmacia':
            return jsonify({'success': False, 'message': 'Acesso restrito à Farmácia.'}), 403

        dados = request.get_json()

        # Validações
        if not dados:
            return jsonify({'success': False, 'message': 'Dados não fornecidos.'}), 400

        campos_obrigatorios = ['id_med_classe', 'lote', 'validade', 'local', 'quantidade']
        for campo in campos_obrigatorios:
            if campo not in dados or not dados[campo] or str(dados[campo]).strip() == '':
                return jsonify({'success': False, 'message': f'Campo {campo} é obrigatório.'}), 400

        # Verificar se a classe existe
        from app.models import MedicacaoClasse, MedicacaoItem
        classe = MedicacaoClasse.query.get(dados['id_med_classe'])
        if not classe:
            return jsonify({'success': False, 'message': 'Classe de medicação não encontrada.'}), 404

        # Verificar se lote já existe para esta classe
        item_existente = MedicacaoItem.query.filter_by(
            id_med_classe=dados['id_med_classe'],
            lote=dados['lote'].strip()
        ).first()
        if item_existente:
            return jsonify({'success': False, 'message': 'Lote já cadastrado para esta medicação.'}), 400

        # Validar quantidade
        try:
            quantidade = int(dados['quantidade'])
            if quantidade <= 0:
                return jsonify({'success': False, 'message': 'Quantidade deve ser maior que zero.'}), 400
        except (ValueError, TypeError):
            return jsonify({'success': False, 'message': 'Quantidade deve ser um número válido.'}), 400

        # Validar data de validade
        from datetime import datetime
        try:
            validade_date = datetime.strptime(dados['validade'], '%Y-%m-%d').date()
            if validade_date <= datetime.now().date():
                return jsonify({'success': False, 'message': 'Data de validade deve ser futura.'}), 400
        except ValueError:
            return jsonify({'success': False, 'message': 'Formato de data inválido. Use YYYY-MM-DD.'}), 400

        # Criar novo item
        novo_item = MedicacaoItem(
            id_med_classe=dados['id_med_classe'],
            lote=dados['lote'].strip(),
            validade=validade_date,
            local=dados['local'].strip(),
            quantidade=quantidade
        )

        db.session.add(novo_item)
        db.session.commit()

        return jsonify({
            'success': True,
            'message': 'Item de medicação cadastrado com sucesso!',
            'data': {
                'id': novo_item.id,
                'id_med_classe': novo_item.id_med_classe,
                'lote': novo_item.lote,
                'validade': novo_item.validade.strftime('%Y-%m-%d'),
                'local': novo_item.local,
                'quantidade': novo_item.quantidade,
                'medicacao_classe': {
                    'nome': classe.nome,
                    'apresentacao': classe.apresentacao,
                    'codigo': classe.codigo
                }
            }
        }), 201

    except Exception as e:
        db.session.rollback()
        logging.error(f"Erro ao criar item de medicação: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro interno do servidor.', 'error': str(e)}), 500

@bp.route('/api/medicamentos/itens/<int:id_classe>', methods=['GET'])
@login_required
def listar_medicacao_itens(id_classe):
    """
    Endpoint para listar todos os itens de uma classe de medicação
    """
    try:
        usuario_atual = get_current_user()
        if not usuario_atual:
            return jsonify({'success': False, 'message': 'Sessão expirada. Faça login novamente.'}), 401

        if usuario_atual.cargo.strip().lower() != 'farmacia':
            return jsonify({'success': False, 'message': 'Acesso restrito à Farmácia.'}), 403

        from app.models import MedicacaoClasse, MedicacaoItem

        # Verificar se a classe existe
        classe = MedicacaoClasse.query.get(id_classe)
        if not classe:
            return jsonify({'success': False, 'message': 'Classe de medicação não encontrada.'}), 404

        # Buscar itens da classe
        itens = MedicacaoItem.query.filter_by(id_med_classe=id_classe).order_by(MedicacaoItem.validade).all()

        itens_data = []
        from datetime import datetime
        hoje = datetime.now().date()

        for item in itens:
            status = 'ativo'
            if item.validade < hoje:
                status = 'vencido'
            elif (item.validade - hoje).days <= 30:
                status = 'proximo_vencimento'

            itens_data.append({
                'id': item.id,
                'id_med_classe': item.id_med_classe,
                'lote': item.lote,
                'validade': item.validade.strftime('%Y-%m-%d'),
                'local': item.local,
                'quantidade': item.quantidade,
                'status': status,
                'dias_para_vencer': (item.validade - hoje).days if item.validade > hoje else -(hoje - item.validade).days
            })

        return jsonify({
            'success': True,
            'data': {
                'classe': {
                    'id': classe.id,
                    'nome': classe.nome,
                    'apresentacao': classe.apresentacao,
                    'codigo': classe.codigo
                },
                'itens': itens_data,
                'total_itens': len(itens_data),
                'itens_ativos': len([i for i in itens_data if i['status'] == 'ativo']),
                'itens_proximos_vencimento': len([i for i in itens_data if i['status'] == 'proximo_vencimento']),
                'itens_vencidos': len([i for i in itens_data if i['status'] == 'vencido'])
            }
        }), 200

    except Exception as e:
        logging.error(f"Erro ao listar itens de medicação: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro interno do servidor.', 'error': str(e)}), 500

@bp.route('/api/medicamentos/itens/item/<int:item_id>', methods=['GET'])
@login_required
def obter_medicacao_item(item_id):
    """
    Endpoint para obter detalhes de um item (lote) específico
    """
    try:
        usuario_atual = get_current_user()
        if not usuario_atual:
            return jsonify({'success': False, 'message': 'Sessão expirada. Faça login novamente.'}), 401

        if usuario_atual.cargo.strip().lower() != 'farmacia':
            return jsonify({'success': False, 'message': 'Acesso restrito à Farmácia.'}), 403

        from app.models import MedicacaoClasse, MedicacaoItem
        item = MedicacaoItem.query.get(item_id)
        if not item:
            return jsonify({'success': False, 'message': 'Item de medicação não encontrado.'}), 404

        classe = MedicacaoClasse.query.get(item.id_med_classe)

        return jsonify({
            'success': True,
            'data': {
                'id': item.id,
                'id_med_classe': item.id_med_classe,
                'lote': item.lote,
                'validade': item.validade.strftime('%Y-%m-%d') if item.validade else None,
                'local': item.local,
                'quantidade': item.quantidade,
                'classe': {
                    'id': classe.id if classe else None,
                    'nome': classe.nome if classe else None,
                    'apresentacao': classe.apresentacao if classe else None,
                    'codigo': classe.codigo if classe else None
                }
            }
        }), 200

    except Exception as e:
        logging.error(f"Erro ao obter item de medicação: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro interno do servidor.', 'error': str(e)}), 500

@bp.route('/api/medicamentos/itens/<int:item_id>', methods=['PUT'])
@login_required
def atualizar_medicacao_item(item_id):
    """
    Endpoint para atualizar dados de um item (lote) de medicação
    """
    try:
        usuario_atual = get_current_user()
        if not usuario_atual:
            return jsonify({'success': False, 'message': 'Sessão expirada. Faça login novamente.'}), 401

        if usuario_atual.cargo.strip().lower() != 'farmacia':
            return jsonify({'success': False, 'message': 'Acesso restrito à Farmácia.'}), 403

        from app.models import MedicacaoItem
        item = MedicacaoItem.query.get(item_id)
        if not item:
            return jsonify({'success': False, 'message': 'Item de medicação não encontrado.'}), 404

        dados = request.get_json() or {}

        # Validar e preparar campos
        if 'lote' in dados and str(dados['lote']).strip() != '':
            novo_lote = str(dados['lote']).strip()
            # Verificar duplicidade de lote na mesma classe
            existente = MedicacaoItem.query.filter(
                MedicacaoItem.id_med_classe == item.id_med_classe,
                MedicacaoItem.lote == novo_lote,
                MedicacaoItem.id != item.id
            ).first()
            if existente:
                return jsonify({'success': False, 'message': 'Já existe um item com este lote para esta medicação.'}), 400
            item.lote = novo_lote

        if 'validade' in dados and str(dados['validade']).strip() != '':
            from datetime import datetime
            try:
                nova_validade = datetime.strptime(str(dados['validade']).strip(), '%Y-%m-%d').date()
                item.validade = nova_validade
            except ValueError:
                return jsonify({'success': False, 'message': 'Formato de data inválido. Use YYYY-MM-DD.'}), 400

        if 'local' in dados and str(dados['local']).strip() != '':
            item.local = str(dados['local']).strip()

        if 'quantidade' in dados:
            try:
                nova_qtd = int(dados['quantidade'])
                if nova_qtd <= 0:
                    return jsonify({'success': False, 'message': 'Quantidade deve ser maior que zero.'}), 400
                item.quantidade = nova_qtd
            except (ValueError, TypeError):
                return jsonify({'success': False, 'message': 'Quantidade deve ser um número válido.'}), 400

        db.session.commit()

        return jsonify({'success': True, 'message': 'Item atualizado com sucesso.'}), 200

    except Exception as e:
        db.session.rollback()
        logging.error(f"Erro ao atualizar item de medicação: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro interno do servidor.', 'error': str(e)}), 500

@bp.route('/api/medicamentos/itens/<int:item_id>', methods=['DELETE'])
@login_required
def excluir_medicacao_item(item_id):
    """
    Endpoint para excluir um item (lote) de medicação
    """
    try:
        usuario_atual = get_current_user()
        if not usuario_atual:
            return jsonify({'success': False, 'message': 'Sessão expirada. Faça login novamente.'}), 401

        if usuario_atual.cargo.strip().lower() != 'farmacia':
            return jsonify({'success': False, 'message': 'Acesso restrito à Farmácia.'}), 403

        from app.models import MedicacaoItem
        item = MedicacaoItem.query.get(item_id)
        if not item:
            return jsonify({'success': False, 'message': 'Item de medicação não encontrado.'}), 404

        db.session.delete(item)
        db.session.commit()

        return jsonify({'success': True, 'message': 'Item excluído com sucesso.'}), 200

    except Exception as e:
        db.session.rollback()
        logging.error(f"Erro ao excluir item de medicação: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro interno do servidor.', 'error': str(e)}), 500

@bp.route('/api/medicamentos/transferir', methods=['POST'])
@login_required
def transferir_medicamento():
    """
    Endpoint para transferir medicamentos entre farmácias
    """
    try:
        usuario_atual = get_current_user()
        if not usuario_atual:
            return jsonify({'success': False, 'message': 'Sessão expirada. Faça login novamente.'}), 401

        if usuario_atual.cargo.strip().lower() != 'farmacia':
            return jsonify({'success': False, 'message': 'Acesso restrito à Farmácia.'}), 403

        dados = request.get_json()

        # Validações
        if not dados:
            return jsonify({'success': False, 'message': 'Dados não fornecidos.'}), 400

        campos_obrigatorios = ['classe_id', 'lote_id', 'origem', 'destino', 'quantidade']
        for campo in campos_obrigatorios:
            if campo not in dados or not dados[campo]:
                return jsonify({'success': False, 'message': f'Campo {campo} é obrigatório.'}), 400

        # Validar quantidade
        try:
            quantidade = int(dados['quantidade'])
            if quantidade <= 0:
                return jsonify({'success': False, 'message': 'Quantidade deve ser maior que zero.'}), 400
        except (ValueError, TypeError):
            return jsonify({'success': False, 'message': 'Quantidade deve ser um número válido.'}), 400

        from app.models import MedicacaoClasse, MedicacaoItem

        # Verificar se a classe existe
        classe = MedicacaoClasse.query.get(dados['classe_id'])
        if not classe:
            return jsonify({'success': False, 'message': 'Classe de medicação não encontrada.'}), 404

        # Verificar se o lote existe
        lote = MedicacaoItem.query.get(dados['lote_id'])
        if not lote:
            return jsonify({'success': False, 'message': 'Lote não encontrado.'}), 404

        # Verificar se o lote pertence à classe
        if lote.id_med_classe != dados['classe_id']:
            return jsonify({'success': False, 'message': 'Lote não pertence à classe selecionada.'}), 400

        # Verificar se o lote está no local de origem
        if lote.local != dados['origem']:
            return jsonify({'success': False, 'message': 'Lote não está no local de origem especificado.'}), 400

        # Verificar se há quantidade suficiente
        if lote.quantidade < quantidade:
            return jsonify({'success': False, 'message': 'Quantidade insuficiente no lote.'}), 400

        # Verificar se origem e destino são diferentes
        if dados['origem'] == dados['destino']:
            return jsonify({'success': False, 'message': 'Origem e destino devem ser diferentes.'}), 400

        # Verificar se já existe um lote igual no destino
        lote_destino = MedicacaoItem.query.filter_by(
            id_med_classe=dados['classe_id'],
            lote=lote.lote,
            local=dados['destino']
        ).first()

        if (lote_destino):
            # Se existe lote igual no destino, somar as quantidades
            lote_destino.quantidade += quantidade
        else:
            # Criar novo lote no destino
            lote_destino = MedicacaoItem(
                id_med_classe=dados['classe_id'],
                lote=lote.lote,
                validade=lote.validade,
                local=dados['destino'],
                quantidade=quantidade
            )
            db.session.add(lote_destino)

        # Diminuir quantidade do lote de origem
        lote.quantidade -= quantidade

        # Se quantidade do lote de origem zerar, remover o lote
        if lote.quantidade == 0:
            db.session.delete(lote)

        # Registrar observação se fornecida
        if dados.get('observacoes'):
            logging.info(f"Transferência registrada: {dados['observacoes']} - {usuario_atual.nome} - {datetime.now()}")

        db.session.commit()

        return jsonify({
            'success': True,
            'message': f'Transferência realizada com sucesso! {quantidade} unidade(s) transferida(s) de {dados["origem"]} para {dados["destino"]}.',
            'data': {
                'medicamento': classe.nome,
                'lote': lote.lote,
                'origem': dados['origem'],
                'destino': dados['destino'],
                'quantidade': quantidade
            }
        }), 200

    except Exception as e:
        db.session.rollback()
        logging.error(f"Erro ao realizar transferência: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro interno do servidor.', 'error': str(e)}), 500

@bp.route('/api/medicamentos/saida', methods=['POST'])
@login_required
def saida_medicamento():
    """
    Endpoint para registrar saída de medicamentos (dispensação)
    """
    try:
        usuario_atual = get_current_user()
        if not usuario_atual:
            return jsonify({'success': False, 'message': 'Sessão expirada. Faça login novamente.'}), 401

        if usuario_atual.cargo.strip().lower() != 'farmacia':
            return jsonify({'success': False, 'message': 'Acesso restrito à Farmácia.'}), 403

        dados = request.get_json()

        # Validações
        if not dados:
            return jsonify({'success': False, 'message': 'Dados não fornecidos.'}), 400

        campos_obrigatorios = ['classe_id', 'lote_id', 'local', 'quantidade']
        for campo in campos_obrigatorios:
            if campo not in dados or not dados[campo]:
                return jsonify({'success': False, 'message': f'Campo {campo} é obrigatório.'}), 400

        # Validar quantidade
        try:
            quantidade = int(dados['quantidade'])
            if quantidade <= 0:
                return jsonify({'success': False, 'message': 'Quantidade deve ser maior que zero.'}), 400
        except (ValueError, TypeError):
            return jsonify({'success': False, 'message': 'Quantidade deve ser um número válido.'}), 400

        from app.models import MedicacaoClasse, MedicacaoItem

        # Verificar se a classe existe
        classe = MedicacaoClasse.query.get(dados['classe_id'])
        if not classe:
            return jsonify({'success': False, 'message': 'Classe de medicação não encontrada.'}), 404

        # Verificar se o lote existe
        lote = MedicacaoItem.query.get(dados['lote_id'])
        if not lote:
            return jsonify({'success': False, 'message': 'Lote não encontrado.'}), 404

        # Verificar se o lote pertence à classe
        if lote.id_med_classe != dados['classe_id']:
            return jsonify({'success': False, 'message': 'Lote não pertence à classe selecionada.'}), 400

        # Verificar se o lote está no local especificado
        if lote.local != dados['local']:
            return jsonify({'success': False, 'message': 'Lote não está no local especificado.'}), 400

        # Verificar se há quantidade suficiente
        if lote.quantidade < quantidade:
            return jsonify({'success': False, 'message': 'Quantidade insuficiente no lote.'}), 400

        # Diminuir quantidade do lote
        lote.quantidade -= quantidade

        # Se quantidade do lote zerar, remover o lote
        if lote.quantidade == 0:
            db.session.delete(lote)

        # Registrar saída no log com ALERTA de retirada forçada
        logging.warning(f"🚨 RETIRADA FORÇADA: {quantidade} un de '{classe.nome}' (Lote {lote.lote}) - Local: {dados['local']} - Farmaceuta: {usuario_atual.nome} - {datetime.now()}")
        logging.warning(f"⚠️ ATENÇÃO: Esta retirada foi registrada para análise posterior")

        db.session.commit()

        return jsonify({
            'success': True,
            'message': f'Retirada forçada realizada! {quantidade} unidade(s) de "{classe.nome}" removida(s) permanentemente do estoque.',
            'data': {
                'medicamento': classe.nome,
                'lote': lote.lote,
                'local': dados['local'],
                'quantidade': quantidade,
                'farmaceuta': usuario_atual.nome,
                'tipo': 'retirada_forcada'
            }
        }), 200

    except Exception as e:
        db.session.rollback()
        logging.error(f"Erro ao registrar saída: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro interno do servidor.', 'error': str(e)}), 500

### Observacao

@bp.route('/observacao')
@login_required
def observacao_lobby():
    try:
        current_user = get_current_user()
        if current_user.cargo.lower() not in ['medico', 'enfermeiro', 'multi']:
            flash('Acesso restrito a médicos, enfermeiros e profissionais multi.', 'danger')
            return redirect(url_for('main.index'))
        
        return render_template('observacao.html')
        
    except Exception as e:
        logging.error(f"Erro ao acessar página da clínica: {str(e)}")
        logging.error(traceback.format_exc())
        flash('Erro ao acessar a página. Por favor, tente novamente.', 'danger')
        return redirect(url_for('main.index'))


@bp.route('/recepcionista/status-pacientes')
@login_required
def status_pacientes_recepcionista():
    """
    Página para recepcionistas verem o status de pacientes internados e em observação.
    """
    try:
        current_user = get_current_user()
        if current_user.cargo.lower() != 'recepcionista':
            flash('Acesso restrito a recepcionistas.', 'danger')
            return redirect(url_for('main.index'))

        return render_template('recepcionista_status_pacientes.html')

    except Exception as e:
        logging.error(f"Erro ao acessar status de pacientes: {str(e)}")
        logging.error(traceback.format_exc())
        flash('Erro ao acessar a página. Por favor, tente novamente.', 'danger')
        return redirect(url_for('main.recepcionista'))


@bp.route('/api/pacientes/status/internados')
@login_required
def api_pacientes_internados():
    """
    API para obter lista de pacientes internados (para recepcionistas).
    """
    try:
        current_user = get_current_user()
        if current_user.cargo.lower() != 'recepcionista':
            return jsonify({'success': False, 'message': 'Acesso não autorizado'}), 403

        # Buscar apenas pacientes cujo atendimento esteja com status exatamente "Internado"
        internacoes_ativas = db.session.query(Internacao).join(
            Atendimento, Internacao.atendimento_id == Atendimento.id
        ).filter(
            Atendimento.status == 'Internado'
        ).all()

        pacientes_data = []
        for internacao in internacoes_ativas:
            paciente = Paciente.query.get(internacao.paciente_id)
            if paciente:
                medico = Funcionario.query.get(internacao.medico_id)

                pacientes_data.append({
                    'id': paciente.id,
                    'nome': paciente.nome,
                    'cpf': paciente.cpf,
                    'data_nascimento': paciente.data_nascimento.strftime('%Y-%m-%d') if paciente.data_nascimento else None,
                    'leito': internacao.leito,
                    'data_entrada': internacao.data_internacao.isoformat() if internacao.data_internacao else None,
                    'medico_nome': medico.nome if medico else None,
                    'diagnostico': internacao.diagnostico
                })

        return jsonify({
            'success': True,
            'pacientes': pacientes_data,
            'total': len(pacientes_data)
        })

    except Exception as e:
        logging.error(f"Erro ao buscar pacientes internados: {str(e)}")
        return jsonify({'success': False, 'message': 'Erro interno do servidor'}), 500


@bp.route('/api/pacientes/status/observacao')
@login_required
def api_pacientes_observacao():
    """
    API para obter lista de pacientes em observação (para recepcionistas).
    """
    try:
        current_user = get_current_user()
        if current_user.cargo.lower() != 'recepcionista':
            return jsonify({'success': False, 'message': 'Acesso não autorizado'}), 403

        # Buscar pacientes em observação (internações com caráter "Observação")
        observacoes_ativas = Internacao.query.filter(
            Internacao.carater_internacao.ilike('%observação%'),
            Internacao.data_alta.is_(None)
        ).all()

        pacientes_data = []
        for internacao in observacoes_ativas:
            paciente = Paciente.query.get(internacao.paciente_id)
            if paciente:
                medico = Funcionario.query.get(internacao.medico_id)

                pacientes_data.append({
                    'id': paciente.id,
                    'nome': paciente.nome,
                    'cpf': paciente.cpf,
                    'data_nascimento': paciente.data_nascimento.strftime('%Y-%m-%d') if paciente.data_nascimento else None,
                    'leito': internacao.leito,
                    'data_entrada': internacao.data_internacao.isoformat() if internacao.data_internacao else None,
                    'medico_nome': medico.nome if medico else None,
                    'diagnostico': internacao.diagnostico
                })

        return jsonify({
            'success': True,
            'pacientes': pacientes_data,
            'total': len(pacientes_data)
        })

    except Exception as e:
        logging.error(f"Erro ao buscar pacientes em observação: {str(e)}")
        return jsonify({'success': False, 'message': 'Erro interno do servidor'}), 500


@bp.route('/api/pacientes/status/aguardando-medico')
@login_required
def api_pacientes_aguardando_medico():
    """
    API para obter lista de pacientes aguardando atendimento médico (para recepcionistas).
    """
    try:
        current_user = get_current_user()
        if current_user.cargo.lower() != 'recepcionista':
            return jsonify({'success': False, 'message': 'Acesso não autorizado'}), 403

        atendimentos = Atendimento.query.filter(
            db.or_(
                Atendimento.status.ilike('%aguardando%medico%'),
                Atendimento.status.ilike('%aguardando%médico%')
            )
        ).all()

        pacientes_data = []
        for a in atendimentos:
            paciente = Paciente.query.get(a.paciente_id)
            if not paciente:
                continue
            medico = Funcionario.query.get(a.medico_id) if hasattr(a, 'medico_id') else None

            pacientes_data.append({
                'id': paciente.id,
                'nome': paciente.nome,
                'cpf': paciente.cpf,
                'data_nascimento': paciente.data_nascimento.strftime('%Y-%m-%d') if paciente.data_nascimento else None,
                'data_entrada': a.data_atendimento.isoformat() if a.data_atendimento else None,
                'hora_triagem': a.horario_triagem.strftime('%H:%M') if getattr(a, 'horario_triagem', None) else None,
                'classificacao_risco': getattr(a, 'classificacao_risco', None),
                'medico_nome': medico.nome if medico else None,
                'atendimento_id': a.id
            })

        return jsonify({'success': True, 'pacientes': pacientes_data, 'total': len(pacientes_data)})

    except Exception as e:
        logging.error(f"Erro ao buscar pacientes aguardando médico: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro interno do servidor'}), 500


@bp.route('/api/pacientes/status/aguardando-triagem')
@login_required
def api_pacientes_aguardando_triagem():
    """
    API para obter lista de pacientes aguardando triagem (para recepcionistas).
    """
    try:
        current_user = get_current_user()
        if current_user.cargo.lower() != 'recepcionista':
            return jsonify({'success': False, 'message': 'Acesso não autorizado'}), 403

        # Buscar pacientes aguardando triagem (status = 'Aguardando Triagem')
        atendimentos_aguardando = Atendimento.query.filter(
            Atendimento.status.ilike('%aguardando triagem%')
        ).all()

        pacientes_data = []
        for atendimento in atendimentos_aguardando:
            paciente = Paciente.query.get(atendimento.paciente_id)
            if paciente:
                recepcionista = Funcionario.query.get(atendimento.funcionario_id)

                pacientes_data.append({
                    'id': paciente.id,
                    'nome': paciente.nome,
                    'cpf': paciente.cpf,
                    'data_nascimento': paciente.data_nascimento.strftime('%Y-%m-%d') if paciente.data_nascimento else None,
                    'data_entrada': atendimento.data_atendimento.isoformat() if atendimento.data_atendimento else None,
                    'hora_entrada': atendimento.hora_atendimento.strftime('%H:%M') if atendimento.hora_atendimento else None,
                    'recepcionista_nome': recepcionista.nome if recepcionista else None,
                    'atendimento_id': atendimento.id
                })

        return jsonify({
            'success': True,
            'pacientes': pacientes_data,
            'total': len(pacientes_data)
        })

    except Exception as e:
        logging.error(f"Erro ao buscar pacientes aguardando triagem: {str(e)}")
        return jsonify({'success': False, 'message': 'Erro interno do servidor'}), 500


@bp.route('/api/enfermeiro/pacientes-triagem')
@login_required
def api_enfermeiro_pacientes_triagem():
    """
    API para enfermeiros obterem lista de pacientes aguardando triagem.
    """
    try:
        current_user = get_current_user()
        if current_user.cargo.lower() not in ['enfermeiro', 'multi']:
            return jsonify({'success': False, 'message': 'Acesso não autorizado'}), 403

        # Buscar pacientes aguardando triagem (status = 'Aguardando Triagem')
        atendimentos_aguardando = Atendimento.query.filter(
            Atendimento.status.ilike('%aguardando triagem%')
        ).all()

        pacientes_data = []
        for atendimento in atendimentos_aguardando:
            paciente = Paciente.query.get(atendimento.paciente_id)
            if paciente:
                recepcionista = Funcionario.query.get(atendimento.funcionario_id)

                # Calcular idade em anos
                idade = None
                try:
                    if paciente.data_nascimento:
                        hoje = date.today()
                        dn = paciente.data_nascimento
                        idade = hoje.year - dn.year - ((hoje.month, hoje.day) < (dn.month, dn.day))
                except Exception:
                    idade = None

                # Combinar data e hora de criação da ficha
                criado_em = None
                try:
                    if atendimento.data_atendimento and atendimento.hora_atendimento:
                        criado_dt = datetime.combine(atendimento.data_atendimento, atendimento.hora_atendimento)
                        criado_em = criado_dt.isoformat()
                except Exception:
                    criado_em = None

                pacientes_data.append({
                    'id': paciente.id,
                    'nome': paciente.nome,
                    'cpf': paciente.cpf,
                    'sexo': paciente.sexo,
                    'idade': idade,
                    'data_nascimento': paciente.data_nascimento.strftime('%Y-%m-%d') if paciente.data_nascimento else None,
                    'data_entrada': atendimento.data_atendimento.isoformat() if atendimento.data_atendimento else None,
                    'hora_entrada': atendimento.hora_atendimento.strftime('%H:%M') if atendimento.hora_atendimento else None,
                    'criado_em': criado_em,
                    'recepcionista_nome': recepcionista.nome if recepcionista else None,
                    'atendimento_id': atendimento.id,
                    'prioridade': paciente.prioridade,
                    'desc_prioridade': paciente.desc_prioridade
                })

        return jsonify({
            'success': True,
            'pacientes': pacientes_data,
            'total': len(pacientes_data)
        })

    except Exception as e:
        logging.error(f"Erro ao buscar pacientes aguardando triagem para enfermeiro: {str(e)}")
        return jsonify({'success': False, 'message': 'Erro interno do servidor'}), 500


@bp.route('/api/enfermeiro/paciente_detalhado')
@login_required
def api_enfermeiro_paciente_detalhado():
    """
    API para enfermeiros obterem dados detalhados de um paciente.
    """
    try:
        current_user = get_current_user()
        if current_user.cargo.lower() not in ['enfermeiro', 'multi']:
            return jsonify({'success': False, 'message': 'Acesso não autorizado'}), 403

        paciente_id = request.args.get('paciente_id')
        if not paciente_id:
            return jsonify({'success': False, 'message': 'ID do paciente não fornecido'}), 400

        paciente = Paciente.query.get(paciente_id)
        if not paciente:
            return jsonify({'success': False, 'message': 'Paciente não encontrado'}), 404

        # Calcular idade se data de nascimento estiver disponível
        idade = None
        if paciente.data_nascimento:
            hoje = date.today()
            idade = hoje.year - paciente.data_nascimento.year - ((hoje.month, hoje.day) < (paciente.data_nascimento.month, paciente.data_nascimento.day))

        paciente_data = {
            'id': paciente.id,
            'nome': paciente.nome,
            'nome_social': paciente.nome_social,
            'cartao_sus': paciente.cartao_sus,
            'data_nascimento': paciente.data_nascimento.strftime('%d/%m/%Y') if paciente.data_nascimento else None,
            'idade': idade,
            'sexo': paciente.sexo,
            'cpf': paciente.cpf,
            'endereco': paciente.endereco,
            'bairro': paciente.bairro,
            'municipio': paciente.municipio,
            'telefone': paciente.telefone,
            'filiacao': paciente.filiacao,
            'alergias': paciente.alergias,
            'prioridade': paciente.prioridade,
            'desc_prioridade': paciente.desc_prioridade
        }

        return jsonify({
            'success': True,
            'paciente': paciente_data
        })

    except Exception as e:
        logging.error(f"Erro ao buscar dados detalhados do paciente: {str(e)}")
        return jsonify({'success': False, 'message': 'Erro interno do servidor'}), 500


@bp.route('/api/enfermeiro/atendimento')
@login_required
def api_enfermeiro_atendimento():
    """
    API para enfermeiros obterem o atendimento de um paciente.
    """
    try:
        current_user = get_current_user()
        if current_user.cargo.lower() not in ['enfermeiro', 'multi']:
            return jsonify({'success': False, 'message': 'Acesso não autorizado'}), 403

        paciente_id = request.args.get('paciente_id')
        if not paciente_id:
            return jsonify({'success': False, 'message': 'ID do paciente não fornecido'}), 400

        # Buscar o atendimento mais recente do paciente com status "Aguardando Triagem"
        atendimento = Atendimento.query.filter(
            Atendimento.paciente_id == paciente_id,
            Atendimento.status.ilike('%aguardando triagem%')
        ).order_by(Atendimento.data_atendimento.desc()).first()

        if not atendimento:
            return jsonify({'success': False, 'message': 'Atendimento não encontrado'}), 404

        atendimento_data = {
            'id': atendimento.id,
            'atendimento_id': atendimento.id,
            'data_atendimento': atendimento.data_atendimento.isoformat() if atendimento.data_atendimento else None,
            'hora_atendimento': atendimento.hora_atendimento.strftime('%H:%M') if atendimento.hora_atendimento else None,
            'status': atendimento.status
        }

        return jsonify({
            'success': True,
            'atendimento_id': atendimento.id,
            'atendimento': atendimento_data
        })

    except Exception as e:
        logging.error(f"Erro ao buscar atendimento do paciente: {str(e)}")
        return jsonify({'success': False, 'message': 'Erro interno do servidor'}), 500


@bp.route('/api/enfermeiro/atendimentos/<string:atendimento_id>/evasao', methods=['PUT'])
@login_required
def api_enfermeiro_marcar_evasao(atendimento_id):
    """
    Permite ao enfermeiro marcar o atendimento como Evasão.
    Efeito: altera o status para "Evasão" e define horário de alta (saindo da fila).
    A ação é permitida apenas enquanto o atendimento estiver aguardando triagem.
    """
    try:
        user = get_current_user()
        if not user or user.cargo.lower() not in ['enfermeiro', 'multi']:
            return jsonify({'success': False, 'message': 'Acesso não autorizado'}), 403

        atendimento = Atendimento.query.get(atendimento_id)
        if not atendimento:
            return jsonify({'success': False, 'message': 'Atendimento não encontrado'}), 404

        # Garantir que está em fase de triagem para o enfermeiro poder marcar evasão
        status_atual = (atendimento.status or '').strip().lower()
        if 'aguardando triagem' not in status_atual:
            return jsonify({'success': False, 'message': 'Ação permitida apenas durante a triagem'}), 400

        atendimento.status = 'Evasão'
        atendimento.conduta_final = 'EVASAO POR CONTA PROPRIA'

        # Registrar horário de alta caso não exista
        if not getattr(atendimento, 'horario_alta', None):
            atendimento.horario_alta = datetime.utcnow() - timedelta(hours=3)

        db.session.commit()

        logging.info(
            f"Atendimento {atendimento_id} marcado como Evasão por {user.nome} (Enfermeiro ID: {user.id})"
        )

        return jsonify({'success': True, 'message': 'Atendimento marcado como Evasão'})

    except Exception as e:
        db.session.rollback()
        logging.error(f"Erro ao marcar evasão: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'Erro interno do servidor'}), 500


@bp.route('/api/enfermeiro/info')
@login_required
def api_enfermeiro_info():
    """
    API para obter informações do enfermeiro logado.
    """
    try:
        current_user = get_current_user()
        if current_user.cargo.lower() not in ['enfermeiro', 'multi']:
            return jsonify({'success': False, 'message': 'Acesso não autorizado'}), 403

        funcionario_data = {
            'id': current_user.id,
            'nome': current_user.nome,
            'numero_profissional': getattr(current_user, 'numero_profissional', None),
            'cargo': current_user.cargo
        }

        return jsonify({
            'success': True,
            'funcionario': funcionario_data
        })

    except Exception as e:
        logging.error(f"Erro ao buscar informações do enfermeiro: {str(e)}")
        return jsonify({'success': False, 'message': 'Erro interno do servidor'}), 500


@bp.route('/ficha_paciente_triagem')
@login_required
def ficha_paciente_triagem():
    """
    Página de triagem de pacientes para enfermeiros.
    """
    try:
        current_user = get_current_user()
        if current_user.cargo.lower() not in ['enfermeiro', 'multi']:
            flash('Acesso restrito a enfermeiros.', 'danger')
            return redirect(url_for('main.index'))

        # Verificar se o paciente_id foi fornecido
        paciente_id = request.args.get('paciente_id')
        if not paciente_id:
            flash('ID do paciente não fornecido.', 'warning')
            return redirect(url_for('main.enfermeiro'))

        # Verificar se o paciente existe
        paciente = Paciente.query.get(paciente_id)
        if not paciente:
            flash('Paciente não encontrado.', 'danger')
            return redirect(url_for('main.enfermeiro'))

        # Verificar se existe um atendimento ativo para este paciente com status "Aguardando Triagem"
        atendimento = Atendimento.query.filter(
            Atendimento.paciente_id == paciente_id,
            Atendimento.status.ilike('%aguardando triagem%')
        ).first()

        if not atendimento:
            flash('Não há atendimento aguardando triagem para este paciente.', 'warning')
            return redirect(url_for('main.enfermeiro'))

        return render_template('ficha_paciente_triagem.html')

    except Exception as e:
        logging.error(f"Erro ao acessar ficha de triagem: {str(e)}")
        flash('Erro interno do servidor.', 'danger')
        return redirect(url_for('main.enfermeiro'))


@bp.route('/clinica/observacao')
@login_required
def clinica_observacao():
    try:
        current_user = get_current_user()
        if current_user.cargo.lower() not in ['medico', 'enfermeiro', 'multi']:
            flash('Acesso restrito a médicos, enfermeiros e profissionais multi.', 'danger')
            return redirect(url_for('main.index'))
        
        # Obter o atendimento_id da URL
        atendimento_id = request.args.get('atendimento_id')
        
        logging.info(f"Acessando clinica_observacao com atendimento_id: {atendimento_id}")
        
        if not atendimento_id:
            flash('ID do atendimento não fornecido.', 'warning')
            return redirect(url_for('main.observacao_lobby'))
        
        # Buscar dados do paciente e internação
        internacao = Internacao.query.filter_by(atendimento_id=atendimento_id).first()
        logging.info(f"Internação encontrada: {internacao}")
        
        if not internacao:
            flash('Internação não encontrada.', 'danger')
            return redirect(url_for('main.observacao_lobby'))
        
        paciente = Paciente.query.get(internacao.paciente_id)
        logging.info(f"Paciente encontrado: {paciente}")
        
        if not paciente:
            flash('Paciente não encontrado.', 'danger')
            return redirect(url_for('main.observacao_lobby'))
        
        # Renderizar o template clinica_observacao para controlar pacientes em observação
        return render_template('clinica_evolucao_paciente_medico.html', 
                            paciente=paciente, 
                            internacao=internacao)
        
    except Exception as e:
        logging.error(f"Erro ao acessar página de observação clínica: {str(e)}")
        logging.error(traceback.format_exc())
        flash('Erro ao acessar a página. Por favor, tente novamente.', 'danger')
        return redirect(url_for('main.observacao_lobby'))


@bp.route('/api/lista-observacao/')
@login_required
def get_lista_observacao():
    try:
        current_user = get_current_user()
        if current_user.cargo.lower() not in ['medico', 'enfermeiro', 'multi']:
            return jsonify({'error': 'Acesso restrito a médicos, enfermeiros e profissionais multi.'}), 403
        
        # Buscar apenas observações SEM conduta_final (pacientes que ainda estão em observação)
        observacoes = db.session.query(ListaObservacao, Paciente, Atendimento).join(
            Paciente, ListaObservacao.id_paciente == Paciente.id
        ).join(
            Atendimento, ListaObservacao.id_atendimento == Atendimento.id
        ).filter(
            db.or_(
                ListaObservacao.conduta_final == None,
                ListaObservacao.conduta_final == '',
                ListaObservacao.conduta_final.is_(None)
            )
        ).all()
        
        resultado = []
        atendimento_ids_incluidos = set()

        for obs, paciente, atendimento in observacoes:
            # Buscar dados do médico de entrada se disponível
            medico_entrada_nome = None
            if obs.medico_entrada:
                medico = Funcionario.query.filter_by(nome=obs.medico_entrada).first()
                medico_entrada_nome = medico.nome if medico else obs.medico_entrada
            
            resultado.append({
                'id': obs.id,
                'id_atendimento': obs.id_atendimento,
                'id_paciente': obs.id_paciente,
                'nome': paciente.nome,
                'cpf': paciente.cpf,
                'cartao_sus': paciente.cartao_sus,
                'medico_entrada': obs.medico_entrada,
                'medico_conduta': obs.medico_conduta,
                'medico_nome': medico_entrada_nome or obs.medico_entrada,
                'data_entrada': obs.data_entrada.strftime('%Y-%m-%d %H:%M:%S') if obs.data_entrada else None,
                'data_saida': obs.data_saida.strftime('%Y-%m-%d %H:%M:%S') if obs.data_saida else None,
                'horario_observacao': atendimento.horario_observacao.strftime('%Y-%m-%d %H:%M:%S') if atendimento.horario_observacao else None,
                'conduta_final': obs.conduta_final,
                'atendimento_id': obs.id_atendimento
            })
            atendimento_ids_incluidos.add(obs.id_atendimento)

        # INCLUIR também todos os atendimentos com status "Em Observação" que ainda não estejam na lista
        atendimentos_obs = db.session.query(Atendimento, Paciente, Funcionario).join(
            Paciente, Atendimento.paciente_id == Paciente.id
        ).outerjoin(
            Funcionario, Atendimento.medico_id == Funcionario.id
        ).filter(
            Atendimento.status == 'Em Observação'
        ).all()

        for atendimento, paciente, medico in atendimentos_obs:
            if atendimento.id in atendimento_ids_incluidos:
                continue

            resultado.append({
                'id': None,
                'id_atendimento': atendimento.id,
                'id_paciente': paciente.id,
                'nome': paciente.nome,
                'cpf': paciente.cpf,
                'cartao_sus': paciente.cartao_sus,
                'medico_entrada': medico.nome if medico else None,
                'medico_conduta': None,
                'medico_nome': (medico.nome if medico else None),
                'data_entrada': None,
                'data_saida': None,
                'horario_observacao': atendimento.horario_observacao.strftime('%Y-%m-%d %H:%M:%S') if atendimento.horario_observacao else None,
                'conduta_final': None,
                'atendimento_id': atendimento.id
            })
        
        return jsonify({'pacientes': resultado}), 200
        
    except Exception as e:
        logging.error(f"Erro ao buscar lista de observações: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'error': 'Erro interno do servidor'}), 500


@bp.route('/api/lista-observacao/atendimento/<int:id_atendimento>')
@login_required
def get_observacao_por_atendimento(id_atendimento):
    try:
        current_user = get_current_user()
        if current_user.cargo.lower() not in ['medico', 'enfermeiro', 'multi']:
            return jsonify({'error': 'Acesso restrito a médicos, enfermeiros e profissionais multi.'}), 403
        
        # Buscar observações por ID do atendimento
        observacoes = ListaObservacao.query.filter_by(id_atendimento=id_atendimento).all()
        
        resultado = []
        for obs in observacoes:
            resultado.append({
                'id': obs.id,
                'id_atendimento': obs.id_atendimento,
                'id_paciente': obs.id_paciente,
                'medico_entrada': obs.medico_entrada,
                'medico_conduta': obs.medico_conduta,
                'data_entrada': obs.data_entrada.strftime('%Y-%m-%d %H:%M:%S') if obs.data_entrada else None,
                'data_saida': obs.data_saida.strftime('%Y-%m-%d %H:%M:%S') if obs.data_saida else None,
                'conduta_final': obs.conduta_final
            })
        
        return jsonify({'observacoes': resultado}), 200
        
    except Exception as e:
        logging.error(f"Erro ao buscar observações por atendimento {id_atendimento}: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'error': 'Erro interno do servidor'}), 500


@bp.route('/observacao/evolucao-paciente-medico/<string:atendimento_id>')
@login_required
def observacao_paciente_medico(atendimento_id):
    try:
        # Verificar se o usuário é médico
        current_user = get_current_user()
        if current_user.cargo.lower() not in ['medico', 'multi']:
            flash('Acesso restrito a médicos.', 'danger')
            return redirect(url_for('main.observacao_lobby'))
        
        # Buscar atendimento primeiro
        atendimento = Atendimento.query.filter_by(id=atendimento_id).first()
        if not atendimento:
            flash('Atendimento não encontrado.', 'danger')
            return redirect(url_for('main.observacao_lobby'))
        
        # Verificar se o atendimento está em observação
        if atendimento.status != 'Em Observação':
            flash('Este atendimento não está em observação.', 'warning')
            return redirect(url_for('main.observacao_lobby'))
        
        # Buscar paciente
        paciente = Paciente.query.get(atendimento.paciente_id)
        if not paciente:
            flash('Paciente não encontrado.', 'danger')
            return redirect(url_for('main.observacao_lobby'))
        
        # Buscar ou criar registro de observação
        observacao = ListaObservacao.query.filter_by(id_atendimento=atendimento_id).first()
        if not observacao:
            # Criar registro de observação automaticamente se não existir
            observacao = ListaObservacao(
                id_atendimento=atendimento_id,
                id_paciente=paciente.id,
                medico_entrada=current_user.nome,
                data_entrada=datetime.now()
            )
            db.session.add(observacao)
            db.session.commit()
            logging.info(f"Registro de observação criado automaticamente para atendimento {atendimento_id}")
        
        # Buscar dados da internação para compatibilidade com o template
        internacao = Internacao.query.filter_by(atendimento_id=atendimento_id).first()
        
        return render_template('clinica_observacao.html', 
                            paciente=paciente, 
                            internacao=internacao,
                            observacao=observacao)
        
    except Exception as e:
        logging.error(f"Erro ao acessar evolução do paciente (médico): {str(e)}")
        logging.error(traceback.format_exc())
        flash('Erro ao acessar a evolução do paciente. Por favor, tente novamente.', 'danger')
        return redirect(url_for('main.observacao_lobby'))

@bp.route('/observacao/evolucao-paciente-enfermeiro/<string:atendimento_id>')
@login_required
def observacao_paciente_enfermeiro(atendimento_id):
    try:
        # Verificar se o usuário é enfermeiro
        current_user = get_current_user()
        if current_user.cargo.lower() not in ['enfermeiro', 'multi']:
            flash('Acesso restrito a enfermeiros.', 'danger')
            return redirect(url_for('main.observacao_lobby'))
        
        # Buscar atendimento primeiro
        atendimento = Atendimento.query.filter_by(id=atendimento_id).first()
        if not atendimento:
            flash('Atendimento não encontrado.', 'danger')
            return redirect(url_for('main.observacao_lobby'))
        
        # Verificar se o atendimento está em observação
        if atendimento.status != 'Em Observação':
            flash('Este atendimento não está em observação.', 'warning')
            return redirect(url_for('main.observacao_lobby'))
        
        # Buscar paciente
        paciente = Paciente.query.get(atendimento.paciente_id)
        if not paciente:
            flash('Paciente não encontrado.', 'danger')
            return redirect(url_for('main.observacao_lobby'))
        
        # Buscar ou criar registro de observação
        observacao = ListaObservacao.query.filter_by(id_atendimento=atendimento_id).first()
        if not observacao:
            # Criar registro de observação automaticamente se não existir
            # Para enfermeiro, usamos o médico do atendimento se disponível
            medico_entrada = None
            if atendimento.medico_id:
                medico = Funcionario.query.get(atendimento.medico_id)
                if medico:
                    medico_entrada = medico.nome
            
            observacao = ListaObservacao(
                id_atendimento=atendimento_id,
                id_paciente=paciente.id,
                medico_entrada=medico_entrada,
                data_entrada=datetime.now()
            )
            db.session.add(observacao)
            db.session.commit()
            logging.info(f"Registro de observação criado automaticamente para atendimento {atendimento_id} por enfermeiro")
        
        # Buscar dados da internação para compatibilidade com o template
        internacao = Internacao.query.filter_by(atendimento_id=atendimento_id).first()
        
        return render_template('clinica_observacao_enfermagem.html', 
                            paciente=paciente, 
                            internacao=internacao,
                            observacao=observacao)
        
    except Exception as e:
        logging.error(f"Erro ao acessar evolução do paciente (enfermeiro): {str(e)}")
        logging.error(traceback.format_exc())
        flash('Erro ao acessar a evolução do paciente. Por favor, tente novamente.', 'danger')
        return redirect(url_for('main.observacao_lobby'))
    

@bp.route('/api/definir-conduta', methods=['POST'])
@login_required
def definir_conduta():
    """
    Define a conduta final para um paciente em observação.
    Se for Alta, Transferido, A pedido ou Óbito: atualiza status e registra conduta_final
    Se for Internar: muda status para Internado e segue protocolo de internação
    
    IMPORTANTE: Esta rota depende que exista um registro de Internacao (atendimento_clinica).
    A rota /api/observacao-paciente foi corrigida para SEMPRE criar esse registro.
    
    MELHORIAS: Processo mais robusto com melhor tratamento de erros e logs detalhados
    """
    try:
        # Verificar se o usuário é médico
        current_user = get_current_user()
        if not current_user:
            logging.error("Usuário não autenticado tentou definir conduta")
            return jsonify({
                'success': False,
                'message': 'Usuário não autenticado'
            }), 401
            
        if current_user.cargo.lower() != 'medico':
            logging.warning(f"Usuário {current_user.nome} (cargo: {current_user.cargo}) tentou definir conduta sem permissão")
            return jsonify({
                'success': False,
                'message': 'Apenas médicos podem definir conduta'
            }), 403
        
        # Validar dados recebidos
        dados = request.get_json()
        if not dados:
            logging.error("Requisição sem dados JSON")
            return jsonify({
                'success': False,
                'message': 'Dados não fornecidos'
            }), 400
        
        atendimento_id = dados.get('atendimento_id')
        conduta_raw = dados.get('conduta', '')
        
        # Log dos dados recebidos
        logging.info(f"🔍 Definir conduta - Dados recebidos: {dados}")
        
        # Normalizar conduta (aceitar variações sem acento/minúsculas)
        conduta_map = {
            'alta': 'Alta',
            'transferido': 'Transferido',
            'transferencia': 'Transferido',
            'transferência': 'Transferido',
            'a pedido': 'A pedido',
            'alta a pedido': 'A pedido',
            'obito': 'Óbito',
            'óbito': 'Óbito',
            'internar': 'Internar',
            'evasao': 'Evasão',
            'evasão': 'Evasão',
        }
        conduta = conduta_map.get(str(conduta_raw).strip().lower(), conduta_raw)
        evolucao_medica_final = dados.get('observacao', '')  # Campo renomeado para evolução médica final
        leito_selecionado = dados.get('leito', '')  # Leito selecionado para internação
        
        # Validações básicas
        if not atendimento_id:
            logging.error("Atendimento ID não fornecido")
            return jsonify({
                'success': False,
                'message': 'ID do atendimento é obrigatório'
            }), 400
            
        if not conduta:
            logging.error("Conduta não fornecida")
            return jsonify({
                'success': False,
                'message': 'Conduta é obrigatória'
            }), 400
        
        # Validar leito se conduta for Internar
        if conduta == 'Internar' and not leito_selecionado:
            logging.error(f"Tentativa de internar sem leito - Atendimento: {atendimento_id}")
            return jsonify({
                'success': False,
                'message': 'Leito é obrigatório para internação'
            }), 400
        
        # Validar conduta
        condutas_validas = ['Alta', 'Transferido', 'A pedido', 'Óbito', 'Internar', 'Evasão']
        if conduta not in condutas_validas:
            logging.error(f"Conduta inválida recebida: {conduta}")
            return jsonify({
                'success': False,
                'message': f'Conduta inválida. Valores aceitos: {", ".join(condutas_validas)}'
            }), 400
        
        logging.info(f"✅ Validações iniciais OK - Atendimento: {atendimento_id}, Conduta: {conduta}")
        
        # Buscar atendimento
        atendimento = Atendimento.query.get(atendimento_id)
        if not atendimento:
            logging.error(f"Atendimento {atendimento_id} não encontrado no banco")
            return jsonify({
                'success': False,
                'message': f'Atendimento {atendimento_id} não encontrado'
            }), 404
        
        logging.info(f"✅ Atendimento encontrado: {atendimento.paciente.nome if atendimento.paciente else 'Sem paciente'}")
        
        # Buscar observação
        observacao = ListaObservacao.query.filter_by(id_atendimento=atendimento_id).first()
        if not observacao:
            logging.error(f"Observação não encontrada para atendimento {atendimento_id}")
            return jsonify({
                'success': False,
                'message': 'Registro de observação não encontrado. Verifique se o paciente está em observação.'
            }), 404
        
        logging.info(f"✅ Observação encontrada")
        
        # Buscar internação (CRÍTICO: deve existir para observações)
        internacao = Internacao.query.filter_by(atendimento_id=atendimento_id).first()
        if not internacao:
            logging.error(f"❌ CRÍTICO: Internação não encontrada para atendimento {atendimento_id}")
            logging.error(f"   Isso indica que a observação não foi criada corretamente")
            logging.error(f"   Tentando criar registro de emergência...")
            
            # Tentar criar um registro de internação de emergência
            try:
                internacao = Internacao(
                    atendimento_id=atendimento_id,
                    paciente_id=atendimento.paciente_id,
                    medico_id=current_user.id,
                    enfermeiro_id=None,
                    hda=observacao.medico_entrada if observacao else 'Registro criado automaticamente',
                    diagnostico_inicial='Observação sem registro de internação - criado automaticamente',
                    data_internacao=now_brasilia(),
                    leito='Observação',
                    carater_internacao='Observação',
                    dieta=None
                )
                db.session.add(internacao)
                db.session.flush()
                logging.info(f"✅ Registro de internação de emergência criado (ID: {internacao.id})")
            except Exception as emergency_error:
                logging.error(f"❌ Falha ao criar registro de emergência: {str(emergency_error)}")
                return jsonify({
                    'success': False,
                    'message': 'Registro de internação não encontrado e não foi possível criar automaticamente. Entre em contato com o suporte.',
                    'details': f'Atendimento {atendimento_id} sem registro de Internacao'
                }), 404
        
        logging.info(f"✅ Internação encontrada (ID: {internacao.id})")
        
        # Data/hora atual
        agora = now_brasilia()
        
        # Criar evolução médica final SEMPRE que houver observação (campo obrigatório no frontend)
        if evolucao_medica_final and evolucao_medica_final.strip():
            try:
                # Salvar a evolução médica completa na tabela EvolucaoAtendimentoClinica
                # Formato: "CONDUTA FINAL: [conduta]\n\n[texto da evolução]"
                texto_evolucao = f"CONDUTA FINAL: {conduta}\n\n{evolucao_medica_final.strip()}"
                
                nova_evolucao = EvolucaoAtendimentoClinica(
                    atendimentos_clinica_id=internacao.id,
                    funcionario_id=current_user.id,
                    data_evolucao=agora,
                    evolucao=texto_evolucao
                )
                db.session.add(nova_evolucao)
                db.session.flush()  # Flush para detectar erros antes do commit final
                logging.info(f"✅ Evolução médica final criada na tabela EvolucaoAtendimentoClinica")
                logging.info(f"   Conteúdo: {texto_evolucao[:100]}...")
            except Exception as e:
                logging.error(f"❌ Erro ao criar evolução médica: {str(e)}")
                logging.error(traceback.format_exc())
                db.session.rollback()
                return jsonify({
                    'success': False,
                    'message': f'Erro ao registrar evolução médica: {str(e)}'
                }), 500
        
        if conduta == 'Internar':
            logging.info(f"🏥 Processando internação no leito: {leito_selecionado}")
            
            # Validar se o leito selecionado existe e está disponível
            try:
                leito = Leito.query.filter_by(nome=leito_selecionado).first()
                if not leito:
                    logging.error(f"Leito {leito_selecionado} não encontrado no banco")
                    return jsonify({
                        'success': False,
                        'message': f'Leito "{leito_selecionado}" não existe no sistema'
                    }), 404
                
                logging.info(f"✅ Leito encontrado: {leito.nome} (Status: {leito.status}, Ocupação: {leito.ocupacao_atual}/{leito.capacidade_maxima})")
                
                # Verificar se o leito está disponível
                if leito.status != 'Disponível':
                    logging.warning(f"Tentativa de usar leito indisponível: {leito_selecionado} (Status: {leito.status})")
                    return jsonify({
                        'success': False,
                        'message': f'Leito "{leito_selecionado}" não está disponível no momento'
                    }), 400
                
                # Verificar capacidade do leito
                internacoes_no_leito = Internacao.query.filter_by(
                    leito=leito_selecionado,
                    data_alta=None
                ).count()
                
                logging.info(f"📊 Internações ativas no leito: {internacoes_no_leito}/{leito.capacidade_maxima}")
                
                if internacoes_no_leito >= leito.capacidade_maxima:
                    logging.warning(f"Leito {leito_selecionado} com capacidade máxima atingida")
                    return jsonify({
                        'success': False,
                        'message': f'Leito "{leito_selecionado}" está com capacidade máxima ({leito.capacidade_maxima} pacientes)'
                    }), 400
                
            except Exception as e:
                logging.error(f"❌ Erro ao validar leito: {str(e)}")
                db.session.rollback()
                return jsonify({
                    'success': False,
                    'message': f'Erro ao validar leito: {str(e)}'
                }), 500
            
            # Conduta: Internar - muda status para Internado e aloca leito
            try:
                atendimento.status = 'Internado'
                # Salvar apenas o tipo de conduta (sem o texto da evolução médica)
                # A evolução médica completa está na tabela EvolucaoAtendimentoClinica
                atendimento.conduta_final = f"INTERNADO NO LEITO {leito_selecionado}"
                
                # Registrar data de internação se ainda não foi registrada
                if not internacao.data_internacao:
                    internacao.data_internacao = agora
                    logging.info(f"📅 Data de internação registrada: {agora}")
                
                # Atualizar leito da internação
                internacao.leito = leito_selecionado
                
                # NÃO definir dieta aqui - ela deve ser NULL durante a internação
                # A dieta só será definida como '1' quando o prontuário for fechado
                
                # Atualizar ocupação do leito
                leito.ocupacao_atual += 1
                if leito.ocupacao_atual >= leito.capacidade_maxima:
                    leito.status = 'Ocupado'
                    logging.info(f"🔒 Leito {leito_selecionado} marcado como Ocupado")
                db.session.add(leito)
                
                logging.info(f"✅ Ocupação do leito atualizada: {leito.ocupacao_atual}/{leito.capacidade_maxima}")
                
            except Exception as e:
                logging.error(f"❌ Erro ao atualizar status de internação: {str(e)}")
                db.session.rollback()
                return jsonify({
                    'success': False,
                    'message': f'Erro ao processar internação: {str(e)}'
                }), 500
            
            # NOVO: Criar registro na ListaInternacao
            try:
                lista_internacao = ListaInternacao(
                    id_atendimento=atendimento_id,
                    id_paciente=atendimento.paciente_id,
                    medico_entrada=current_user.nome,
                    data_entrada=agora,
                    medico_conduta=None,  # Será preenchido quando der alta
                    data_saida=None,      # Será preenchido quando der alta
                    conduta_final=None    # Será preenchido quando der alta
                )
                db.session.add(lista_internacao)
                db.session.flush()
                logging.info(f"✅ Registro em ListaInternacao criado")
            except Exception as e:
                logging.error(f"❌ Erro ao criar registro em ListaInternacao: {str(e)}")
                # Não falhar por causa disso - apenas logar o erro
                logging.error(traceback.format_exc())
            
            # Atualizar observação com conduta final (mantém o registro ativo)
            try:
                observacao.medico_conduta = current_user.nome
                observacao.data_saida = agora
                # Salvar apenas o tipo de conduta (sem o texto da evolução médica)
                # A evolução médica completa está na tabela EvolucaoAtendimentoClinica
                observacao.conduta_final = f"INTERNADO NO LEITO {leito_selecionado}"
                logging.info(f"✅ Observação atualizada")
            except Exception as e:
                logging.error(f"❌ Erro ao atualizar observação: {str(e)}")
                db.session.rollback()
                return jsonify({
                    'success': False,
                    'message': f'Erro ao atualizar registro de observação: {str(e)}'
                }), 500
            
            # Log da ação
            logging.info(f"✅ Paciente {atendimento.paciente.nome} (Atendimento: {atendimento_id}) teve conduta definida como INTERNADO no leito {leito_selecionado} pelo médico {current_user.nome}")
            
        else:
            # Conduta: Alta, Transferido, A pedido, Óbito ou Evasão
            logging.info(f"📋 Processando conduta: {conduta}")
            
            try:
                # Atualizar status do atendimento para refletir a conduta e remover da lista de "Em Observação"
                # Armazenar conduta pendente no campo dieta
                atendimento.status = conduta
                internacao.dieta = f'PENDENTE:{conduta}'
                # Salvar apenas o tipo de conduta (sem o texto da evolução médica)
                # A evolução médica completa está na tabela EvolucaoAtendimentoClinica
                atendimento.conduta_final = conduta.upper()
                logging.info(f"✅ Conduta pendente armazenada: {conduta} (será aplicada ao fechar prontuário)")
                
                # Atualizar observação com conduta final
                observacao.medico_conduta = current_user.nome
                observacao.data_saida = agora
                # Salvar apenas o tipo de conduta (sem o texto da evolução médica)
                # A evolução médica completa está na tabela EvolucaoAtendimentoClinica
                observacao.conduta_final = conduta.upper()
                logging.info(f"✅ Observação atualizada")
                
                # Registrar data de alta na internação
                internacao.data_alta = agora
                logging.info(f"📅 Data de alta registrada: {agora}")
                
            except Exception as e:
                logging.error(f"❌ Erro ao atualizar registros de conduta: {str(e)}")
                db.session.rollback()
                return jsonify({
                    'success': False,
                    'message': f'Erro ao processar conduta: {str(e)}'
                }), 500
            
            # Atualizar ocupação do leito se houver
            if internacao.leito:
                try:
                    leito = Leito.query.filter_by(nome=internacao.leito).first()
                    if leito:
                        leito_anterior = leito.nome
                        if leito.ocupacao_atual > 0:
                            leito.ocupacao_atual -= 1
                        if leito.status == 'Ocupado' and leito.ocupacao_atual < leito.capacidade_maxima:
                            leito.status = 'Disponível'
                        db.session.add(leito)
                        logging.info(f"✅ Leito {leito_anterior} liberado (Nova ocupação: {leito.ocupacao_atual}/{leito.capacidade_maxima})")
                except Exception as e:
                    logging.error(f"❌ Erro ao atualizar ocupação do leito: {str(e)}")
                    # Não falhar por causa disso - apenas logar
                    logging.error(traceback.format_exc())
            
            # Log da ação principal
            logging.info(f"✅ Paciente {atendimento.paciente.nome} (Atendimento: {atendimento_id}) teve conduta definida como {conduta.upper()} pelo médico {current_user.nome}")

            # Registrar fluxo do paciente com status exatamente igual à Conduta Final
            try:
                fluxo = FluxoPaciente(
                    id_atendimento=atendimento.id,
                    id_medico=current_user.id if current_user else None,
                    id_enfermeiro=None,
                    nome_paciente=atendimento.paciente.nome if atendimento.paciente else '',
                    mudanca_status=atendimento.conduta_final,
                    mudanca_hora=agora
                )
                db.session.add(fluxo)
                db.session.flush()
                logging.info(f"✅ FluxoPaciente registrado")
            except Exception as e:
                logging.error(f"❌ Falha ao registrar FluxoPaciente: {str(e)}")
                logging.error(traceback.format_exc())
                # Não falhar por causa disso - apenas logar
        
        # Salvar todas as alterações com tratamento robusto
        try:
            db.session.commit()
            logging.info(f"💾 ✅ COMMIT REALIZADO COM SUCESSO - Conduta: {conduta}")
            
            return jsonify({
                'success': True,
                'message': f'Conduta "{conduta}" definida com sucesso',
                'conduta': conduta,
                'atendimento_id': atendimento_id
            })
            
        except Exception as commit_error:
            db.session.rollback()
            logging.error(f"❌ ERRO CRÍTICO no commit final: {str(commit_error)}")
            logging.error(traceback.format_exc())
            return jsonify({
                'success': False,
                'message': f'Erro ao salvar alterações no banco de dados: {str(commit_error)}',
                'details': 'Todas as alterações foram revertidas. Tente novamente.'
            }), 500
        
    except Exception as e:
        db.session.rollback()
        logging.error(f"❌ ERRO GERAL ao definir conduta: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': 'Erro interno do servidor ao processar conduta',
            'error': str(e),
            'details': 'Verifique os logs do servidor para mais informações'
        }), 500

@bp.route('/api/internar-paciente', methods=['POST'])
@login_required
def internar_paciente():
    """
    Registra um novo paciente e cria uma internação associada.
    """
    try:
        current_user = get_current_user()
        if current_user.cargo.lower() != 'medico':
            return jsonify({
                'success': False,
                'message': 'Apenas médicos podem realizar internações'
            }), 403

        dados = request.get_json()

        # Campos obrigatórios da internação (sempre exigidos)
        campos_internacao_obrigatorios = ['diagnostico_inicial', 'leito']
        for campo in campos_internacao_obrigatorios:
            if campo not in dados or not dados[campo]:
                return jsonify({
                    'success': False,
                    'message': f'Campo obrigatório não informado: {campo}'
                }), 400

        # Tentar usar um atendimento existente se fornecido
        atendimento_id = (dados.get('atendimento_id') or '').strip()
        atendimento = Atendimento.query.get(atendimento_id) if atendimento_id else None

        if atendimento is not None:
            # Usar o paciente já existente do atendimento; não exigir dados de paciente
            paciente = Paciente.query.get(atendimento.paciente_id)
            # Atualizar status/horário de internação se aplicável
            agora = datetime.now(ZoneInfo("America/Sao_Paulo"))
            atendimento.status = 'Internado'
            atendimento.horario_internacao = agora
            db.session.add(atendimento)
        else:
            # Não há atendimento existente: criar paciente e atendimento novos
            campos_paciente_obrigatorios = ['nome', 'cpf', 'data_nascimento', 'sexo']
            for campo in campos_paciente_obrigatorios:
                if campo not in dados or not dados[campo]:
                    return jsonify({
                        'success': False,
                        'message': f'Campo obrigatório não informado: {campo}'
                    }), 400

            paciente_existente = Paciente.query.filter_by(cpf=dados['cpf']).first()

            if paciente_existente:
                paciente = paciente_existente
                # Atualizar alergias do paciente existente se fornecidas
                if dados.get('alergias'):
                    paciente.alergias = dados.get('alergias')
            else:
                try:
                    data_nascimento = datetime.strptime(dados['data_nascimento'], '%Y-%m-%d').date()
                except ValueError:
                    return jsonify({
                        'success': False,
                        'message': 'Formato de data de nascimento inválido. Use YYYY-MM-DD.'
                    }), 400

                # Tratar cartão SUS - pode ser vazio se o checkbox estiver marcado
                cartao_sus = dados.get('cartao_sus')
                if dados.get('sem_cartao_sus', False) or not cartao_sus:
                    cartao_sus = None

                paciente = Paciente(
                    nome=dados['nome'],
                    filiacao=dados.get('filiacao', 'Não informado'),
                    cpf=dados['cpf'],
                    data_nascimento=data_nascimento,
                    sexo=dados['sexo'],
                    telefone=dados.get('telefone', 'Não informado'),
                    endereco=dados.get('endereco', 'Não informado'),
                    municipio=dados.get('municipio', 'Não informado'),
                    bairro=dados.get('bairro', 'Não informado'),
                    cartao_sus=cartao_sus,
                    nome_social=dados.get('nome_social', ''),
                    cor=dados.get('cor', 'Não informada'),
                    alergias=dados.get('alergias', ''),
                    identificado=True
                )
                db.session.add(paciente)
                db.session.flush()

            agora = datetime.now(ZoneInfo("America/Sao_Paulo"))

            # Se cliente forneceu um atendimento_id para criação, validar; senão gerar
            atendimento_id = (dados.get('atendimento_id') or '').strip()
            if atendimento_id:
                if len(atendimento_id) > 8:
                    return jsonify({
                        'success': False,
                        'message': 'ID de atendimento excede o limite máximo de 8 dígitos.'
                    }), 400
                if Atendimento.query.get(atendimento_id):
                    return jsonify({
                        'success': False,
                        'message': 'ID de atendimento já existe. Tente novamente.'
                    }), 400
            else:
                prefixo_data = agora.strftime('%y%m%d')
                numero_unico = str(paciente.id)[-2:].zfill(2)
                atendimento_id = f"{prefixo_data}{numero_unico}"

            # Atualizar alergias do paciente se fornecidas
            if dados.get('alergias'):
                paciente.alergias = dados.get('alergias')

            atendimento = Atendimento(
                id=atendimento_id,
                paciente_id=paciente.id,
                funcionario_id=current_user.id,
                medico_id=current_user.id,
                data_atendimento=date.today(),
                hora_atendimento=time(agora.hour, agora.minute, agora.second),
                status='Internado',
                horario_internacao=agora
            )
            db.session.add(atendimento)
            db.session.flush()

        # Evitar duplicidade de internação para o mesmo atendimento
        existente = Internacao.query.filter_by(atendimento_id=atendimento_id).first()
        if existente:
            # Persistir atualização de status/horário no atendimento
            db.session.commit()
            return jsonify({
                'success': True,
                'message': 'Internação já existente para este atendimento',
                'paciente_id': existente.paciente_id,
                'internacao_id': existente.id,
                'atendimento_id': atendimento_id
            }), 200

        internacao = Internacao(
            atendimento_id=atendimento_id,
            paciente_id=paciente.id,
            medico_id=current_user.id,
            data_internacao=datetime.now(ZoneInfo("America/Sao_Paulo")),
            hda=dados.get('hda', ''),
            justificativa_internacao_sinais_e_sintomas=f"{(dados.get('hda') or '').strip()}\n\n{(dados.get('folha_anamnese') or '').strip()}",
            justificativa_internacao_condicoes="RISCO DE COMPLICAÇÃO",
            justificativa_internacao_principais_resultados_diagnostico="ANMNESE + EXAME FISICO",
            diagnostico_inicial=dados.get('diagnostico_inicial', ''),
            folha_anamnese=dados.get('folha_anamnese', ''),
            conduta=dados.get('conduta_inicial', 'Não informada'),
            carater_internacao=dados.get('carater_internacao', 'Não informado'),
            cid_principal=dados.get('cid_principal', ''),
            cid_10_secundario=dados.get('cid_10_secundario', ''),
            cid_10_causas_associadas=dados.get('cid_10_causas_associadas', ''),
            exames_laboratoriais=dados.get('exames_realizados', ''),
            leito=dados.get('leito', '')
        )

        db.session.add(internacao)
        
        # Criar registro na ListaInternacao
        lista_internacao = ListaInternacao(
            id_atendimento=atendimento_id,
            id_paciente=paciente.id,
            medico_entrada=current_user.nome,
            data_entrada=agora,
            medico_conduta=None,  # Deixar em aberto
            data_saida=None,      # Deixar em aberto
            conduta_final=None    # Deixar em aberto
        )
        db.session.add(lista_internacao)
        
        # Atualizar ocupação do leito
        leito_registrado = Leito.query.filter_by(nome=internacao.leito).first()
        if leito_registrado:
            leito_registrado.ocupacao_atual += 1

            # Se a ocupação chegou no limite, atualiza o status
            if leito_registrado.ocupacao_atual >= leito_registrado.capacidade_maxima:
                leito_registrado.status = 'Ocupado'
            else:
                leito_registrado.status = 'Disponível'

            db.session.add(leito_registrado)

        db.session.commit()

        return jsonify({
            'success': True,
            'message': 'Paciente internado com sucesso',
            'paciente_id': paciente.id,
            'internacao_id': internacao.id,
            'atendimento_id': atendimento_id,
            'lista_internacao_id': lista_internacao.id
        }), 201

    except Exception as e:
        db.session.rollback()
        logging.error(f'Erro ao internar paciente: {str(e)}')
        logging.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': f'Erro interno do servidor: {str(e)}'
        }), 500

@bp.route('/api/rn-internacao', methods=['POST'])
@login_required
def adicionar_rn_internacao():
    """
    Registra um novo RN, cria atendimento, internação e primeira evolução médica
    """
    try:
        current_user = get_current_user()
        if current_user.cargo.lower() != 'medico':
            return jsonify({
                'success': False,
                'message': 'Apenas médicos podem adicionar RN em internação'
            }), 403

        dados = request.get_json()
        
        # Campos obrigatórios do RN
        campos_rn_obrigatorios = ['nome', 'data_nascimento', 'sexo', 'peso_ao_nascer', 'responsavel_id']
        for campo in campos_rn_obrigatorios:
            if campo not in dados or not dados[campo]:
                return jsonify({
                    'success': False,
                    'message': f'Campo obrigatório não informado: {campo}'
                }), 400

        # Campos obrigatórios da internação
        campos_internacao_obrigatorios = ['hda', 'diagnostico_inicial']
        for campo in campos_internacao_obrigatorios:
            if campo not in dados or not dados[campo]:
                return jsonify({
                    'success': False,
                    'message': f'Campo obrigatório não informado: {campo}'
                }), 400

        # Verificar se o responsável existe
        responsavel = Paciente.query.get(dados['responsavel_id'])
        if not responsavel:
            return jsonify({
                'success': False,
                'message': 'Responsável não encontrado'
            }), 404

        # Processar data de nascimento
        try:
            data_nascimento = datetime.strptime(dados['data_nascimento'], '%Y-%m-%d').date()
        except ValueError:
            return jsonify({
                'success': False,
                'message': 'Formato de data de nascimento inválido. Use YYYY-MM-DD.'
            }), 400

        # Gerar CPF temporário para RN se não fornecido
        if not dados.get('cpf'):
            import random
            data_str = datetime.now(ZoneInfo("America/Sao_Paulo")).strftime('%y%m%d')
            random_digits = str(random.randint(0, 9999)).zfill(4)
            dados['cpf'] = f'RN{data_str}{random_digits}'[:14]

        # Verificar se o paciente RN já existe
        paciente_existente = Paciente.query.filter_by(cpf=dados['cpf']).first()

        if paciente_existente:
            paciente = paciente_existente
            # Atualizar alergias do paciente existente se fornecidas
            if dados.get('alergias'):
                paciente.alergias = dados.get('alergias')
        else:
            # Criar novo paciente RN
            paciente = Paciente(
                nome=dados['nome'],
                filiacao=dados.get('filiacao', responsavel.nome),
                cpf=dados['cpf'],
                data_nascimento=data_nascimento,
                sexo=dados['sexo'],
                telefone=dados.get('telefone', responsavel.telefone),
                endereco=dados.get('endereco', responsavel.endereco),
                municipio=dados.get('municipio', responsavel.municipio),
                bairro=dados.get('bairro', responsavel.bairro),
                cartao_sus=dados.get('cartao_sus'),
                nome_social=dados.get('nome_social', ''),
                cor=dados.get('cor', 'Não informada'),
                alergias=dados.get('alergias', ''),
                identificado=True
            )
            db.session.add(paciente)
            db.session.flush()

        agora = datetime.now(ZoneInfo("America/Sao_Paulo"))
        
        # Gerar ID de atendimento para RN (começando com 999)
        def gerar_id_rn():
            import random
            while True:
                sufixo = str(random.randint(0, 99999)).zfill(5)
                novo_id = f"999{sufixo}"
                if not Atendimento.query.get(novo_id):
                    return novo_id

        atendimento_id = gerar_id_rn()

        # Criar atendimento
        atendimento = Atendimento(
            id=atendimento_id,
            paciente_id=paciente.id,
            funcionario_id=current_user.id,
            medico_id=current_user.id,
            data_atendimento=date.today(),
            hora_atendimento=time(agora.hour, agora.minute, agora.second),
            status='Internado',
            horario_internacao=agora,
            peso=dados.get('peso_ao_nascer'),
            altura=dados.get('altura')
        )
        db.session.add(atendimento)
        db.session.flush()

        # Criar entrada na tabela Internacao (atendimento_clinica)
        internacao = Internacao(
            atendimento_id=atendimento_id,
            paciente_id=paciente.id,
            medico_id=current_user.id,
            enfermeiro_id=dados.get('enfermeiro_id'),
            hda=dados.get('hda', ''),
            diagnostico_inicial=dados.get('diagnostico_inicial', ''),
            folha_anamnese=dados.get('exame_fisico', ''),
            cid_principal=dados.get('cid_principal', ''),
            cid_10_secundario=dados.get('cid_secundario', ''),
            data_internacao=agora,
            leito=dados.get('leito', 'Berçário'),
            carater_internacao='Internação RN'
            # dieta NÃO é mais definida aqui - fica NULL até fechar prontuário
        )
        db.session.add(internacao)
        db.session.flush()

        # Criar registro na tabela PacienteRN
        paciente_rn = PacienteRN(
            paciente_id=paciente.id,
            responsavel_id=dados['responsavel_id'],
            data_nascimento=data_nascimento,
            tipo_parto=dados.get('tipo_parto', ''),
            idade_gestacional=dados.get('idade_gestacional', ''),
            peso_ao_nascer=dados.get('peso_ao_nascer'),
            observacoes=dados.get('observacoes_rn', '')
        )
        db.session.add(paciente_rn)
        db.session.flush()

        # Criar primeira evolução médica com dados específicos do RN
        dados_rn_para_evolucao = []
        
        if dados.get('peso_ao_nascer'):
            dados_rn_para_evolucao.append(f"Peso ao nascer: {dados['peso_ao_nascer']}g")
        
        if dados.get('altura'):
            dados_rn_para_evolucao.append(f"Altura: {dados['altura']}cm")
        
        if dados.get('apgar'):
            dados_rn_para_evolucao.append(f"Índice de Apgar: {dados['apgar']}")
        
        if dados.get('tipo_parto'):
            dados_rn_para_evolucao.append(f"Tipo de parto: {dados['tipo_parto']}")
        
        if dados.get('idade_gestacional'):
            dados_rn_para_evolucao.append(f"Idade gestacional: {dados['idade_gestacional']}")

        evolucao_texto = f"Primeira evolução médica - RN\n\n"
        evolucao_texto += f"Dados do nascimento:\n"
        evolucao_texto += "\n".join(dados_rn_para_evolucao)
        
        if dados.get('primeira_evolucao'):
            evolucao_texto += f"\n\nEvolução médica:\n{dados['primeira_evolucao']}"

        primeira_evolucao = EvolucaoAtendimentoClinica(
            atendimentos_clinica_id=internacao.id,
            funcionario_id=current_user.id,
            data_evolucao=agora,
            hda=dados.get('hda', ''),
            evolucao=evolucao_texto,
            conduta=dados.get('conduta', 'Internação em berçário - Acompanhamento médico'),
            parametros=dados.get('parametros', '')
        )
        db.session.add(primeira_evolucao)

        db.session.commit()

        return jsonify({
            'success': True,
            'message': 'RN adicionado à internação com sucesso',
            'paciente_id': paciente.id,
            'paciente_rn_id': paciente_rn.id,
            'atendimento_id': atendimento_id,
            'internacao_id': internacao.id
        }), 201

    except Exception as e:
        db.session.rollback()
        logging.error(f'Erro ao adicionar RN em internação: {str(e)}')
        logging.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': f'Erro interno do servidor: {str(e)}'
        }), 500


@bp.route('/api/atendimentos/<string:atendimento_id>/triagem', methods=['PUT'])
@login_required
def salvar_triagem_normal(atendimento_id):
    """
    Salva a triagem de um paciente não gestante.
    Atualiza os dados da triagem no atendimento e muda status para "Aguardando Avaliação Médico".
    """
    try:
        current_user = get_current_user()
        if current_user.cargo.lower() not in ['enfermeiro', 'multi']:
            return jsonify({'success': False, 'message': 'Acesso não autorizado'}), 403

        # Buscar o atendimento
        atendimento = Atendimento.query.get(atendimento_id)
        if not atendimento:
            return jsonify({'success': False, 'message': 'Atendimento não encontrado'}), 404

        # Verificar se o atendimento está aguardando triagem (robusto a nulos e variações)
        status_atual = (atendimento.status or '').strip().lower()
        if 'aguardando triagem' not in status_atual:
            return jsonify({'success': False, 'message': 'Este atendimento não está aguardando triagem'}), 400

        dados = request.get_json()
        if not dados:
            return jsonify({'success': False, 'message': 'Dados não fornecidos'}), 400

        # Atualizar os dados da triagem no atendimento, convertendo strings vazias em None
        def _to_none_if_empty(value):
            if isinstance(value, str) and value.strip() == '':
                return None
            return value

        atendimento.pressao = _to_none_if_empty(dados.get('pressao'))
        atendimento.pulso = _to_none_if_empty(dados.get('pulso'))
        atendimento.sp02 = _to_none_if_empty(dados.get('sp02'))
        atendimento.temp = _to_none_if_empty(dados.get('temp'))
        atendimento.peso = _to_none_if_empty(dados.get('peso'))
        atendimento.altura = _to_none_if_empty(dados.get('altura'))
        atendimento.fr = _to_none_if_empty(dados.get('fr'))
        atendimento.dx = _to_none_if_empty(dados.get('dx'))
        atendimento.triagem = _to_none_if_empty(dados.get('triagem'))
        atendimento.classificacao_risco = _to_none_if_empty(dados.get('classificacao_risco'))
        atendimento.anamnese_exame_fisico = _to_none_if_empty(dados.get('anamnese_exame_fisico'))
        atendimento.observacao = _to_none_if_empty(dados.get('observacao'))

        # Atualizar alergias do paciente (não do atendimento)
        if atendimento.paciente:
            atendimento.paciente.alergias = _to_none_if_empty(dados.get('alergias'))

        # Definir enfermeiro responsável e horário da triagem
        atendimento.enfermeiro_id = current_user.id
        # Horário do servidor menos 3 horas
        atendimento.horario_triagem = datetime.utcnow() - timedelta(hours=3)

        # Mudar status para aguardando avaliação médica (limite 20 chars na coluna)
        atendimento.status = 'Aguardando Medico'

        db.session.commit()

        logging.info(f'Triagem salva para atendimento {atendimento_id} por enfermeiro {current_user.nome}')

        return jsonify({
            'success': True,
            'message': 'Triagem salva com sucesso',
            'atendimento_id': atendimento_id
        })

    except Exception as e:
        db.session.rollback()
        logging.error(f'Erro ao salvar triagem normal: {str(e)}')
        return jsonify({
            'success': False,
            'message': 'Erro interno do servidor'
        }), 500


@bp.route('/api/atendimentos/<string:atendimento_id>/transformar_gestante', methods=['PUT'])
@login_required
def salvar_triagem_gestante(atendimento_id):
    """
    Salva a triagem de um paciente gestante.
    Atualiza os dados da triagem no atendimento, cria registro em AtendimentosGestante
    e muda status para "Aguardando Avaliação Médico".
    """
    try:
        current_user = get_current_user()
        if current_user.cargo.lower() not in ['enfermeiro', 'multi']:
            return jsonify({'success': False, 'message': 'Acesso não autorizado'}), 403

        # Buscar o atendimento
        atendimento = Atendimento.query.get(atendimento_id)
        if not atendimento:
            return jsonify({'success': False, 'message': 'Atendimento não encontrado'}), 404

        # Verificar se o atendimento está aguardando triagem (robusto a nulos e variações)
        status_atual = (atendimento.status or '').strip().lower()
        if 'aguardando triagem' not in status_atual:
            return jsonify({'success': False, 'message': 'Este atendimento não está aguardando triagem'}), 400

        dados = request.get_json()
        if not dados:
            return jsonify({'success': False, 'message': 'Dados não fornecidos'}), 400

        # Atualizar os dados da triagem no atendimento, convertendo strings vazias em None
        def _to_none_if_empty(value):
            if isinstance(value, str) and value.strip() == '':
                return None
            return value

        atendimento.pressao = _to_none_if_empty(dados.get('pressao'))
        atendimento.pulso = _to_none_if_empty(dados.get('pulso'))
        atendimento.sp02 = _to_none_if_empty(dados.get('sp02'))
        atendimento.temp = _to_none_if_empty(dados.get('temp'))
        atendimento.peso = _to_none_if_empty(dados.get('peso'))
        atendimento.altura = _to_none_if_empty(dados.get('altura'))
        atendimento.fr = _to_none_if_empty(dados.get('fr'))
        atendimento.dx = _to_none_if_empty(dados.get('dx'))
        atendimento.triagem = _to_none_if_empty(dados.get('triagem'))
        atendimento.classificacao_risco = _to_none_if_empty(dados.get('classificacao_risco'))
        atendimento.anamnese_exame_fisico = _to_none_if_empty(dados.get('anamnese_exame_fisico'))
        atendimento.observacao = _to_none_if_empty(dados.get('observacao'))

        # Atualizar alergias do paciente (não do atendimento)
        if atendimento.paciente:
            atendimento.paciente.alergias = _to_none_if_empty(dados.get('alergias'))

        # Definir enfermeiro responsável e horário da triagem
        atendimento.enfermeiro_id = current_user.id
        # Horário do servidor menos 3 horas
        atendimento.horario_triagem = datetime.utcnow() - timedelta(hours=3)

        # Criar registro de gestante
        gestante = AtendimentosGestante(
            id_atendimentos=atendimento_id,
            id_paciente=atendimento.paciente_id,
            semanas=_to_none_if_empty(dados.get('idade_gestacional_semanas')),
            dias=_to_none_if_empty(dados.get('idade_gestacional_dias')),
            altura_uterina=_to_none_if_empty(dados.get('altura_uterina')),
            quantidade_gestacoes=_to_none_if_empty(dados.get('quantidade_de_gestacoes')),
            ultima_menstruacao=_to_none_if_empty(dados.get('ultima_menstruacao')),
            bcf=_to_none_if_empty(dados.get('bcf')),
            data_primeiro_ultrassom=_to_none_if_empty(dados.get('data_primeira_ultrassom')),
            abo_rh=_to_none_if_empty(dados.get('abo_rh'))
        )

        # Adicionar condições clínicas ao atendimento como texto estruturado
        condicoes_clinicas = []
        if dados.get('itu'):
            condicoes_clinicas.append("ITU")
        if dados.get('sheg'):
            condicoes_clinicas.append("SHEG")
        if dados.get('diabetes'):
            condicoes_clinicas.append("Diabetes")
        if dados.get('cardiopata'):
            condicoes_clinicas.append("Cardiopata")
        if dados.get('tromboembolismo'):
            condicoes_clinicas.append("Tromboembolismo")
        if dados.get('hipertensao_arterial'):
            condicoes_clinicas.append("Hipertensão Arterial")
        if dados.get('cirugia'):
            condicoes_clinicas.append("Cirurgia")
        if dados.get('usa_insulina'):
            condicoes_clinicas.append("Usa Insulina")
        if dados.get('anemia'):
            condicoes_clinicas.append("Anemia")
        if dados.get('toxoplasmose'):
            condicoes_clinicas.append("Toxoplasmose")

        # Adicionar DST se informada
        if dados.get('dst'):
            condicoes_clinicas.append(f"DST: {dados.get('dst')}")

        if condicoes_clinicas:
            atendimento.observacao = (atendimento.observacao or '') + f"\n\nCondições Clínicas: {', '.join(condicoes_clinicas)}"

        # Mudar status para aguardando avaliação médica (limite 20 chars na coluna)
        atendimento.status = 'Aguardando Medico'

        db.session.add(gestante)
        db.session.commit()

        logging.info(f'Triagem gestante salva para atendimento {atendimento_id} por enfermeiro {current_user.nome}')

        return jsonify({
            'success': True,
            'message': 'Triagem gestante salva com sucesso',
            'atendimento_id': atendimento_id,
            'novo_id': atendimento_id  # Para compatibilidade com o frontend
        })

    except Exception as e:
        db.session.rollback()
        logging.error(f'Erro ao salvar triagem gestante: {str(e)}')
        return jsonify({
            'success': False,
            'message': 'Erro interno do servidor'
        }), 500


@bp.route('/api/atendimentos/<string:atendimento_id>/assinatura', methods=['PUT'])
@login_required
def assinar_triagem_normal(atendimento_id):
    """
    Assina a triagem normal, confirmando que foi realizada pelo enfermeiro.
    """
    try:
        current_user = get_current_user()
        if current_user.cargo.lower() not in ['enfermeiro', 'multi']:
            return jsonify({'success': False, 'message': 'Acesso não autorizado'}), 403

        # Buscar o atendimento
        atendimento = Atendimento.query.get(atendimento_id)
        if not atendimento:
            return jsonify({'success': False, 'message': 'Atendimento não encontrado'}), 404

        # Verificar se já foi assinado
        if hasattr(atendimento, 'assinatura_enfermeiro_triagem') and atendimento.assinatura_enfermeiro_triagem:
            return jsonify({'success': False, 'message': 'Triagem já foi assinada'}), 400

        # Verificar se tem enfermeiro responsável
        if not atendimento.enfermeiro_id:
            atendimento.enfermeiro_id = current_user.id

        # Adicionar assinatura (usando um campo de texto para registrar quem assinou)
        assinatura_texto = f"Assinado por: {current_user.nome} (Enfermeiro) - {datetime.now(ZoneInfo('America/Sao_Paulo')).strftime('%d/%m/%Y %H:%M')}"

        # Se não tem campo específico, adicionar na observação
        if atendimento.observacao:
            atendimento.observacao += f"\n\n{assinatura_texto}"
        else:
            atendimento.observacao = assinatura_texto

        # Garantir status correto após finalizar triagem
        status_atual = (atendimento.status or '').strip().lower()
        if 'aguardando medico' not in status_atual:
            atendimento.status = 'Aguardando Medico'

        db.session.commit()

        logging.info(f'Triagem normal assinada para atendimento {atendimento_id} por {current_user.nome}')

        return jsonify({
            'success': True,
            'message': 'Triagem assinada com sucesso'
        })

    except Exception as e:
        db.session.rollback()
        logging.error(f'Erro ao assinar triagem normal: {str(e)}')
        return jsonify({
            'success': False,
            'message': 'Erro interno do servidor'
        }), 500


@bp.route('/api/atendimentos_gestantes/<string:atendimento_id>/assinatura', methods=['PUT'])
@login_required
def assinar_triagem_gestante(atendimento_id):
    """
    Assina a triagem gestante, confirmando que foi realizada pelo enfermeiro.
    """
    try:
        current_user = get_current_user()
        if current_user.cargo.lower() not in ['enfermeiro', 'multi']:
            return jsonify({'success': False, 'message': 'Acesso não autorizado'}), 403

        # Buscar o atendimento gestante
        gestante = AtendimentosGestante.query.filter_by(id_atendimentos=atendimento_id).first()
        if not gestante:
            return jsonify({'success': False, 'message': 'Registro de gestante não encontrado'}), 404

        # Buscar o atendimento
        atendimento = Atendimento.query.get(atendimento_id)
        if not atendimento:
            return jsonify({'success': False, 'message': 'Atendimento não encontrado'}), 404

        # Verificar se já foi assinado (usando um campo adicional se existir, ou verificando na observação)
        if hasattr(gestante, 'assinatura_enfermeiro') and gestante.assinatura_enfermeiro:
            return jsonify({'success': False, 'message': 'Triagem gestante já foi assinada'}), 400

        # Verificar se tem enfermeiro responsável
        if not atendimento.enfermeiro_id:
            atendimento.enfermeiro_id = current_user.id

        # Adicionar assinatura
        assinatura_texto = f"Assinado por: {current_user.nome} (Enfermeiro) - {datetime.now(ZoneInfo('America/Sao_Paulo')).strftime('%d/%m/%Y %H:%M')}"

        # Adicionar na observação do atendimento
        if atendimento.observacao:
            atendimento.observacao += f"\n\n{assinatura_texto}"
        else:
            atendimento.observacao = assinatura_texto

        # Garantir status correto após finalizar triagem (gestante)
        status_atual = (atendimento.status or '').strip().lower()
        if 'aguardando medico' not in status_atual:
            atendimento.status = 'Aguardando Medico'

        db.session.commit()

        logging.info(f'Triagem gestante assinada para atendimento {atendimento_id} por {current_user.nome}')

        return jsonify({
            'success': True,
            'message': 'Triagem gestante assinada com sucesso'
        })

    except Exception as e:
        db.session.rollback()
        logging.error(f'Erro ao assinar triagem gestante: {str(e)}')
        return jsonify({
            'success': False,
            'message': 'Erro interno do servidor'
        }), 500


# ============================================
# ROTAS DO CHAMADOR (DISPLAY PARA TVS)
# ============================================

@bp.route('/chamador')
@login_required
def chamador():
    """
    Interface full-screen do chamador para exibição em TVs.
    Mostra o último chamado e lista dos chamados recentes.
    Requer login com cargo 'chamador'.
    """
    try:
        # Verifica se o usuário tem cargo de chamador
        user_id = session.get('user_id')
        funcionario = Funcionario.query.get(user_id)
        
        if not funcionario or funcionario.cargo.lower() != 'chamador':
            flash('Acesso negado. Esta página é exclusiva para o perfil Chamador.', 'danger')
            return redirect(url_for('main.index'))
        
        return render_template('chamador.html')
    except Exception as e:
        logging.error(f'Erro ao carregar página do chamador: {str(e)}')
        return "Erro ao carregar o chamador", 500


@bp.route('/api/chamados/criar', methods=['POST'])
@login_required
def api_criar_chamado():
    """
    API para criar um novo chamado.
    Recebe: id_paciente, local, id_atendimento (opcional)
    """
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({
                'success': False,
                'message': 'Dados não fornecidos'
            }), 400
        
        id_paciente = data.get('id_paciente')
        local = data.get('local', 'Não especificado')
        id_atendimento = data.get('id_atendimento')
        
        if not id_paciente:
            return jsonify({
                'success': False,
                'message': 'ID do paciente é obrigatório'
            }), 400
        
        # Verifica se o paciente existe
        paciente = Paciente.query.get(id_paciente)
        if not paciente:
            return jsonify({
                'success': False,
                'message': 'Paciente não encontrado'
            }), 404
        
        # Obtém o usuário atual (funcionário que está chamando)
        user_id = session.get('user_id')
        funcionario = Funcionario.query.get(user_id)
        
        # Cria o novo chamado
        novo_chamado = Chamado(
            id_paciente=id_paciente,
            id_atendimento=id_atendimento,
            local=local,
            hora=datetime.now(ZoneInfo('America/Sao_Paulo')).time(),
            data=datetime.now(ZoneInfo('America/Sao_Paulo')).date(),
            status='ativo'
        )
        
        # Atribui o profissional responsável baseado no cargo
        if funcionario:
            if funcionario.cargo.lower() == 'enfermeiro':
                novo_chamado.id_enfermeiro = user_id
            elif funcionario.cargo.lower() == 'medico':
                novo_chamado.id_medico = user_id
        
        db.session.add(novo_chamado)
        db.session.commit()
        
        logging.info(f'Chamado criado: ID={novo_chamado.id}, Paciente={paciente.nome}, Local={local}')
        
        return jsonify({
            'success': True,
            'message': 'Chamado criado com sucesso',
            'chamado_id': novo_chamado.id,
            'paciente_nome': paciente.nome
        })
        
    except Exception as e:
        db.session.rollback()
        logging.error(f'Erro ao criar chamado: {str(e)}')
        logging.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': 'Erro ao criar chamado'
        }), 500


@bp.route('/api/chamados/recentes')
def api_chamados_recentes():
    """
    API que retorna os chamados mais recentes para atualização em tempo real.
    Retorna os últimos 10 chamados ordenados por data e hora (mais recente primeiro).
    """
    try:
        # Busca os últimos 10 chamados, ordenados por data e hora (mais recente primeiro)
        chamados = Chamado.query.options(
            joinedload(Chamado.paciente),
            joinedload(Chamado.medico),
            joinedload(Chamado.enfermeiro)
        ).order_by(
            Chamado.data.desc(),
            Chamado.hora.desc()
        ).limit(10).all()
        
        # Formata os dados para JSON
        chamados_data = []
        for chamado in chamados:
            # Formata hora e data
            hora_str = chamado.hora.strftime('%H:%M') if chamado.hora else '--:--'
            data_str = chamado.data.strftime('%d/%m/%Y') if chamado.data else '--/--/----'
            
            # Pega o nome do paciente
            nome_paciente = chamado.paciente.nome if chamado.paciente else 'Paciente não identificado'
            
            # Pega o profissional responsável (médico ou enfermeiro)
            profissional = None
            if chamado.medico:
                profissional = f"Dr(a). {chamado.medico.nome}"
            elif chamado.enfermeiro:
                profissional = f"Enf. {chamado.enfermeiro.nome}"
            
            chamados_data.append({
                'id': chamado.id,
                'paciente': nome_paciente,
                'local': chamado.local,
                'hora': hora_str,
                'data': data_str,
                'status': chamado.status,
                'profissional': profissional
            })
        
        return jsonify({
            'success': True,
            'chamados': chamados_data,
            'total': len(chamados_data)
        })
        
    except Exception as e:
        logging.error(f'Erro ao buscar chamados recentes: {str(e)}')
        logging.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': 'Erro ao buscar chamados',
            'chamados': []
        }), 500